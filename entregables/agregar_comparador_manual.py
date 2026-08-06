# -*- coding: utf-8 -*-
"""Agrega la sección del Comparador de Inversores (Página 4b) al Manual de Usuario v3."""
import copy

from docx import Document
from docx.shared import Pt

ORIGEN = "attached_assets/MANUAL_CALCULADORA_BIPV_v3_agosto2026_(7)_1786026454397.docx"
DESTINO = "entregables/MANUAL_CALCULADORA_BIPV_v3.1_agosto2026.docx"

doc = Document(ORIGEN)

# ── localizar el ancla: el heading "7. Página 5b — Motor Óptico" ──────────────
ancla = None
for p in doc.paragraphs:
    if p.style.name == "Heading 2" and p.text.strip().startswith("7. Página 5b"):
        ancla = p
        break
assert ancla is not None, "No se encontró el heading de la Página 5b"


def antes(texto, estilo=None, negrita=False):
    """Inserta un párrafo antes del ancla y lo devuelve."""
    nuevo = ancla.insert_paragraph_before(texto, style=estilo)
    if negrita and nuevo.runs:
        nuevo.runs[0].bold = True
    return nuevo


def vineta(texto):
    return antes(texto, estilo="List Bullet")


# ══ Sección nueva 6b ══════════════════════════════════════════════════════════
antes("6b. Página 4b — Comparador de Inversores ⚖️  NUEVO", "Heading 2")
antes(
    "Nueva página (agosto 2026) ubicada entre Dimensionamiento y Mismatch. Compara varias "
    "configuraciones de inversor (modelo × unidades) usando la simulación horaria REAL ya "
    "corrida en 📊 Producción, aplicando el recorte de potencia (clipping) que cada límite AC "
    "produciría, y entrega los indicadores financieros a 25 años de cada opción."
)

antes("Prerrequisitos", "Heading 3")
vineta("Panel seleccionado en 📐 Dimensionamiento (con Voc, Vmp, Isc y coeficientes térmicos completos).")
vineta("Simulación horaria corrida en 📊 Producción EN LA MISMA SESIÓN (la serie horaria no se guarda en el proyecto).")
vineta("Si el proyecto usa multi-superficie, la página se bloquea: dimensiona los inversores por superficie desde la Página 9.")
vineta("Si hay corrección por diodos de bypass, la página la aplica automáticamente y lo avisa en pantalla.")

antes("Las 3 secciones de la página", "Heading 3")
vineta(
    "1️⃣ Compatibilidad: evalúa TODO el catálogo contra tu string (Voc en frío ≤ V máx., ventana "
    "MPPT, corriente Isc×1,25 por tracker). La columna \"modo\" indica cómo conectar: normal, o "
    "\"1 string/tracker\" cuando el panel tiene tanta corriente que solo cabe un string por entrada "
    "MPPT (típico con paneles de 700+ W). La columna \"motivo\" explica cada rechazo."
)
vineta(
    "2️⃣ Comparativa: eliges 2–4 modelos; la app calcula sola las unidades necesarias según las "
    "entradas de cada equipo y muestra E_ac con clipping, CAPEX, TIR, VPN, Payback y LCOE. "
    "Exportable a CSV. Incluye el botón \"✅ Adoptar esta configuración\"."
)
vineta(
    "3️⃣ Barrido DC/AC: curva completa de ratio DC/AC (1,0 a 2,2) con el óptimo por LCOE marcado "
    "con ⭐. Te dice cuántos kW AC realmente necesita tu campo solar."
)

antes("Pautas para elegir bien antes de oprimir \"Adoptar\"", "Heading 3")
antes("Aplícalas en este orden — primero descartar, luego comparar, luego desempatar:", negrita=True)
vineta(
    "① Descarta por clipping y ratio: solo considera configuraciones con clipping ≤ 2% "
    "(0,5–2% es la zona sana; 0% suele significar inversor sobredimensionado). El ratio DC/AC "
    "recomendable en Colombia es 1,2–1,4: debajo de 1,1 pagas capacidad AC de más; arriba de "
    "1,5 regalas energía."
)
vineta(
    "② Elige por LCOE, no por TIR: el LCOE (COP/kWh) es cuánto cuesta producir cada kWh en 25 "
    "años — la configuración de menor LCOE es casi siempre la mejor. TIR y VPN son desempates. "
    "Ojo: si un modelo no tiene costo en el catálogo, la app lo avisa y su LCOE NO es comparable "
    "(compáralo solo por E_ac y clipping)."
)
vineta(
    "③ Diferencias pequeñas son empates: menos de ~2% en E_ac o ~3% en LCOE es empate numérico. "
    "Desempata con criterios prácticos: redundancia (2 equipos medianos superan a 1 grande — si "
    "uno falla sigues produciendo), y disponibilidad real en Colombia (repuestos, garantía y "
    "soporte local de la marca)."
)
vineta(
    "④ Revisa el modo de conexión: si dice \"1 str/MPPT\", confirma que las unidades calculadas "
    "no te obliguen a comprar un equipo extra casi vacío solo por falta de entradas."
)
vineta(
    "⑤ Corrige las advertencias ANTES de adoptar: si aparece el aviso de string parcial (módulos "
    "no múltiplo del N en serie), ajusta N o el número de módulos — un string incompleto es "
    "eléctricamente inválido."
)
vineta(
    "⑥ Contrasta con el barrido: la potencia AC total de tu elección debería quedar cerca de los "
    "kW AC del punto ⭐ de la sección 3. Si queda lejos, revisa por qué."
)

antes("Qué pasa al oprimir \"✅ Adoptar esta configuración\"", "Heading 3")
vineta("Se fija el inversor, el número de unidades y el N en serie como configuración oficial del proyecto.")
vineta(
    "Se INVALIDAN automáticamente los resultados guardados de Producción, Bypass, Financiero y CO₂ "
    "(estaban calculados con el inversor anterior) — es la misma invalidación en cadena del resto "
    "de la app."
)
vineta(
    "Paso obligatorio después de adoptar: vuelve a correr 📊 Producción y 💰 Financiero. "
    "Sin eso el proyecto queda sin energía oficial."
)
vineta(
    "El CAPEX y el costo USD/kW del comparador son supuestos editables: cuando tengas cotizaciones "
    "reales de los finalistas, actualízalos y confirma que el ganador sigue ganando."
)

# ══ Entrada en la tabla de contenido ══════════════════════════════════════════
for p in doc.paragraphs[:30]:
    if p.text.strip().startswith("7. Página 5b"):
        toc = p.insert_paragraph_before("6b. Página 4b — Comparador de Inversores ⚖️  NUEVO", style=p.style)
        # copiar formato de runs del vecino si aplica
        break

doc.save(DESTINO)
print("Guardado:", DESTINO)
