# -*- coding: utf-8 -*-
"""Alternativas de configuración eléctrica preliminar — Proyecto agrivoltaico Urabá 220,32 kWp."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

AZUL = RGBColor(0x1B, 0x4F, 0x72); VERDE = RGBColor(0x1E, 0x84, 0x49); GRIS = RGBColor(0x55, 0x55, 0x55)
doc = Document()
doc.styles['Normal'].font.name = 'Calibri'; doc.styles['Normal'].font.size = Pt(10.5)
for s in doc.sections:
    s.top_margin = Cm(1.6); s.bottom_margin = Cm(1.6); s.left_margin = Cm(1.8); s.right_margin = Cm(1.8)

def p(text='', bold=False, size=None, color=None, align=None, after=6, italic=False):
    par = doc.add_paragraph(); par.paragraph_format.space_after = Pt(after)
    if align: par.alignment = align
    r = par.add_run(text); r.bold = bold; r.italic = italic
    if size: r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return par

def h(t):
    par = doc.add_paragraph(); par.paragraph_format.space_before = Pt(12); par.paragraph_format.space_after = Pt(5)
    r = par.add_run(t); r.bold = True; r.font.size = Pt(12.5); r.font.color.rgb = AZUL

def tabla(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, htxt in enumerate(headers):
        c = t.rows[0].cells[j]; c.paragraphs[0].text = ''
        r = c.paragraphs[0].add_run(htxt); r.bold = True; r.font.size = Pt(9)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.rows[i+1].cells[j]; c.paragraphs[0].text = ''
            r = c.paragraphs[0].add_run(str(val)); r.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

p('ALTERNATIVAS DE CONFIGURACIÓN ELÉCTRICA PRELIMINAR', bold=True, size=16, color=AZUL, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
p('Proyecto agrivoltaico Urabá · 306 × JA Solar JAM66D46-720/LB bifacial · 220,32 kWp DC', bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
p('Documento de trabajo para simulación comparativa en la Calculadora BIPV · agosto 2026', size=9.5, color=GRIS, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)

h('0. Datos eléctricos del panel que gobiernan el diseño')
tabla(['Parámetro (STC)', 'Valor', 'Implicación de diseño'], [
    ['Potencia', '720 Wp', '306 paneles = 220,32 kWp DC'],
    ['Voc', '≈ 49,0 V', 'String de 18 → 882 V; de 20 → 980 V; de 22 → 1.078 V'],
    ['Vmp', '≈ 41,0 V', 'String de 18 → Vmp ≈ 738 V (dentro de ventana MPPT típica 500–850 V)'],
    ['Imp / Isc', '≈ 17,6 / 18,6 A', '⚠️ CRITERIO CLAVE: muchos inversores string aceptan máx. 15–16 A por entrada; se necesitan entradas de 20 A o más'],
    ['Tensión máx. sistema', '1.500 V', 'El panel no limita; limita el inversor (1.100 V string / 1.500 V central)'],
    ['Clima Apartadó', 'T. mín. ≈ 22 °C', 'Corrección de Voc por frío casi nula → se pueden usar strings más largos que en clima frío sin riesgo de sobretensión'],
])

h('Alternativa A — Caso base actual: 1 inversor ~200 kW, 17 strings de 18')
tabla(['Aspecto', 'Detalle'], [
    ['Arreglo', '17 strings de 18 módulos (1 string por matriz de 2×9) · 882 V Voc'],
    ['Inversor tipo', 'String grande 185–215 kW AC, 1.100 V (ej. Huawei SUN2000-215KTL-H0, Sungrow SG250HX operado a ~200 kW)'],
    ['Ratio DC/AC', '≈ 1,10'],
    ['Ventajas', 'Menor costo por W; un solo punto de conexión; cableado AC mínimo'],
    ['Riesgos a verificar', 'Corriente por entrada (17,6 A) vs límite del tracker; un solo punto de falla: si sale de servicio, se pierde el 100% de la producción'],
])

h('Alternativa B — Redundancia: 2 inversores de 100–110 kW')
tabla(['Aspecto', 'Detalle'], [
    ['Arreglo', '2 × (8 y 9 strings de 18) · mitades independientes de la granja'],
    ['Inversor tipo', '100–110 kW, 1.100 V, 10 MPPT (ej. Huawei SUN2000-100KTL, Growatt MAX 100KTL3-X LV, Sungrow SG110CX)'],
    ['Ratio DC/AC', '≈ 1,00–1,10'],
    ['Ventajas', 'Si falla uno, se conserva ~50% de la producción; repuestos más comunes en Colombia; mantenimiento sin parar toda la granja'],
    ['Riesgos a verificar', 'Corriente por entrada; costo por W algo mayor que el central'],
])

h('Alternativa C — Granularidad: 4 inversores de 50–60 kW (1 por grupo de ~4 matrices)')
tabla(['Aspecto', 'Detalle'], [
    ['Arreglo', '4 × (4–5 strings de 18) distribuidos por fila de matrices'],
    ['Inversor tipo', '50–60 kW, 1.100 V (ej. Solis S5-GC60K, Deye SUN-50K, Growatt MAX 60KTL3)'],
    ['Ratio DC/AC', '≈ 0,92–1,10 según modelo'],
    ['Ventajas', 'MPPT casi por matriz → menor pérdida si una zona se sombrea o ensucia distinto (cultivo, polvo); falla afecta solo 25%; cableado DC corto (inversor junto a cada grupo)'],
    ['Riesgos a verificar', 'Mayor costo por W y más puntos de conexión AC; verificar 17,6 A por entrada (algunos 50–60 kW aceptan solo 15 A)'],
])

h('Alternativa D — Strings largos: 15 strings de 20 + 1 de 6 → mejor usar 18 strings de 17')
p('Con clima cálido se puede subir la tensión de string para reducir corriente y cableado. Dos variantes enteras con 306 paneles:', after=4)
tabla(['Variante', 'Arreglo', 'Voc string', 'Comentario'], [
    ['D1', '18 strings de 17 módulos', '≈ 833 V', 'Más strings y menos tensión — útil si el inversor tiene muchas entradas de baja corriente'],
    ['D2', '15 strings de 20 + 1 string de 6 (descartar el corto) o ajustar a 300 paneles = 15×20', '≈ 980 V', 'Menos strings, menos combiners y menos cobre; exige revisar Voc en frío vs 1.100 V (en Apartadó sobra margen)'],
    ['D3', '14 strings de 22 → 308 paneles (agregar 2)', '≈ 1.078 V', 'Al límite de 1.100 V: solo viable en clima cálido como Urabá; máximo ahorro de cableado'],
])

h('Alternativa E — Ratio DC/AC: sobredimensionar el campo FV')
p('Con HSP 5,3 y ganancia bifacial, conviene simular cuánta energía se recorta (clipping) a distintos ratios:', after=4)
tabla(['Ratio DC/AC', 'AC instalado para 220,3 kWp', 'Qué esperar'], [
    ['1,00', '≈ 220 kW', 'Cero clipping, mayor CAPEX de inversores'],
    ['1,10 (base)', '≈ 200 kW', 'Clipping mínimo (<1%) — punto de partida típico'],
    ['1,20', '≈ 184 kW', 'Clipping moderado al mediodía; menor costo de inversor por kWh'],
    ['1,30', '≈ 170 kW', 'Solo si la tarifa premia energía en horas hombro; verificar con la simulación horaria'],
])

h('Criterios de decisión para comparar en la Calculadora')
tabla(['#', 'Criterio', 'Dónde se ve en la app'], [
    ['1', 'Compatibilidad eléctrica (ventana MPPT, Voc máx., corriente por entrada ≥ 18 A)', 'Dimensionamiento — emparejamiento panel-inversor'],
    ['2', 'Energía anual E_ac y clipping por ratio DC/AC', 'Producción (simulación horaria 8.760 h)'],
    ['3', 'Costo del inversor y CAPEX total (USD/Wp)', 'Presupuesto'],
    ['4', 'TIR, VPN, Payback y LCOE por alternativa', 'Financiero'],
    ['5', 'Riesgo operativo (puntos de falla, repuestos en Colombia)', 'Criterio cualitativo — anotar en el reporte'],
])

h('Plan de trabajo paso a paso')
tabla(['Paso', 'Acción'], [
    ['1', 'Cargar en el Catálogo de Inversores (PDF) las fichas de los modelos candidatos de cada alternativa'],
    ['2', 'Simular Alternativa A (caso base) y guardar el proyecto como "Urabá — Alt A"'],
    ['3', 'Duplicar el proyecto y simular B, C y D cambiando solo inversor y arreglo de strings'],
    ['4', 'Para la mejor de A–D, barrer el ratio DC/AC (Alternativa E) y medir clipping vs CAPEX'],
    ['5', 'Comparar E_ac, LCOE y TIR de todas las corridas y elegir la configuración preliminar óptima'],
    ['6', 'Con la configuración elegida, actualizar la Ficha Técnica Preliminar y el formulario de estructura'],
])

p()
p('Nota: los modelos de inversor citados son referencias de mercado para orientar la búsqueda de fichas '
  'técnicas; la verificación final de corrientes por entrada, ventana MPPT y Voc en frío la hace la '
  'Calculadora con la ficha PDF real de cada equipo.', size=9, color=GRIS, italic=True)
p('Documento de trabajo — agosto 2026 · Proyecto agrivoltaico Urabá · Innovación Química', size=9, color=GRIS, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.save('entregables/Alternativas_Inversores_Uraba.docx')
print('OK')
