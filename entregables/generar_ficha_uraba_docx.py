# -*- coding: utf-8 -*-
"""Ficha Técnica Preliminar Urabá v2 — versión Word editable, sin solapes."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

AZUL = RGBColor(0x1B, 0x4F, 0x72); VERDE = RGBColor(0x1E, 0x84, 0x49); GRIS = RGBColor(0x55, 0x55, 0x55)
doc = Document()
doc.styles['Normal'].font.name = 'Calibri'; doc.styles['Normal'].font.size = Pt(10)
for s in doc.sections:
    s.top_margin = Cm(1.8); s.bottom_margin = Cm(1.8); s.left_margin = Cm(2.0); s.right_margin = Cm(2.0)

def p(text='', bold=False, size=None, color=None, align=None, after=6, italic=False):
    par = doc.add_paragraph(); par.paragraph_format.space_after = Pt(after)
    if align: par.alignment = align
    r = par.add_run(text); r.bold = bold; r.italic = italic
    if size: r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return par

def h1(t):
    par = doc.add_paragraph(); par.paragraph_format.space_before = Pt(14); par.paragraph_format.space_after = Pt(6)
    r = par.add_run(t); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = AZUL

def h2(t):
    par = doc.add_paragraph(); par.paragraph_format.space_before = Pt(10); par.paragraph_format.space_after = Pt(4)
    r = par.add_run(t); r.bold = True; r.font.size = Pt(11.5); r.font.color.rgb = VERDE

def tabla(rows):
    t = doc.add_table(rows=len(rows), cols=2); t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(rows):
        c0, c1 = t.rows[i].cells
        c0.width = Cm(6.2); c1.width = Cm(10.5)
        c0.paragraphs[0].text = ''; r0 = c0.paragraphs[0].add_run(k); r0.font.size = Pt(9.5)
        c1.paragraphs[0].text = ''; r1 = c1.paragraphs[0].add_run(v); r1.bold = True; r1.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def nota(lineas):
    par = doc.add_paragraph(); par.paragraph_format.space_after = Pt(8)
    r = par.add_run(lineas); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GRIS

# ══ Portada / encabezado ══
p('FICHA TÉCNICA PRELIMINAR', bold=True, size=18, color=AZUL, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
p('Proyecto Agrivoltaico Urabá — 220,32 kWp DC', bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
p('Versión 2.1 · corrige desfase de timezone del TMY y agrega pérdida IAM · validada contra PVsyst · agosto 2026',
  size=9.5, color=GRIS, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)

# ══ 1. Sitio, generador y estructura ══
h1('1. Sitio, generador FV y estructura')
h2('Emplazamiento')
tabla([
    ('Ubicación', 'Apartadó, Urabá antioqueño, Colombia'),
    ('Coordenadas / altitud', '7.884° N, −76.635° O · 30 m s.n.m.'),
    ('Terreno', '3.200 m² (32 m N-S × 100 m E-O)'),
    ('Uso del suelo', 'Agrivoltaico — cultivo bajo y entre paneles'),
    ('Viento / nieve de diseño', '30 m/s (por confirmar NSR-10) · 0 kN/m²'),
])
h2('Generador fotovoltaico')
tabla([
    ('Módulo', 'JA Solar JAM66D46-720/LB · bifacial n-type'),
    ('Potencia / cantidad', '720 Wp × 306 módulos = 220,32 kWp DC'),
    ('Dimensiones / peso', '2384 × 1303 × 33 mm · ~38,5 kg'),
    ('Degradación garantizada', '≤1% año 1 · ≤0,4%/año · ~87,8% a 30 años'),
])
h2('Configuración eléctrica preliminar (optimizada por simulación)')
tabla([
    ('Arreglo DC', '17 strings de 18 módulos (1 por matriz)'),
    ('Tensión de string', 'Voc ≈ 882 V · Vmp ≈ 738 V (límite de equipo 1.100–1.500 V)'),
    ('Corriente por string', 'Imp ≈ 17,6 A → se requieren entradas MPPT ≥18 A, con 1 string por tracker'),
    ('Inversores', '2 × 80–90 kW AC (clase string, 1.500 V)'),
    ('Ratio DC/AC', '1,22–1,38 · clipping simulado ≤0,11%'),
    ('Candidatos verificados (~100 kW)', 'Huawei SUN2000-100KTL-M1 · Sungrow SG110CX · Growatt MAX 100KTL3 LV — penalización <1% frente al óptimo de 80–90 kW'),
])
nota('El barrido horario de ratio DC/AC mostró que el pico real del campo es ~184 kW AC; '
     '2×80–90 kW maximiza la TIR y minimiza el LCOE sin pérdida apreciable de energía. '
     'La redundancia de 2 equipos conserva ~50% de la producción ante la falla de uno.')
h2('Estructura y disposición agrivoltaica')
tabla([
    ('Matrices', '17 matrices de 2×9 módulos apaisados'),
    ('Altura libre bajo panel', '3,0 m (maquinaria y cultivo)'),
    ('Inclinación / orientación', '10° · azimut Sur'),
    ('Factor de ocupación', '30% — corredores de cultivo ~4 m, pasillos de 2,8 m'),
    ('Cimentación', 'Tornillo de tierra + acero ZAM (requiere estudio de suelo)'),
])

doc.add_page_break()

# ══ 2. Producción ══
h1('2. Producción estimada (simulación horaria)')
h2('Método')
p('Simulación horaria de 8.760 horas con año meteorológico típico PVGIS para el punto exacto '
  '(7.884, −76.635), hora local correcta (America/Bogotá): transposición Hay-Davies al plano de 10° '
  'Sur, pérdida por reflexión angular IAM (ASHRAE, vidrio estándar) sobre la componente directa y '
  'difusa, temperatura de celda por modelo Faiman, coeficiente de potencia −0,30%/°C, pérdidas DC '
  'combinadas del 8% (soiling, mismatch, cableado), eficiencia de inversor 98,2% y recorte '
  '(clipping) AC real.')
p('Esta cifra corrige la versión anterior de este documento (≈278.600 kWh/año): el script de '
  'simulación tenía un desfase de 5 horas entre la irradiancia del TMY (en UTC) y la posición solar '
  '(en hora local), y no modelaba la pérdida por reflexión angular (IAM). Ambos se corrigieron y se '
  'validaron contra una corrida real de PVsyst para el mismo proyecto: sin ganancia bifacial, la '
  'cifra corregida (310.037 kWh/año) queda a 1,6% de PVsyst (315.074 kWh/año) — diferencia menor, ya '
  'explicada por pérdidas que PVsyst modela y este motor aún no (nivel de irradiancia, ~0,7%).', italic=True)
h2('Resultados anuales — base monofacial (defendible ante banca)')
tabla([
    ('Energía AC año 1', '≈ 310.000 kWh/año (2×90 kW) / ≈ 309.700 kWh/año (2×80 kW)'),
    ('Yield específico', '≈ 1.407 kWh/kWp·año'),
    ('Performance Ratio (IEC 61724)', '≈ 85,3%'),
    ('Pico AC real del campo', '≈ 184 kW (nunca alcanza los 220,32 kWp nominales)'),
    ('Clipping con 2×90 kW', '0,00% · con 2×80 kW: 0,11%'),
    ('Producción año 25 (con degradación)', '≈ 281.600 kWh/año'),
    ('Energía acumulada en 25 años', '≈ 7,39 GWh'),
])
h2('Sinergia agrivoltaica — ganancia bifacial (validada contra PVsyst)')
tabla([
    ('Suelo libre para cultivo', '≈ 2.250 m² (70% del terreno)'),
    ('Ganancia bifacial validada', '+7,6% real (PVsyst modo bifacial: altura 3,0 m, pitch 6,6 m, GCR≈0,39, albedo 0,20 pasto verde, φ=0,80)'),
    ('Energía AC año 1 con bifacial', '≈ 334.800 kWh/año (2×90 kW) — +24.800 kWh/año sobre la base monofacial'),
])
nota('Nota bifacial: se corrió PVsyst en modo bifacial ("Fixed Tilted Plane, Unlimited Sheds") con la '
     'geometría real del plano de disposición (altura de montaje 3,0 m, pitch entre filas 6,6 m, '
     'GCR≈0,39, albedo 0,20 de pasto verde, factor de bifacialidad del módulo φ=0,80). Resultado: '
     '339.033 kWh/año, ganancia bifacial real +7,6% sobre el caso monofacial de PVsyst (315.074 kWh/año) '
     '— muy cerca del +8% que ya usaba este motor. La calculadora con ese +8% da 334.846 kWh/año para el '
     'mismo caso, a solo -1,2% de PVsyst. La ganancia bifacial queda validada y puede incluirse en el '
     'caso base con confianza.')
nota('Nota: la cifra contractual definitiva debe salir de la Calculadora BIPV (Motor IV con la curva '
     'del módulo + diodos de bypass si hay sombras) una vez cerrado el layout final.')

doc.add_page_break()

# ══ 3. Financieros ══
h1('3. Estimación financiera preliminar')
h2('Supuestos declarados')
tabla([
    ('TRM', 'COP 3.118,24/USD — Banco de la República, TRM oficial vigente (datos.gov.co)'),
    ('Tarifa evitada', '950 COP/kWh (EPM, 100% autoconsumo)'),
    ('Vida útil / degradación', '25 años · 0,4% anual'),
    ('OPEX', '10 USD/kWp·año'),
    ('Precios de inversor', 'Referencia de mercado — pendiente cotización local'),
])
h2('Inversión (sin BOM oficial — rangos de mercado)')
tabla([
    ('Costos duros', '≈ 0,68 USD/Wp — módulos, estructura elevada 3 m, 2 inversores 80–90 kW, BOS y montaje'),
    ('Costos blandos (17%)', 'Ingeniería, trámites UPME/RETIE, interventoría e imprevistos'),
    ('CAPEX central', '≈ USD 177.200 ≈ 0,80 USD/Wp ≈ COP 552,5 millones (a TRM del día)'),
    ('Rango (±16%)', 'USD 148.000 – 205.000'),
])
h2('Indicadores (simulación horaria + flujo de caja a 25 años) — base monofacial')
tabla([
    ('Ahorro año 1', '≈ COP 294,5 millones (310.037 kWh × 950 COP)'),
    ('TIR', '≈ 51,7%'),
    ('VPN (tasa 10%)', '≈ USD 635.200'),
    ('Payback simple', '≈ 1,9 años'),
    ('LCOE', '≈ 0,072 USD/kWh ≈ 225 COP/kWh (vs tarifa de 950)'),
])
nota('Indicadores del caso 2×90 kW AC (CAPEX ≈ USD 177.200), con TRM oficial vigente al generar este '
     'documento (COP 3.118,24/USD) — al ser una tasa de mercado, reconfirmarla antes de una decisión de '
     'cierre. El caso 2×80 kW da resultados casi idénticos con clipping ligeramente mayor. Si se valida '
     'el upside bifacial (+8%), estos indicadores mejoran adicionalmente (ver Informe Final para Evaluador).')
h2('Beneficios Ley 1715/2014 (no incluidos arriba — mejoran los indicadores)')
tabla([
    ('Art. 11', 'Deducción del 50% de la inversión en el impuesto de renta (hasta 15 años)'),
    ('Art. 12 / 13', 'Exclusión de IVA y exención arancelaria de equipos'),
    ('Art. 14', 'Depreciación acelerada hasta 33,3% anual'),
])
nota('Los indicadores asumen 100% de autoconsumo; si parte de la energía se exporta como excedentes '
     '(Res. CREG 174/2021), el retorno se reduce según la tarifa de venta. Pendientes para pasar a '
     'ingeniería: estudio de suelo (tornillo de tierra), confirmación de viento NSR-10 y cotizaciones '
     'reales de inversores de 80–90 kW y del BOM completo.')
p()
p('Documento preliminar para cotización — no constituye ingeniería de detalle · Innovación Química · agosto 2026',
  size=8.5, color=GRIS, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.save('entregables/Ficha_Tecnica_Preliminar_Agrivoltaico_Uraba_v2.docx')
print('OK')
