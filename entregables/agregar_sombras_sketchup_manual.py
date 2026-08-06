# -*- coding: utf-8 -*-
"""Agrega la sección de Sombras desde SketchUp (Página 5a) al Manual v3.1 → v3.2."""
from docx import Document

ORIGEN = "entregables/MANUAL_CALCULADORA_BIPV_v3.1_agosto2026.docx"
DESTINO = "entregables/MANUAL_CALCULADORA_BIPV_v3.2_agosto2026.docx"

doc = Document(ORIGEN)

ancla = None
for p in doc.paragraphs:
    if p.style.name == "Heading 2" and p.text.strip().startswith("8b."):
        ancla = p
        break
assert ancla is not None, "No se encontró el heading 8b (Página 9)"


def antes(texto, estilo=None, negrita=False):
    nuevo = ancla.insert_paragraph_before(texto, style=estilo)
    if negrita and nuevo.runs:
        nuevo.runs[0].bold = True
    return nuevo


def vineta(texto):
    return antes(texto, estilo="List Bullet")


antes("8a. Página 5a — Sombras desde SketchUp 🌳  NUEVO", "Heading 2")
antes(
    "Nueva página (agosto 2026) que calcula el Factor de Sombreado horario automáticamente "
    "a partir de un modelo 3D del sitio hecho en SketchUp. Es la segunda puerta de entrada al "
    "modelo de bypass diodes: la Calculadora de Sombreado web (sección 14) sigue funcionando "
    "igual — las dos rutas conviven y producen el mismo CSV. Esta funcionalidad cierra la "
    "brecha frente a PVsyst en escenas 3D de sombras cercanas, con un modelador mejor."
)

antes("Cómo preparar el modelo en SketchUp", "Heading 3")
vineta("Modela en METROS y con el norte real en el eje verde (Y). Si el modelo quedó girado, la página tiene un campo de corrección de norte (° horario).")
vineta("Incluye SOLO los obstáculos que producen sombra: edificios vecinos, árboles, tanques, la propia edificación si sombrea la fachada. NO incluyas los paneles (se sombrearían a sí mismos).")
vineta("Árboles: modélalos como volúmenes simples (cilindro + esfera). En la página hay un deslizador de transparencia (0,3–0,6 típico de follaje). Si mezclas edificios y árboles, calcula en dos pasadas.")
vineta("Exporta con Archivo → Exportar → Modelo 3D en formato OBJ o STL (también acepta DAE, PLY, GLB).")
vineta("Modelos muy pesados: la página rechaza mallas de más de 300.000 triángulos — borra mobiliario y detalle, deja solo los volúmenes que dan sombra.")

antes("Flujo en la página (3 pasos)", "Heading 3")
vineta("1️⃣ Modelo 3D: sube el archivo, elige las unidades y la corrección de norte. La página muestra triángulos y dimensiones — si mide más de 2 km, las unidades no son metros.")
vineta("2️⃣ Puntos de análisis: una fila de módulos = un punto, con coordenadas (x=Este, y=Norte, z=altura) tomadas del propio SketchUp con la herramienta de medición. La columna Fachada permite filtrar después en la Página 5.")
vineta("3️⃣ Calcular: la app lanza un rayo hacia el sol por cada punto y cada hora del año usando el MISMO TMY del proyecto. Muestra % de horas con sombra, FS medio y una gráfica de verificación (FS a mediodía por mes).")

antes("Requisito obligatorio: el TMY del proyecto", "Heading 3")
antes(
    "La página se bloquea si no has corrido antes ☀️ Recurso Solar. No es un capricho: el TMY "
    "de PVGIS viene en hora UTC y el CSV se alinea con Producción por (mes, día, hora) — sin el "
    "mismo TMY, la sombra quedaría corrida unas 5 horas. Con el TMY cargado, la coincidencia con "
    "📊 Producción es hora a hora, 1:1."
)

antes("Conexión con el resto de la calculadora", "Heading 3")
vineta("Botón «📤 Enviar a la Página 5»: deja el CSV listo en la sesión. Al abrir 🔀 Mismatch aparece el botón «🌳 Usar el CSV generado en Sombras SketchUp» — de ahí en adelante la cadena es la de siempre: bypass → E_ac corregida → Producción → Financiero → Reporte.")
vineta("También puedes descargar el CSV y guardarlo: tiene el mismo formato de la Calculadora web (Mes, Dia, Hora, FS_geometrico, FS, Fachada), con FS_geometrico = sombra física pura (convención 0 = sin sombra, 1 = sombra total — sin riesgo de FS invertido).")
vineta("Si cambias el modelo, las unidades, el norte, los puntos o la transparencia, el resultado anterior se invalida automáticamente — no puedes enviar por error un cálculo viejo.")

antes("Avisos y validaciones automáticas", "Heading 3")
vineta("Punto DENTRO de un edificio: la página lo detecta y avisa (daría sombra total falsa).")
vineta("Punto pegado al obstáculo (menos de ~10 cm): aviso de resultado ambiguo.")
vineta("Modelo sin ninguna sombra sobre los puntos: aviso para revisar unidades, norte y coordenadas.")
vineta("La Página 5 promedia los puntos de cada hora con IGUAL peso: usa un punto por fila de módulos y procura que las filas tengan un número similar de módulos.")

antes("Requisito de instalación (una sola vez en el servidor)", "Heading 3")
antes(
    "La página usa la librería trimesh (ya incluida en requirements.txt). Si el servidor no la "
    "tiene, la propia página muestra el comando: venv/bin/pip install trimesh."
)

# ── Entrada en la tabla de contenido ─────────────────────────────────────────
for p in doc.paragraphs[:40]:
    if p.text.strip().startswith("8b.") or "Vista 3D y Multi-Superficie" in p.text:
        p.insert_paragraph_before("8a. Página 5a — Sombras desde SketchUp 🌳  NUEVO", style=p.style)
        break

doc.save(DESTINO)
print("Guardado:", DESTINO)
