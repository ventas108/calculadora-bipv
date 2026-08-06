# -*- coding: utf-8 -*-
"""Genera la Ficha Técnica de la Calculadora BIPV complementada (agosto 2026)."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AZUL = RGBColor(0x1B, 0x4F, 0x72)
VERDE = RGBColor(0x1E, 0x84, 0x49)
GRIS = RGBColor(0x55, 0x55, 0x55)

doc = Document()
st_n = doc.styles['Normal']
st_n.font.name = 'Calibri'
st_n.font.size = Pt(10.5)
for sec in doc.sections:
    sec.top_margin = Cm(1.8); sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)

def p(text='', bold=False, size=None, color=None, align=None, space_after=6, italic=False):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(space_after)
    if align: par.alignment = align
    r = par.add_run(text)
    r.bold = bold; r.italic = italic
    if size: r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return par

def h(num, text):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(14); par.paragraph_format.space_after = Pt(6)
    r = par.add_run(f"{num}. {text}" if num else text)
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = AZUL

def sub(text):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(8); par.paragraph_format.space_after = Pt(4)
    r = par.add_run(text); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = VERDE

def tabla(headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, htxt in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.paragraphs[0].text = ''
        r = cell.paragraphs[0].add_run(htxt); r.bold = True; r.font.size = Pt(9.5)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.paragraphs[0].text = ''
            r = cell.paragraphs[0].add_run(str(val)); r.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# ══════════ ENCABEZADO ══════════
p('FICHA TÉCNICA', bold=True, size=22, color=AZUL, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
p('Calculadora BIPV — Plataforma de simulación fotovoltaica integrada en edificios y agrivoltaica',
  bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
p('Innovación Química · calc.innovacionquimica.com.co · Versión agosto 2026',
  size=10, color=GRIS, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
p('¿Qué es? Simula 8 760 horas del año con datos climáticos reales del sitio, resuelve la curva '
  'corriente-voltaje de cada panel, modela la activación de los diodos de bypass bajo sombra parcial '
  'y entrega TIR, VPN, Payback y LCOE en pesos colombianos con la Ley 1715/2014 aplicada. '
  'Cada número del reporte es trazable y auditable — del rayo de sol al flujo de caja.', space_after=10)

# ══════════ 1. TECNOLOGÍAS ══════════
h(1, 'Seis tecnologías de integración, cada una con su física correcta')
p('La Calculadora no usa un modelo genérico: cada tipo de instalación carga automáticamente su densidad '
  'de potencia, su Performance Ratio típico y su inclinación físicamente correcta, con alertas reactivas '
  'si un valor sale del rango recomendado.')
tabla(['Tipo de instalación', 'Densidad (W/m²)', 'PR típico', 'Tilt defecto', 'Modelo térmico'], [
    ['🏢 Fachada BIPV', '80–180', '0.65', '90°', 'Confinado (k=1.3)'],
    ['🏠 Techo inclinado BIPV', '100–200', '0.75', '15°', 'Ventilado (k=1.0)'],
    ['⛱️ Pérgola BIPV', '60–150', '0.70', '10°', 'Semi-ventilado (k=1.1)'],
    ['Marquesina BIPV', '70–160', '0.68', '30°', 'Semi-ventilado (k=1.1)'],
    ['🏚️ Techo plano', '120–220', '0.78', '10°', 'Ventilado (k=1.0)'],
    ['🌱 Granja FV / agrivoltaica', '130–250', '0.80', '15° (≈5°N Colombia)', 'Ventilado + bifacial'],
])
p('🌱 Modo agrivoltaico completo: factor de ocupación configurable (5–100%) que dimensiona paneles y '
  'presupuesto sobre el área útil, sincroniza la separación entre filas (GCR) y muestra en Vista 3D las '
  'filas elevadas sobre el cultivo. Validado con proyectos reales de cultivo bajo paneles en Colombia.')

# ══════════ 2. FLUJO DE TRABAJO (NUEVO) ══════════
h(2, 'Flujo de trabajo — de la idea al reporte bancable en una sola sesión')
p('La plataforma guía el proyecto por módulos encadenados: cada página alimenta a la siguiente y los '
  'resultados aguas abajo siempre reflejan los datos aguas arriba.')
tabla(['Paso', 'Módulo', 'Qué se obtiene'], [
    ['1', '🏠 Proyecto', 'Sitio, tipo de instalación, área, tarifa eléctrica y gestión de proyectos guardados'],
    ['2', '☀️ Recurso Solar', 'TMY/EPW del sitio, transposición al plano (Hay-Davies / infinite_sheds bifacial)'],
    ['3', '🔬 Motor IV', 'Curva corriente-voltaje real del panel seleccionado, hora a hora'],
    ['4', '📐 Dimensionamiento', 'N.º de paneles, strings y emparejamiento eléctrico con el inversor'],
    ['5', '🔀 Mismatch + 🔆 Motor Óptico', 'Bypass diodes con sombras reales · cascada IAM/soiling/térmica BIPV'],
    ['6', '📊 Producción', 'E_ac anual y mensual con diagnóstico de PR real vs esperado'],
    ['7', '💰 Financiero + 💼 Presupuesto', 'TIR, VPN, Payback, LCOE con Ley 1715 · CAPEX/OPEX detallado'],
    ['8', '🗺️ Vista 3D · 🔋 Baterías · 🌿 CO₂ · 📄 Reporte', 'Visualización, balance con almacenamiento, impacto ambiental y PDF ejecutivo'],
])
sub('🔒 Consistencia garantizada — invalidación en cadena')
p('Si el usuario cambia un dato estructural (área útil, tipo de instalación, coordenadas, inclinación, '
  'azimut o albedo), la plataforma detecta el cambio e invalida automáticamente todos los resultados '
  'derivados — POA, producción, bypass, financiero y CO₂ — con un aviso explícito. Es imposible generar '
  'un reporte donde la producción corresponda a una geometría vieja: o los números están sincronizados, '
  'o la herramienta obliga a recalcular.')
sub('📁 Multi-proyecto')
p('Cada proyecto se guarda con nombre propio y puede recuperarse, duplicarse o eliminarse; el consumo '
  'energético se auto-guarda sin intervención del usuario. Un mismo asesor maneja su cartera completa '
  'de clientes desde la misma sesión.')

# ══════════ 3. MOTORES DE CÁLCULO ══════════
h(3, 'Motores de cálculo — la diferencia entre estimar y simular')
sub('☀️ Recurso solar hora a hora (8 760 iteraciones/año)')
p('Archivos climáticos TMY/EPW reales del sitio + transposición Hay-Davies al plano exacto de instalación. '
  'Para bifaciales y agrivoltaica, modelo infinite_sheds de pvlib (estándar NREL) que calcula la luz que '
  'llega a la cara trasera según la separación real entre filas. El reporte declara la ganancia bifacial '
  'anual (%) con altura de montaje y albedo — el banco ve de dónde sale cada kWh extra.')
sub('⚡ Motor IV — curva corriente-voltaje real, no regla de tres')
p('Resuelve el modelo de diodo único del panel en cada condición de irradiancia y temperatura, con Voc, '
  'Isc, Vmp, Imp y coeficientes térmicos de la ficha real del fabricante. Se activa automáticamente con '
  'ficha completa e incluye triple defensa contra el error más común del mercado: el conteo de celdas en '
  'paneles half-cut (verificación física Ns ≈ Voc/0.74).')
sub('🌗 Diodos de bypass — el cálculo que casi nadie hace')
p('Cuando un obstáculo sombrea 2 de 8 módulos de un string, la pérdida real NO es 25%: puede ser 40–60% '
  'de la producción en esas horas, porque los diodos de bypass eliminan la tensión completa de los módulos '
  'sombreados. La Calculadora resuelve el circuito IV del string hora a hora con el factor de sombra '
  'geométrico del sitio (CSV de la Calculadora de Sombreado 3D) y entrega pérdida anual, horas de bypass '
  'activo y tabla mensual. Referencia técnica: Deline et al. 2013 (NREL).')
tabla(['Pérdida bypass anual', 'Diagnóstico', 'Acción sugerida'], [
    ['< 2%', '🟢 Bajo', 'Sombras leves — diseño aprobado'],
    ['2–5%', '🟡 Moderado', 'Considerar redistribución de strings'],
    ['5–10%', '🔴 Alto', 'Evaluar cambio de orientación o layout'],
    ['> 10%', '⛔ Muy alto', 'Rediseño del sombreado — detectado ANTES de construir'],
])
sub('🔬 Motor Óptico BIPV')
p('Cascada completa de pérdidas que solo existen en integración arquitectónica: reflexión por ángulo de '
  'incidencia (IAM), suciedad (soiling) y efecto térmico del confinamiento en fachada (el panel confinado '
  'opera ~30% más caliente que en rack ventilado, k=1.3).')
sub('🏢 Multi-superficie — el edificio completo en una simulación')
p('Combina fachada sur + techo plano + pérgola + marquesina en un solo sistema: POA y producción por cada '
  'orientación, bypass individual por superficie y una E_ac total trazable que alimenta el financiero, '
  'las baterías y el cálculo de CO₂ evitado.')
sub('🗺️ Vista 3D interactiva')
p('Modelo tridimensional del sistema directamente en el navegador: edificio con las superficies BIPV '
  'ubicadas en su orientación real, o granja agrivoltaica con las matrices de paneles elevadas sobre el '
  'terreno, corredores de cultivo, pasillos de mantenimiento y porcentaje de suelo libre calculado a partir '
  'de la geometría real. Si el terreno no aloja los paneles del dimensionamiento, la vista lo advierte — '
  'la coherencia entre diseño eléctrico y espacio físico se verifica visualmente antes de ir a campo.')
p('🎯 Argumento clave: El Performance Ratio no se asume: se obtiene. La cadena óptica → térmica → '
  'eléctrica produce el PR como resultado físico del sitio. En Bogotá o Medellín puede superar el 100% '
  '(altitud y baja temperatura) — y la herramienta lo demuestra, capa por capa.', italic=True)

# ══════════ 4. FINANCIERO ══════════
h(4, 'Análisis financiero bancable — Colombia, no genérico')
p('TIR, VPN, Payback y LCOE calculados sobre flujo de caja en COP con TRM en línea, precio real del '
  'inversor del catálogo y los tres beneficios de la Ley 1715/2014: deducción de renta (Art. 11), '
  'exclusión de IVA sobre equipos (Art. 12) y depreciación acelerada (Art. 14).')
tabla(['Indicador', '🟢 Atractivo', '🟡 Aceptable', '🔴 Revisar'], [
    ['TIR', '> 12%', '8–12%', '< 8%'],
    ['VPN', '> 0 USD', '≈ 0', '< 0 USD'],
    ['Payback simple', '< 10 años', '10–15 años', '> 15 años'],
    ['LCOE', '< tarifa red', '≈ tarifa red', '> tarifa red'],
])
sub('Estimación Rápida — rango de inversión en menos de 2 minutos')
p('Benchmarks del mercado colombiano (julio 2026) por tipo de instalación, escenario económico '
  '(optimista/base/conservador) y factor de zona geográfica (Bogotá ×1.00 hasta Urabá/Chocó ×1.17). '
  'Desglosa CAPEX en equipos (55–65%), construcción (20–28%), costos blandos (8–14%) y contingencias, '
  'más OPEX anual de 5 componentes (O&M, limpieza, fondo de reposición de inversor, monitoreo y seguro). '
  'Precisión declarada honestamente: ±25–35% — perfecta para decidir la factibilidad; las cotizaciones '
  'reales siempre tienen prioridad automática cuando se ingresan.')
tabla(['Tipo', 'CAPEX referencia (USD/Wp)', 'OPEX (USD/kWp·año)'], [
    ['Granja FV en campo', '0.70–1.10', '8–12'],
    ['Techo industrial', '0.90–1.45', '6–12'],
    ['BIPV fachada / pérgola', '1.60–3.20', '10–16'],
])

# ══════════ 5. BATERÍAS (NUEVO como sección propia) ══════════
h(5, 'Almacenamiento — baterías con balance horario, no promedios')
p('El módulo de baterías cruza la generación simulada hora a hora contra el perfil de consumo del cliente '
  'y entrega los indicadores que definen si el almacenamiento se justifica:')
tabla(['Capacidad', 'Detalle'], [
    ['Balance horario 8 760 h', 'Autogeneración, autosuficiencia y excedentes calculados hora a hora, no con promedios mensuales'],
    ['Dimensionamiento recomendado', 'Capacidad de batería sugerida a partir del balance real del sitio'],
    ['Catálogo con extractor PDF', 'Las fichas de baterías se cargan arrastrando el PDF del fabricante, igual que paneles e inversores'],
    ['Compatibilidad batería-inversor', 'Verificación automática de tensiones y química contra el inversor del proyecto, con alertas'],
    ['Integración financiera', 'El balance con baterías alimenta directamente el flujo de caja y el reporte ejecutivo'],
])

# ══════════ 6. IMPACTO CO2 (NUEVO) ══════════
h(6, 'Impacto ambiental — CO₂ evitado con factores oficiales de Colombia')
p('El CO₂ evitado se calcula sobre la E_ac trazable de la simulación (no sobre capacidad instalada) con '
  'dos escenarios declarados: factor promedio del SIN Colombia (XM/UPME, GHG Protocol location-based) y '
  'factor marginal combinado (metodología CDM AMS-I.D). Los resultados se traducen a equivalencias '
  'verificables — árboles plantados (IDEAM) y kilómetros de vehículo a gasolina (FECOC 2022) — y se '
  'contextualizan frente a la meta NDC Colombia 2030 (Ley 1931/2018). Un argumento ambiental defendible '
  'ante banca verde y certificaciones, no un número de folleto.')

# ══════════ 7. CATALOGOS ══════════
h(7, 'Catálogos inteligentes — cero transcripción manual')
p('Paneles, inversores y baterías se cargan arrastrando la ficha técnica PDF del fabricante. El extractor '
  'lee potencias, tensiones MPPT, corrientes por tracker, coeficientes térmicos, celdas y bifacialidad — '
  'incluso de fichas escaneadas (OCR), en español o inglés, con tablas multi-modelo (un PDF, varios equipos).')
tabla(['Capacidad', 'Detalle'], [
    ['Fabricantes de inversores reconocidos', '20+: Growatt, Deye, Solis, Huawei, SMA, Fronius, GoodWe, Sungrow, Victron, SolarEdge, MUST, SolaX, LuxPower, POWEST y más'],
    ['Control de calidad automático', 'Validador físico que rechaza valores imposibles + detección de fallos silenciosos: ningún dato incompleto entra al catálogo sin aviso'],
    ['Verificación permanente', 'Banco de regresión de 84 fichas reales de paneles + 26 casos de inversores que corren antes de cada mejora'],
    ['Emparejamiento panel-inversor', 'Motor automático que verifica ventanas MPPT, ratios DC/AC y compatibilidad eléctrica del arreglo'],
])

# ══════════ 8. CALIDAD Y DIAGNOSTICO (NUEVO) ══════════
h(8, 'Calidad continua — la herramienta se audita a sí misma')
tabla(['Mecanismo', 'Qué garantiza'], [
    ['🔍 Diagnóstico integrado', 'Página dedicada que revisa la coherencia del proyecto completo y señala datos faltantes o inconsistentes antes de generar el reporte'],
    ['📈 Histórico de diagnósticos', 'Cada corrida queda registrada en el servidor: se puede demostrar si el diseño mejora o empeora entre iteraciones'],
    ['🧪 Bancos de regresión', 'Suites automáticas de extractores (110 fichas reales), invalidación en cadena, bypass y financiero que corren antes de cada mejora'],
    ['⚠️ Fallos ruidosos, no silenciosos', 'Filosofía de diseño: ante un dato dudoso la herramienta avisa o se detiene; nunca rellena con un valor inventado'],
])

# ══════════ 9. ENTREGABLES ══════════
h(9, 'Entregables con trazabilidad total')
p('El Reporte PDF ejecutivo incluye hasta 11 secciones seleccionables: proyecto, recurso solar, motor '
  'óptico, producción con diagnóstico PR real vs esperado, bypass diodes con semáforo de impacto, desglose '
  'multi-superficie, financiero, costos con KPIs de bancabilidad (USD/Wp, USD/m², OPEX/CAPEX), balance con '
  'baterías y CO₂ evitado.')
p('🔍 Trazabilidad: Cada reporte declara explícitamente qué energía se usó y por qué — por ejemplo: '
  '"E_ac usada: 88 150 kWh/año (corregida por bypass diodes; pérdida descontada: 2 850 kWh/año)". '
  'El cliente, el banco o la UPME pueden verificar que las cifras son conservadoras, no optimistas.')

# ══════════ 10. ESPECIFICACIONES ══════════
h(10, 'Especificaciones generales')
tabla(['Característica', 'Detalle'], [
    ['Resolución temporal', 'Horaria — 8 760 horas por año meteorológico típico (TMY/EPW)'],
    ['Tipos de instalación', '6: fachada, techo inclinado, techo plano, pérgola, marquesina, granja FV/agrivoltaica'],
    ['Multi-superficie', 'Ilimitadas orientaciones combinadas con bypass individual por superficie'],
    ['Multi-proyecto', 'Proyectos guardados con nombre; carga, duplicado y eliminación desde la app'],
    ['Baterías', 'Balance horario consumo vs generación: autogeneración, autosuficiencia, excedentes y dimensionamiento recomendado'],
    ['Impacto ambiental', 'CO₂ evitado con factor SIN Colombia (XM/UPME) + factor marginal · equivalencias IDEAM · meta NDC 2030'],
    ['Consistencia de datos', 'Invalidación en cadena automática: ningún resultado derivado sobrevive a un cambio de geometría o sitio'],
    ['Moneda y normativa', 'COP con TRM en línea · Ley 1715/2014 (Arts. 11, 12, 14) · RETIE/UPME en costos blandos'],
    ['Base científica', 'pvlib (NREL, estándar de la industria) + modelos propios BIPV + Deline et al. 2013 para bypass'],
    ['Acceso', 'Aplicación web, sin instalación — calc.innovacionquimica.com.co'],
    ['Ecosistema', 'Integrada con la Calculadora de Sombreado 3D (bipv.innovacionquimica.com.co)'],
])

p()
p('👥 Diseñada para: desarrolladores inmobiliarios, arquitectos, EPCs, banca verde y proyectos '
  'agrivoltaicos que necesitan cifras defendibles ante un comité de inversión — no promesas de folleto.',
  bold=True)
p('Ficha técnica — agosto de 2026 · Innovación Química · calc.innovacionquimica.com.co',
  size=9, color=GRIS, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.save('entregables/FICHA_TECNICA_CALCULADORA_BIPV_agosto2026_complementada.docx')
print('OK')
