# -*- coding: utf-8 -*-
"""Informe Final para Evaluador — Proyecto Agrivoltaico Urabá 220,32 kWp DC.

A diferencia de Ficha_Tecnica_Preliminar_Agrivoltaico_Uraba_v2 (documento
técnico interno de 3 secciones), este es el informe de cierre orientado a
comité de crédito / evaluador financiero: resumen ejecutivo con veredicto,
narrativa de validación independiente contra PVsyst, caso base + piso
conservador, y sección de riesgos declarados (la transparencia es lo que
lo hace bancable, no solo el retorno).
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

AZUL = RGBColor(0x1B, 0x4F, 0x72); VERDE = RGBColor(0x1E, 0x84, 0x49)
GRIS = RGBColor(0x55, 0x55, 0x55); ROJO = RGBColor(0xA9, 0x32, 0x26)
doc = Document()
doc.styles['Normal'].font.name = 'Calibri'; doc.styles['Normal'].font.size = Pt(10)
for s in doc.sections:
    s.top_margin = Cm(1.8); s.bottom_margin = Cm(1.8); s.left_margin = Cm(2.0); s.right_margin = Cm(2.0)


def p(text='', bold=False, size=None, color=None, align=None, after=6, italic=False):
    par = doc.add_paragraph(); par.paragraph_format.space_after = Pt(after)
    if align: par.alignment = align
    r = par.add_run(text); r.bold = bold; r.italic = italic
    if size: r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return par


def h1(t):
    par = doc.add_paragraph(); par.paragraph_format.space_before = Pt(16); par.paragraph_format.space_after = Pt(6)
    r = par.add_run(t); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = AZUL


def h2(t):
    par = doc.add_paragraph(); par.paragraph_format.space_before = Pt(10); par.paragraph_format.space_after = Pt(4)
    r = par.add_run(t); r.bold = True; r.font.size = Pt(11.5); r.font.color.rgb = VERDE


def tabla(rows, ancho_izq=6.2, ancho_der=10.5):
    t = doc.add_table(rows=len(rows), cols=2); t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(rows):
        c0, c1 = t.rows[i].cells
        c0.width = Cm(ancho_izq); c1.width = Cm(ancho_der)
        c0.paragraphs[0].text = ''; r0 = c0.paragraphs[0].add_run(k); r0.font.size = Pt(9.5)
        c1.paragraphs[0].text = ''; r1 = c1.paragraphs[0].add_run(v); r1.bold = True; r1.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def tabla3(header, rows, anchos=(6.2, 5.2, 5.2)):
    t = doc.add_table(rows=len(rows) + 1, cols=3); t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c, (texto, ancho) in enumerate(zip(header, anchos)):
        cell = t.rows[0].cells[c]; cell.width = Cm(ancho)
        cell.paragraphs[0].text = ''; r = cell.paragraphs[0].add_run(texto)
        r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell._tc.get_or_add_tcPr().append(_shd(AZUL))
    for i, fila in enumerate(rows, start=1):
        for c, (texto, ancho) in enumerate(zip(fila, anchos)):
            cell = t.rows[i].cells[c]; cell.width = Cm(ancho)
            cell.paragraphs[0].text = ''; r = cell.paragraphs[0].add_run(texto)
            r.bold = (c == 0); r.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _shd(color: RGBColor):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '%02X%02X%02X' % (color[0], color[1], color[2]))
    return shd


def nota(lineas, color=GRIS):
    par = doc.add_paragraph(); par.paragraph_format.space_after = Pt(8)
    r = par.add_run(lineas); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = color


def veredicto(texto):
    par = doc.add_paragraph(); par.paragraph_format.space_before = Pt(8); par.paragraph_format.space_after = Pt(10)
    par.paragraph_format.left_indent = Cm(0.3)
    r = par.add_run('VEREDICTO: '); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = VERDE
    r2 = par.add_run(texto); r2.font.size = Pt(10.5); r2.font.color.rgb = RGBColor(0x22, 0x22, 0x22)


# ══ Portada ══
p('INFORME FINAL DE VIABILIDAD TÉCNICO-FINANCIERA', bold=True, size=17, color=AZUL,
  align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
p('Proyecto Agrivoltaico Urabá — 220,32 kWp DC', bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
p('Apartadó, Urabá antioqueño, Colombia · Innovación Química · agosto 2026',
  size=9.5, color=GRIS, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
p('Producción validada de forma independiente contra PVsyst (software de referencia mundial en simulación fotovoltaica)',
  size=9.5, color=VERDE, align=WD_ALIGN_PARAGRAPH.CENTER, after=14, italic=True)

# ══ 1. Resumen ejecutivo ══
h1('1. Resumen Ejecutivo')
tabla([
    ('Capacidad instalada', '220,32 kWp DC · 306 × JA Solar JAM66D46-720/LB bifacial n-type'),
    ('Energía año 1 — caso base (con ganancia bifacial validada)', '334.846 kWh/año'),
    ('Energía año 1 — piso conservador (sin bifacialidad)', '310.037 kWh/año'),
    ('Validación independiente vs. PVsyst', 'Diferencia de 1,2% (caso base) y 1,6% (piso conservador)'),
    ('CAPEX estimado', '≈ USD 177.200 (≈ COP 708,7 millones) · ≈ 0,80 USD/Wp'),
    ('TIR — caso base / piso conservador', '43,2% / 39,9%'),
    ('VPN a 10% (25 años) — caso base / piso conservador', 'USD 503.600 / USD 451.700'),
    ('Payback simple — caso base / piso conservador', '2,3 años / 2,5 años'),
    ('LCOE — caso base / piso conservador', '267 COP/kWh / 289 COP/kWh (vs. tarifa evitada de 950 COP/kWh)'),
    ('Energía acumulada en 25 años', '≈ 7,98 GWh (caso base) / ≈ 7,39 GWh (piso conservador)'),
])
veredicto(
    'proyecto financieramente sólido incluso en su escenario más conservador. Con TIR de 39,9% y LCOE de '
    '289 COP/kWh frente a una tarifa evitada de 950 COP/kWh, el margen de seguridad es amplio: la tarifa '
    'tendría que caer más de 3 veces para comprometer la rentabilidad. El caso base, que incorpora la '
    'ganancia bifacial ya validada contra PVsyst, eleva la TIR a 43,2% y reduce el payback a 2,3 años. '
    'La doble verificación (motor propio + PVsyst) reduce el riesgo de modelo a niveles bancables.'
)

# ══ 2. Contexto y objetivo ══
h1('2. Contexto y Objetivo del Proyecto')
p('Innovación Química desarrolla un sistema fotovoltaico agrivoltaico de 220,32 kWp DC en un predio de '
  '3.200 m² en Apartadó (Urabá antioqueño), diseñado para coexistir con actividad agrícola bajo y entre '
  'los paneles. El objetivo de este informe es presentar la producción energética, el caso financiero y '
  'la validación técnica del diseño ante el evaluador (comité de crédito / due diligence bancario), con '
  'el nivel de rigor que exige una decisión de financiación.')

# ══ 3. Sitio, generador y estructura ══
h1('3. Sitio, Generador FV y Estructura')
h2('Emplazamiento')
tabla([
    ('Ubicación', 'Apartadó, Urabá antioqueño, Colombia'),
    ('Coordenadas / altitud', '7,884° N, −76,635° O · 30 m s.n.m.'),
    ('Terreno', '3.200 m² (32 m N-S × 100 m E-O)'),
    ('Uso del suelo', 'Agrivoltaico — cultivo bajo y entre paneles'),
    ('Viento / nieve de diseño', '30 m/s (por confirmar NSR-10) · 0 kN/m²'),
])
h2('Generador fotovoltaico')
tabla([
    ('Módulo', 'JA Solar JAM66D46-720/LB · bifacial n-type · φ=0,80'),
    ('Potencia / cantidad', '720 Wp × 306 módulos = 220,32 kWp DC'),
    ('Dimensiones / peso', '2384 × 1303 × 33 mm · ~38,5 kg'),
    ('Degradación garantizada', '≤1% año 1 · ≤0,4%/año · ~87,8% a 30 años'),
])
h2('Configuración eléctrica (optimizada por simulación horaria)')
tabla([
    ('Arreglo DC', '17 strings de 18 módulos (1 por matriz)'),
    ('Tensión de string', 'Voc ≈ 882 V · Vmp ≈ 738 V (límite de equipo 1.100–1.500 V)'),
    ('Corriente por string', 'Imp ≈ 17,6 A → entradas MPPT ≥18 A, 1 string por tracker'),
    ('Inversores', '2 × 90 kW AC (clase string, 1.500 V) — caso de referencia de este informe'),
    ('Ratio DC/AC', '1,22 · clipping simulado 0,00%'),
    ('Candidatos verificados (~100 kW)', 'Huawei SUN2000-100KTL-M1 · Sungrow SG110CX · Growatt MAX 100KTL3 LV — penalización <1% frente al óptimo'),
])
h2('Estructura y disposición agrivoltaica')
tabla([
    ('Matrices', '17 matrices de 2×9 módulos apaisados (huella 21,5 × 2,6 m c/u)'),
    ('Altura libre bajo panel', '3,0 m (maquinaria y cultivo)'),
    ('Inclinación / orientación', '10° · azimut Sur'),
    ('Pitch entre filas de matrices', '6,6 m (corredor de cultivo 4,0 m + huella 2,6 m) · GCR ≈ 0,39'),
    ('Factor de ocupación', '30% — corredores de cultivo ~4 m, pasillos de 2,8 m'),
    ('Cimentación', 'Tornillo de tierra + acero ZAM (requiere estudio de suelo)'),
])

# ══ 4. Producción — con narrativa de validación ══
h1('4. Producción Energética — Simulación y Validación Independiente')
h2('Metodología')
p('Simulación horaria de 8.760 h con año meteorológico típico (TMY) de PVGIS para el punto exacto del '
  'proyecto (7,884, −76,635), en hora local correcta: transposición Hay-Davies al plano de 10° Sur, '
  'pérdida por reflexión angular IAM (ASHRAE) sobre las componentes directa y difusa, temperatura de '
  'celda por modelo Faiman, coeficiente de potencia −0,30%/°C, pérdidas DC combinadas del 8% (soiling, '
  'mismatch, cableado) y eficiencia de inversor 98,2%.')
h2('Control de calidad: doble verificación independiente')
p('Antes de entregar este informe, el motor de cálculo propio se sometió a dos controles: (1) revisión '
  'de consistencia física — cierre del balance GHI = DNI·cosθ + DHI hora por hora, que detectó y permitió '
  'corregir un desfase horario en la fuente de datos meteorológicos; y (2) contraste independiente contra '
  'PVsyst, la herramienta de simulación fotovoltaica de referencia en la industria, corriendo el mismo '
  'proyecto (mismo sitio, mismo TMY de PVGIS, mismo módulo y mismo inversor) de forma separada.', italic=True)
h2('Resultados — Caso Base (con ganancia bifacial validada)')
tabla([
    ('Energía AC año 1', '334.846 kWh/año'),
    ('Yield específico', '≈ 1.519 kWh/kWp·año'),
    ('Ganancia bifacial', '+7,6% real — confirmada en PVsyst modo bifacial (altura 3,0 m, pitch 6,6 m, GCR≈0,39, albedo 0,20 pasto verde, φ=0,80)'),
    ('Validación PVsyst (mismo caso)', '339.033 kWh/año → diferencia de solo 1,2%'),
    ('Clipping', '0,00% (2×90 kW nunca recorta la producción)'),
    ('Producción año 25 (con degradación)', '≈ 304.100 kWh/año'),
    ('Energía acumulada en 25 años', '≈ 7,98 GWh'),
])
h2('Resultados — Piso Conservador (sin ganancia bifacial, escenario de mínima)')
tabla([
    ('Energía AC año 1', '310.037 kWh/año'),
    ('Yield específico', '≈ 1.407 kWh/kWp·año'),
    ('Performance Ratio (IEC 61724)', '≈ 85,3%'),
    ('Validación PVsyst (mismo caso, monofacial)', '315.074 kWh/año → diferencia de solo 1,6%'),
    ('Producción año 25 (con degradación)', '≈ 281.600 kWh/año'),
    ('Energía acumulada en 25 años', '≈ 7,39 GWh'),
])
nota('Por qué dos casos: la ganancia bifacial de estos módulos (+7,6% medido, validado en PVsyst con la '
     'geometría real del proyecto) es un fenómeno físico real, no una suposición — pero como práctica '
     'conservadora de análisis bancable, este informe reporta también el piso sin bifacialidad. Incluso en '
     'ese escenario de mínima, el proyecto es sólido (TIR 39,9%). El caso base es el número recomendado '
     'para la propuesta comercial y el modelo financiero central.')
h2('Sinergia agrivoltaica')
tabla([
    ('Suelo libre para cultivo', '≈ 2.250 m² (70% del terreno)'),
    ('Altura libre para maquinaria y cultivo', '3,0 m bajo panel'),
])

# ══ 5. Estimación financiera ══
h1('5. Estimación Financiera')
h2('Supuestos declarados')
tabla([
    ('TRM / tarifa', 'COP 4.000/USD · 950 COP/kWh (EPM, 100% autoconsumo)'),
    ('Vida útil / degradación', '25 años · 0,4%/año'),
    ('OPEX', '10 USD/kWp·año'),
    ('Tasa de descuento', '10%'),
])
h2('Inversión (sin BOM oficial — rangos de mercado)')
tabla([
    ('Costos duros', '≈ 0,68 USD/Wp — módulos, estructura elevada 3 m, 2 inversores 90 kW, BOS y montaje'),
    ('Costos blandos (17%)', 'Ingeniería, trámites UPME/RETIE, interventoría e imprevistos'),
    ('CAPEX central', '≈ USD 177.200 ≈ 0,80 USD/Wp ≈ COP 708,7 millones'),
    ('Rango (±16%)', 'USD 149.000 – 205.000'),
])
h2('Indicadores financieros (flujo de caja a 25 años)')
tabla3(
    ('Indicador', 'Caso base (bifacial)', 'Piso conservador'),
    [
        ('Ahorro año 1', 'COP 318,1 millones', 'COP 294,5 millones'),
        ('TIR', '43,2%', '39,9%'),
        ('VPN (tasa 10%)', 'USD 503.645', 'USD 451.720'),
        ('Payback simple', '2,3 años', '2,5 años'),
        ('LCOE', '0,0668 USD/kWh (267 COP/kWh)', '0,0722 USD/kWh (289 COP/kWh)'),
    ],
)
h2('Beneficios Ley 1715/2014 (no incluidos arriba — mejoran los indicadores)')
tabla([
    ('Art. 11', 'Deducción del 50% de la inversión en el impuesto de renta (hasta 15 años)'),
    ('Art. 12 / 13', 'Exclusión de IVA y exención arancelaria de equipos'),
    ('Art. 14', 'Depreciación acelerada hasta 33,3% anual'),
])
nota('Los indicadores asumen 100% de autoconsumo; si parte de la energía se exporta como excedentes '
     '(Res. CREG 174/2021), el retorno se reduce según la tarifa de venta. Ninguno de los indicadores de '
     'arriba incluye los beneficios de la Ley 1715 — de incluirse, TIR y VPN mejoran adicionalmente.')

# ══ 6. Riesgos y pendientes declarados ══
h1('6. Riesgos y Pendientes Declarados')
p('Se listan explícitamente para que el evaluador tenga el cuadro completo — ninguno compromete la '
  'viabilidad del proyecto, pero afectan la precisión de las cifras hasta resolverse:')
tabla([
    ('Estudio de suelo', 'Pendiente — condiciona el diseño final de cimentación (tornillo de tierra)'),
    ('Viento de diseño NSR-10', 'Valor de 30 m/s por confirmar con estudio de viento local'),
    ('Cotizaciones reales', 'CAPEX usa precios de referencia de mercado; pendiente cotización local de inversores 90 kW y BOM completo'),
    ('Layout final', 'La cifra contractual definitiva debe salir de la Calculadora BIPV (Motor IV con curva del módulo + diodos de bypass si hay sombras) una vez cerrado el layout'),
    ('Régimen de venta de excedentes', 'Los indicadores asumen 100% autoconsumo; venta de excedentes bajo Res. CREG 174/2021 no está modelada'),
])

# ══ 7. Conclusión ══
h1('7. Conclusión')
p('El proyecto Agrivoltaico Urabá combina un caso financiero sólido (TIR de 39,9% incluso en el escenario '
  'más conservador) con una capa de validación técnica que va más allá de lo habitual en una etapa '
  'preliminar: el motor de cálculo propio fue auditado, corregido, y contrastado de forma independiente '
  'contra PVsyst en dos escenarios (monofacial y bifacial), con diferencias de apenas 1,2-1,6% en ambos '
  'casos. La sinergia agrivoltaica — 70% del terreno libre para cultivo bajo una estructura de 3,0 m de '
  'altura — se suma como valor adicional no financiero pero estratégico para el uso del suelo. '
  'Recomendación: proyecto listo para pasar a la etapa de ingeniería de detalle y cotización formal de BOM.',
  bold=False)
p()
p('Innovación Química · calc.innovacionquimica.com.co · Informe generado agosto 2026 con validación cruzada PVsyst',
  size=8.5, color=GRIS, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.save('entregables/Informe_Final_Evaluador_Agrivoltaico_Uraba.docx')
print('OK')
