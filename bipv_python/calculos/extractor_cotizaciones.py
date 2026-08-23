# -*- coding: utf-8 -*-
"""Extractor genérico de cotizaciones de proveedores (PDF o Word) para
alimentar 💼 Presupuesto — Perfilería y Estructura.

A diferencia de los extractores de fichas técnicas (paneles/inversores, que
usan patrones específicos por marca porque los campos son siempre los
mismos: Voc, Isc, Vmp...), una cotización comercial NO tiene un formato
estándar entre proveedores -- cada uno arma su propio documento, en su
propio idioma, con su propio orden de columnas. Por eso este extractor
combina dos mecanismos deliberadamente genéricos, ninguno atado a un
proveedor puntual:

1. Regex de "etiqueta: valor" sobre el texto plano, usando un diccionario de
   SINÓNIMOS por campo (español + inglés) -- para agregar cobertura de un
   proveedor nuevo basta con sumar términos a `_CAMPOS_TEXTO`, sin tocar la
   lógica de extracción.
2. Lectura de tablas por ENCABEZADO DE COLUMNA (`pdfplumber.extract_tables`
   / tablas de Word): detecta cuál columna es "Capacidad/Cantidad", cuál es
   "Precio unitario", cuál es "Total" -- sin importar el orden ni el idioma
   de las columnas -- y usa el texto de cada FILA (p. ej. "Ocean freight")
   para decidir a qué campo pertenece ese total.

Si con eso queda algún campo sin encontrar, y el servidor tiene una clave de
IA configurada (mismo mecanismo que 🧭 Asistente, ver `calculos.ia_proveedor`),
se intenta un tercer paso: pedirle al modelo el valor CITANDO el fragmento
textual exacto de donde lo sacó. Ese fragmento se verifica programáticamente
contra el texto real del documento -- si el modelo no puede citar un
fragmento que exista literalmente, el campo se descarta como no verificado
(nunca se confía "a ciegas" en la palabra del modelo).

Ningún valor entra a Presupuesto automáticamente: la página siempre muestra
cada campo propuesto junto a su evidencia citada del documento, para que el
usuario confirme visualmente antes de aplicarlo.
"""
from __future__ import annotations

import io
import json
import re
import unicodedata

try:
    import pdfplumber
    _HAS_PDF = True
except ImportError:
    _HAS_PDF = False

try:
    import docx  # python-docx
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False


# ═══════════════════════════ Lectura de archivos ════════════════════════════

def leer_texto_y_tablas_pdf(datos: bytes) -> tuple[str, list]:
    """Devuelve (texto_plano, tablas) de un PDF. `tablas` es una lista de
    tablas; cada tabla es una lista de filas; cada fila, una lista de celdas
    (str, puede ser vacía)."""
    if not _HAS_PDF:
        raise RuntimeError("Falta la librería pdfplumber en el servidor "
                           "(pip install pdfplumber).")
    texto_partes: list[str] = []
    tablas: list = []
    with pdfplumber.open(io.BytesIO(datos)) as pdf:
        for pagina in pdf.pages:
            t = pagina.extract_text() or ""
            if t:
                texto_partes.append(t)
            for tabla in (pagina.extract_tables() or []):
                tablas.append([[(c or "").strip() for c in fila] for fila in tabla])
    return "\n".join(texto_partes), tablas


def leer_texto_y_tablas_docx(datos: bytes) -> tuple[str, list]:
    """Devuelve (texto_plano, tablas) de un archivo Word (.docx)."""
    if not _HAS_DOCX:
        raise RuntimeError("Falta la librería python-docx en el servidor.")
    doc = docx.Document(io.BytesIO(datos))
    texto_partes = [p.text for p in doc.paragraphs if p.text.strip()]
    tablas: list = []
    for tabla in doc.tables:
        filas = [[celda.text.strip() for celda in fila.cells] for fila in tabla.rows]
        tablas.append(filas)
        # Las tablas de Word también aportan texto corrido -- útil para el
        # regex de "etiqueta: valor" si algún campo vive dentro de una celda.
        for fila in filas:
            texto_partes.append(" ".join(fila))
    return "\n".join(texto_partes), tablas


def leer_cotizacion(datos: bytes, nombre_archivo: str) -> tuple[str, list]:
    """Despacha a PDF o Word según la extensión del archivo."""
    ext = (nombre_archivo or "").lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        return leer_texto_y_tablas_pdf(datos)
    if ext == "docx":
        return leer_texto_y_tablas_docx(datos)
    raise ValueError(f"Formato no soportado: .{ext} (usa PDF o Word .docx)")


# ═══════════════════════════ Utilidades de texto ═════════════════════════════

def _sin_acentos(s: str) -> str:
    t = unicodedata.normalize("NFKD", (s or "").lower())
    return "".join(c for c in t if not unicodedata.combining(c))


def _parsear_numero(texto: str) -> float | None:
    """Convierte un número escrito en cualquiera de las dos convenciones de
    separador (1,234.56 formato US / 1.234,56 formato latino) a float.
    Ambigüedad conocida: un único separador con exactamente 3 dígitos después
    (p. ej. "1.234") se interpreta como decimal, no como miles -- caso poco
    común en precios/capacidades, que suelen traer 2 decimales o ninguno."""
    s = re.sub(r"[^\d.,]", "", texto or "")
    if not s:
        return None
    i_coma = s.rfind(",")
    i_punto = s.rfind(".")
    if i_coma == -1 and i_punto == -1:
        try:
            return float(s)
        except ValueError:
            return None
    if i_coma > i_punto:
        entero, dec = s[:i_coma], s[i_coma + 1:]
        entero = entero.replace(".", "").replace(",", "")
    else:
        entero, dec = s[:i_punto], s[i_punto + 1:]
        entero = entero.replace(",", "").replace(".", "")
    s_norm = f"{entero}.{dec}" if dec else (entero or "0")
    try:
        return float(s_norm)
    except ValueError:
        return None


def _detectar_moneda(texto: str) -> str | None:
    t = texto.upper()
    if "US$" in t or "USD" in t:
        return "USD"
    if "COP" in t or "COL$" in t:
        return "COP"
    if "EUR" in t or "€" in texto:
        return "EUR"
    if "CNY" in t or "RMB" in t or "¥" in texto:
        return "CNY"
    if "$" in texto:
        return "USD"  # convención de la app: "$" solo, sin más contexto, se asume USD
    return None


# ═══════════════ Campos por regex de "etiqueta: valor" (texto plano) ════════

# Campos de metadata / totales que casi siempre aparecen como una línea de
# "etiqueta ... valor" al aplanar el documento a texto -- incluso si en el
# PDF original vivían en una tabla angosta de 2 columnas. Extensible: agregar
# un proveedor nuevo es agregar sinónimos aquí, no escribir código nuevo.
_CAMPOS_TEXTO: dict[str, list[str]] = {
    "proveedor": [
        "beneficiary name", "beneficiary's name", "supplier name", "seller",
        "company name", "proveedor", "fabricante", "manufacturer", "vendedor",
    ],
    "numero_cotizacion": [
        "quotation no", "quotation number", "quote no", "cotizacion no",
        "cotización no", "numero de cotizacion", "número de cotización",
        "reference no", "ref no",
    ],
    "fecha_cotizacion": [
        "quotation date", "quote date", "fecha de cotizacion",
        "fecha de cotización", "date", "fecha",
    ],
    "incoterm": [
        "delivery terms", "terminos de entrega", "términos de entrega",
        "incoterm", "condiciones de entrega",
    ],
    "condiciones_pago": [
        "payment terms", "condiciones de pago", "forma de pago",
        "terminos de pago", "términos de pago",
    ],
    "total_cif": [
        "total amount (cif", "total cif", "valor total cif", "monto total cif",
    ],
    "total_fob": [
        "total amount (fob", "total fob", "subtotal fob",
    ],
    "flete": [
        "ocean freight", "flete maritimo", "flete marítimo",
        "freight charges", "shipping cost", "costo de flete",
    ],
}

# Campos NO numéricos (se devuelve el texto tal cual, sin parsear número).
_CAMPOS_TEXTO_NO_NUM = {"proveedor", "numero_cotizacion", "fecha_cotizacion",
                        "incoterm", "condiciones_pago"}

# Sinónimo demasiado genérico ("No.") -- solo se intenta como último recurso
# y excluyendo líneas que claramente son datos bancarios/contacto, no el
# número de la cotización.
_NUMERO_GENERICO = "no."
_EXCLUIR_NUMERO_GENERICO = ("swift", "account", "bank", "iban", "phone",
                            "tel", "fax", "email", "e-mail")


_PAT_NUM_CON_MONEDA = re.compile(r"(?:US\$|USD|COP|EUR|CNY|RMB|\$|¥|€)\s?[\d][\d.,]*")
_PAT_NUM_SIMPLE = re.compile(r"[\d][\d.,]*")


def _primer_numero(candidato: str) -> str | None:
    """Devuelve el primer número monetario reconocible en `candidato`. Un
    número con símbolo/código de moneda pegado (US$8,000.00, USD 6.200,00)
    tiene prioridad sobre cualquier otro dígito suelto de la misma línea
    (p. ej. el "1" de "1*40HQ" en "Ocean freight charges (1*40HQ) US$8,000.00")
    -- así se evita confundir un conteo/código con el monto real."""
    m = _PAT_NUM_CON_MONEDA.search(candidato)
    if m:
        m2 = _PAT_NUM_SIMPLE.search(m.group(0))
        if m2:
            return m2.group(0)
    m3 = _PAT_NUM_SIMPLE.search(candidato)
    return m3.group(0) if m3 else None


def _buscar_valor_texto(texto: str, sinonimos: list[str], es_numero: bool,
                         excluir: tuple[str, ...] = ()) -> dict | None:
    """Busca la primera línea que contenga alguno de los sinónimos (sin
    distinguir mayúsculas/acentos) y extrae el valor que sigue -- un número
    si `es_numero`, o el resto de la línea como texto si no. Si la línea no
    trae nada después de la etiqueta, prueba la línea siguiente (frecuente
    cuando el PDF/Word separa etiqueta y valor en celdas distintas). Devuelve
    {"valor":..., "evidencia": "<línea completa>"} o None."""
    lineas = texto.split("\n")
    lineas_norm = [_sin_acentos(l) for l in lineas]
    for syn in sinonimos:
        syn_n = _sin_acentos(syn)
        for i, ln in enumerate(lineas_norm):
            pos = ln.find(syn_n)
            if pos == -1:
                continue
            if excluir and any(x in ln for x in excluir):
                continue
            resto = lineas[i][pos + len(syn):]
            candidato = resto.lstrip(" :\t-–—")
            evidencia = lineas[i].strip()
            if not candidato.strip() and i + 1 < len(lineas):
                candidato = lineas[i + 1].strip()
                evidencia = f"{lineas[i].strip()} {candidato}".strip()
            if es_numero:
                num_str = _primer_numero(candidato)
                if num_str is None:
                    continue
                val = _parsear_numero(num_str)
                if val is None:
                    continue
                return {"valor": val, "evidencia": evidencia}
            candidato = candidato.strip()
            if candidato:
                return {"valor": candidato, "evidencia": evidencia}
    return None


def _extraer_texto(texto: str) -> dict:
    resultado = {}
    for campo, sinonimos in _CAMPOS_TEXTO.items():
        es_num = campo not in _CAMPOS_TEXTO_NO_NUM
        r = _buscar_valor_texto(texto, sinonimos, es_num)
        if r:
            resultado[campo] = {**r, "metodo": "patron"}
    if "numero_cotizacion" not in resultado:
        r = _buscar_valor_texto(texto, [_NUMERO_GENERICO], es_numero=False,
                                 excluir=_EXCLUIR_NUMERO_GENERICO)
        if r:
            resultado["numero_cotizacion"] = {**r, "metodo": "patron"}
    return resultado


# ═══════════════ Campos por encabezado de columna (tablas) ══════════════════

_ENCABEZADOS_TABLA: dict[str, list[str]] = {
    "descripcion_item": ["description", "descripcion", "descripción", "item",
                         "product", "producto", "concepto"],
    "capacidad_w": ["install capacity", "installed capacity", "capacity (w)",
                    "capacidad instalada", "potencia instalada", "potencia (w)",
                    "power (w)", "watts", "capacity", "capacidad", "potencia"],
    "precio_unitario_w": ["price/watt", "price per watt", "unit price",
                          "precio/watt", "precio unitario", "usd/w"],
    "monto_item": ["total amount", "amount", "total", "monto", "importe",
                   "valor total"],
}

_KW_FLETE = ("freight", "flete", "shipping", "envio", "envío")
_KW_CIF = ("cif",)
_KW_FOB = ("fob",)


def _clasificar_columnas(encabezado: list[str]) -> dict[int, str]:
    """Dado el texto de cada celda de una posible fila de encabezado,
    devuelve {índice_columna: campo_canónico} para las columnas reconocidas.
    Cada campo se asigna como máximo a una columna."""
    mapa: dict[int, str] = {}
    usados = set()
    for idx, celda in enumerate(encabezado):
        c = _sin_acentos(celda or "")
        if not c:
            continue
        for campo, sinonimos in _ENCABEZADOS_TABLA.items():
            if campo in usados:
                continue
            if any(_sin_acentos(syn) in c for syn in sinonimos):
                mapa[idx] = campo
                usados.add(campo)
                break
    return mapa


def _extraer_de_tablas(tablas: list) -> dict:
    resultado: dict = {}
    for tabla in tablas:
        if len(tabla) < 2:
            continue
        idx_hdr, mapa_cols = None, {}
        for i, fila in enumerate(tabla):
            m = _clasificar_columnas(fila)
            if len(m) >= 2:
                idx_hdr, mapa_cols = i, m
                break
        if idx_hdr is None:
            continue
        for fila in tabla[idx_hdr + 1:]:
            fila_txt = " ".join((c or "") for c in fila).strip()
            if not fila_txt:
                continue
            fila_norm = _sin_acentos(fila_txt)
            es_flete = any(k in fila_norm for k in _KW_FLETE)
            es_cif = any(k in fila_norm for k in _KW_CIF)
            es_fob = any(k in fila_norm for k in _KW_FOB)
            for idx, campo in mapa_cols.items():
                if idx >= len(fila):
                    continue
                valor_celda = (fila[idx] or "").strip()
                if not valor_celda:
                    continue
                if campo == "descripcion_item":
                    # Nunca tomar la descripción de una fila de flete/CIF/FOB --
                    # esas son resúmenes ("Ocean freight...", "Total Amount (CIF...")
                    # que a veces caen en la columna Descripción, no el ítem real.
                    if ("descripcion_item" not in resultado and not es_flete
                            and not es_cif and not es_fob):
                        resultado["descripcion_item"] = {
                            "valor": valor_celda, "evidencia": fila_txt,
                            "metodo": "patron"}
                    continue
                num = _parsear_numero(valor_celda)
                if num is None:
                    continue
                if campo == "monto_item":
                    if es_flete:
                        campo_final = "flete"
                    elif es_cif:
                        campo_final = "total_cif"
                    elif es_fob:
                        campo_final = "total_fob"
                    else:
                        campo_final = "total_fob"  # línea del ítem principal
                    if campo_final not in resultado:
                        resultado[campo_final] = {
                            "valor": num, "evidencia": fila_txt, "metodo": "patron"}
                elif campo not in resultado and not es_flete:
                    resultado[campo] = {
                        "valor": num, "evidencia": fila_txt, "metodo": "patron"}
    return resultado


def extraer_por_patrones(texto: str, tablas: list) -> dict:
    """Combina el paso de texto (etiqueta:valor) con el de tablas
    (encabezado de columna). El texto tiene prioridad para los campos que
    ambos podrían encontrar (suele traer la etiqueta más explícita); las
    tablas rellenan lo que el texto no pudo (típicamente los campos de la
    fila de ítem: capacidad, precio unitario, descripción)."""
    resultado = _extraer_texto(texto)
    for campo, r in _extraer_de_tablas(tablas).items():
        if campo not in resultado:
            resultado[campo] = r
    return resultado


# ══════════════════ Respaldo con IA (evidencia verificada) ══════════════════

_PROMPT_EXTRACCION = """Eres un asistente que EXTRAE datos de una cotización comercial \
de un proveedor. NUNCA inventes ni calcules valores. Para cada campo solicitado, si \
puedes encontrarlo LITERALMENTE en el texto, responde con el valor Y el fragmento \
EXACTO (copiado tal cual, sin parafrasear ni corregir) de donde lo sacaste. Si un \
campo no aparece en el texto, ponlo en null -- NUNCA lo inventes ni lo calcules a \
partir de otros campos.

Responde ÚNICAMENTE con un JSON válido (sin texto adicional, sin bloque de código \
markdown), con esta forma exacta:
{"nombre_campo": {"valor": "...", "evidencia": "fragmento exacto del texto"}, \
"otro_campo": null}"""


def extraer_con_ia(texto: str, campos_faltantes: list[str]) -> dict:
    """Pide al proveedor de IA configurado (ver calculos.ia_proveedor) los
    campos que el paso de patrones no encontró. Cada valor se descarta si el
    modelo no puede citar un fragmento que exista LITERALMENTE en el texto
    del documento -- es la salvaguarda contra invención."""
    from calculos.ia_proveedor import llamar_ia, proveedor_disponible

    if not campos_faltantes or not proveedor_disponible():
        return {}
    contenido = (
        f"CAMPOS A EXTRAER: {', '.join(campos_faltantes)}\n\n"
        f"TEXTO DE LA COTIZACIÓN:\n{texto[:12000]}"
    )
    try:
        out = llamar_ia(_PROMPT_EXTRACCION, contenido, timeout=45)
    except RuntimeError:
        return {}

    bruto = re.sub(r"^```(?:json)?\s*|\s*```$", "", out["texto"].strip())
    try:
        datos = json.loads(bruto)
    except (ValueError, TypeError):
        return {}
    if not isinstance(datos, dict):
        return {}

    texto_norm = re.sub(r"\s+", " ", texto)
    resultado = {}
    for campo in campos_faltantes:
        entrada = datos.get(campo)
        if not isinstance(entrada, dict):
            continue
        valor, evidencia = entrada.get("valor"), entrada.get("evidencia")
        if valor is None or not evidencia:
            continue
        evidencia_norm = re.sub(r"\s+", " ", str(evidencia)).strip()
        # ── Salvaguarda anti-invención: la evidencia citada debe existir tal
        # cual en el documento real. Si el modelo no puede citar un
        # fragmento verdadero, el campo se descarta -- no se confía "a
        # ciegas" en su palabra.
        if not evidencia_norm or evidencia_norm not in texto_norm:
            continue
        if campo in _CAMPOS_TEXTO_NO_NUM or campo == "descripcion_item":
            valor_final = str(valor).strip()
            if not valor_final:
                continue
        else:
            valor_final = _parsear_numero(str(valor))
            if valor_final is None:
                continue
        resultado[campo] = {"valor": valor_final, "evidencia": str(evidencia),
                             "metodo": "ia"}
    return resultado


# ═══════════════════════════ Punto de entrada ════════════════════════════════

# ═══════════ Clasificador de categoría de costo (Presupuesto) ══════════════

# Sinónimos por CATEGORÍA de costo -- mismo espíritu que _CAMPOS_TEXTO
# (diccionario extensible, sin ramas de código por proveedor), aplicado a un
# nivel más grueso: no "qué campo es este número" sino "de qué sección de
# 💼 Presupuesto es esta cotización". Las claves coinciden EXACTAMENTE con
# las claves de sección de pages/8_💼_Presupuesto.py (las mismas que usa
# `pstore.SECCIONES_PERSISTIBLES` para perfileria/mano_obra/sistema_fv/
# inversor). "inversor"/"inverter" se deja FUERA a propósito: en esa pestaña
# significa tableros/breakers, no el modelo de inversor (que tiene su propio
# flujo de precio desde el catálogo de 📐 Dimensionamiento) -- incluir la
# palabra genérica "inversor" generaría falsos positivos frecuentes.
_CATEGORIAS_COSTO: dict[str, list[str]] = {
    "perfileria": [
        "structure", "mounting", "estructura", "montaje", "perfileria",
        "perfilería", "racking", "rack", "soporte", "screw foundation",
        "ground mount", "roof mount", "clamp", "rail", "carril", "anclaje",
        "cimentacion", "cimentación",
    ],
    "mano_obra": [
        "installation", "instalacion", "instalación", "labor", "mano de obra",
        "commissioning", "puesta en marcha", "certificacion", "certificación",
        "retie", "ritel", "contratista", "cuadrilla",
    ],
    "sistema_fv": [
        "cable", "conductor", "combiner box", "caja de conexiones",
        "puesta a tierra", "grounding", "monitoring", "monitoreo", "conduit",
        "canalizacion", "canalización", "string box", "conector mc4",
    ],
    "inversor": [
        "breaker", "tablero", "switchgear", "distribution board",
        "panel electrico", "panel eléctrico", "proteccion ac",
        "protección ac", "transfer switch",
    ],
    "catalogo": [
        "panel solar", "solar panel", "modulo fotovoltaico",
        "módulo fotovoltaico", "modulo bipv", "pv module", "bateria",
        "batería", "battery",
    ],
    "soft": [
        "ingenieria", "ingeniería", "engineering", "diseno", "diseño",
        "consultoria", "consultoría", "legal", "seguro", "poliza", "póliza",
        "insurance", "project management", "gerencia de proyecto",
        "auditoria", "auditoría", "licencia", "tramite", "trámite",
    ],
}

CATEGORIA_LABELS: dict[str, str] = {
    "perfileria": "🔩 Perfilería y Estructura",
    "mano_obra":  "👷 Mano de Obra",
    "sistema_fv": "⚡ Sistema FV",
    "inversor":   "🔌 Inversor y Equipos Eléctricos",
    "catalogo":   "📦 Equipos del Catálogo",
    "soft":       "🧾 Costos Blandos",
}


def clasificar_categoria_costo(texto: str) -> tuple[str | None, dict[str, int]]:
    """Sugiere a qué sección de 💼 Presupuesto pertenece una cotización,
    contando coincidencias de palabras clave por categoría. Devuelve
    (categoria_sugerida, {categoria: puntaje, ...}) -- ambos ordenados de
    mayor a menor puntaje en el dict. `categoria_sugerida` es None si el
    texto no coincide con ninguna categoría (el usuario elige a mano)."""
    t = _sin_acentos(texto)
    puntajes = {}
    for cat, palabras in _CATEGORIAS_COSTO.items():
        score = sum(t.count(_sin_acentos(p)) for p in palabras)
        if score:
            puntajes[cat] = score
    if not puntajes:
        return None, {}
    puntajes = dict(sorted(puntajes.items(), key=lambda kv: -kv[1]))
    mejor = next(iter(puntajes))
    return mejor, puntajes


CAMPOS_COTIZACION = (
    "proveedor", "numero_cotizacion", "fecha_cotizacion", "descripcion_item",
    "capacidad_w", "precio_unitario_w", "total_fob", "flete", "total_cif",
    "incoterm", "condiciones_pago",
)


def extraer_cotizacion(datos: bytes, nombre_archivo: str) -> dict:
    """Punto de entrada principal. Lee el archivo (PDF o Word), extrae los
    campos con los patrones genéricos y, si falta alguno y el servidor tiene
    una clave de IA configurada, completa con extracción asistida por IA
    (con verificación de evidencia).

    Devuelve {campo: {"valor", "evidencia", "metodo"}, ...} más las claves
    especiales "_advertencias" (lista de avisos, p. ej. moneda distinta de
    USD) y "_texto_crudo" (para auditoría)."""
    texto, tablas = leer_cotizacion(datos, nombre_archivo)
    resultado = extraer_por_patrones(texto, tablas)
    faltantes = [c for c in CAMPOS_COTIZACION if c not in resultado]
    if faltantes:
        resultado.update(extraer_con_ia(texto, faltantes))

    advertencias = []
    moneda = _detectar_moneda(texto)
    if moneda and moneda != "USD":
        advertencias.append(
            f"La cotización parece estar en {moneda}, no en USD -- convierte los "
            "valores manualmente antes de aplicarlos a Presupuesto (los montos de "
            "Presupuesto son siempre en USD)."
        )
    resultado["_advertencias"] = advertencias
    resultado["_texto_crudo"] = texto
    return resultado
