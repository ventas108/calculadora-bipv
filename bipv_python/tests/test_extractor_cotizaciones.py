# -*- coding: utf-8 -*-
"""Extractor genérico de cotizaciones (2026-08-22).

Pedido del usuario: poder cargar la cotización de un proveedor (PDF o Word)
y que sus valores entren a 💼 Presupuesto de forma automática pero VERÍDICA
(nunca inventada), y que el mecanismo sea GENÉRICO -- no atado a la
plantilla de un solo proveedor (el caso real que lo disparó fue una
cotización de Xiamen Mibet New Energy para estructura de montaje, pero debe
funcionar igual con la cotización de otro proveedor con layout distinto).

Estos tests verifican:
- El parseo robusto de números en ambas convenciones de separador.
- La extracción por patrones de texto + tablas contra el texto REAL de la
  cotización de Mibet (reproducido tal como lo pegó el usuario).
- Que el MISMO extractor -- sin ningún cambio ni rama de código por marca --
  también funciona contra una cotización SINTÉTICA de un proveedor
  completamente distinto, en español, con layout y formato de número
  diferentes (prueba de que es genérico, no solo "funciona para Mibet").
- La salvaguarda anti-invención del respaldo por IA: si el modelo cita un
  fragmento que no existe en el documento real, el campo se descarta.
"""
import io

import pytest

from calculos.extractor_cotizaciones import (
    _detectar_moneda,
    _parsear_numero,
    extraer_con_ia,
    extraer_por_patrones,
    leer_cotizacion,
)

# ── Texto de la cotización REAL de Mibet Energy, tal como lo pegó el usuario,
# aproximando cómo pdfplumber aplanaría el PDF a texto plano. ────────────────
_TEXTO_MIBET = """TEL: +86-592-3754999 ext. 688
TEL:+57 318 7241820
Email: irene.li@mbt-energy.com
Email:
Contact: Irene Li    Contact: Mauricio
No.: PJ-260807-04 V01    Date: Aug 11, 2026
Bank Information
Beneficiary Name: Xiamen Mibet New Energy Co., Ltd
Beneficiary's address: NO.45,Sushan Road,Jimei District,Xiamen,China.
Account Number: 000000501511035997
SWIFT NO: SCBLCNSXIMN
Bank Name: Standard Chartered Bank (China) Limited
Bank Address: Unit EFGH, 18/F, International Plaza No. 8 Lu Jiang Road, Xiamen Fujian 361001, China
Description    Install Capacity (W)    Price/Watt (USD, FOB Tianjin)    Total Amount (USD, FOB Tianjin)
220.32kw Ground Mounting Structure
(Ground Srew Foundation，C4)
220320    US$0.09051    US$19,940.33
    Ocean freight charges ( 1*40HQ)    US$8,000.00
    Total Amount (CIF Cartagena)    US$27,940.33
Terms and Conditions
1. This the price is based on direct supply agreements being in place between the Purchaser and Mibet Energy.
Valid 10 working days from quoation date.
2. Payment Terms: T/T 30% in advance, and the balance before shipment.
3. Delivery terms: CIF Cartagena
4. Lead time: 20 working days after received payment.
"""

_TABLAS_MIBET = [
    [
        ["Description", "Install Capacity (W)", "Price/Watt (USD, FOB Tianjin)",
         "Total Amount (USD, FOB Tianjin)"],
        ["220.32kw Ground Mounting Structure\n(Ground Srew Foundation，C4)",
         "220320", "US$0.09051", "US$19,940.33"],
        ["", "Ocean freight charges ( 1*40HQ)", "", "US$8,000.00"],
        ["", "Total Amount (CIF Cartagena)", "", "US$27,940.33"],
    ]
]

# ── Cotización SINTÉTICA de un proveedor totalmente distinto: en español,
# con formato de número latino (punto=miles, coma=decimales), layout de
# tabla con columnas en OTRO orden, y otras etiquetas. Prueba de que el
# extractor generaliza sin ningún código específico de proveedor. ───────────
_TEXTO_OTRO_PROVEEDOR = """Estructuras del Caribe S.A.S.
Contacto: Ana Torres
Cotización No.: EC-2026-0731
Fecha: 15/09/2026
Cliente: Innovación Química
Condiciones de pago: 50% anticipo, 50% contra entrega
Términos de entrega: CIF Barranquilla

Concepto                          Precio Unitario (USD/W)   Potencia (W)   Total
Estructura galvanizada tipo C      USD 0,0875                 180500        USD 15.793,75
Flete marítimo (contenedor 40 pies)                                          USD 6.200,00
Total CIF Barranquilla                                                       USD 21.993,75
"""

_TABLAS_OTRO_PROVEEDOR = [
    [
        ["Concepto", "Precio Unitario (USD/W)", "Potencia (W)", "Total"],
        ["Estructura galvanizada tipo C", "USD 0,0875", "180500", "USD 15.793,75"],
        ["Flete marítimo (contenedor 40 pies)", "", "", "USD 6.200,00"],
        ["Total CIF Barranquilla", "", "", "USD 21.993,75"],
    ]
]


# ═══════════════════════════ _parsear_numero ═══════════════════════════════

@pytest.mark.parametrize("texto,esperado", [
    ("19,940.33", 19940.33),        # formato US
    ("27.940,33", 27940.33),        # formato latino
    ("220320", 220320.0),
    ("0.09051", 0.09051),
    ("0,09051", 0.09051),
    ("US$8,000.00", 8000.00),
    ("USD 6.200,00", 6200.00),
    ("", None),
    ("sin numeros", None),
])
def test_parsear_numero_ambas_convenciones(texto, esperado):
    assert _parsear_numero(texto) == esperado


# ═══════════════════════════ Cotización de Mibet ═══════════════════════════

def test_mibet_extrae_proveedor_y_numero_de_cotizacion():
    r = extraer_por_patrones(_TEXTO_MIBET, _TABLAS_MIBET)
    assert "Mibet" in r["proveedor"]["valor"]
    assert r["numero_cotizacion"]["valor"].strip().startswith("PJ-260807-04")


def test_mibet_extrae_totales_y_flete_por_patron_de_texto():
    r = extraer_por_patrones(_TEXTO_MIBET, _TABLAS_MIBET)
    assert r["total_cif"]["valor"] == pytest.approx(27940.33)
    assert r["flete"]["valor"] == pytest.approx(8000.00)
    assert r["total_cif"]["metodo"] == "patron"


def test_mibet_extrae_capacidad_y_precio_por_tabla():
    r = extraer_por_patrones(_TEXTO_MIBET, _TABLAS_MIBET)
    assert r["capacidad_w"]["valor"] == pytest.approx(220320)
    assert r["precio_unitario_w"]["valor"] == pytest.approx(0.09051)
    assert "220320" in r["capacidad_w"]["evidencia"] or "220320" in r["precio_unitario_w"]["evidencia"]


def test_mibet_extrae_condiciones_de_entrega_y_pago():
    r = extraer_por_patrones(_TEXTO_MIBET, _TABLAS_MIBET)
    assert "CIF Cartagena" in r["incoterm"]["valor"]
    assert "30%" in r["condiciones_pago"]["valor"]


def test_mibet_numero_generico_no_confunde_con_datos_bancarios():
    # El sinónimo genérico "No." NO debe terminar devolviendo el SWIFT o el
    # número de cuenta bancaria -- debe encontrar "PJ-260807-04 V01".
    r = extraer_por_patrones(_TEXTO_MIBET, _TABLAS_MIBET)
    assert "SCBLCNSXIMN" not in r["numero_cotizacion"]["valor"]
    assert "000000501511035997" not in r["numero_cotizacion"]["valor"]


def test_mibet_moneda_detectada_es_usd():
    assert _detectar_moneda(_TEXTO_MIBET) == "USD"


# ═════════════ Mismo extractor, proveedor y layout DISTINTOS (genérico) ═════

def test_otro_proveedor_extrae_proveedor_y_numero_sin_codigo_especifico():
    r = extraer_por_patrones(_TEXTO_OTRO_PROVEEDOR, _TABLAS_OTRO_PROVEEDOR)
    assert "EC-2026-0731" in r["numero_cotizacion"]["valor"]


def test_otro_proveedor_extrae_totales_con_formato_de_numero_latino():
    r = extraer_por_patrones(_TEXTO_OTRO_PROVEEDOR, _TABLAS_OTRO_PROVEEDOR)
    assert r["total_cif"]["valor"] == pytest.approx(21993.75)
    assert r["flete"]["valor"] == pytest.approx(6200.00)


def test_otro_proveedor_extrae_capacidad_y_precio_con_columnas_en_otro_orden():
    # En esta plantilla "Precio Unitario" viene ANTES que "Potencia" (al
    # revés que en Mibet) -- el extractor debe leer por encabezado, no por
    # posición fija de columna.
    r = extraer_por_patrones(_TEXTO_OTRO_PROVEEDOR, _TABLAS_OTRO_PROVEEDOR)
    assert r["capacidad_w"]["valor"] == pytest.approx(180500)
    assert r["precio_unitario_w"]["valor"] == pytest.approx(0.0875)


def test_otro_proveedor_extrae_condiciones_en_espanol():
    r = extraer_por_patrones(_TEXTO_OTRO_PROVEEDOR, _TABLAS_OTRO_PROVEEDOR)
    assert "CIF Barranquilla" in r["incoterm"]["valor"]
    assert "50%" in r["condiciones_pago"]["valor"]


def test_otro_proveedor_moneda_detectada_es_usd():
    assert _detectar_moneda(_TEXTO_OTRO_PROVEEDOR) == "USD"


# ═══════════════ Respaldo por IA: salvaguarda anti-invención ═══════════════

def test_ia_descarta_campo_si_la_evidencia_citada_no_existe_en_el_documento(monkeypatch):
    import calculos.extractor_cotizaciones as mod

    def _fake_llamar_ia(prompt_sistema, contenido_usuario, timeout=60):
        return {"texto": '{"proveedor": {"valor": "Proveedor Inventado", '
                          '"evidencia": "esto no aparece en el documento real"}}',
                "proveedor": "Gemini"}

    monkeypatch.setattr(mod, "proveedor_disponible", lambda: "Gemini", raising=False)
    monkeypatch.setattr("calculos.ia_proveedor.llamar_ia", _fake_llamar_ia)
    monkeypatch.setattr("calculos.ia_proveedor.proveedor_disponible", lambda: "Gemini")

    r = extraer_con_ia("Este es el texto real de la cotización.", ["proveedor"])
    assert "proveedor" not in r  # la evidencia no existe en el texto -> se descarta


def test_ia_acepta_campo_si_la_evidencia_citada_existe_literalmente(monkeypatch):
    texto_real = "Beneficiary Name: Acme Structures Ltd, un proveedor certificado."

    def _fake_llamar_ia(prompt_sistema, contenido_usuario, timeout=60):
        return {"texto": '{"proveedor": {"valor": "Acme Structures Ltd", '
                          '"evidencia": "Beneficiary Name: Acme Structures Ltd"}}',
                "proveedor": "Gemini"}

    monkeypatch.setattr("calculos.ia_proveedor.llamar_ia", _fake_llamar_ia)
    monkeypatch.setattr("calculos.ia_proveedor.proveedor_disponible", lambda: "Gemini")

    r = extraer_con_ia(texto_real, ["proveedor"])
    assert r["proveedor"]["valor"] == "Acme Structures Ltd"
    assert r["proveedor"]["metodo"] == "ia"


def test_ia_no_se_llama_si_no_hay_clave_configurada(monkeypatch):
    monkeypatch.setattr("calculos.ia_proveedor.proveedor_disponible", lambda: None)
    r = extraer_con_ia("cualquier texto", ["proveedor"])
    assert r == {}


# ═══════════════════ Lectura de Word (.docx) real, ida y vuelta ═════════════

def test_leer_docx_real_extrae_texto_y_tabla_correctamente():
    import docx as pydocx

    doc = pydocx.Document()
    doc.add_paragraph("Cotización No.: WD-2026-001")
    doc.add_paragraph("Beneficiary Name: Word Structures Inc")
    tabla = doc.add_table(rows=2, cols=3)
    tabla.rows[0].cells[0].text = "Description"
    tabla.rows[0].cells[1].text = "Capacity (W)"
    tabla.rows[0].cells[2].text = "Total Amount"
    tabla.rows[1].cells[0].text = "Estructura de montaje"
    tabla.rows[1].cells[1].text = "150000"
    tabla.rows[1].cells[2].text = "US$12,500.00"

    buf = io.BytesIO()
    doc.save(buf)
    datos = buf.getvalue()

    texto, tablas = leer_cotizacion(datos, "cotizacion.docx")
    assert "Word Structures Inc" in texto
    assert len(tablas) == 1
    assert tablas[0][0][1] == "Capacity (W)"

    r = extraer_por_patrones(texto, tablas)
    assert "Word Structures" in r["proveedor"]["valor"]
    assert r["capacidad_w"]["valor"] == pytest.approx(150000)


def test_leer_cotizacion_formato_no_soportado_da_error_claro():
    with pytest.raises(ValueError, match="no soportado"):
        leer_cotizacion(b"contenido", "cotizacion.xlsx")
