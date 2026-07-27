"""
Genera: Guia_Despliegue_BIPV_DigitalOcean.docx
Guía completa paso a paso para subir la calculadora BIPV Python
al servidor Digital Ocean de Mauricio.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def h1(t):
    p = doc.add_heading(t, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
def h2(t):
    p = doc.add_heading(t, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
def h3(t):
    doc.add_heading(t, level=3)
def par(t='', bold=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(t); r.bold = bold; r.font.size = Pt(10)
    if color: r.font.color.rgb = RGBColor(*color)
def cmd(t):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(t); r.font.name = 'Courier New'; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
def warn(t):
    p = doc.add_paragraph()
    r = p.add_run("⚠️  " + t)
    r.bold = True; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xCC, 0x55, 0x00)
def ok(t):
    p = doc.add_paragraph()
    r = p.add_run("✅  " + t); r.font.size = Pt(10)
def tabla(filas, h1t='Paso', h2t='Comando / Acción'):
    t = doc.add_table(rows=1, cols=2); t.style = 'Table Grid'
    t.rows[0].cells[0].text = h1t; t.rows[0].cells[1].text = h2t
    for a in t.rows[0].cells:
        a.paragraphs[0].runs[0].bold = True
    for a, b in filas:
        r = t.add_row().cells; r[0].text = a; r[1].text = b

# ═══════════════════════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════════════════════
titulo = doc.add_heading('', 0)
r = titulo.add_run('GUÍA DE DESPLIEGUE\nCalculadora BIPV Python → Digital Ocean')
r.font.size = Pt(22); r.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run(
    'Servidor: Ubuntu 22.04 — Digital Ocean\n'
    'App actual: Node.js en bipv.innovacionquimica.com.co\n'
    'App nueva: Streamlit en calc.innovacionquimica.com.co\n\n'
    'Mauricio Acevedo — Ingeniería BIPV Colombia\n'
    'Julio 2026'
)
r2.font.size = Pt(12); r2.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 0 — ALERTA DE SEGURIDAD
# ═══════════════════════════════════════════════════════════════════════════
h1('0. ⚠️ ACCIÓN URGENTE: CREDENCIALES EXPUESTAS')
par(
    'El documento OJO_API_de_GitHub.docx que subiste a Replit contiene en texto plano '
    'tu token de GitHub y la contraseña SSH de tu servidor. Replit puede guardar estos '
    'archivos. Debes actuar ANTES de continuar con esta guía.',
    bold=True, color=(0xCC, 0x00, 0x00)
)
doc.add_paragraph()
h2('0.1 Revocar y regenerar el token de GitHub')
par('Pasos en github.com (desde el navegador):')
for paso in [
    'Ir a github.com → menú superior derecho → Settings',
    'Menú izquierdo → Developer settings (al fondo)',
    'Personal access tokens → Tokens (classic)',
    'Encontrar el token "ghp_Wz41..." → botón Delete',
    'Generar nuevo token: Generate new token (classic)',
    'Scopes necesarios: marcar "repo" (todo el repo)',
    'Copiar el nuevo token en un lugar seguro (gestión de contraseñas)',
    'NUNCA subirlo a ningún archivo de texto o documento',
]:
    doc.add_paragraph(paso, style='List Number')

doc.add_paragraph()
h2('0.2 Cambiar contraseña SSH del servidor')
par('Conectarte al servidor y cambiar la contraseña root:')
cmd('ssh root@198.199.160.X   # usa la IP real de tu servidor')
cmd('passwd                    # escribe la nueva contraseña dos veces')
par('O desde el panel de Digital Ocean: Droplets → tu servidor → Access → Reset Root Password')
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 1 — ARQUITECTURA
# ═══════════════════════════════════════════════════════════════════════════
h1('1. ARQUITECTURA DEL SISTEMA')
par(
    'Tu servidor Digital Ocean ya tiene la calculadora TypeScript corriendo. '
    'Vamos a agregar la calculadora Python/Streamlit en el mismo servidor, '
    'en un subdominio separado. No hay que crear un servidor nuevo ni pagar más.'
)
doc.add_paragraph()
cmd('''ANTES (lo que ya tienes):
  nginx (puerto 80/443)
    └── bipv.innovacionquimica.com.co → PM2 → Node.js (puerto 3000)
         /var/www/bipv/calculadora/

DESPUÉS (lo que agregaremos):
  nginx (puerto 80/443)
    ├── bipv.innovacionquimica.com.co  → PM2 → Node.js (puerto 3000)   [sin cambios]
    └── calc.innovacionquimica.com.co  → PM2 → Streamlit (puerto 8501) [NUEVO]
         /var/www/bipv/calculadora_bipv/
''')
par('Ventajas de esta arquitectura:', bold=True)
for v in [
    'Sin costo adicional — mismo servidor, mismo IP, misma suscripción DO',
    'Subdominio separado — las dos apps son independientes, una no afecta la otra',
    'PM2 gestiona ambos procesos — si el servidor se reinicia, ambos arrancan solos',
    'RAM disponible — pvlib + numpy pueden usar hasta 1-2 GB sin problema en DO',
    'HTTPS independiente — Certbot agrega el certificado al subdominio con 1 comando',
    'Flujo de actualización idéntico — git pull + pm2 restart (igual que el Node.js)',
]:
    doc.add_paragraph(v, style='List Bullet')
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 2 — PASO A PASO COMPLETO
# ═══════════════════════════════════════════════════════════════════════════
h1('2. INSTALACIÓN PASO A PASO')
warn('Ejecutar todos los comandos como root en el servidor (vía SSH)')

# ─── PASO 1 ───
h2('PASO 1 — Conectarte al servidor')
par('Desde PowerShell o Terminal de tu computador:')
cmd('ssh root@<IP_DE_TU_SERVIDOR>')
par('(usa la IP real de tu servidor — la que ya sabes)')
doc.add_paragraph()

# ─── PASO 2 ───
h2('PASO 2 — Actualizar el sistema e instalar Python')
cmd('apt update && apt upgrade -y')
cmd('apt install -y python3.11 python3.11-venv python3-pip git')
cmd('python3.11 --version   # debe mostrar Python 3.11.x')
doc.add_paragraph()

# ─── PASO 3 ───
h2('PASO 3 — Configurar el nuevo token de GitHub y clonar')
par('Primero agrega la carpeta bipv_python/ al repo existente desde Replit:')
cmd('# En Replit (ya está hecho): la carpeta bipv_python/ está lista')
cmd('# Súbela a tu GitHub con el flujo normal de Replit → git push')
doc.add_paragraph()
par('En el servidor, obtener el código:')
cmd('mkdir -p /var/www/bipv/calculadora_bipv')
cmd('cd /var/www/bipv')
cmd('# Si ya tienes el repo clonado:')
cmd('cd /var/www/bipv/calculadora-bipv && git pull origin main')
cmd('')
cmd('# Copiar la carpeta bipv_python al directorio de la app:')
cmd('cp -r /var/www/bipv/calculadora-bipv/bipv_python/* /var/www/bipv/calculadora_bipv/')
doc.add_paragraph()

# ─── PASO 4 ───
h2('PASO 4 — Crear entorno virtual e instalar dependencias')
cmd('cd /var/www/bipv/calculadora_bipv')
cmd('python3.11 -m venv venv')
cmd('source venv/bin/activate')
cmd('pip install --upgrade pip')
cmd('pip install -r requirements.txt')
doc.add_paragraph()
par('Verificar instalación:')
cmd('python3 -c "import pvlib; print(\'pvlib\', pvlib.__version__)"')
cmd('python3 -c "import streamlit; print(\'streamlit\', streamlit.__version__)"')
cmd('python3 -c "import numpy; print(\'numpy\', numpy.__version__)"')
doc.add_paragraph()

# ─── PASO 5 ───
h2('PASO 5 — Ejecutar los tests de validación')
par('Antes de arrancar la app, verificar que el motor SDM está correcto:')
cmd('cd /var/www/bipv/calculadora_bipv')
cmd('source venv/bin/activate')
cmd('python -m pytest tests/test_validacion_vba.py -v')
doc.add_paragraph()
par('Resultado esperado:')
cmd('''PASSED test_ff_vs_irradiancia[100-69.75]
PASSED test_ff_vs_irradiancia[200-76.28]
PASSED test_ff_vs_irradiancia[400-74.51]
PASSED test_ff_vs_irradiancia[600-72.87]
PASSED test_ff_vs_irradiancia[800-71.55]
PASSED test_ff_vs_irradiancia[1000-64.92]
PASSED test_maximo_ff_en_bajo_G
PASSED test_validacion_stc_vs_ficha
PASSED test_optimizar_n_serie[8-True]
PASSED test_optimizar_n_serie[9-False]
10 passed in 3.2s''')
warn('Si algún test FALLA, no continuar. Reportar el error para corregir el código.')
doc.add_paragraph()

# ─── PASO 6 ───
h2('PASO 6 — Probar Streamlit manualmente')
par('Antes de configurar PM2, probar que la app arranca:')
cmd('cd /var/www/bipv/calculadora_bipv')
cmd('source venv/bin/activate')
cmd('streamlit run app.py --server.port 8501 --server.address 0.0.0.0')
doc.add_paragraph()
par('Desde otro terminal SSH, verificar que responde:')
cmd('curl -s http://localhost:8501/healthz')
par('Debe responder: ok')
par('Cuando funcione, presionar Ctrl+C para detenerlo — PM2 lo iniciará permanentemente.')
doc.add_paragraph()

# ─── PASO 7 ───
h2('PASO 7 — Configurar PM2')
cmd('cd /var/www/bipv/calculadora_bipv')
cmd('pm2 start ecosystem.config.js')
cmd('pm2 save                    # guardar para que arranque al reiniciar servidor')
cmd('pm2 status                  # verificar que aparece "online"')
doc.add_paragraph()
par('Verificar logs en tiempo real:')
cmd('pm2 logs calculadora-bipv-python --lines 30')
par('Debe mostrar: "You can now view your Streamlit app in your browser"')
doc.add_paragraph()

# ─── PASO 8 ───
h2('PASO 8 — Configurar nginx (subdominio)')
cmd('cp /var/www/bipv/calculadora_bipv/nginx_bipv_python.conf /etc/nginx/sites-available/bipv-python')
cmd('ln -sf /etc/nginx/sites-available/bipv-python /etc/nginx/sites-enabled/bipv-python')
cmd('nginx -t                    # verificar que no hay errores de sintaxis')
cmd('systemctl reload nginx')
doc.add_paragraph()
warn('Antes de este paso, debes agregar el subdominio DNS (ver Paso 8B).')
doc.add_paragraph()

# ─── PASO 8B ───
h2('PASO 8B — Agregar subdominio DNS')
par('En el panel donde tienes el dominio innovacionquimica.com.co:')
tabla([
    ('Tipo', 'A'),
    ('Nombre / Host', 'calc'),
    ('Valor / Destino', 'IP_DE_TU_SERVIDOR_DO'),
    ('TTL', '3600 (o "Automático")'),
], 'Campo', 'Valor')
par('Guardar y esperar 5-15 minutos para que el DNS se propague.')
doc.add_paragraph()

# ─── PASO 9 ───
h2('PASO 9 — Certificado HTTPS para el subdominio')
cmd('certbot --nginx -d calc.innovacionquimica.com.co')
par('Certbot pedirá tu email (ya lo tiene del certificado anterior). '
    'Selecciona "Expand" cuando pregunte si expandir el certificado existente.')
cmd('nginx -t && systemctl reload nginx')
doc.add_paragraph()

# ─── PASO 10 ───
h2('PASO 10 — Verificación final')
par('Desde tu navegador:')
for url in [
    'https://calc.innovacionquimica.com.co  →  debe mostrar la calculadora Streamlit',
    'https://bipv.innovacionquimica.com.co  →  debe seguir funcionando el app TypeScript',
]:
    doc.add_paragraph(url, style='List Bullet')
cmd('pm2 status    # ambos procesos deben mostrar "online"')
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 3 — FLUJO DE ACTUALIZACIÓN DIARIA
# ═══════════════════════════════════════════════════════════════════════════
h1('3. FLUJO DE ACTUALIZACIÓN (una vez instalado)')
par('Exactamente igual a los 4 comandos del proyecto Node.js, solo cambia el nombre:')
cmd('''# ─── Los 4 comandos para actualizar la calculadora Python ─────────────────

cd /var/www/bipv/calculadora_bipv
git pull
pm2 restart calculadora-bipv-python
pm2 status

# Eso es todo. Si agregaste nuevas dependencias, agregar entre paso 2 y 3:
# source venv/bin/activate && pip install -r requirements.txt''')
doc.add_paragraph()
par('Comparación con el flujo Node.js que ya conoces:', bold=True)
tabla([
    ('cd /var/www/bipv/calculadora', 'cd /var/www/bipv/calculadora_bipv'),
    ('git pull', 'git pull'),
    ('(no necesario)', 'source venv/bin/activate && pip install -r requirements.txt  (solo si cambió requirements)'),
    ('pm2 restart calculadora-bipv', 'pm2 restart calculadora-bipv-python'),
], 'Node.js (ya conoces)', 'Python (nuevo)')
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 4 — RESOLUCIÓN DE PROBLEMAS COMUNES
# ═══════════════════════════════════════════════════════════════════════════
h1('4. PROBLEMAS COMUNES Y SOLUCIONES')

problemas = [
    ('pm2 muestra "errored" o "stopped"',
     'pm2 logs calculadora-bipv-python --lines 50\n'
     'Revisar el error. Causa común: Python no encuentra un módulo → '
     'source venv/bin/activate && pip install -r requirements.txt'),
    ('La página muestra "502 Bad Gateway"',
     'nginx está corriendo pero Streamlit no.\n'
     'Verificar: pm2 status\n'
     'Si está "stopped": pm2 start calculadora-bipv-python\n'
     'Si está "online": curl http://localhost:8501/healthz'),
    ('WebSocket error (página se queda cargando)',
     'Verificar nginx: el config debe incluir los headers de Upgrade.\n'
     'Revisar el archivo nginx_bipv_python.conf (ya está correcto en el repo).\n'
     'nginx -t && systemctl reload nginx'),
    ('Los cálculos son muy lentos',
     'Verificar RAM disponible: free -h\n'
     'Si el servidor tiene < 1GB libre, considerar upgrading el Droplet de DO.\n'
     'Para cálculos pesados (TMY completo), se recomienda 2GB RAM mínimo.'),
    ('"ModuleNotFoundError: No module named pvlib"',
     'El proceso PM2 no está usando el virtualenv.\n'
     'Verificar ecosystem.config.js: la ruta de streamlit debe ser\n'
     '/var/www/bipv/calculadora_bipv/venv/bin/streamlit'),
    ('"git pull" pide usuario y contraseña',
     'Configurar el remote con el nuevo token:\n'
     'git remote set-url origin https://ventas108:<NUEVO_TOKEN>@github.com/ventas108/calculadora-bipv.git\n'
     'NUNCA uses el token anterior (ya fue revocado).'),
]

for problema, solucion in problemas:
    h3(f'❓ {problema}')
    par(solucion)
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 5 — ESTRUCTURA DEL CÓDIGO
# ═══════════════════════════════════════════════════════════════════════════
h1('5. ESTRUCTURA DEL CÓDIGO PYTHON')
cmd('''bipv_python/                     ← carpeta en el repo GitHub
│
├── app.py                          ← Página principal Streamlit
├── requirements.txt                ← Dependencias Python
├── ecosystem.config.js             ← Configuración PM2
├── nginx_bipv_python.conf          ← Configuración nginx
├── instalar_servidor.sh            ← Script de instalación
├── actualizar_en_servidor.sh       ← Script de actualización
│
├── .streamlit/
│   └── config.toml                 ← Puerto, tema, headless
│
├── datos/
│   ├── tecnologias_bipv.py         ← ASP-ST1-T40 y catálogo SolTech
│   ├── catalogo_inversores.py      ← Growatt MID15KTL3-X y otros
│   └── ciudades_colombia.py        ← 12 ciudades + Ley 1715 + CO₂
│
├── calculos/
│   ├── modelo_iv.py                ← Motor De Soto 2006 + pvlib (NÚCLEO)
│   ├── dimensionamiento.py         ← Semáforo string sizing
│   └── temperatura.py              ← NOCT dinámico
│
├── pages/
│   ├── 1_🏠_Proyecto.py
│   ├── 3_🔬_Motor_IV.py            ← Curva I-V interactiva + validación VBA
│   └── 4_📐_Dimensionamiento.py    ← Tabla semáforo OK/ALERTA/FALLA
│
└── tests/
    └── test_validacion_vba.py      ← 10 tests contra datos del XLSM''')

doc.add_paragraph()
par('Páginas pendientes de implementar (en las próximas sesiones):', bold=True)
tabla([
    ('2_☀️_Recurso_Solar.py',  'TMY de PVGIS + cálculo irradiancia POA por orientación'),
    ('5_⚡_Mismatch.py',       'Análisis de mismatch MPPT (validado vs XLSM)'),
    ('6_📊_Produccion.py',     'Simulación IEC 61724 hora a hora completa'),
    ('7_💰_Financiero.py',     'VPN, TIR, payback, Ley 1715/2014, CO₂'),
    ('8_📄_Reporte.py',        'Reporte Word + Excel descargable'),
], 'Archivo', 'Función')
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 6 — COMPARACIÓN DIGITAL OCEAN VS ALTERNATIVAS
# ═══════════════════════════════════════════════════════════════════════════
h1('6. POR QUÉ DIGITAL OCEAN ES LA ELECCIÓN CORRECTA')

tabla([
    ('Digital Ocean (tu elección) ✅',
     'Control total de RAM/CPU, sin límites de tiempo, pvlib corre sin restricciones, '
     'ya tienes el servidor configurado con nginx+PM2+HTTPS, costo fijo mensual'),
    ('Streamlit Community Cloud (gratis)',
     'Limitado a 1GB RAM (pvlib puede excederlo con TMY completo), '
     'requiere repo público en GitHub, hibernación tras inactividad, sin base de datos'),
    ('Heroku / Railway',
     'Limites de RAM en plan gratuito, cold starts, más caro que DO para lo mismo'),
    ('Google Cloud Run / AWS Lambda',
     'Serverless: cada request arranca frío, mal para Streamlit que mantiene estado, '
     'configuración compleja'),
], 'Plataforma', 'Características para BIPV/pvlib')

doc.add_paragraph()
par('Tu servidor DO actual — especificaciones relevantes para pvlib:', bold=True)
par(
    'pvlib procesando un TMY anual completo (8760 horas) del panel ASP-ST1-T40 '
    'consume aproximadamente:\n'
    '• RAM: 150-400 MB (numpy arrays de G y T × 8760 puntos)\n'
    '• CPU: 2-8 segundos (cálculo completo con singlediode())\n'
    '• Con 1-2GB RAM en tu droplet, tienes margen suficiente para hasta 5 usuarios simultáneos'
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# RESUMEN EJECUTIVO
# ═══════════════════════════════════════════════════════════════════════════
h1('7. RESUMEN — LO QUE HAY QUE HACER')

tabla([
    ('HOY — Urgente (seguridad)', 'Revocar token GitHub → generar nuevo\nCambiar contraseña SSH del servidor'),
    ('Paso 1 — Subir código a GitHub', 'Desde Replit: la carpeta bipv_python/ ya está lista → git push al repo existente'),
    ('Paso 2 — Instalar en servidor', 'SSH al servidor → ejecutar Pasos 2-7 de esta guía (20-30 min)'),
    ('Paso 3 — DNS + HTTPS', 'Agregar registro A "calc" en tu proveedor de dominio → certbot'),
    ('Paso 4 — Verificar', 'https://calc.innovacionquimica.com.co debe mostrar Streamlit funcionando'),
    ('Siguiente — Completar páginas', 'Implementar páginas 2 (Solar), 5 (Mismatch), 6 (Producción), 7 (Financiero)'),
], 'Qué', 'Cómo')

doc.add_paragraph()
par(
    'Una vez completados los 4 pasos, tendrás:\n'
    '✅ bipv.innovacionquimica.com.co — Calculadora TypeScript (existente)\n'
    '✅ calc.innovacionquimica.com.co — Calculadora BIPV Python + pvlib (nueva)\n'
    '✅ Motor De Soto 2006 validado contra tu XLSM (10 tests pasan)\n'
    '✅ Panel ASP-ST1-T40 con parámetros SDM calibrados\n'
    '✅ Semáforo de dimensionamiento (N=8 paneles/string validado)\n'
    '✅ Flujo de actualización: git pull + pm2 restart (2 comandos)'
)

fname = 'Guia_Despliegue_BIPV_DigitalOcean.docx'
doc.save(fname)
print(f'✅ Documento generado: {fname}')
