# -*- coding: utf-8 -*-
"""
Genera entregables/MANUAL_CALCULADORA_BIPV_v3.4_agosto2026.docx a partir de la
v3.3, agregando el anexo 20: la ruta externa de sombreado con Site Designer
(Andrew Marsh) — cómo dibujar la escena, exportar el JSON, cargarlo en la
página 🌳 Sombras y de ahí seguir la cadena de siempre (Mismatch → bypass →
Producción → Financiero), con la trazabilidad de fuente en los informes.
Mismo formato del manual (estilos, separadores, colores).

Ejecutar:  python3 scripts/actualizar_manual_v34_agosto2026.py
Luego:     python3 bipv_python/scripts/generar_base_conocimiento.py
"""
import copy
import shutil

from docx import Document
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph

SRC = "entregables/MANUAL_CALCULADORA_BIPV_v3.3_agosto2026.docx"
DST = "entregables/MANUAL_CALCULADORA_BIPV_v3.4_agosto2026.docx"
shutil.copy(SRC, DST)
doc = Document(DST)

SEP = "────────────────────────────────────────────────────────────"
AZUL = RGBColor(0x1A, 0x53, 0x76)
NARANJA = RGBColor(0xE6, 0x51, 0x00)
VERDE = RGBColor(0x1B, 0x5E, 0x20)


def find_para(pred):
    for p in doc.paragraphs:
        if pred(p.text.strip()):
            return p
    raise RuntimeError("párrafo no encontrado")


# ── 1. TOC: nueva entrada después de "Preguntas frecuentes" ────────────────
toc_faq = find_para(lambda t: t == "Preguntas frecuentes")
new_toc = copy.deepcopy(toc_faq._p)
toc_faq._p.addnext(new_toc)
p_toc = Paragraph(new_toc, toc_faq._parent)
for r in list(p_toc.runs):
    r._r.getparent().remove(r._r)
p_toc.add_run("Anexo — Sombras desde Site Designer / Andrew Marsh "
              "(ruta externa, agosto 2026)  ")
rn = p_toc.add_run("NUEVO")
rn.bold = True
rn.font.color.rgb = NARANJA

# ── 2. Anexo antes del bloque de cierre ────────────────────────────────────
cierre = find_para(lambda t: t.startswith("Manual actualizado el"))
prev = cierre._p.getprevious()
if prev is not None and prev.tag.endswith("}p") and Paragraph(prev, cierre._parent).text.strip() == SEP:
    anchor = Paragraph(prev, cierre._parent)
else:
    anchor = cierre


def before(text="", style=None, bold=False, color=None):
    p = anchor.insert_paragraph_before(text, style=style)
    if p.runs:
        r = p.runs[0]
        r.bold = bold
        if color:
            r.font.color.rgb = color
    return p


def sep():
    before(SEP)


def h2(t):
    p = before(t, style="Heading 2")
    if p.runs:
        p.runs[0].font.color.rgb = AZUL


def sub(t):
    before(t, bold=True)


def body(t):
    before(t)


def bullet(t):
    before(t, style="List Bullet")


def warn(t):
    p = before(t)
    if p.runs:
        p.runs[0].bold = True
        p.runs[0].font.color.rgb = NARANJA


def tip(t):
    p = before(t)
    if p.runs:
        p.runs[0].font.color.rgb = VERDE


sep()
h2("20. Anexo — Sombras desde Site Designer (Andrew Marsh): la ruta externa")
body("Además de SketchUp y de la Calculadora de Sombreado 3D web, la app acepta escenas "
     "dibujadas en Site Designer, la herramienta gratuita de Andrew Marsh que corre en el "
     "navegador (buscar «Andrew Marsh Site Designer» o «3D Site Designer drajmarsh»). "
     "Es la opción más rápida para modelar los edificios vecinos como cajas simples: se "
     "dibujan los volúmenes, se exporta un archivo JSON y la calculadora hace el resto "
     "con su propio motor solar — el mismo ray-casting oficial de la ruta SketchUp. "
     "Site Designer solo aporta la geometría; los números de sombra los calcula siempre "
     "la calculadora, por eso ambas rutas dan resultados idénticos para la misma escena.")

sep()
sub("20.1 Qué dibujar en Site Designer  NUEVO")
bullet("Primero fija la ubicación del proyecto en Site Designer (latitud/longitud o "
       "buscando la ciudad): esa ubicación viaja dentro del JSON y la calculadora la "
       "compara contra el proyecto activo — si no coincide, avisa.")
bullet("Dibuja SOLO los obstáculos que producen sombra: edificios vecinos, muros, "
       "volúmenes de la propia edificación si sombrean la fachada. Cada obstáculo es un "
       "bloque (caja) con su posición y altura reales.")
bullet("NO dibujes los paneles ni la fachada de estudio: los puntos de análisis se "
       "definen después, dentro de la calculadora (igual que en la ruta SketchUp).")
bullet("Si Site Designer muestra el norte girado (northOffset), déjalo tal cual: el "
       "archivo lo registra y la calculadora aplica la corrección automáticamente.")

sep()
sub("20.2 Exportar el JSON  NUEVO")
bullet("En Site Designer usa la opción de guardar/exportar el proyecto: descarga un "
       "archivo con nombre tipo «site-designer-AAAA-MM-DD-HHMM-SS.json».")
bullet("No edites el JSON a mano. Si le falta información (por ejemplo el norte), la "
       "calculadora lo rechaza con un mensaje claro en vez de asumir valores.")
body("El archivo contiene la ubicación (latitud, longitud, zona horaria, elevación, "
     "corrección de norte) y los obstáculos como cajas en milímetros. La conversión a "
     "metros es automática y fija — no hay selector de unidades que configurar.")

sep()
sub("20.3 Cargarlo en la calculadora  NUEVO")
bullet("Corre primero ☀️ Recurso Solar (el TMY del proyecto es obligatorio para alinear "
       "las horas de sombra con Producción).")
bullet("Abre 🌳 Sombras y sube el archivo .json en el mismo cargador donde va el modelo "
       "de SketchUp. La app confirma cuántos obstáculos leyó, sus dimensiones en metros, "
       "el norte corregido y la ubicación del archivo.")
bullet("Define los puntos de análisis (una fila de módulos = un punto, con sus "
       "coordenadas x, y, z en metros en el mismo sistema de la escena) y pulsa "
       "▶️ Calcular sombras.")
bullet("Envía el resultado a la Página 5 con «📤 Enviar a Mismatch»: de ahí en adelante "
       "la cadena es la de siempre — bypass → E_ac corregida → Producción → Financiero.")
warn("Si la app avisa que la ubicación del archivo no coincide con la del proyecto, "
     "verifica que la escena sea del sitio correcto antes de continuar: una escena de "
     "otro proyecto produce sombras sin sentido físico.")

sep()
sub("20.4 Trazabilidad: el informe dice de dónde salieron las sombras  NUEVO")
body("La fuente del sombreado queda registrada y visible en toda la cadena: en el "
     "resumen del modelo bypass (Página 5), en Producción y en el Reporte PDF aparece "
     "«Fuente del sombreado: SketchUp (interno)», «Site Designer + TMY (externo)» o "
     "«CSV externo». Es un dato informativo — no cambia ningún cálculo — pero le da "
     "credibilidad al informe que se entrega al cliente.")
tip("¿Cuándo usar cada ruta? Site Designer: escenas rápidas de cajas (edificios "
    "vecinos) sin instalar nada. SketchUp: geometrías detalladas, aleros, árboles con "
    "transparencia. Calculadora web: cuando ya existe el análisis punto a punto. Las "
    "tres desembocan en el mismo CSV de FS horario y el mismo modelo bypass.")

doc.save(DST)
print(f"OK → {DST}")
