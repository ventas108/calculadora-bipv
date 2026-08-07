# -*- coding: utf-8 -*-
"""
Genera entregables/MANUAL_CALCULADORA_BIPV_v3.3_agosto2026.docx a partir de la
v3.2, agregando el anexo con las actualizaciones del 6-7 de agosto de 2026:
Asistente, login/planes/pagos, multi-proyectos con privacidad por usuario,
persistencia de resultados, perfil de carga horario, motor óptico-térmico,
coherencia CAPEX Presupuesto→Financiero, reporte PDF con gráficas, robustez
del catálogo de inversores y el nuevo tab solar de la Vista 3D.
Mismo formato del manual (estilos, separadores, colores).

Ejecutar:  python3 scripts/actualizar_manual_v33_agosto2026.py
"""
import copy
import shutil

from docx import Document
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph

SRC = "entregables/MANUAL_CALCULADORA_BIPV_v3.2_agosto2026.docx"
DST = "entregables/MANUAL_CALCULADORA_BIPV_v3.3_agosto2026.docx"
shutil.copy(SRC, DST)
doc = Document(DST)

SEP = "────────────────────────────────────────────────────────────"
AZUL = RGBColor(0x1A, 0x53, 0x76)
NARANJA = RGBColor(0xE6, 0x51, 0x00)
VERDE = RGBColor(0x1B, 0x5E, 0x20)


def find_para(pred):
    for p in doc.paragraphs:
        if pred(p.text.strip()):
            return p
    raise RuntimeError("párrafo no encontrado")


# ── 1. TOC: nueva entrada después de "Preguntas frecuentes" ────────────────
toc_faq = find_para(lambda t: t == "Preguntas frecuentes")
new_toc = copy.deepcopy(toc_faq._p)
toc_faq._p.addnext(new_toc)
p_toc = Paragraph(new_toc, toc_faq._parent)
for r in list(p_toc.runs):
    r._r.getparent().remove(r._r)
p_toc.add_run("Anexo — Actualizaciones 6-7 de agosto 2026 "
              "(Asistente, cuentas, proyectos y Vista 3D solar)  ")
rn = p_toc.add_run("NUEVO")
rn.bold = True
rn.font.color.rgb = NARANJA

# ── 2. Anexo antes del bloque de cierre ────────────────────────────────────
cierre = find_para(lambda t: t.startswith("Manual actualizado el"))
prev = cierre._p.getprevious()
if prev is not None and prev.tag.endswith("}p") and Paragraph(prev, cierre._parent).text.strip() == SEP:
    anchor = Paragraph(prev, cierre._parent)
else:
    anchor = cierre


def before(text="", style=None, bold=False, color=None):
    p = anchor.insert_paragraph_before(text, style=style)
    if p.runs:
        r = p.runs[0]
        r.bold = bold
        if color:
            r.font.color.rgb = color
    return p


def sep():
    before(SEP)


def h2(t):
    p = before(t, style="Heading 2")
    if p.runs:
        p.runs[0].font.color.rgb = AZUL


def sub(t):
    before(t, bold=True)


def body(t):
    before(t)


def bullet(t):
    before(t, style="List Bullet")


def warn(t):
    p = before(t)
    if p.runs:
        p.runs[0].bold = True
        p.runs[0].font.color.rgb = NARANJA


def tip(t):
    p = before(t)
    if p.runs:
        p.runs[0].font.color.rgb = VERDE


sep()
h2("18. Anexo — Actualizaciones del 6 y 7 de agosto de 2026")
body("Esta entrega convierte la calculadora en una aplicación multiusuario: cada persona "
     "entra con su cuenta, guarda varios proyectos privados y no pierde el trabajo al cerrar "
     "la pestaña. Además se suma el Asistente 🧭 que guía el flujo paso a paso, un tab de "
     "análisis solar en la Vista 3D para elegir la mejor orientación, y mejoras de robustez "
     "en catálogos, presupuesto, financiero y reporte PDF.")

# 18.1 Login
sep()
sub("18.1 Acceso con cuenta: login, planes y panel de administración  NUEVO")
body("La app ahora pide iniciar sesión con correo y contraseña antes de usar cualquier "
     "página. Cada cuenta tiene un plan con fecha de vencimiento (por ejemplo, prueba de "
     "14 días, mensual o anual); al vencer, la app muestra la pantalla de renovación.")
bullet("El administrador gestiona todo desde la página 17 — Administración: crear usuarios, "
       "asignar plan y vigencia, extender o revocar accesos y cerrar sesiones activas.")
bullet("En la pantalla de renovación aparecen los botones de pago configurados por el "
       "administrador: link de pago Wompi (mensual o anual) y/o datos de transferencia "
       "bancaria. Solo se aceptan links de dominios oficiales de Wompi, como protección "
       "contra suplantación.")
bullet("La primera vez que se instala en un servidor nuevo, el primer administrador se crea "
       "con un código de configuración de un solo uso (no viaja en el repositorio).")
warn("⚠️ Para no cometer errores: cada persona debe usar su propia cuenta. Los proyectos y "
     "resultados guardados son privados por usuario — si dos personas comparten una cuenta, "
     "se pisarán los datos entre sí.")

# 18.2 Asistente
sep()
sub("18.2 Página 0 — 🧭 Asistente: guía paso a paso y chat con el manual  NUEVO")
body("Nueva primera página del menú. Tiene dos partes:")
bullet("Guía del flujo: un checklist en vivo que detecta qué pasos ya completaste en esta "
       "sesión (proyecto definido, recurso solar descargado, dimensionamiento, producción, "
       "financiero, presupuesto, reporte) y te dice cuál es el siguiente paso y en qué página "
       "está.")
bullet("Chat del manual: puedes preguntar en lenguaje natural (\"¿cómo cargo el CSV de "
       "sombras?\", \"¿qué significa el PR?\") y el asistente responde usando este mismo "
       "Manual de Usuario como fuente. Requiere que el administrador haya configurado una "
       "clave de IA en el servidor; si no hay clave, la guía paso a paso funciona igual.")
tip("Consejo: si te pierdes en el flujo, abre el Asistente — el checklist te muestra "
    "exactamente qué falta y en qué orden.")

# 18.3 Multi-proyectos
sep()
sub("18.3 Página 1 — Proyecto: varios proyectos guardados, privados por usuario  NUEVO")
body("Ahora puedes guardar varios proyectos con nombre y alternar entre ellos sin perder "
     "nada. En la parte superior de la Página 1 está el selector: Guardar proyecto actual, "
     "Cargar y Eliminar.")
bullet("El proyecto activo se muestra en la barra lateral de todas las páginas, para que "
       "siempre sepas sobre qué proyecto estás trabajando.")
bullet("Privacidad: cada proyecto queda amarrado a la cuenta que lo guardó. Otros usuarios "
       "no pueden verlo, cargarlo ni borrarlo. Los proyectos guardados antes de esta versión "
       "solo los ve el administrador; al volver a guardarlos quedan asociados a su dueño.")
bullet("Al cargar un proyecto, la app pide re-ejecutar Recurso Solar (banner de pasos "
       "pendientes). Si la ciudad y las coordenadas no cambiaron, la descarga se revalida "
       "sola desde el caché en segundos.")
warn("⚠️ Para no cometer errores: cargar un proyecto NO revive simulaciones viejas — es a "
     "propósito, para que Producción y Financiero nunca muestren números de otra corrida. "
     "Sigue el banner de pasos pendientes en orden.")

# 18.4 Persistencia
sep()
sub("18.4 Tu trabajo sobrevive recargas y pestañas nuevas  NUEVO")
body("Los resultados importantes ya no viven solo en la pestaña del navegador: se guardan en "
     "el servidor, en archivos privados de tu cuenta, y se restauran al volver a entrar.")
bullet("Producción Anual y Análisis Financiero: al abrir la página en una sesión nueva, si "
       "hay resultados guardados aparece un banner \"restaurado de la sesión anterior\" con "
       "la fecha. Antes de restaurar, la app verifica que la ciudad y las coordenadas "
       "coincidan con las guardadas.")
bullet("Consumo energético: el perfil de consumo y el modo de entrada que usaste se "
       "recuerdan automáticamente.")
bullet("Presupuesto: las tablas editadas (ítems, precios, activos/inactivos, costos blandos, "
       "OPEX) se guardan en disco por proyecto y por usuario, y vuelven al recargar.")
tip("Consejo: igual conviene oprimir los botones de guardar del Presupuesto después de "
    "ediciones grandes — el guardado explícito es inmediato.")

# 18.5 Perfil de carga horario
sep()
sub("18.5 Página 11 — Baterías: perfil de carga horario real  NUEVO")
body("El balance energético con batería ahora puede usar un perfil de consumo hora a hora "
     "(8 760 valores) en lugar de un promedio plano. Esto cambia mucho el resultado en "
     "edificios con consumo concentrado en el día o en la noche: el autoconsumo, los ciclos "
     "de la batería y el ahorro se calculan contra el consumo real de cada hora.")
bullet("Puedes elegir entre perfiles típicos (residencial, comercial, industrial) o subir tu "
       "propio CSV horario.")
warn("⚠️ Para no cometer errores: si tienes la curva real de tu operador de red, úsala — el "
     "perfil típico es una aproximación y puede sobrestimar el autoconsumo.")

# 18.6 Motor óptico-térmico
sep()
sub("18.6 Motor Óptico: efecto térmico BIPV integrado sin doble conteo  ACTUALIZADO")
body("El sobrecalentamiento típico de los paneles integrados a fachada (k_bipv) ahora entra "
     "directamente al modelo de temperatura de celda y al modelo eléctrico del panel, en un "
     "solo lugar. Antes existía el riesgo de descontar el efecto térmico dos veces.")
bullet("Además, la POA corregida por el Motor Óptico (IAM + soiling) fluye automáticamente a "
       "Dimensionamiento y a Producción — ya no hay que activar nada manualmente.")

# 18.7 Financiero
sep()
sub("18.7 Página 7 — Financiero: CAPEX del Presupuesto y ahorro de batería  ACTUALIZADO")
bullet("CAPEX como fuente única: el análisis financiero toma automáticamente el CAPEX Total "
       "del Presupuesto (con su nivel de contingencia), y muestra la fuente y la fecha de "
       "última actualización. Un toggle permite desvincularlo si quieres probar un CAPEX "
       "manual.")
bullet("El OPEX anual del Presupuesto entra como valor absoluto coherente en el VPN y la "
       "TIR.")
bullet("El ahorro que genera la batería (autoconsumo nocturno) ahora sí se incluye en la "
       "TIR y el Payback, no solo en la gráfica de balance.")
warn("⚠️ Para no cometer errores: si cambias precios en el Presupuesto, vuelve a abrir "
     "Financiero para que el CAPEX vinculado se refresque, y revisa la fecha de última "
     "actualización que aparece junto al valor.")

# 18.8 Presupuesto/cotización
sep()
sub("18.8 Página 8 — Cotización exportada: cifras coherentes y celdas seguras  ACTUALIZADO")
bullet("El total en USD de la cotización (Excel y PDF) ahora se deriva exactamente del mismo "
       "total en COP de los ítems, con la TRM vigente — ya no pueden salir dos cifras de "
       "bases distintas en el mismo documento.")
bullet("El Excel exportado neutraliza textos que empiecen con símbolos de fórmula (=, +, -, "
       "@): protección estándar contra archivos maliciosos al compartir cotizaciones.")

# 18.9 Reporte PDF
sep()
sub("18.9 Página 10 — Reporte PDF: gráficas, logo y trazabilidad  ACTUALIZADO")
bullet("Nueva gráfica de producción mensual (barras) y nueva curva de flujo de caja "
       "acumulado con el año de payback marcado.")
bullet("Encabezado con el logo y los datos de contacto de tu empresa (se configuran una vez "
       "en el Presupuesto y se reutilizan).")
bullet("El reporte indica qué tasa de degradación se usó (historial PR real vs slider "
       "manual) y la zona horaria del análisis.")

# 18.10 Catálogo de inversores
sep()
sub("18.10 Catálogo de Inversores: diagnóstico, confianza y recarga automática  ACTUALIZADO")
bullet("Dimensionamiento muestra un semáforo de salud del catálogo de inversores (filas "
       "válidas, campos críticos vacíos, duplicados) con botón de recarga.")
bullet("Al extraer una ficha PDF, cada campo queda marcado con su nivel de confianza; antes "
       "de sobrescribir un inversor existente la app muestra un diff campo por campo y pide "
       "confirmación.")
bullet("Si reemplazas el Excel del catálogo en el servidor, la app lo detecta y recarga sola "
       "— ya no hay que esperar una hora ni reiniciar.")
bullet("En Dimensionamiento también aparece un banner cuando el panel elegido del catálogo "
       "tiene confianza distinta de Alta: sus dimensiones son estimadas y conviene "
       "confirmarlas con el fabricante antes de cerrar el diseño.")

# 18.11 Vista 3D solar
sep()
sub("18.11 Página 9 — Vista 3D: nuevo análisis solar y de orientación  NUEVO")
body("El tab solar de la Vista 3D ahora responde la pregunta clave: ¿está bien orientada mi "
     "fachada, y cuánto ganaría si la girara?")
bullet("Diagrama solar con líneas iso-hora: sobre las curvas del recorrido del sol se dibujan "
       "líneas punteadas de 07:00 a 17:00 hora local, para leer de un vistazo a qué hora el "
       "sol pasa por cada posición.")
bullet("Comparación de AOI mensual entre superficies: un multiselect permite poner lado a "
       "lado el ángulo de incidencia promedio de varias fachadas, con colores semáforo "
       "(verde < 40°, naranja < 60°, rojo ≥ 60°). Menor AOI = luz más perpendicular = más "
       "energía.")
bullet("🧭 Orientación de mejor incidencia (geométrica): barrido de azimuth cada 5° que "
       "sugiere hacia dónde girar la superficie para mejorar el ángulo de incidencia, "
       "comparando solo orientaciones con horas de sol equiparables e indicando las horas de "
       "sol directo al año.")
bullet("⚡ Orientación de máxima energía real (con TMY): barrido de 72 orientaciones que "
       "calcula la irradiación anual real (kWh/m²) de cada azimuth usando el clima del TMY y "
       "el horizonte de obstáculos del proyecto, y muestra el azimuth óptimo, la mejora "
       "porcentual y cuántos grados habría que girar. Incluye gráfica POA vs azimuth con la "
       "orientación actual y la óptima marcadas.")
tip("Consejo: usa primero la métrica de energía real (⚡) para decidir — es la que manda, "
    "porque incluye nubes y horizonte. La geométrica (🧭) sirve para entender el porqué.")
warn("⚠️ Para no cometer errores: estos barridos evalúan la orientación manteniendo la misma "
     "inclinación. En BIPV de fachada muchas veces la orientación viene dada por el edificio; "
     "usa el resultado como criterio para elegir ENTRE fachadas candidatas.")

# 18.12 CSV meses en texto
sep()
sub("18.12 Página 5 — Mismatch: el CSV acepta meses en texto  ACTUALIZADO")
body("El CSV de Factor de Sombreado ahora puede traer los meses escritos (Ene, Feb, Mar… o "
     "January, February…) además de números — es el formato que exporta la Calculadora de "
     "Sombreado 3D web. La app los reconoce automáticamente.")

sep()

# ── 3. Actualizar bloque de cierre ─────────────────────────────────────────
cierre.text = "Manual actualizado el 7 de agosto de 2026"
nov = Paragraph(cierre._p.getnext(), cierre._parent)
if nov.text.startswith("Novedades"):
    nov.text = ("Novedades de esta versión: Asistente 🧭 con guía paso a paso y chat del manual, "
                "login con planes y pagos (Wompi/transferencia), múltiples proyectos privados por "
                "usuario, persistencia de resultados y presupuesto, perfil de carga horario, motor "
                "óptico-térmico sin doble conteo, CAPEX vinculado al Financiero, reporte PDF con "
                "gráficas y logo, diagnóstico del catálogo de inversores y análisis de orientación "
                "en la Vista 3D (incidencia geométrica y energía real con TMY).")

doc.save(DST)
print("Generado", DST)
