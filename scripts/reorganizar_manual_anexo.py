"""
Reorganiza el manual: mueve las secciones del Anexo 18 al final del capítulo
correspondiente, elimina el anexo y su entrada en la tabla de contenido, y
marca los capítulos tocados como ACTUALIZADO en el TOC.

Ejecutar:  python3 scripts/reorganizar_manual_anexo.py
"""
import shutil
from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import RGBColor

SRC = "entregables/MANUAL_CALCULADORA_BIPV_v3_agosto2026.docx"
shutil.copy(SRC, SRC + ".bak3")
doc = Document(SRC)
SEP = "────────────────────────────────────────────────────────────"
NARANJA = RGBColor(0xE6, 0x51, 0x00)

body = doc.paragraphs[0]._parent  # documento

def paras():
    return doc.paragraphs

def find(pred, style=None):
    for p in paras():
        if (style is None or p.style.name == style) and pred(p.text.strip()):
            return p
    raise RuntimeError("no encontrado")

def block_of(prefix):
    """Devuelve [SEP, sub, ...cuerpo] del bloque 18.x (hasta el próximo SEP, excl.)."""
    sub = find(lambda t: t.startswith(prefix))
    els = []
    prev = sub._p.getprevious()
    if prev is not None and Paragraph(prev, body).text.strip() == SEP:
        els.append(prev)
    els.append(sub._p)
    nxt = sub._p.getnext()
    while nxt is not None and Paragraph(nxt, body).text.strip() != SEP:
        els.append(nxt)
        nxt = nxt.getnext()
    return els

def chapter_end_anchor(next_heading_prefix):
    """Elemento ante el cual insertar: el SEP inmediatamente antes del
    encabezado del capítulo siguiente (o el encabezado mismo)."""
    h = find(lambda t: t.startswith(next_heading_prefix), style="Heading 2")
    prev = h._p.getprevious()
    if prev is not None and Paragraph(prev, body).text.strip() == SEP:
        return prev
    return h._p

MOVES = [
    # (prefijo bloque, prefijo del capítulo SIGUIENTE al destino, nuevo título)
    ("18.1", "4. Página 2",  "Factor de ocupación con paneles — agrivoltaica  NUEVO (5-ago-2026)"),
    ("18.2", "5. Página 3",  "GCR sincronizado con el factor de ocupación  NUEVO (5-ago-2026)"),
    ("18.5", "6. Página 4",  "Activación automática y defensa Ns half-cut  NUEVO (5-ago-2026)"),
    ("18.3", "9. Página 6",  "Modo Granja agrivoltaica  NUEVO (5-ago-2026)"),
    ("18.4", "11. Página 8", "Precio real del inversor del catálogo  NUEVO (5-ago-2026)"),
    ("18.6", "3. Página 1",  "Flujo recomendado para proyectos agrivoltaicos  NUEVO (5-ago-2026)"),
]

for pref, next_ch, titulo in MOVES:
    els = block_of(pref)
    anchor = chapter_end_anchor(next_ch)
    # retitular el sub-encabezado (segundo elemento)
    sub = Paragraph(els[1], body)
    sub.runs[0].text = titulo
    for r in sub.runs[1:]:
        r.text = ""
    for el in els:
        el.getparent().remove(el)
        anchor.addprevious(el)

# ── eliminar lo que queda del anexo: SEP + H2 18 + intro ──────────────────
h18 = find(lambda t: t.startswith("18. Anexo"), style="Heading 2")
els = []
prev = h18._p.getprevious()
if prev is not None and Paragraph(prev, body).text.strip() == SEP:
    els.append(prev)
els.append(h18._p)
nxt = h18._p.getnext()
while nxt is not None and Paragraph(nxt, body).text.strip() != SEP:
    els.append(nxt)
    nxt = nxt.getnext()
for el in els:
    el.getparent().remove(el)

# ── TOC: quitar entrada del anexo y marcar capítulos actualizados ─────────
toc_anexo = find(lambda t: t.startswith("Anexo — Actualizaciones"))
toc_anexo._p.getparent().remove(toc_anexo._p)

def toc_mark(texto_toc):
    try:
        p = find(lambda t: t == texto_toc, style="List Number")
    except RuntimeError:
        return
    r = p.add_run("  ACTUALIZADO")
    r.bold = True
    r.font.color.rgb = NARANJA

toc_mark("Página 3 — Motor IV")
toc_mark("Página 7 — Análisis Financiero")   # por si no tenía etiqueta
toc_mark("Flujo de trabajo recomendado")

doc.save(SRC)
print("OK — anexo reorganizado en capítulos")
