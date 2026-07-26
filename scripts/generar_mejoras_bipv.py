
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Paleta de colores ──
AZUL      = RGBColor(0x1A, 0x5C, 0x8A)
VERDE     = RGBColor(0x17, 0x6B, 0x17)
MORADO    = RGBColor(0x6E, 0x27, 0x94)
ROJO      = RGBColor(0xC0, 0x39, 0x2B)
NARANJA   = RGBColor(0xD4, 0x7A, 0x00)
AZUL_COD  = RGBColor(0x10, 0x10, 0x60)

# ═══════════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════════
t = doc.add_heading('DIAGNÓSTICO Y MEJORAS — CALCULADORA BIPV', 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.runs[0].font.color.rgb = AZUL

st = doc.add_paragraph('Análisis técnico basado en el Algoritmo paso a paso BIPV')
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
st.runs[0].font.bold = True
st.runs[0].font.size = Pt(12)
st.runs[0].font.color.rgb = MORADO

repo = doc.add_paragraph('Repositorio: github.com/ventas108/calculadora-bipv  |  Stack: React + TypeScript + Vite')
repo.alignment = WD_ALIGN_PARAGRAPH.CENTER
repo.runs[0].font.size = Pt(10)
repo.runs[0].font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
repo.runs[0].italic = True
doc.add_paragraph('')

intro = doc.add_paragraph(
    'Se realizó una lectura completa del código fuente del repositorio. '
    'Este documento registra: (1) qué módulos del algoritmo BIPV ya están '
    'implementados y a qué nivel de detalle, (2) qué módulos faltan o están '
    'incompletos, y (3) el código TypeScript/JavaScript exacto para cada mejora, '
    'indicando en qué archivo del proyecto debe integrarse.'
)
intro.runs[0].font.size = Pt(10)
intro.runs[0].italic = True
intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
doc.add_paragraph('')

# ── Helpers ──
def sep():
    p = doc.add_paragraph('─' * 78)
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

def h1(txt, color=AZUL):
    doc.add_paragraph('')
    h = doc.add_heading(txt, level=1)
    h.runs[0].font.size = Pt(14)
    h.runs[0].font.color.rgb = color
    doc.add_paragraph('')

def h2(txt, color=MORADO):
    h = doc.add_heading(txt, level=2)
    h.runs[0].font.size = Pt(12)
    h.runs[0].font.color.rgb = color

def h3(txt, color=NARANJA):
    h = doc.add_heading(txt, level=3)
    h.runs[0].font.size = Pt(11)
    h.runs[0].font.color.rgb = color

def body(txt):
    p = doc.add_paragraph(txt)
    p.runs[0].font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def badge(label, txt, color):
    p = doc.add_paragraph()
    r1 = p.add_run(f'{label}  ')
    r1.bold = True
    r1.font.color.rgb = color
    p.add_run(txt).font.size = Pt(10)

def codigo(txt, lang='typescript'):
    p = doc.add_paragraph()
    cr = p.add_run(txt)
    cr.font.name = 'Courier New'
    cr.font.size = Pt(9)
    cr.font.color.rgb = AZUL_COD
    doc.add_paragraph('')

def archivo(ruta):
    p = doc.add_paragraph()
    r = p.add_run('📁 Archivo:  ')
    r.bold = True
    r.font.color.rgb = VERDE
    cr = p.add_run(ruta)
    cr.font.name = 'Courier New'
    cr.font.size = Pt(9.5)
    cr.font.color.rgb = VERDE

# ═══════════════════════════════════════════════════════════════
# SECCIÓN 1 — QUÉ TIENE LA CALCULADORA
# ═══════════════════════════════════════════════════════════════
h1('SECCIÓN 1 — LO QUE LA CALCULADORA YA TIENE (MUY BIEN IMPLEMENTADO)')

body(
    'Tras leer el código completo, la calculadora supera en varios aspectos '
    'al algoritmo básico propuesto. Los siguientes módulos están implementados '
    'con un nivel de detalle profesional:'
)
doc.add_paragraph('')

items_ok = [
    ('✅ MÓDULO 5 — Producción energética IEC 61724-1:2021',
     'energyProduction.ts',
     'Implementado completamente: Yr, Yf, Ya, Lc, Ls, PR, PR_T (corregido temperatura), '
     'CF (Factor de Capacidad), BOS efficiency, EPI (Energy Performance Index). '
     'Incluye desglose de capture losses por categoría. Referencia normativa correcta.'),

    ('✅ MÓDULO 4 — Temperatura BIPV con modelo térmico confinado',
     'iamSoilingEngine.ts',
     'Implementado con k_bipv = 1.3 para fachadas confinadas (sin ventilación). '
     'El factor multiplica el incremento de temperatura sobre el NOCT estándar. '
     'Es exactamente la penalidad de +10-15°C del algoritmo, modelada con precisión.'),

    ('✅ MÓDULO 6 — Análisis de sombras avanzado',
     'ShadingCalculator.tsx + facadeShadingAnalysis.ts',
     'Importación de modelos 3D (GLTF, OBJ, Marsh Site Designer), diagrama solar '
     'con polígonos de obstáculos interactivos, trayectoria solar hora a hora, '
     'crossing modal para sombras mutuas entre fachadas. Supera al algoritmo propuesto.'),

    ('✅ MÓDULO 2 — Recurso solar con transposición anisótropa',
     'iamSoilingEngine.ts + OrientationOptimizer.tsx',
     'Implementado modelo de Perez (anisótropo) y Liu-Jordan (isótropo) para '
     'transposición de irradiancia al plano inclinado. Integración directa con datos '
     'horarios EPW, PVWatts v8 (NREL) y PVGIS. Supera al algoritmo propuesto.'),

    ('✅ IAM-ASHRAE + Soiling estacional',
     'iamSoilingEngine.ts',
     'IAM = 1 − b0×(1/cos(AOI)−1) con corrección hora a hora. Soiling mensual '
     'con autolavado por agua precipitable. No estaba en el algoritmo — es una mejora propia.'),

    ('✅ MÓDULO 9 — Análisis financiero completo',
     'bipvROIOptimizer.ts',
     'LCOE, VAN (NPV), TIR (IRR), payback, ahorro HVAC por reducción SHGC, '
     'incentivos Ley 1715 Colombia (deducción 50%, depreciación acelerada, '
     'exención IVA y aranceles), análisis de sensibilidad multi-variable.'),

    ('✅ MÓDULO 3 — Catálogo de tecnologías BIPV',
     'bipvGlassCatalog.ts + panelTechnologies.ts',
     'Catálogo 1G/2G/3G con vidrio semitransparente CdTe, HJT, mono-Si, poli-Si, CIGS. '
     'Niveles de transparencia (10-60%). Importador de fichas técnicas PDF. '
     'Paneles personalizados por el usuario.'),

    ('✅ Diagnóstico de rendimiento IEC 61724',
     'performanceDiagnostic.ts',
     'Sistema de alertas ok/leve/moderada/severa/crítica basado en desviación del PR. '
     'Causas probables ponderadas con categoría (ambiental/equipo/instalación). '
     'Historial de mediciones de campo.'),

    ('✅ Modelo Mulcue-Llanos',
     'mulcueLlanos.ts + shared/mulcueLlanos.ts',
     'PR_max = K_sist × (1 + γ × (1.12×T_a − 10)), PR_C corregido por temperatura, '
     'T_ref = 21°C (no STC 25°C), P_exp con ajuste por irradiancia. '
     'Validado contra ejercicios del curso BIPV Global.'),
]

for titulo, archivo_ref, descripcion in items_ok:
    h2(titulo, color=VERDE)
    p_arch = doc.add_paragraph()
    r_a = p_arch.add_run('📁 ')
    r_a.font.color.rgb = VERDE
    cr_a = p_arch.add_run(archivo_ref)
    cr_a.font.name = 'Courier New'
    cr_a.font.size = Pt(9.5)
    cr_a.font.color.rgb = VERDE
    body(descripcion)
    doc.add_paragraph('')

# ═══════════════════════════════════════════════════════════════
# SECCIÓN 2 — QUÉ FALTA
# ═══════════════════════════════════════════════════════════════
sep()
h1('SECCIÓN 2 — MÓDULOS FALTANTES O INCOMPLETOS (Con código de mejora)')

body(
    'Los siguientes módulos del algoritmo BIPV no se encontraron en el repositorio '
    'o están implementados de forma parcial. Para cada uno se entrega: '
    'descripción del problema, impacto en los cálculos, y el código TypeScript '
    'exacto a agregar, con el nombre del archivo destino.'
)
doc.add_paragraph('')

# ──────────────────────────────────────────────────────────────
# MEJORA 1 — DIMENSIONADO ELÉCTRICO DE STRINGS
# ──────────────────────────────────────────────────────────────
sep()
h1('MEJORA 1 — Dimensionado eléctrico de strings (MÓDULO 7 faltante)', ROJO)

badge('⚠ PROBLEMA:', 
      'No existe ningún módulo de sizing eléctrico. El usuario puede definir '
      'cuántos paneles tiene, pero el sistema nunca valida si la combinación '
      'serie/paralelo es compatible con el inversor seleccionado. '
      'Un error de sizing puede dañar el inversor o reducir la producción hasta un 30%.', ROJO)
doc.add_paragraph('')

badge('📊 IMPACTO:', 
      'Sin string sizing, el ratio DC/AC no se verifica, las tensiones Voc en frío '
      'pueden superar el límite del inversor (riesgo de falla), y el MPPT puede quedar '
      'fuera de rango. PVsyst siempre valida esto — la calculadora BIPV no.', NARANJA)
doc.add_paragraph('')

h2('Código a agregar — nuevo archivo:')
archivo('client/src/lib/stringSizing.ts')
codigo(
    '// ═══════════════════════════════════════════════════════════\n'
    '// MÓDULO: Dimensionado eléctrico de strings BIPV\n'
    '// Archivo: client/src/lib/stringSizing.ts\n'
    '// ═══════════════════════════════════════════════════════════\n\n'
    'export interface PanelElectricalParams {\n'
    '  Voc_stc: number;       // Tensión circuito abierto STC (V)\n'
    '  Vmp_stc: number;       // Tensión máxima potencia STC (V)\n'
    '  Isc_stc: number;       // Corriente cortocircuito STC (A)\n'
    '  Imp_stc: number;       // Corriente máxima potencia STC (A)\n'
    '  coef_v: number;        // Coef. temperatura tensión (%/°C, negativo)\n'
    '  coef_i: number;        // Coef. temperatura corriente (%/°C, positivo)\n'
    '}\n\n'
    'export interface InverterElectricalParams {\n'
    '  P_ac_kW: number;       // Potencia AC nominal (kW)\n'
    '  Vdc_max: number;       // Tensión DC máxima admitida (V)\n'
    '  Vmppt_min: number;     // Tensión mínima rango MPPT (V)\n'
    '  Vmppt_max: number;     // Tensión máxima rango MPPT (V)\n'
    '  Idc_max: number;       // Corriente DC máxima (A)\n'
    '  efic: number;          // Eficiencia del inversor (0-1)\n'
    '}\n\n'
    'export interface StringSizingResult {\n'
    '  n_serie_min: number;   // Mínimo paneles en serie (MPPT_min / Vmp_caliente)\n'
    '  n_serie_max: number;   // Máximo paneles en serie (Vdc_max / Voc_frio)\n'
    '  n_serie_opt: number;   // Óptimo dentro del rango MPPT\n'
    '  n_strings_max: number; // Máximo strings en paralelo\n'
    '  Voc_frio: number;      // Voc corregido a T_min (V)\n'
    '  Vmp_caliente: number;  // Vmp corregido a T_max (V)\n'
    '  ratio_dc_ac: number;   // Ratio DC/AC del sistema\n'
    '  ratio_ok: boolean;     // true si 1.10 ≤ ratio ≤ 1.30\n'
    '  warnings: string[];    // Advertencias de diseño\n'
    '}\n\n'
    'export function calcularStringSizing(\n'
    '  panel: PanelElectricalParams,\n'
    '  inversor: InverterElectricalParams,\n'
    '  potencia_dc_kWp: number,\n'
    '  T_min: number = -5,    // Temperatura mínima registrada del lugar (°C)\n'
    '  T_max: number = 70     // Temperatura máxima de celda en verano (°C)\n'
    '): StringSizingResult {\n'
    '  const warnings: string[] = [];\n\n'
    '  // Corrección de tensiones por temperatura\n'
    '  const Voc_frio     = panel.Voc_stc * (1 + (panel.coef_v / 100) * (T_min - 25));\n'
    '  const Vmp_caliente = panel.Vmp_stc * (1 + (panel.coef_v / 100) * (T_max - 25));\n\n'
    '  // Límites de strings en serie\n'
    '  const n_serie_max = Math.floor(inversor.Vdc_max    / Voc_frio);\n'
    '  const n_serie_min = Math.ceil(inversor.Vmppt_min   / Vmp_caliente);\n'
    '  const n_serie_opt = Math.floor(inversor.Vmppt_max  / panel.Vmp_stc);\n\n'
    '  // Strings en paralelo\n'
    '  const n_strings_max = Math.floor(inversor.Idc_max  / panel.Isc_stc);\n\n'
    '  // Ratio DC/AC\n'
    '  const ratio_dc_ac = potencia_dc_kWp / inversor.P_ac_kW;\n'
    '  const ratio_ok    = ratio_dc_ac >= 1.10 && ratio_dc_ac <= 1.30;\n\n'
    '  // Advertencias\n'
    '  if (n_serie_min > n_serie_max)\n'
    '    warnings.push("CRÍTICO: No hay rango válido de paneles en serie. Cambiar inversor.");\n'
    '  if (!ratio_ok)\n'
    '    warnings.push(`Ratio DC/AC = ${ratio_dc_ac.toFixed(2)} fuera del rango óptimo 1.10-1.30.`);\n'
    '  if (Voc_frio * n_serie_max > inversor.Vdc_max)\n'
    '    warnings.push("Riesgo: Voc en frío supera Vdc_max del inversor.");\n\n'
    '  return {\n'
    '    n_serie_min, n_serie_max, n_serie_opt,\n'
    '    n_strings_max, Voc_frio, Vmp_caliente,\n'
    '    ratio_dc_ac, ratio_ok, warnings\n'
    '  };\n'
    '}'
)

h2('Dónde integrar en la interfaz:')
archivo('client/src/components/EnergyProductionSimulator.tsx')
body(
    'Agregar una nueva sección "Dimensionado Eléctrico" dentro del Simulador de Energía, '
    'después del selector de paneles y antes del cálculo de producción. '
    'Mostrar una tarjeta de validación con semáforo: verde si el string sizing es correcto, '
    'rojo si hay advertencias. Usar los datos del panel seleccionado (panelTechnologies.ts) '
    'y añadir campos de entrada para los parámetros del inversor.'
)
doc.add_paragraph('')

# ──────────────────────────────────────────────────────────────
# MEJORA 2 — FACTOR DE COBERTURA
# ──────────────────────────────────────────────────────────────
sep()
h1('MEJORA 2 — Factor de cobertura activa vs área total de fachada', ROJO)

badge('⚠ PROBLEMA:',
      'El sistema calcula la potencia como Área × eficiencia, pero en BIPV real '
      'no toda el área de la fachada tiene celdas activas. Los marcos, juntas, '
      'secciones opacas y módulos con transparencia reducen el área generadora real. '
      'Sin este factor, la potencia calculada puede estar sobrestimada un 10-20%.', ROJO)
doc.add_paragraph('')

h2('Código a agregar — modificar función existente:')
archivo('client/src/lib/iamSoilingEngine.ts  →  función calculateBIPVPower()')
codigo(
    '// ─── AGREGAR este parámetro a BIPVSimulationConfig ───────────────────\n'
    'export interface BIPVSimulationConfig {\n'
    '  // ... campos existentes ...\n'
    '  /** Factor de cobertura: fracción del área total con celdas activas (0-1).\n'
    '   *  Fachada continua → 0.90  |  Ventanas/vidrio → 0.60-0.75\n'
    '   *  Muro cortina con marcos → 0.70-0.85  |  Por defecto: 0.85 */\n'
    '  factorCobertura?: number;   // NUEVO — agregar aquí\n'
    '}\n\n'
    '// ─── MODIFICAR calculateBIPVPower() ──────────────────────────────────\n'
    'export function calculateBIPVPower(\n'
    '  poaTotalOptica: number,     // W/m² — POA después de IAM+soiling\n'
    '  areaM2: number,             // m² — área TOTAL de fachada\n'
    '  factorCobertura: number = 0.85,  // NUEVO parámetro\n'
    '  eficienciaAjustada: number, // eficiencia real (0-1)\n'
    '  transparencia: number       // fracción de transparencia (0-1)\n'
    '): number {\n'
    '  // Área activa real con celdas generadoras\n'
    '  const areaActiva = areaM2 * factorCobertura;  // ← NUEVO\n\n'
    '  // La fracción transparente no genera — solo la opaca genera energía\n'
    '  const fraccionOpaca = 1 - transparencia;\n\n'
    '  // Potencia DC = POA × área_activa × eficiencia × fracción_opaca\n'
    '  return poaTotalOptica * areaActiva * eficienciaAjustada * fraccionOpaca;\n'
    '}'
)
doc.add_paragraph('')

# ──────────────────────────────────────────────────────────────
# MEJORA 3 — BALANCE MENSUAL GENERACIÓN VS CONSUMO
# ──────────────────────────────────────────────────────────────
sep()
h1('MEJORA 3 — Balance energético mensual (generación vs consumo del edificio)', ROJO)

badge('⚠ PROBLEMA:',
      'El ROI Optimizer usa un porcentaje fijo de autoconsumo (ej: 70%). '
      'Pero en la realidad el autoconsumo varía mes a mes: en verano el BIPV '
      'genera más de lo que necesita el edificio (excedente a la red), y en invierno '
      'puede faltar energía (compra de red). Sin este cálculo, el VAN y el payback '
      'pueden estar distorsionados hasta en un 15%.', ROJO)
doc.add_paragraph('')

h2('Código a agregar — nuevo archivo:')
archivo('client/src/lib/energyBalance.ts')
codigo(
    '// ═══════════════════════════════════════════════════════════\n'
    '// MÓDULO: Balance energético mensual BIPV\n'
    '// Archivo: client/src/lib/energyBalance.ts\n'
    '// ═══════════════════════════════════════════════════════════\n\n'
    'export interface MonthlyEnergyBalance {\n'
    '  month: string;\n'
    '  generacion_kwh: number;    // kWh generados por BIPV ese mes\n'
    '  consumo_kwh: number;       // kWh consumidos por el edificio ese mes\n'
    '  autoconsumo_kwh: number;   // kWh usados directamente del BIPV\n'
    '  excedente_kwh: number;     // kWh inyectados a la red\n'
    '  compra_red_kwh: number;    // kWh comprados a la red\n'
    '  tasa_autoconsumo_pct: number;  // % de la generación auto-consumida\n'
    '  tasa_autosuficiencia_pct: number; // % del consumo cubierto por BIPV\n'
    '  ahorro_mes: number;        // $ ahorrados ese mes\n'
    '}\n\n'
    'export interface EnergyBalanceResult {\n'
    '  monthly: MonthlyEnergyBalance[];\n'
    '  annual: {\n'
    '    generacion_total_kwh: number;\n'
    '    consumo_total_kwh: number;\n'
    '    autoconsumo_pct_promedio: number;\n'
    '    autosuficiencia_pct_promedio: number;\n'
    '    excedente_total_kwh: number;\n'
    '    compra_red_total_kwh: number;\n'
    '    ahorro_total: number;\n'
    '    clasificacion_energetica: "A+" | "A" | "B" | "C" | "D";\n'
    '  };\n'
    '}\n\n'
    'const MONTHS = ["Ene","Feb","Mar","Abr","May","Jun",\n'
    '                "Jul","Ago","Sep","Oct","Nov","Dic"];\n\n'
    'export function calcularBalanceEnergetico(\n'
    '  generacionMensualKwh: number[],  // 12 valores del simulador BIPV\n'
    '  consumoEdificioMensualKwh: number[], // 12 valores del edificio\n'
    '  tarifaCompraCLP: number,         // $/kWh tarifa eléctrica\n'
    '  tarifaVentaCLP: number           // $/kWh precio inyección a red\n'
    '): EnergyBalanceResult {\n'
    '  const monthly: MonthlyEnergyBalance[] = [];\n'
    '  let totalGen = 0, totalCons = 0, totalExc = 0, totalRed = 0, totalAhorro = 0;\n\n'
    '  for (let i = 0; i < 12; i++) {\n'
    '    const gen  = generacionMensualKwh[i]  ?? 0;\n'
    '    const cons = consumoEdificioMensualKwh[i] ?? 0;\n\n'
    '    const autoconsumo = Math.min(gen, cons);\n'
    '    const excedente   = Math.max(0, gen - cons);\n'
    '    const compra_red  = Math.max(0, cons - gen);\n\n'
    '    const ahorro_mes  = autoconsumo * tarifaCompraCLP\n'
    '                      + excedente   * tarifaVentaCLP;\n\n'
    '    monthly.push({\n'
    '      month: MONTHS[i],\n'
    '      generacion_kwh: gen,\n'
    '      consumo_kwh: cons,\n'
    '      autoconsumo_kwh: autoconsumo,\n'
    '      excedente_kwh: excedente,\n'
    '      compra_red_kwh: compra_red,\n'
    '      tasa_autoconsumo_pct: gen > 0 ? (autoconsumo / gen)   * 100 : 0,\n'
    '      tasa_autosuficiencia_pct: cons > 0 ? (autoconsumo / cons) * 100 : 0,\n'
    '      ahorro_mes,\n'
    '    });\n\n'
    '    totalGen += gen; totalCons += cons;\n'
    '    totalExc += excedente; totalRed += compra_red;\n'
    '    totalAhorro += ahorro_mes;\n'
    '  }\n\n'
    '  const autoconsumo_pct = totalGen > 0\n'
    '    ? (monthly.reduce((s, m) => s + m.autoconsumo_kwh, 0) / totalGen) * 100\n'
    '    : 0;\n'
    '  const autosuficiencia_pct = totalCons > 0\n'
    '    ? (monthly.reduce((s, m) => s + m.autoconsumo_kwh, 0) / totalCons) * 100\n'
    '    : 0;\n\n'
    '  // Clasificación energética según cobertura promedio del BIPV\n'
    '  let clasificacion: "A+" | "A" | "B" | "C" | "D";\n'
    '  if      (autosuficiencia_pct >= 80) clasificacion = "A+";\n'
    '  else if (autosuficiencia_pct >= 60) clasificacion = "A";\n'
    '  else if (autosuficiencia_pct >= 40) clasificacion = "B";\n'
    '  else if (autosuficiencia_pct >= 20) clasificacion = "C";\n'
    '  else                                 clasificacion = "D";\n\n'
    '  return {\n'
    '    monthly,\n'
    '    annual: {\n'
    '      generacion_total_kwh: totalGen,\n'
    '      consumo_total_kwh: totalCons,\n'
    '      autoconsumo_pct_promedio: autoconsumo_pct,\n'
    '      autosuficiencia_pct_promedio: autosuficiencia_pct,\n'
    '      excedente_total_kwh: totalExc,\n'
    '      compra_red_total_kwh: totalRed,\n'
    '      ahorro_total: totalAhorro,\n'
    '      clasificacion_energetica: clasificacion,\n'
    '    },\n'
    '  };\n'
    '}'
)

h2('Dónde integrar:')
archivo('client/src/components/BIPVROIOptimizer.tsx')
body(
    'Agregar una nueva pestaña "Balance Mensual" junto a las pestañas existentes '
    '"scenarios / sensitivity / recommendations". Mostrar una tabla con los 12 meses '
    'y una barra de progreso para la tasa de autosuficiencia. '
    'Agregar un campo de entrada para el consumo mensual del edificio (kWh/mes). '
    'Mostrar la clasificación energética A+/A/B/C/D en una tarjeta destacada.'
)
doc.add_paragraph('')

# ──────────────────────────────────────────────────────────────
# MEJORA 4 — DEGRADACIÓN AÑO A AÑO + PRODUCCIÓN 25 AÑOS
# ──────────────────────────────────────────────────────────────
sep()
h1('MEJORA 4 — Tabla y gráfico de producción con degradación en 25 años', ROJO)

badge('⚠ PROBLEMA:',
      'El ROI Optimizer aplica degradación anual (0.5%) para calcular el VAN, '
      'pero no muestra al usuario la curva de producción año a año. '
      'El usuario no puede ver en qué año la producción cae bajo cierto umbral, '
      'ni cuándo se necesita el primer reemplazo de componentes.', ROJO)
doc.add_paragraph('')

h2('Código a agregar — nueva función en archivo existente:')
archivo('client/src/lib/bipvROIOptimizer.ts  →  agregar al final')
codigo(
    '// ─── NUEVA FUNCIÓN: Curva de producción con degradación ──────────\n'
    'export interface YearlyProjection {\n'
    '  year: number;\n'
    '  production_kwh: number;     // kWh ese año con degradación\n'
    '  cumulative_kwh: number;     // kWh acumulados desde año 1\n'
    '  energy_revenue: number;     // $ ingreso por energía ese año\n'
    '  cumulative_revenue: number; // $ ingresos acumulados\n'
    '  cumulative_cashflow: number; // $ flujo acumulado (negativo hasta payback)\n'
    '  co2_avoided_ton: number;    // Toneladas CO2 evitadas ese año\n'
    '}\n\n'
    'export function calcularProyeccion25Anos(\n'
    '  produccionAno1Kwh: number,\n'
    '  costoTotalSistema: number,\n'
    '  ahorroAno1: number,\n'
    '  degradacionAnual: number  = 0.005,  // 0.5% por año\n'
    '  crecimientoTarifa: number = 0.035,  // 3.5% alza anual tarifa\n'
    '  factorCO2_kgKwh: number   = 0.294,  // kg CO2/kWh (SEN Chile)\n'
    '                                        // Colombia: 0.126 kg CO2/kWh\n'
    '  horizonte: number = 25\n'
    '): YearlyProjection[] {\n'
    '  const proyeccion: YearlyProjection[] = [];\n'
    '  let cumKwh = 0, cumRevenue = 0, cumCashflow = -costoTotalSistema;\n\n'
    '  for (let n = 1; n <= horizonte; n++) {\n'
    '    const prod   = produccionAno1Kwh * Math.pow(1 - degradacionAnual, n - 1);\n'
    '    const ingreso = ahorroAno1 * Math.pow(1 + crecimientoTarifa, n - 1);\n'
    '    const co2    = prod * factorCO2_kgKwh / 1000;  // toneladas\n\n'
    '    cumKwh      += prod;\n'
    '    cumRevenue  += ingreso;\n'
    '    cumCashflow += ingreso;\n\n'
    '    proyeccion.push({\n'
    '      year: n,\n'
    '      production_kwh:      Math.round(prod),\n'
    '      cumulative_kwh:      Math.round(cumKwh),\n'
    '      energy_revenue:      Math.round(ingreso),\n'
    '      cumulative_revenue:  Math.round(cumRevenue),\n'
    '      cumulative_cashflow: Math.round(cumCashflow),\n'
    '      co2_avoided_ton:     Math.round(co2 * 100) / 100,\n'
    '    });\n'
    '  }\n'
    '  return proyeccion;\n'
    '}'
)

h2('Dónde integrar:')
archivo('client/src/components/BIPVROIOptimizer.tsx')
body(
    'Agregar una pestaña "Proyección 25 años" con dos gráficos Recharts: '
    '(1) línea de producción anual con área rellena mostrando la degradación, '
    '(2) barras de flujo de caja acumulado (rojo = antes del payback, verde = después). '
    'Añadir una tarjeta con el CO₂ evitado acumulado en 25 años.'
)
doc.add_paragraph('')

# ──────────────────────────────────────────────────────────────
# MEJORA 5 — DIMENSIONADO DE BATERÍAS
# ──────────────────────────────────────────────────────────────
sep()
h1('MEJORA 5 — Cálculo de capacidad del banco de baterías (sistemas híbridos)', ROJO)

badge('⚠ PROBLEMA:',
      'Para edificios BIPV con almacenamiento (cada vez más frecuente), '
      'no existe ninguna herramienta para dimensionar el banco de baterías. '
      'El usuario no puede calcular cuántos kWh de batería necesita, '
      'el voltaje del banco, ni la capacidad en Ah según la tecnología (litio vs plomo).', ROJO)
doc.add_paragraph('')

h2('Código a agregar — nueva función:')
archivo('client/src/lib/stringSizing.ts  →  agregar al final del mismo archivo')
codigo(
    '// ─── DIMENSIONADO DE BANCO DE BATERÍAS ──────────────────────\n'
    'export interface BatteryBankResult {\n'
    '  capacidad_kwh: number;      // Energía útil total necesaria (kWh)\n'
    '  capacidad_ah: number;       // Capacidad en Ah (C_banco)\n'
    '  voltaje_banco: number;      // Voltaje del banco (V)\n'
    '  n_baterias: number;         // Número de baterías unitarias\n'
    '  dod_usado: number;          // DOD aplicado (0-1)\n'
    '  tecnologia: string;         // Tecnología seleccionada\n'
    '  advertencias: string[];     // Advertencias de diseño\n'
    '}\n\n'
    'export type BatteryTechnology = "litio-LFP" | "litio-NMC" | "plomo-acido" | "flujo";\n\n'
    'const DOD_POR_TECNOLOGIA: Record<BatteryTechnology, number> = {\n'
    '  "litio-LFP"  : 0.90,  // LiFePO4 — hasta 90% DOD\n'
    '  "litio-NMC"  : 0.80,  // NMC — hasta 80% DOD\n'
    '  "plomo-acido": 0.50,  // Plomo ácido — max 50% DOD\n'
    '  "flujo"      : 0.80,  // Flujo vanadio — hasta 80% DOD\n'
    '};\n\n'
    'export function dimensionarBaterias(\n'
    '  consumoDiarioKwh: number,      // kWh/día consumo del edificio\n'
    '  diasAutonomia: number = 1,     // Días de autonomía sin sol\n'
    '  voltajeBancoV: number = 48,    // Voltaje del banco (V)\n'
    '  tecnologia: BatteryTechnology = "litio-LFP",\n'
    '  capacidadBateriaUnitariaAh: number = 100  // Ah de cada batería individual\n'
    '): BatteryBankResult {\n'
    '  const advertencias: string[] = [];\n'
    '  const dod = DOD_POR_TECNOLOGIA[tecnologia];\n\n'
    '  // Energía bruta necesaria (considerando DOD y eficiencia carga ~0.95)\n'
    '  const energia_kwh_util = consumoDiarioKwh * diasAutonomia;\n'
    '  const energia_bruta    = energia_kwh_util / (dod * 0.95);\n\n'
    '  // Capacidad en Ah = Energía_Wh / Voltaje_banco\n'
    '  const capacidad_ah = (energia_bruta * 1000) / voltajeBancoV;\n\n'
    '  // Número de baterías unitarias (redondear hacia arriba)\n'
    '  const n_baterias = Math.ceil(capacidad_ah / capacidadBateriaUnitariaAh);\n\n'
    '  if (diasAutonomia > 3)\n'
    '    advertencias.push("Más de 3 días de autonomía encarece el sistema significativamente.");\n'
    '  if (tecnologia === "plomo-acido" && diasAutonomia > 1)\n'
    '    advertencias.push("Plomo-ácido con alta autonomía reduce vida útil. Considerar LFP.");\n\n'
    '  return {\n'
    '    capacidad_kwh: Math.round(energia_kwh_util * 10) / 10,\n'
    '    capacidad_ah:  Math.round(capacidad_ah),\n'
    '    voltaje_banco: voltajeBancoV,\n'
    '    n_baterias,\n'
    '    dod_usado: dod,\n'
    '    tecnologia,\n'
    '    advertencias,\n'
    '  };\n'
    '}'
)
doc.add_paragraph('')

# ──────────────────────────────────────────────────────────────
# MEJORA 6 — CONVERTIDOR DE MÉTRICAS
# ──────────────────────────────────────────────────────────────
sep()
h1('MEJORA 6 — Módulo de conversiones de métricas y HSP', ROJO)

badge('⚠ PROBLEMA:',
      'No existe ningún convertidor de unidades. Los usuarios que reciben datos de '
      'estaciones meteorológicas en W/m²/día, Wh/m²/día, kWh/m²/mes, BTU/ft², °F, '
      'o HP no pueden convertirlos antes de ingresarlos. '
      'Errores de unidad son la causa más común de resultados incorrectos en la práctica.', ROJO)
doc.add_paragraph('')

h2('Código a agregar — nuevo archivo:')
archivo('client/src/lib/unitConversions.ts')
codigo(
    '// ═══════════════════════════════════════════════════════════\n'
    '// MÓDULO: Conversiones de métricas para cálculos BIPV\n'
    '// Archivo: client/src/lib/unitConversions.ts\n'
    '// ═══════════════════════════════════════════════════════════\n\n'
    '// ── TEMPERATURA ──────────────────────────────────────────\n'
    'export const celsiusToFahrenheit = (c: number): number => (c * 9/5) + 32;\n'
    'export const fahrenheitToCelsius = (f: number): number => (f - 32) * 5/9;\n'
    'export const celsiusToKelvin     = (c: number): number => c + 273.15;\n\n'
    '// ── POTENCIA ─────────────────────────────────────────────\n'
    'export const wattsToKilowatts    = (w:  number): number => w  / 1000;\n'
    'export const kilowattsToWatts    = (kw: number): number => kw * 1000;\n'
    'export const wattsToHP           = (w:  number): number => w  / 745.7;\n'
    'export const hpToWatts           = (hp: number): number => hp * 745.7;\n\n'
    '// ── ENERGÍA ──────────────────────────────────────────────\n'
    'export const whToKwh   = (wh:  number): number => wh  / 1000;\n'
    'export const kwhToWh   = (kwh: number): number => kwh * 1000;\n'
    'export const kwhToMwh  = (kwh: number): number => kwh / 1000;\n'
    'export const whToJoules= (wh:  number): number => wh  * 3600;\n\n'
    '// ── ÁREA ─────────────────────────────────────────────────\n'
    'export const m2ToFt2      = (m2: number): number => m2 * 10.764;\n'
    'export const ft2ToM2      = (ft2:number): number => ft2 * 0.0929;\n'
    'export const m2ToHectareas= (m2: number): number => m2 / 10000;\n'
    'export const km2ToM2      = (km2:number): number => km2 * 1_000_000;\n\n'
    '// ── IRRADIANCIA / HSP ────────────────────────────────────\n'
    '/**\n'
    ' * HSP del lugar = kWh/m²/día (numéricamente equivalentes).\n'
    ' * Wh/m²/día debe dividirse entre 1000 para obtener HSP.\n'
    ' */\n'
    'export const whM2DiaToHSP      = (wh_m2: number): number => wh_m2 / 1000;\n'
    'export const kwhM2DiaToHSP     = (kwh_m2:number): number => kwh_m2;      // iguales\n'
    'export const ghiAnualToHSPDia  = (ghiAnual: number): number => ghiAnual / 365;\n\n'
    '/**\n'
    ' * Convierte irradiación mensual (kWh/m²/mes) a HSP diaria promedio del mes.\n'
    ' */\n'
    'export function irradiacionMensualToHSP(\n'
    '  kwhM2Mes: number,\n'
    '  diasDelMes: number\n'
    '): number {\n'
    '  return kwhM2Mes / diasDelMes;\n'
    '}\n\n'
    '/**\n'
    ' * Tabla de factores de emisión CO2 por país (kg CO2/kWh).\n'
    ' * Fuente: IEA 2024 / SEN 2024 / UPME 2024\n'
    ' */\n'
    'export const FACTORES_CO2_KG_KWH: Record<string, number> = {\n'
    '  "Chile"     : 0.294,   // SEN 2024\n'
    '  "Colombia"  : 0.126,   // UPME / IDEAM 2024 (alta hidro)\n'
    '  "México"    : 0.398,   // CENACE 2024\n'
    '  "Argentina" : 0.341,   // CAMMESA 2024\n'
    '  "Perú"      : 0.214,   // COES 2024\n'
    '  "España"    : 0.187,   // REE 2024\n'
    '  "Global"    : 0.436,   // IEA World Average 2024\n'
    '};\n\n'
    'export const co2EvitadoTon = (\n'
    '  energiaKwh: number,\n'
    '  pais: string = "Chile"\n'
    '): number => {\n'
    '  const factor = FACTORES_CO2_KG_KWH[pais] ?? FACTORES_CO2_KG_KWH["Global"];\n'
    '  return (energiaKwh * factor) / 1000;  // toneladas\n'
    '};'
)

h2('Dónde integrar:')
archivo('client/src/components/EnergyProductionSimulator.tsx')
body(
    'Agregar una sección de ayuda "Convertidor de Unidades" desplegable (collapsible) '
    'al inicio del Simulador de Energía. El usuario puede ingresar un valor en cualquier '
    'unidad y ver la conversión en tiempo real. También integrar co2EvitadoTon() '
    'en el ReportGenerator para mostrar el CO₂ evitado con el factor del país del proyecto.'
)
doc.add_paragraph('')

# ──────────────────────────────────────────────────────────────
# MEJORA 7 — TEMPERATURA DE CELDA CON PENALIDAD EXPLÍCITA
# ──────────────────────────────────────────────────────────────
sep()
h1('MEJORA 7 — Mostrar al usuario la penalidad de temperatura de fachada confinada', ROJO)

badge('⚠ PROBLEMA:',
      'El k_bipv = 1.3 está implementado en el motor de cálculo (iamSoilingEngine.ts) '
      'pero el usuario nunca ve cuánto impacta en la temperatura de celda '
      'ni cuánta producción pierde por ese efecto. '
      'En fachadas confinadas (sin ventilación trasera), la celda puede operar '
      '15-20°C más caliente que un panel ventilado, reduciendo la producción 5-8%.', ROJO)
doc.add_paragraph('')

h2('Código a agregar — nueva función de visualización:')
archivo('client/src/lib/iamSoilingEngine.ts  →  agregar función comparativa')
codigo(
    '// ─── COMPARATIVA DE IMPACTO TÉRMICO BIPV ─────────────────────\n'
    'export interface ThermalImpactComparison {\n'
    '  tCell_ventilado: number;    // T celda con montaje ventilado (°C)\n'
    '  tCell_confinado: number;    // T celda con montaje confinado (°C)\n'
    '  diferencia_C: number;       // Diferencia de temperatura (°C)\n'
    '  efic_ventilado: number;     // Eficiencia real ventilado (0-1)\n'
    '  efic_confinado: number;     // Eficiencia real confinado (0-1)\n'
    '  perdida_produccion_pct: number; // % de producción perdida por confinamiento\n'
    '}\n\n'
    'export function compararImpactoTermico(\n'
    '  tAmb: number,           // Temperatura ambiente (°C)\n'
    '  irradiancia: number,    // Irradiancia W/m²\n'
    '  noct: number,           // NOCT del panel (°C)\n'
    '  coefTemp: number,       // Coef. temperatura (%/°C, negativo)\n'
    '  eficNom: number         // Eficiencia nominal STC (0-1)\n'
    '): ThermalImpactComparison {\n'
    '  // k_bipv = 1.0 → montaje ventilado (panel convencional)\n'
    '  // k_bipv = 1.3 → montaje confinado (fachada sin ventilación trasera)\n'
    '  const calcTCell = (k_bipv: number): number =>\n'
    '    tAmb + k_bipv * ((noct - 20) / 800) * irradiancia;\n\n'
    '  const tCell_v = calcTCell(1.0);\n'
    '  const tCell_c = calcTCell(1.3);\n\n'
    '  const calcEfic = (tCell: number): number =>\n'
    '    eficNom * (1 + (coefTemp / 100) * (tCell - 25));\n\n'
    '  const efic_v = calcEfic(tCell_v);\n'
    '  const efic_c = calcEfic(tCell_c);\n\n'
    '  const perdida_pct = efic_v > 0\n'
    '    ? ((efic_v - efic_c) / efic_v) * 100\n'
    '    : 0;\n\n'
    '  return {\n'
    '    tCell_ventilado:       Math.round(tCell_v * 10) / 10,\n'
    '    tCell_confinado:       Math.round(tCell_c * 10) / 10,\n'
    '    diferencia_C:          Math.round((tCell_c - tCell_v) * 10) / 10,\n'
    '    efic_ventilado:        Math.round(efic_v * 10000) / 10000,\n'
    '    efic_confinado:        Math.round(efic_c * 10000) / 10000,\n'
    '    perdida_produccion_pct: Math.round(perdida_pct * 10) / 10,\n'
    '  };\n'
    '}'
)

h2('Dónde integrar:')
archivo('client/src/components/BIPVGlassSimulator.tsx')
body(
    'Agregar una tarjeta de alerta cuando el tipo de montaje sea "fachada_confinada". '
    'Mostrar un comparativo: temperatura de celda ventilada vs confinada, diferencia en °C, '
    'y pérdida estimada de producción. Esto ayuda al usuario a justificar '
    'el diseño de cámara de ventilación trasera en fachadas BIPV.'
)
doc.add_paragraph('')

# ═══════════════════════════════════════════════════════════════
# SECCIÓN 3 — TABLA RESUMEN
# ═══════════════════════════════════════════════════════════════
sep()
h1('SECCIÓN 3 — TABLA RESUMEN DE DIAGNÓSTICO')

filas_tabla = [
    ('Módulo 1 — Datos del edificio',       '✅ Completo',  'BIPVGlassSimulator.tsx — superficies BIPV por tipo'),
    ('Módulo 2 — Recurso solar',            '✅ Superior',  'Perez + Liu-Jordan + PVWatts + PVGIS hora a hora'),
    ('Módulo 3 — Módulo BIPV',              '✅ Completo',  'bipvGlassCatalog.ts + importador PDF'),
    ('Módulo 4 — Temperatura operación',    '⚠ Mejorar',   'k_bipv=1.3 existe, falta comparativa visible (MEJORA 7)'),
    ('Módulo 5 — Energía generada',         '✅ Superior',  'IEC 61724-1:2021 completo con PR_T, EPI, BOS'),
    ('Módulo 6 — Análisis de sombras',      '✅ Superior',  '3D GLTF/OBJ, polígonos, crossing modal, trayectoria solar'),
    ('Módulo 7 — Dimensionado eléctrico',   '❌ Falta',     'No existe string sizing → MEJORA 1 obligatoria'),
    ('Módulo 8 — Balance energético',       '⚠ Parcial',   'Autoconsumo % fijo → necesita MEJORA 3 (balance mensual)'),
    ('Módulo 9 — Análisis financiero',      '✅ Completo',  'LCOE, VAN, TIR, payback, Ley 1715, HVAC savings'),
    ('Módulo 10 — Gráficos e informe',      '✅ Completo',  'Recharts, PDF report, CSV export, global report'),
    ('Factor de cobertura activa',          '❌ Falta',     'Sin factor_cobertura la potencia está sobrestimada (MEJORA 2)'),
    ('Degradación 25 años — gráfico',       '⚠ Parcial',   'Se calcula en VAN pero no se muestra al usuario (MEJORA 4)'),
    ('Banco de baterías',                   '❌ Falta',     'No hay cálculo de capacidad Ah/kWh (MEJORA 5)'),
    ('Conversiones de métricas',            '❌ Falta',     'No hay convertidor de unidades (MEJORA 6)'),
    ('CO₂ evitado por país',               '⚠ Parcial',   'No usa factor de emisión por país (MEJORA 6 lo incluye)'),
]

tabla = doc.add_table(rows=1, cols=3)
tabla.style = 'Table Grid'
hdr = tabla.rows[0].cells
hdr[0].text = 'Módulo / Funcionalidad'
hdr[1].text = 'Estado'
hdr[2].text = 'Observación'
for c in hdr:
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            r.font.size = Pt(9)
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE
    shd = _OE('w:shd')
    shd.set(_qn('w:val'), 'clear')
    shd.set(_qn('w:color'), 'auto')
    shd.set(_qn('w:fill'), '1A5C8A')
    c._tc.get_or_add_tcPr().append(shd)

for f in filas_tabla:
    row = tabla.add_row().cells
    for i, val in enumerate(f):
        row[i].text = val
        run = row[i].paragraphs[0].runs[0] if row[i].paragraphs[0].runs else row[i].paragraphs[0].add_run(val)
        run.font.size = Pt(9)
        if i == 1:
            if '✅' in val: run.font.color.rgb = VERDE
            elif '⚠' in val: run.font.color.rgb = NARANJA
            elif '❌' in val: run.font.color.rgb = ROJO

doc.add_paragraph('')

# ── Prioridad de implementación ──
h1('PRIORIDAD DE IMPLEMENTACIÓN', AZUL)

prioridades = [
    ('🔴 ALTA — Implementar primero',
     'MEJORA 1: String sizing (riesgo de daño al inversor sin esta validación)\n'
     'MEJORA 2: Factor de cobertura (potencia sobrestimada en todos los proyectos)\n'
     'MEJORA 3: Balance energético mensual (VAN y payback más precisos)'),
    ('🟡 MEDIA — Implementar en segunda fase',
     'MEJORA 7: Comparativa de temperatura confinada vs ventilada (mejora la UX)\n'
     'MEJORA 4: Gráfico de degradación 25 años (el dato ya existe, solo falta mostrarlo)\n'
     'MEJORA 6: Convertidor de unidades (reduce errores de entrada del usuario)'),
    ('🟢 BAJA — Implementar en tercera fase',
     'MEJORA 5: Banco de baterías (solo para proyectos híbridos off-grid)'),
]

for titulo_p, desc_p in prioridades:
    h2(titulo_p)
    p_d = doc.add_paragraph(desc_p)
    p_d.runs[0].font.size = Pt(10)
    doc.add_paragraph('')

# Pie de página
doc.add_paragraph('')
footer = doc.add_paragraph(
    'Diagnóstico elaborado sobre código fuente real — github.com/ventas108/calculadora-bipv  |  2026'
)
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.runs[0].font.size = Pt(9)
footer.runs[0].font.color.rgb = RGBColor(0x7F,0x7F,0x7F)
footer.runs[0].italic = True

doc.save('Diagnostico_Mejoras_Calculadora_BIPV.docx')
print("Documento creado correctamente.")
