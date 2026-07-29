"""
Genera el Manual de Operación BIPV Colombia v2026 — FASE 4 incluida.
Salida: attached_assets/MANUAL_OPERACION_BIPV_v2026_FASE4.docx

Ejecutar:
    python scripts/generar_manual_bipv_v2026.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os

OUT = "attached_assets/MANUAL_OPERACION_BIPV_v2026_FASE4.docx"

# ── Paleta ─────────────────────────────────────────────────────────────────
C_VERDE_OSC  = RGBColor(0x1B, 0x5E, 0x20)   # encabezado verde oscuro
C_VERDE_MED  = RGBColor(0x2E, 0x7D, 0x32)   # filas impares verde
C_VERDE_CLR  = RGBColor(0xE8, 0xF5, 0xE9)   # fondo filas pares verde
C_AZUL       = RGBColor(0x0D, 0x47, 0xA1)   # títulos sección azul
C_AZUL_CLR   = RGBColor(0xE3, 0xF2, 0xFD)   # fondo info box
C_NARANJA    = RGBColor(0xE6, 0x51, 0x00)   # alertas naranja
C_NARANJA_CLR= RGBColor(0xFF, 0xF3, 0xE0)   # fondo alerta leve
C_ROJO       = RGBColor(0xB7, 0x1C, 0x1C)   # error crítico
C_ROJO_CLR   = RGBColor(0xFF, 0xEB, 0xEE)   # fondo error
C_GRIS       = RGBColor(0x37, 0x47, 0x4F)   # gris oscuro
C_GRIS_CLR   = RGBColor(0xEC, 0xEF, 0xF1)   # fondo gris claro
C_AMARILLO   = RGBColor(0xFF, 0xF9, 0xC4)   # advertencia FASE4
C_MORADO     = RGBColor(0x4A, 0x14, 0x8C)   # FASE 4 baterías
C_MORADO_CLR = RGBColor(0xF3, 0xE5, 0xF5)   # fondo morado claro
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helpers ────────────────────────────────────────────────────────────────
def rgb_hex(rgb: RGBColor) -> str:
    """Convierte RGBColor a string hexadecimal sin depender de .red/.green/.blue"""
    # RGBColor es subclase de tuple: (r, g, b)
    r, g, b = int(rgb[0]), int(rgb[1]), int(rgb[2])
    return f"{r:02X}{g:02X}{b:02X}"

def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  rgb_hex(rgb))
    tcPr.append(shd)

def set_cell_borders(cell, color="CCCCCC"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top","left","bottom","right"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def style_header_cell(cell, rgb_bg: RGBColor, text=None, bold=True, sz=9, color=None):
    set_cell_bg(cell, rgb_bg)
    set_cell_borders(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if text is not None:
        p = cell.paragraphs[0]
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(sz)
        run.font.color.rgb = color if color else WHITE

def add_heading(doc, text, level=1, color=None, sz=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    if sz:      run.font.size = Pt(sz)
    elif level == 1: run.font.size = Pt(16)
    elif level == 2: run.font.size = Pt(13)
    elif level == 3: run.font.size = Pt(11)
    else:            run.font.size = Pt(10)
    run.font.color.rgb = color or C_AZUL
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_body(doc, text, italic=False, bold=False, color=None, sz=9, indent=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(sz)
    run.italic = italic
    run.bold   = bold
    if color: run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    if indent: p.paragraph_format.left_indent = Cm(indent)
    return p

def add_note_box(doc, icon, title, text, bg: RGBColor, title_color: RGBColor):
    """Agrega un recuadro de nota estilizado usando una tabla 1x1."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg)
    set_cell_borders(cell, color=rgb_hex(title_color))
    p = cell.paragraphs[0]
    p.clear()
    r1 = p.add_run(f"{icon}  {title}  ")
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = title_color
    r2 = p.add_run(text)
    r2.font.size = Pt(9)
    r2.font.color.rgb = C_GRIS
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_alert(doc, icon, text, critical=False):
    bg     = C_ROJO_CLR   if critical else C_NARANJA_CLR
    color  = C_ROJO       if critical else C_NARANJA
    symbol = "✗" if critical else "⚠"
    add_note_box(doc, f"{symbol} {icon}", "", text, bg, color)

def add_info(doc, text):
    add_note_box(doc, "ℹ", "", text, C_AZUL_CLR, C_AZUL)

def add_success(doc, icon, text):
    add_note_box(doc, f"✔ {icon}", "Señal de éxito:", text, C_VERDE_CLR, C_VERDE_OSC)

def add_fase4_note(doc, text):
    """Recuadro especial para notas críticas de FASE 4 baterías."""
    add_note_box(doc, "🔋", "NOTA FASE 4:", text, C_MORADO_CLR, C_MORADO)

def simple_table(doc, headers, rows, col_widths=None):
    n = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"
    # Header row
    for ci, h in enumerate(headers):
        style_header_cell(tbl.cell(0, ci), C_VERDE_OSC, h, sz=8)
        if col_widths:
            tbl.cell(0, ci).width = Cm(col_widths[ci])
    # Data rows
    for ri, row in enumerate(rows):
        bg = C_VERDE_CLR if ri % 2 == 0 else WHITE
        for ci, val in enumerate(row):
            cell = tbl.cell(ri + 1, ci)
            set_cell_bg(cell, bg)
            set_cell_borders(cell)
            p = cell.paragraphs[0]
            p.clear()
            r = p.add_run(str(val))
            r.font.size = Pt(8)
            r.font.color.rgb = C_GRIS
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return tbl

def page_break(doc):
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
def build_manual():
    doc = Document()

    # ── Márgenes ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)

    # ══ PORTADA ══════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("INNOVACIÓN QUÍMICA")
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = C_VERDE_OSC
    p.paragraph_format.space_before = Pt(20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("calc.innovacionquimica.com.co")
    r.font.size = Pt(10); r.font.color.rgb = C_GRIS; r.italic = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MANUAL DE OPERACIÓN")
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = C_AZUL

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CALCULADORA BIPV COLOMBIA")
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = C_AZUL

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Versión 2026  ·  Incluye FASE 4: Baterías y Balance Energético")
    r.font.size = Pt(11); r.font.color.rgb = C_MORADO; r.bold = True

    doc.add_paragraph()

    add_note_box(doc,
        "⚠", "BORRADOR:",
        " Este manual es un BORRADOR preliminar y será actualizado cuando la "
        "calculadora sea la versión definitiva. No entregarlo como documento "
        "final al cliente.",
        C_NARANJA_CLR, C_NARANJA)

    # ── Índice de modos ──────────────────────────────────────────────────
    doc.add_paragraph()
    add_heading(doc, "Seleccione su punto de partida según el caso:", level=3, color=C_GRIS)

    simple_table(doc,
        ["Modo", "Cuándo usarlo", "Sección del manual"],
        [
            ("[A] Modo ÁREA disponible",
             "Conozco los m² disponibles para instalar los paneles",
             "PARTE A — Pasos 1 al 10"),
            ("[B] Modo CONSUMO / Factura",
             "Tengo la factura eléctrica o sé los kWh/mes de consumo",
             "PARTE B — Paso 1 específico"),
        ],
        col_widths=[5, 8, 5]
    )

    # ── Flujo de páginas ─────────────────────────────────────────────────
    doc.add_paragraph()
    add_heading(doc, "Flujo obligatorio de páginas (seguir en orden):", level=3, color=C_GRIS)

    simple_table(doc,
        ["Pág.", "Nombre", "Tipo"],
        [
            ("1",  "Proyecto",                 "OBLIGATORIA — siempre primero"),
            ("2",  "Recurso Solar",             "OBLIGATORIA — antes de Mismatch y Producción"),
            ("3",  "Motor IV",                  "Opcional — herramienta técnica de validación"),
            ("4",  "Dimensionamiento",          "OBLIGATORIA — seleccionar panel e inversor"),
            ("5",  "Mismatch",                  "Opcional pero muy recomendada"),
            ("5b", "Motor Óptico",              "Opcional pero muy recomendada — correcciones reales IAM + Soiling + Térmico"),
            ("6",  "Producción",                "OBLIGATORIA — antes del Financiero"),
            ("7",  "Financiero",                "OBLIGATORIA — análisis económico final"),
            ("8",  "Presupuesto",               "Opcional pero recomendada — CAPEX real"),
            ("11", "🔋 Baterías y Balance",     "NUEVO FASE 4 — dimensionar batería y analizar balance energético"),
            ("10", "Reporte PDF",               "Opcional — exportar reporte técnico completo del proyecto"),
        ],
        col_widths=[1.2, 5.5, 9.5]
    )

    add_info(doc,
        "NO salte páginas. Cada página alimenta a la siguiente. Si salta una, "
        "los cálculos posteriores quedarán incompletos o usarán valores por "
        "defecto incorrectos.")

    # ══ PARTE A ═══════════════════════════════════════════════════════════
    page_break(doc)
    p = doc.add_paragraph()
    r = p.add_run("PARTE A  —  MODO ÁREA DISPONIBLE  [A]")
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = WHITE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # fondo azul
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, C_AZUL)
    cp = cell.paragraphs[0]
    cp.clear()
    r = cp.add_run("PARTE A  —  MODO ÁREA DISPONIBLE  [A]")
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = WHITE
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    add_body(doc,
        '"Se cuántos metros cuadrados tengo disponibles para instalar los paneles."',
        italic=True, color=C_GRIS)

    # ── PASO 1 ──────────────────────────────────────────────────────────
    add_heading(doc, "PASO 1  —  Página 1: Datos del Proyecto", level=2, color=C_AZUL)
    add_body(doc, "Qué hacer:", bold=True)

    simple_table(doc,
        ["Campo", "Descripción", "Ejemplo"],
        [
            ("Nombre del proyecto", "Identificador del proyecto (solo texto)", "Edificio Chapinero"),
            ("Ciudad", "Seleccionar de la lista desplegable", "Bogotá"),
            ("Área de fachada / techo (m²)", "Metros cuadrados NETOS disponibles para paneles", "50"),
            ("Tarifa eléctrica (COP/kWh)", "Valor del kWh según su factura de energía", "650"),
            ("Performance Ratio — PR (%)", "Eficiencia real del sistema. Dejar en 80% si no sabe.", "80"),
            ("Densidad de potencia (W/m²)", "Potencia por m² del panel elegido. Ver ficha técnica.", "200"),
        ],
        col_widths=[5, 9, 4]
    )

    add_body(doc, "Hacer clic en  Guardar configuración.", bold=True, indent=0.5)

    add_success(doc, "",
        "[OK]  La página muestra un bloque verde con: ciudad, coordenadas, GHI, "
        "temperatura media y una estimación preliminar de energía anual en kWh.")

    add_body(doc, "ALERTAS — No cometer estos errores:", bold=True, color=C_ROJO)
    add_alert(doc, "",
        "NO ingrese el área del lote completo. Solo el área NETA disponible para "
        "paneles: sin pasillos de mantenimiento, sin zonas de sombra de paredes "
        "laterales, sin zonas de exclusión técnica.", critical=True)
    add_alert(doc, "",
        "NO deje la tarifa eléctrica en 0. Con tarifa = 0 el análisis financiero "
        "calculará ahorro = $0 y la TIR saldrá incorrecta o negativa.", critical=True)
    add_alert(doc, "",
        "NO cambie la ciudad después de haber calculado el Recurso Solar. Si la "
        "cambia debe repetir todos los pasos desde el Paso 2.", critical=True)
    add_info(doc,
        "El PR por defecto (80%) es conservador. Para sistemas BIPV en fachada "
        "vertical con sombras, use 70-75%. Para techos sin sombras se puede usar 82-85%.")

    # ── PASO 2 ──────────────────────────────────────────────────────────
    add_heading(doc, "PASO 2  —  Página 2: Recurso Solar", level=2, color=C_AZUL)
    add_body(doc, "Qué hacer:", bold=True)

    simple_table(doc,
        ["Campo", "Descripción", "Valores típicos"],
        [
            ("Azimuth fachada (grados)",
             "Dirección a la que mira la fachada.\nNorte=0 | Sur=180 | Este=90 | Oeste=270",
             "Fachada sur → 180"),
            ("Inclinación — Tilt (grados)",
             "Ángulo del panel respecto al horizonte.\nFachada vertical → 90  |  Techo plano → 10",
             "Fachada vertical → 90"),
            ("Albedo del suelo",
             "Reflectividad del suelo frente a la fachada.\n0.2=concreto/asfalto | 0.8=nieve",
             "0.20"),
        ],
        col_widths=[4.5, 9, 4.7]
    )
    add_body(doc, "Hacer clic en  Calcular Recurso Solar.  Esperar 30-90 segundos.", bold=True, indent=0.5)

    add_success(doc, "",
        "[OK]  Aparece 'Recurso solar calculado' con: gráfica POA mensual, "
        "heatmap horario anual y POA total en kWh/m²/año.")

    add_alert(doc, "",
        "NO continúe si aparece un error de PVGIS. Indica falla en la descarga "
        "de datos climáticos. Sin este paso el resto del flujo NO tiene datos válidos.", critical=True)
    add_alert(doc, "",
        "NO use Tilt = 0 grados. Tilt cero = panel completamente horizontal. "
        "Para fachadas use 90. Para techos inclinados use el ángulo real.", critical=True)
    add_info(doc,
        "El azimuth afecta drásticamente la producción. Una fachada norte (azimuth=0) "
        "en Colombia puede producir entre 40% y 60% menos que una fachada sur (azimuth=180).")

    # ── PASO 3 ──────────────────────────────────────────────────────────
    add_heading(doc, "PASO 3  —  Página 4: Dimensionamiento de Strings", level=2, color=C_AZUL)
    add_info(doc,
        "Salte la Página 3 (Motor IV) por ahora. Es una herramienta técnica de "
        "validación y NO es obligatoria para el flujo de cálculo principal.")

    add_body(doc, "Qué hacer:", bold=True)
    add_body(doc,
        "1.  Abrir  [4] Dimensionamiento.\n"
        "2.  Seleccionar Panel solar (usar ÚNICAMENTE panel con ficha completa — semáforo VERDE) "
        "e Inversor según la potencia total del sistema.\n"
        "3.  Revisar parámetros de temperatura de diseño.", indent=0.3)

    simple_table(doc,
        ["Campo", "Descripción", "Valor defecto"],
        [
            ("T mínima de diseño (°C)", "Temperatura nocturna más fría del sitio", "Viene de Página 1"),
            ("T celda realista (°C)", "Temperatura de operación normal real", "50"),
            ("T celda extremo (°C)", "Temperatura máxima en día de verano sin viento", "70"),
            ("N strings por tracker", "Strings en paralelo por entrada MPPT del inversor", "1"),
        ],
        col_widths=[5.5, 8.5, 4.2]
    )
    add_body(doc, "Hacer clic en  Calcular Dimensionamiento.", bold=True, indent=0.5)

    add_success(doc, "",
        "[OK]  Muestra: N_serie, verificación de tensiones en verde (Voc ≤ Vmax inversor "
        "y Vmpp dentro del rango MPPT) y Potencia DC total del sistema en kWp.")

    add_alert(doc, "",
        "NO elija un panel con ficha incompleta (semáforo ROJO). El motor de simulación "
        "no podrá calcular la curva I-V correctamente.", critical=True)
    add_alert(doc, "",
        "NO ignore las advertencias de tensión. Si aparece 'Voc supera Vmax del inversor', "
        "el sistema eléctrico es INSEGURO para instalación.", critical=True)
    add_info(doc,
        "Si el sistema tiene múltiples orientaciones, use el toggle 'Múltiples orientaciones' "
        "y distribuya los porcentajes. Las fracciones DEBEN sumar exactamente 1.00 = 100%.")

    # ── PASO 4 ──────────────────────────────────────────────────────────
    add_heading(doc, "PASO 4  —  Página 5: Mismatch y Pérdidas", level=2, color=C_AZUL)
    add_body(doc, "Qué hacer:", bold=True)
    add_body(doc,
        "1.  Abrir  [5] Mismatch.\n"
        "2.  Definir el perfil de horizonte (elevación=0 si no hay sombras cercanas).", indent=0.3)

    simple_table(doc,
        ["Parámetro", "Descripción", "Valor típico Colombia"],
        [
            ("Mismatch fabricación (%)", "Variación de parámetros entre paneles del mismo lote", "2%"),
            ("Soiling / suciedad (%)", "Pérdida por polvo, humo y suciedad acumulada", "3%"),
            ("Cableado DC (%)", "Pérdidas resistivas en el cableado de corriente continua", "2%"),
        ],
        col_widths=[5.5, 9, 3.7]
    )
    add_body(doc, "Hacer clic en  Calcular Mismatch.", bold=True, indent=0.5)

    add_success(doc, "",
        "[OK]  Aparece el gráfico de cascada de pérdidas (waterfall) con el Factor "
        "Global de Mismatch y la POA efectiva final en kWh/m²/año.")

    add_info(doc,
        "Esta página es OPCIONAL pero muy recomendada. Si la omite, la Producción "
        "usará la POA bruta sin descuentos de sombra ni mismatch, SOBREESTIMANDO la energía real.")

    # ── PASO 4b ─────────────────────────────────────────────────────────
    add_heading(doc, "PASO 4b  —  Página 5b: Motor Óptico BIPV  (opcional pero muy recomendado)", level=2, color=C_AZUL)
    add_info(doc,
        "Ejecutar DESPUÉS de Mismatch y ANTES de Producción. Si lo omite, la "
        "Producción usará la POA bruta sin las tres correcciones físicas reales del BIPV.")

    add_body(doc,
        "Por qué es importante — Las calculadoras convencionales usan la irradiación bruta "
        "(POA) directamente. El Motor Óptico aplica tres correcciones que sobreestiman la "
        "producción si no se aplican, especialmente en fachadas verticales:\n"
        "  1.  Reflexión del vidrio (IAM ASHRAE): pérdida típica 12-20% anual.\n"
        "  2.  Suciedad estacional (Soiling): pérdida típica 2-5% anual.\n"
        "  3.  Temperatura confinada (Térmico BIPV): pérdida típica 3-7% según montaje.",
        sz=9, indent=0.3)

    simple_table(doc,
        ["Parámetro", "Qué representa", "Fuente del valor", "Cuándo ajustar manualmente"],
        [
            ("b0 (reflexión vidrio)", "Cuánto refleja el vidrio según el ángulo de incidencia.",
             "Auto: CdTe=0.12 / CIGS=0.10 / Si=0.05", "Si tiene dato del fabricante distinto al automático"),
            ("tau — Transparencia (%)", "Fracción del área del vidrio sin semiconductor.",
             "Auto: campo TransparenciaPct del catálogo", "Si el panel tiene transparencia diferente al catálogo"),
            ("k_BIPV (confinamiento)", "Factor de temperatura por ventilación restringida.\n1.0=libre / 1.3=típico / 1.5=sellado",
             "Manual — depende de la arquitectura", "Siempre verificar según detalle constructivo de la fachada"),
            ("NOCT (°C)", "Temperatura nominal de operación del panel a 800 W/m².",
             "Auto: campo NOCT_C del catálogo", "Si el dato del catálogo no coincide con la ficha oficial"),
            ("gamma (%/°C)", "Caída de eficiencia por cada grado sobre 25°C. Siempre negativo.",
             "Auto: campo CoefT_C del catálogo", "Si el catálogo tiene el valor incorrecto"),
        ],
        col_widths=[3.2, 4.8, 4.2, 5.7]
    )

    add_success(doc, "",
        "[OK]  Aparece el waterfall: POA bruta → IAM → Soiling → Térmico → POA efectiva. "
        "El factor global queda guardado en sesión y la Página 6 lo usa automáticamente.")

    add_alert(doc, "",
        "NO use k_BIPV = 1.0 para una fachada BIPV típica. k=1.0 es solo para sistemas "
        "con espacio de ventilación libre mayor a 10 cm.", critical=True)
    add_alert(doc, "",
        "Si el auto-llenado NO aparece (no hay banner verde): regresar a Página 1, "
        "seleccionar el panel y hacer clic en Guardar configuración.", critical=True)
    add_info(doc,
        "Si la sobreestimación supera el 15% (banner naranja), significa que el sistema "
        "sin correcciones estaba calculando una producción significativamente inflada. "
        "Incluir siempre este resultado en la presentación al cliente.")

    # ── PASO 5 ──────────────────────────────────────────────────────────
    add_heading(doc, "PASO 5  —  Página 6: Producción Anual", level=2, color=C_AZUL)
    add_info(doc,
        "IMPORTANTE: Si ejecutó el PASO 4b (Motor Óptico), la Producción usará "
        "automáticamente la POA corregida hora a hora. Un BANNER VERDE confirmará "
        "que las correcciones ópticas están activas antes de calcular. Si el banner "
        "no aparece, regrese al PASO 4b y calcule la cascada óptica primero.")

    simple_table(doc,
        ["Campo", "Descripción"],
        [
            ("Número de paneles", "Viene del Dimensionamiento automáticamente. Editable si necesita ajustar la cantidad."),
            ("Eficiencia del inversor (%)", "Eficiencia CEC o Euro del inversor según ficha técnica. Rango típico: 96-98%."),
        ],
        col_widths=[5.5, 12.7]
    )
    add_body(doc, "Hacer clic en  Calcular Producción.", bold=True, indent=0.5)

    add_success(doc, "",
        "[OK]  Aparece 'Producción calculada' con: Energía AC anual en kWh/año, "
        "Performance Ratio real del sistema, gráficas mensuales y heatmap horario.")

    add_alert(doc, "",
        "NO continúe al Financiero sin completar este paso. Sin producción calculada, "
        "el análisis financiero usará 0 kWh, haciendo la TIR = 0% o negativa.", critical=True)
    add_info(doc,
        "PR > 100% es POSIBLE en climas fríos (Bogotá, Manizales). NO es un error. "
        "El modelo SDM captura la ganancia de producción de los paneles a bajas temperaturas.")

    # ── PASO 6 ──────────────────────────────────────────────────────────
    add_heading(doc, "PASO 6  —  Página 8: Presupuesto  (opcional pero recomendado)", level=2, color=C_AZUL)
    add_info(doc,
        "Por qué hacerlo antes del Financiero: el Presupuesto con costos reales de "
        "cotización da un CAPEX mucho más preciso que el modelo paramétrico general del Financiero.")

    simple_table(doc,
        ["Pestaña", "Qué contiene"],
        [
            ("Materiales", "Estructura de soporte, cableado, protecciones eléctricas, obra civil menor"),
            ("Mano de Obra", "Instalación eléctrica, puesta en marcha y pruebas"),
            ("Sistema FV", "Paneles solares: cantidad (pre-llenada del Dimensionamiento) y precio por unidad"),
            ("Inversor", "Inversor: cantidad y precio unitario en USD"),
            ("Catálogo", "Líneas de costo adicionales personalizadas para el proyecto específico"),
        ],
        col_widths=[3.5, 14.7]
    )

    add_success(doc, "",
        "[OK]  El CAPEX total en USD queda visible al final de la página. Este valor "
        "se transfiere automáticamente al Financiero cuando activa el toggle 'Usar CAPEX del Presupuesto'.")

    add_alert(doc, "",
        "ALERTA CRÍTICA — Costo/Wp > USD 5.0: Si aparece la advertencia, hay un error "
        "de unidades en alguna fila. Las columnas de precio son en USD. Si ingresó un valor "
        "en COP (ej: $18.000.000) en lugar de USD (ej: $5.000), el total queda inflado "
        "por un factor de ~3.600. Revisar cada fila.", critical=True)
    add_alert(doc, "",
        "NO mezcle COP y USD en la misma tabla. TODOS los valores del Presupuesto "
        "deben estar en USD. Divida los valores en COP por la TRM actual.", critical=True)
    add_info(doc,
        "Rango saludable para Colombia: entre USD 1.50 y USD 4.00 por Wp instalado. "
        "Valores fuera de este rango casi siempre indican error en la entrada de datos.")

    # ── PASO 7 ──────────────────────────────────────────────────────────
    add_heading(doc, "PASO 7  —  Página 7: Análisis Financiero", level=2, color=C_AZUL)
    add_body(doc, "Qué hacer:", bold=True)
    add_body(doc,
        "1.  Abrir [7] Financiero.\n"
        "2.  Activar el toggle 'Usar CAPEX del Presupuesto' si completó el Paso 6.\n"
        "3.  Si NO usó el Presupuesto, llenar el CAPEX paramétrico manualmente.", indent=0.3)

    simple_table(doc,
        ["Campo", "Descripción", "Valor típico"],
        [
            ("TRM (COP/USD)", "Tasa de cambio del día. Actualizar siempre.", "4.200 (ajustar al día)"),
            ("Escalación tarifa (%/año)", "Aumento anual esperado del precio del kWh", "5%"),
            ("Degradación módulos (%/año)", "Pérdida anual de producción por envejecimiento", "0.5%"),
            ("OPEX (% del CAPEX/año)", "Costo anual de mantenimiento y operación", "1.0%"),
            ("Tasa de descuento (%)", "Costo de oportunidad del capital del inversionista", "10-12%"),
            ("Horizonte (años)", "Vida útil del proyecto para el análisis", "25"),
            ("Tasa renta corporativa (%)", "Para calcular el beneficio del Art. 11 de la Ley 1715", "35% (empresas)"),
        ],
        col_widths=[5, 9, 4.2]
    )
    add_body(doc, "Hacer clic en  Calcular Análisis Financiero.", bold=True, indent=0.5)

    add_success(doc, "",
        "[OK]  La página muestra: TIR (%), VPN (USD), Payback simple y descontado (años), "
        "LCOE (USD/kWh), Beneficios Ley 1715 (Artículos 11, 12 y 14) y gráfica de flujo de caja acumulado.")

    add_alert(doc, "",
        "Si aparece 'El CAPEX cambió — Recalcula': significa que modificó el Presupuesto "
        "después de calcular el Financiero. Siempre volver a hacer clic en Calcular.", critical=True)
    add_alert(doc, "",
        "TIR negativa o payback > 20 años casi siempre indica: (1) CAPEX inflado por error "
        "COP/USD, (2) tarifa eléctrica en 0, o (3) Producción no calculada.", critical=True)
    add_info(doc,
        "La Ley 1715 aplica SOLO para personas jurídicas (empresas). Para proyectos "
        "residenciales los beneficios del Art. 11 y Art. 12 NO aplican. "
        "Dejar la tasa de renta corporativa = 0% en ese caso.")

    # ══ PASO 8 — FASE 4: BATERÍAS ══════════════════════════════════════
    page_break(doc)

    # Encabezado especial FASE 4
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, C_MORADO)
    cp = cell.paragraphs[0]
    cp.clear()
    r = cp.add_run("🔋  PASO 8  —  Página 11: Baterías y Balance Energético  ·  FASE 4")
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = WHITE
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    add_body(doc,
        "Qué hace esta página:",
        bold=True, color=C_MORADO)
    add_body(doc,
        "Permite dimensionar un sistema de almacenamiento de energía (batería) para el proyecto "
        "y calcular el balance energético mensual: cuánto produce el sistema solar, cuánto consume "
        "el inmueble, cuánto carga/descarga la batería y cuánto sigue dependiendo de la red eléctrica. "
        "El resultado clasifica el proyecto como A+, A, B, C o D según su fracción solar.",
        sz=9, indent=0.3)

    # ╔══ BLOQUE CRÍTICO FASE 4 ══════════════════════════════════════════╗
    doc.add_paragraph()
    add_heading(doc, "⚠  NOTAS CRÍTICAS ANTES DE USAR ESTA PÁGINA", level=3, color=C_ROJO)

    add_fase4_note(doc,
        " PREREQUISITO OBLIGATORIO: La Página 6 (Producción) DEBE estar calculada antes "
        "de abrir la Página 11. La página de Baterías lee directamente la tabla mensual de "
        "producción (df_mensual_produccion) de la sesión. Si llega a esta página sin haber "
        "calculado la Producción, el selector de baterías aparecerá pero el botón "
        "'Calcular balance energético' generará un error o retornará ceros.")

    add_note_box(doc, "⚡", "NOTA TENSIÓN:",
        " TODAS las baterías del catálogo actual son de ALTA TENSIÓN (300-870 V). "
        "Ninguna es compatible con inversores de 48V (DEYE SUN-7.6K, APsystems AHS). "
        "Estas baterías requieren inversores híbridos de alta tensión (>300V) o inversores "
        "comerciales/industriales dedicados. Si su proyecto usa inversor de 48V, las baterías "
        "del catálogo actual NO son compatibles eléctrica ni técnicamente.",
        C_ROJO_CLR, C_ROJO)

    add_note_box(doc, "📋", "NOTA DATOS INCOMPLETOS:",
        " En el catálogo actual, los campos DoD (%), Eficiencia RTE (%), Garantía (años), "
        "Temperatura de operación y Costo (USD) están vacíos para todos los modelos. "
        "La aplicación usa valores por defecto seguros: DoD = 80%, RTE = 95%, Ciclos = 3.000 "
        "(para series BR/BC) o 6.000 (para ATESS ESS). "
        "Los resultados del dimensionamiento son válidos como estimación preliminar, "
        "pero deben revisarse cuando el proveedor entregue los datos reales de su ficha técnica.",
        C_AMARILLO, C_NARANJA)

    add_note_box(doc, "💱", "NOTA FLUJO DE DATOS:",
        " Una vez dimensionada la batería, si se ingresa su costo en el campo 'Costo USD' "
        "dentro de la ficha técnica del catálogo, ese valor se transfiere automáticamente "
        "a la Página 8 (Presupuesto) como una fila adicional 'Batería — [nombre modelo]'. "
        "Si el campo Costo está vacío (situación actual), la fila NO aparece en el Presupuesto.",
        C_AZUL_CLR, C_AZUL)
    # ╚══════════════════════════════════════════════════════════════════╝

    doc.add_paragraph()
    add_heading(doc, "Bloque B-6: Dimensionado de la Batería", level=3, color=C_MORADO)
    add_body(doc, "Qué hacer:", bold=True)
    add_body(doc,
        "1.  Abrir  [11] 🔋 Baterías y Balance  en la barra lateral.\n"
        "2.  Ingresar el consumo mensual del inmueble (kWh/mes) si no viene pre-llenado.\n"
        "3.  Seleccionar el modelo de batería en el selector desplegable.\n"
        "4.  Revisar la ficha técnica que aparece automáticamente al seleccionar el modelo.\n"
        "5.  Hacer clic en  ▶️ Dimensionar batería.", indent=0.3)

    add_success(doc, "",
        "[OK]  Aparecen 8 métricas: Unidades necesarias, Capacidad bruta (kWh), "
        "Capacidad útil real (kWh), DoD aplicado (%), Vida estimada (años), "
        "Potencia continua total (kW), Voltaje nominal (V) y Costo total (USD si está disponible).")

    # Tabla ficha técnica
    add_body(doc, "Campos de la ficha técnica del catálogo de baterías:", bold=True, sz=9)

    simple_table(doc,
        ["Campo", "Descripción", "Valor por defecto si falta"],
        [
            ("Capacidad (kWh)", "Capacidad bruta nominal por unidad de batería", "Requerido — sin este dato no calcula"),
            ("DoD (%)", "Profundidad de descarga máxima recomendada por el fabricante", "80 % (conservador)"),
            ("Eficiencia RTE (%)", "Eficiencia de ida+vuelta del ciclo carga-descarga", "95 %"),
            ("Ciclos de Vida", "Número de ciclos completos de carga/descarga antes del 80% de capacidad", "3.000 (BR series) / 6.000 (ATESS)"),
            ("Potencia Continua (kW)", "Potencia máxima sostenida de descarga", "50% de la capacidad en kWh"),
            ("Voltaje Nominal (V)", "Tensión nominal del banco de baterías", "Según modelo — ver tabla catálogo"),
            ("Costo (USD)", "Precio de lista de la batería. VACÍO en catálogo actual.", "Sin costo → no aparece en Presupuesto"),
        ],
        col_widths=[4.5, 8, 5.7]
    )

    doc.add_paragraph()
    add_heading(doc, "Bloque B-7: Balance Energético Mensual", level=3, color=C_MORADO)
    add_body(doc, "Qué hacer:", bold=True)
    add_body(doc,
        "1.  Ingresar el consumo mensual (kWh/mes) si no está ya ingresado en B-6.\n"
        "2.  Hacer clic en  ▶️ Calcular balance energético.", indent=0.3)

    add_success(doc, "",
        "[OK]  Aparece el gráfico de barras apiladas mensual (Solar directa / Batería / Red) "
        "con la línea de consumo superpuesta, y la Clasificación Energética del proyecto.")

    add_body(doc, "Clasificación Energética del proyecto:", bold=True, sz=9)

    simple_table(doc,
        ["Clase", "Fracción Solar (%)", "Significado"],
        [
            ("A+", "≥ 90%", "Sistema casi completamente autónomo. La batería cubre casi toda la noche."),
            ("A",  "75–90%", "Alta autosuficiencia. Dependencia mínima de la red."),
            ("B",  "50–75%", "Buena cobertura solar. Complemento moderado de la red."),
            ("C",  "25–50%", "Cobertura parcial. La batería reduce pero no elimina la dependencia de la red."),
            ("D",  "< 25%", "Baja autosuficiencia. Revisar capacidad del sistema o del consumo base."),
        ],
        col_widths=[1.5, 4.5, 12.2]
    )

    add_body(doc, "ALERTAS — No cometer estos errores en la Página 11:", bold=True, color=C_ROJO)

    add_alert(doc, "",
        "NO abra la Página 11 sin haber calculado la Producción (Página 6). "
        "El balance mensual requiere la tabla E_ac por mes. Sin ella, el cálculo "
        "retorna ceros o genera un error silencioso.", critical=True)

    add_alert(doc, "",
        "NO seleccione una batería de alta tensión (300-870V) con el inversor DEYE 48V "
        "o APsystems AHS del proyecto. Son incompatibles eléctricamente. "
        "La aplicación mostrará una advertencia en una versión futura; por ahora "
        "el usuario es responsable de verificar la compatibilidad.", critical=True)

    add_alert(doc, "",
        "NO interprete la Vida Estimada como garantía del fabricante. Se calcula como "
        "Ciclos_de_Vida / (365 días/año) asumiendo un ciclo diario. Si el perfil de "
        "uso real tiene menos de un ciclo completo por día, la vida real será mayor.",
        critical=False)

    add_info(doc,
        "Los valores de ahorro por batería (reducción de compra a la red) NO se integran "
        "todavía en la TIR y el Payback de la Página 7 (Financiero). Esta integración "
        "está programada como mejora futura. Mencionarlo al cliente si preguntan por el "
        "retorno de inversión de la batería.")

    add_info(doc,
        "Si el proyecto no requiere batería (cliente sin necesidad de autonomía nocturna), "
        "esta página puede omitirse sin afectar ningún otro cálculo.")

    # Catálogo disponible
    doc.add_paragraph()
    add_heading(doc, "Catálogo de baterías disponible — Versión actual", level=3, color=C_MORADO)
    add_body(doc,
        "El catálogo incluye 26 modelos en dos familias. TODOS son de alta tensión (300-870V):",
        sz=9)

    simple_table(doc,
        ["Familia", "Modelos", "Capacidad (kWh)", "Voltaje nominal", "Estado datos"],
        [
            ("Serie BR — Fabricante pendiente confirmar",
             "BR172R, BR186R, BR200R, BR215R",
             "172 – 215 kWh",  "538–864 V", "Incompleto — DoD/RTE/Costo pendientes"),
            ("ATESS ESS — BC/BR45T a BR145T (7.68 kWh/módulo)",
             "BC45T, BR45T, BC50T, BR50T, BC60T, BR60T, BC75T, BR75T, BC100T, BR100T, BR138T, BR145T",
             "46 – 146 kWh",  "384–730 V", "Incompleto — DoD/RTE/Costo pendientes"),
            ("ATESS ESS — BR114R a BR157R (14.336 kWh/módulo)",
             "BR114R, BR129R, BR143R, BR157R",
             "115 – 158 kWh",  "410–563 V", "Incompleto — DoD/RTE/Costo pendientes"),
            ("ATESS ESS — BC55RPB (5.12 kWh/módulo, IP54 exterior)",
             "BC55RPB-6M a BC55RPB-11M",
             "31 – 56 kWh",   "307–563 V", "Incompleto — DoD/RTE/Costo pendientes"),
        ],
        col_widths=[4, 5.5, 2.7, 2.5, 3.5]
    )

    add_note_box(doc, "📦", "CÓMO AGREGAR NUEVOS MODELOS AL CATÁLOGO:",
        " El catálogo vive en la hoja 'Catalogo_Baterias' del archivo "
        "inversores_catalogo.xlsx en el servidor. Para agregar un nuevo modelo: "
        "abra el Excel del servidor (ruta: /var/www/bipv/calculadora-bipv/bipv_python/datos/), "
        "agregue una fila con los campos obligatorios (Modelo, Capacidad kWh, Voltaje Nominal V, "
        "Ciclos de Vida, Potencia Continua kW) y guarde el archivo. "
        "El cambio se refleja en la app al refrescar la sesión (sin necesidad de reiniciar PM2). "
        "Para verificar que cargó correctamente: ejecutar "
        "python bipv_python/datos/diagnostico_catalogo_baterias.py en el servidor.",
        C_VERDE_CLR, C_VERDE_OSC)

    # ── PASO 9 (Motor IV) ───────────────────────────────────────────────
    page_break(doc)
    add_heading(doc, "PASO 9  —  Página 3: Motor IV  (opcional / técnico)", level=2, color=C_AZUL)
    add_info(doc,
        "Este paso es para validación técnica del panel elegido. NO afecta los "
        "cálculos de producción ni los resultados financieros. Puede hacerse en "
        "cualquier momento del flujo.")

    add_body(doc,
        "1.  Abrir  [3] Motor IV.\n"
        "2.  Seleccionar el mismo panel elegido en el Dimensionamiento.\n"
        "3.  El motor calcula la curva I-V y P-V y valida los resultados contra la ficha técnica.",
        indent=0.3)

    add_success(doc, "",
        "[OK]  Errores de validación menores al 5% en Voc, Isc, Vmp y Pmax. "
        "Aparece el mensaje 'Parámetros validados'.")

    add_info(doc,
        "Si el error de validación es mayor al 5% para algún parámetro, la ficha "
        "técnica del panel puede estar incompleta o los coeficientes de temperatura "
        "son incorrectos. Notificar al administrador del catálogo para corregirlo.")

    # ── PASO 10 (Reporte PDF) ────────────────────────────────────────────
    add_heading(doc, "PASO 10  —  Página 10: Reporte PDF del Proyecto  (opcional)", level=2, color=C_AZUL)
    add_info(doc,
        "Hacer DESPUÉS de completar todos los cálculos. El reporte captura el estado "
        "actual de sesión: si una sección no se ha calculado, no aparecerán sus resultados.")

    add_body(doc,
        "Qué hace esta página: Genera un informe técnico descargable (.html) con todos "
        "los resultados del proyecto. El reporte incluye: resumen del proyecto, recurso solar, "
        "cascada del Motor Óptico, métricas de producción IEC 61724, análisis financiero Ley 1715 "
        "y — si se calculó la Página 11 — el balance energético y clasificación de la batería.",
        sz=9, indent=0.3)

    add_body(doc,
        "1.  Abrir  [10] Reporte PDF  en la barra lateral.\n"
        "2.  Verificar el panel de Estado del proyecto en la parte superior.\n"
        "3.  Llenar: Nombre del proyecto, Nombre de la empresa, checkboxes de secciones.\n"
        "4.  Hacer clic en  Generar Reporte.\n"
        "5.  Hacer clic en  Descargar reporte (.html).\n"
        "6.  Abrir el archivo en Chrome, Edge o Firefox.\n"
        "7.  Para PDF: Ctrl+P → Guardar como PDF → Márgenes: Mínimo → Guardar.",
        indent=0.3)

    add_success(doc, "",
        "[OK]  La descarga del archivo .html se inicia automáticamente. Al abrirlo, "
        "el reporte muestra el encabezado de la empresa, todas las secciones con sus "
        "tablas de resultados y el aviso de BORRADOR en la parte superior.")

    add_alert(doc, "",
        "NO entregue el reporte al cliente sin revisar que todos los checkmarks del "
        "estado estén activos para las secciones críticas (Producción y Financiero al menos).",
        critical=True)
    add_alert(doc, "",
        "NO abrir el archivo HTML en Internet Explorer o Safari antiguo. "
        "Usar Chrome, Edge o Firefox.", critical=False)
    add_info(doc,
        "Si el Motor Óptico fue calculado, el reporte incluirá la tabla de cascada óptica. "
        "Esta sección diferencia el análisis de INNOVACIÓN QUÍMICA de una cotización genérica "
        "del mercado — es un argumento técnico de venta.")

    # ══ PARTE B ═══════════════════════════════════════════════════════════
    page_break(doc)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, C_AZUL)
    cp = cell.paragraphs[0]
    cp.clear()
    r = cp.add_run("PARTE B  —  MODO CONSUMO / FACTURA  [B]")
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = WHITE
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    add_body(doc,
        '"Tengo la factura eléctrica mensual o sé exactamente cuántos kWh/mes consume el inmueble."',
        italic=True, color=C_GRIS)

    add_heading(doc, "PASO 1  —  Página 1: Datos del Proyecto  (diferente al Modo Área)", level=2, color=C_AZUL)

    simple_table(doc,
        ["Campo", "Descripción", "Ejemplo"],
        [
            ("Nombre del proyecto", "Identificador del proyecto", "Fábrica Norte"),
            ("Ciudad", "Seleccionar de la lista", "Medellín"),
            ("Tarifa eléctrica (COP/kWh)", "Valor exacto del kWh según su factura", "580"),
            ("Factura mensual (COP)", "Valor total de la última factura de energía eléctrica", "1.500.000"),
            ("— O bien — Consumo (kWh/mes)", "Si conoce el consumo directo. NO usar junto con la factura.", "2.500"),
            ("% Cobertura solar deseada", "Porcentaje del consumo mensual que quiere cubrir con solar", "80"),
            ("Densidad de potencia (W/m²)", "Del panel elegido. Ver ficha técnica.", "200"),
            ("Performance Ratio — PR (%)", "Eficiencia real. Dejar en 80% si no sabe.", "80"),
        ],
        col_widths=[5.5, 8.5, 4.2]
    )
    add_body(doc, "Hacer clic en  Guardar configuración.", bold=True, indent=0.5)

    add_success(doc, "",
        "[OK]  La calculadora muestra: consumo mensual estimado, Área necesaria para "
        "alcanzar la cobertura deseada y un semáforo VERDE/ROJO según la factibilidad del área.")

    add_alert(doc, "",
        "NO ingrese factura Y consumo al mismo tiempo. Use uno de los dos solamente. "
        "Si ingresa ambos, la calculadora usará la factura.", critical=True)
    add_alert(doc, "",
        "Si la tarifa está en 0 o es incorrecta, el cálculo del consumo derivado de la "
        "factura será erróneo. Verifique el valor exacto del kWh en su factura.", critical=True)

    add_heading(doc, "PASOS 2 al 10  —  Idénticos al Modo Área", level=3, color=C_GRIS)
    add_info(doc,
        "Una vez guardada la configuración del Proyecto en Modo Consumo, el resto del flujo "
        "es EXACTAMENTE IGUAL al Modo Área (Pasos 2 al 10 de la Parte A). La única diferencia "
        "entre los dos modos es el punto de partida: en Modo Área usted define el espacio "
        "disponible y la calculadora estima la producción; en Modo Consumo usted define el "
        "consumo y la calculadora determina el espacio mínimo necesario. "
        "Continuar desde el Paso 2 (Recurso Solar) de la Parte A.")

    # ══ ERRORES FRECUENTES ════════════════════════════════════════════════
    page_break(doc)
    add_heading(doc, "ERRORES FRECUENTES  —  Diagnóstico y solución rápida", level=1, color=C_ROJO)

    simple_table(doc,
        ["Error observado en pantalla", "Causa probable", "Solución inmediata"],
        [
            ("Financiero muestra TIR = 0% o negativa",
             "Producción no calculada o CAPEX inflado por error COP/USD",
             "Completar Página 6 antes de ir a Página 7. Revisar costo/Wp en Presupuesto."),
            ("Costo por Wp = USD 12 o más en Presupuesto",
             "Uno o varios valores de la tabla están en COP en lugar de USD",
             "Identificar la fila anormalmente alta y dividirla por la TRM (~4.200)."),
            ("Inversor aparece a $0/unidad en Presupuesto",
             "Columna 'Costo Inversor' vacía en el Excel del catálogo",
             "Ingresar el precio del inversor manualmente en la pestaña Catálogo del Presupuesto."),
            ("Recurso Solar da error de PVGIS",
             "Sin conexión a internet del servidor o PVGIS temporalmente caído",
             "Verificar conexión. Reintentar pasados unos minutos. Tiempo normal: 30-90 seg."),
            ("Producción parece muy alta (> 1.800 kWh/kWp/año)",
             "Azimuth o Tilt incorrectos en Recurso Solar",
             "Corregir la orientación en Página 2 y recalcular desde ahí."),
            ("Página Mismatch no calcula",
             "Recurso Solar no ha sido calculado previamente",
             "Completar la Página 2 (Recurso Solar) primero."),
            ("Dimensionamiento — advertencia de tensión",
             "Demasiados paneles en serie para el inversor seleccionado",
             "Reducir N_serie o seleccionar inversor con mayor Vmax de entrada."),
            ("'El CAPEX cambió — Recalcula' en Financiero",
             "Se modificó el Presupuesto después de calcular",
             "Hacer clic nuevamente en Calcular en la Página 7."),
            ("TIR muestra 'M COP' en lugar de '%'",
             "Error de visualización del formato numérico",
             "Actualizar la página (F5) y volver a calcular."),
            ("Motor Óptico no muestra banner de auto-llenado",
             "No hay panel guardado en sesión desde la Página 1",
             "Regresar a Página 1, seleccionar el panel y hacer clic en Guardar configuración."),
            ("Banner en Producción dice 'POA bruta' en lugar de 'Motor Óptico activo'",
             "El Motor Óptico no se calculó antes de ir a Producción",
             "Ejecutar el PASO 4b: ir a Página 5b, configurar y calcular cascada óptica. Luego volver a Producción."),
            ("Factor global del Motor Óptico < 70%",
             "Parámetros incorrectos: b0 o k_BIPV demasiado altos",
             "Verificar tipo de vidrio (b0) y tipo de montaje (k_BIPV). Fachada típica Colombia: factor 78-88%."),
            # ── NUEVOS: FASE 4 ─────────────────────────────────────────
            ("🔋 Página 11 — 'Catálogo no encontrado' o dropdown vacío",
             "La hoja 'Catalogo_Baterias' no existe en el Excel del servidor o el loader no detectó la fila de encabezados",
             "En el servidor ejecutar: python bipv_python/datos/diagnostico_catalogo_baterias.py y seguir las instrucciones que muestre."),
            ("🔋 Página 11 — Balance muestra ceros o error al calcular",
             "La Página 6 (Producción) no ha sido calculada en esta sesión",
             "Regresar a Página 6 y hacer clic en Calcular Producción. Luego volver a Página 11."),
            ("🔋 Página 11 — Vida estimada muy corta (< 5 años)",
             "El campo Ciclos de Vida del modelo seleccionado está vacío → usa default 3.000 ciclos",
             "Solicitar al proveedor los ciclos de vida reales e ingresar el dato en la hoja Catalogo_Baterias del Excel del servidor."),
            ("🔋 Página 11 — Batería no aparece en Presupuesto",
             "El campo Costo (USD) del modelo está vacío en el catálogo",
             "Ingresar el precio de la batería en la hoja Catalogo_Baterias del Excel. O bien, agregar la fila manualmente en la pestaña Catálogo del Presupuesto."),
            ("Reporte PDF — secciones vacías o sin datos",
             "Se generó el reporte antes de calcular esas secciones",
             "Completar todas las páginas requeridas y luego volver a Página 10 para regenerar."),
            ("Archivo HTML descargado se ve mal o sin formato",
             "Abrió el archivo en un navegador desactualizado o en el visor de Windows",
             "Abrir con Chrome, Edge o Firefox. Click derecho → Abrir con → Chrome."),
        ],
        col_widths=[5.2, 5.5, 7.5]
    )

    # ══ CHECKLIST ════════════════════════════════════════════════════════
    page_break(doc)
    add_heading(doc, "CHECKLIST  —  Verificación antes de entregar resultados al cliente", level=1, color=C_VERDE_OSC)
    add_body(doc,
        "Marque cada punto ANTES de presentar un análisis al cliente. "
        "Si alguno falla, corríjalo antes de continuar.",
        italic=True, color=C_GRIS, sz=9)

    checks = [
        ("Pág. 1",
         "Ciudad correcta, tarifa eléctrica real, área o consumo correctamente ingresado y guardado."),
        ("Pág. 2",
         "Azimuth y tilt verificados para la orientación real. Recurso Solar calculado sin errores de PVGIS."),
        ("Pág. 4",
         "Panel con ficha técnica completa (semáforo verde), inversor compatible seleccionado, "
         "todas las verificaciones de tensión en verde."),
        ("Pág. 5",
         "Mismatch calculado. Aunque sea con todas las pérdidas en cero, debe estar calculado."),
        ("Pág. 5b",
         "Motor Óptico calculado (RECOMENDADO): factor global entre 75-90%. Si es < 70% revisar b0 y k_BIPV. "
         "Banner de auto-llenado apareció correctamente."),
        ("Pág. 6",
         "Producción calculada: E_ac anual > 0 kWh/año. "
         "Banner confirma 'Motor Óptico activo' si se ejecutó el Paso 4b."),
        ("Pág. 8",
         "Presupuesto revisado línea por línea: costo/Wp entre USD 1.50 y USD 4.00. TRM actualizada al día."),
        ("Pág. 7",
         "TRM actualizada, toggle de Presupuesto activado (si aplica), recalculado después de cualquier cambio. "
         "TIR, VPN y Payback con sentido económico razonable."),
        ("Pág. 11 🔋",
         "[FASE 4 — si el proyecto incluye batería] "
         "Modelo seleccionado revisado: verificar que la tensión de la batería sea compatible con el inversor del proyecto. "
         "Dimensionamiento ejecutado y revisado. Balance energético calculado. "
         "Clasificación energética A+/A/B/C/D visible. "
         "Si el costo de la batería está disponible, verificar que aparece en el Presupuesto."),
        ("Pág. 10",
         "Reporte generado y revisado: checkmarks de Estado activos para secciones críticas. "
         "Leer el PDF generado completo antes de entregar al cliente."),
    ]

    for tag, text in checks:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r1 = p.add_run("☐  ")
        r1.font.size = Pt(11); r1.font.color.rgb = C_VERDE_OSC; r1.bold = True
        r2 = p.add_run(f"{tag}  —  ")
        r2.font.size = Pt(9); r2.font.color.rgb = C_AZUL; r2.bold = True
        r3 = p.add_run(text)
        r3.font.size = Pt(9); r3.font.color.rgb = C_GRIS

    # ══ PIE DE PÁGINA ═════════════════════════════════════════════════════
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "BORRADOR  |  Manual de Operación  |  Calculadora BIPV Colombia v2026  |  "
        "FASE 4: Baterías y Balance Energético  |  "
        "Innovación Química  |  calc.innovacionquimica.com.co"
    )
    r.font.size = Pt(7); r.font.color.rgb = C_GRIS; r.italic = True

    # ── Guardar ──────────────────────────────────────────────────────────
    os.makedirs("attached_assets", exist_ok=True)
    doc.save(OUT)
    print(f"✅ Manual generado: {OUT}")

if __name__ == "__main__":
    build_manual()
