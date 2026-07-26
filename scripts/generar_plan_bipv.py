
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

AZUL    = RGBColor(0x1A, 0x5C, 0x8A)
VERDE   = RGBColor(0x17, 0x6B, 0x17)
MORADO  = RGBColor(0x6E, 0x27, 0x94)
ROJO    = RGBColor(0xC0, 0x39, 0x2B)
NARANJA = RGBColor(0xD4, 0x7A, 0x00)
GRIS    = RGBColor(0x7F, 0x7F, 0x7F)
COD     = RGBColor(0x10, 0x10, 0x60)

def cell_shade(cell, hex_color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)

def sep(color='BBBBBB'):
    p = doc.add_paragraph('─' * 78)
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(0xBB,0xBB,0xBB)

def h1(txt, color=AZUL):
    doc.add_paragraph('')
    h = doc.add_heading(txt, level=1)
    h.runs[0].font.size = Pt(14)
    h.runs[0].font.color.rgb = color

def h2(txt, color=MORADO):
    h = doc.add_heading(txt, level=2)
    h.runs[0].font.size = Pt(12)
    h.runs[0].font.color.rgb = color

def h3(txt, color=NARANJA):
    h = doc.add_heading(txt, level=3)
    h.runs[0].font.size = Pt(11)
    h.runs[0].font.color.rgb = color

def body(txt, justify=True):
    p = doc.add_paragraph(txt)
    p.runs[0].font.size = Pt(10)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def bullet(txt, color=AZUL):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(txt)
    r.font.size = Pt(10)
    r.font.color.rgb = color

def badge(icon, label, txt, color=AZUL):
    p = doc.add_paragraph()
    r1 = p.add_run(f'{icon} {label}  ')
    r1.bold = True
    r1.font.color.rgb = color
    r2 = p.add_run(txt)
    r2.font.size = Pt(10)

def codigo(txt):
    p = doc.add_paragraph()
    cr = p.add_run(txt)
    cr.font.name = 'Courier New'
    cr.font.size = Pt(9)
    cr.font.color.rgb = COD
    doc.add_paragraph('')

# ═══════════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════════
t = doc.add_heading('PLAN MAESTRO — CALCULADORA BIPV EN PYTHON', 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.runs[0].font.color.rgb = AZUL

st = doc.add_paragraph('Hoja de ruta completa: mejoras al repositorio existente + versión Python Streamlit')
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
st.runs[0].font.bold = True
st.runs[0].font.size = Pt(12)
st.runs[0].font.color.rgb = MORADO

intro_p = doc.add_paragraph(
    'Este plan tiene DOS partes independientes:\n'
    '  PARTE A — Mejoras al repositorio TypeScript/React ya existente (7 pasos)\n'
    '  PARTE B — Construcción de la calculadora completa en Python con Streamlit (10 pasos)\n\n'
    'Puedes seguir solo una parte, o ambas en paralelo. '
    'Cada paso incluye: objetivo, archivos involucrados, código de referencia, '
    'y cómo verificar que el paso está completo.'
)
intro_p.runs[0].font.size = Pt(10)
intro_p.runs[0].italic = True
intro_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
doc.add_paragraph('')

# ═══════════════════════════════════════════════════════════════
# RESPUESTA AL RETO — PYTHON VS TYPESCRIPT
# ═══════════════════════════════════════════════════════════════
sep()
h1('RESPUESTA AL RETO: ¿TODA LA CALCULADORA EN PYTHON?')

body(
    'Sí — es técnicamente posible y en ingeniería fotovoltaica es incluso preferible. '
    'Python domina el sector: PVLib, pvfactors, SAM SDK, y la mayoría de las '
    'herramientas profesionales de simulación solar están escritas en Python. '
    'La limitación real es el visualizador 3D de fachadas (Three.js), '
    'que no tiene equivalente directo en Python web. Para todo lo demás, '
    'Python con Streamlit es igual o mejor que TypeScript/React.'
)
doc.add_paragraph('')

# Tabla comparativa
tabla_comp = doc.add_table(rows=1, cols=4)
tabla_comp.style = 'Table Grid'
headers = ['Funcionalidad', 'Calculadora actual\n(TypeScript/React)', 'Python Streamlit', 'Python es mejor?']
hdr_cells = tabla_comp.rows[0].cells
for i, h_txt in enumerate(headers):
    hdr_cells[i].text = h_txt
    cell_shade(hdr_cells[i], '1A5C8A')
    for par in hdr_cells[i].paragraphs:
        for run in par.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            run.font.size = Pt(9)

filas_comp = [
    ('Cálculo de producción IEC 61724', 'energyProduction.ts', 'pvlib + numpy → más preciso', '✅ Sí'),
    ('Modelo Mulcue-Llanos PR', 'mulcueLlanos.ts', 'Función Python idéntica', '= Igual'),
    ('IAM-ASHRAE + Soiling', 'iamSoilingEngine.ts', 'pvlib.irradiance tiene IAM integrado', '✅ Sí'),
    ('Transposición Perez/Liu-Jordan', 'liuJordanModel.ts', 'pvlib.irradiance.perez() oficial', '✅ Sí'),
    ('Análisis financiero (VAN/TIR)', 'bipvROIOptimizer.ts', 'numpy-financial.npv / irr', '✅ Sí'),
    ('Gráficos interactivos', 'Recharts (SVG)', 'Plotly Express / Altair', '= Similar'),
    ('Exportar Word/Excel/PDF', 'bipvReportGenerator.ts', 'python-docx / openpyxl / reportlab', '✅ Sí'),
    ('Datos meteorológicos EPW', 'epwParser.ts', 'pvlib.iotools.read_epw()', '✅ Sí, nativo'),
    ('PVGIS / PVWatts API', 'pvgisApi.ts + pvwattsApi.ts', 'pvlib.iotools.get_pvgis_hourly()', '✅ Sí, nativo'),
    ('Catálogo de paneles', 'panelTechnologies.ts', 'pvlib.pvsystem.retrieve_sam()', '✅ Sí, 10k paneles'),
    ('Diagrama solar interactivo', 'SunPath3D.tsx + Three.js', 'pvlib.solarposition → Plotly 3D', '= Parcial'),
    ('Visualizador 3D GLTF fachadas', 'Three.js + OBJ/GLTF parser', '⚠ No tiene equivalente exacto', '❌ No aplica'),
]

for fila in filas_comp:
    row = tabla_comp.add_row().cells
    for i, val in enumerate(fila):
        row[i].text = val
        if row[i].paragraphs[0].runs:
            run = row[i].paragraphs[0].runs[0]
        else:
            run = row[i].paragraphs[0].add_run(val)
        run.font.size = Pt(9)
        if i == 3:
            if '✅' in val: run.font.color.rgb = VERDE; run.bold = True
            elif '❌' in val: run.font.color.rgb = ROJO; run.bold = True

doc.add_paragraph('')
badge('💡', 'VEREDICTO:',
      'Python con Streamlit puede hacer el 95% de la calculadora. '
      'El 5% restante es el visualizador 3D de obstáculos y la importación de modelos GLTF/OBJ. '
      'Para aprendizaje de Python fotovoltaico, la versión Streamlit es ideal. '
      'Para el producto web final, mantener TypeScript/React.', VERDE)
doc.add_paragraph('')


# ═══════════════════════════════════════════════════════════════
# PARTE A — MEJORAS AL REPOSITORIO TYPESCRIPT
# ═══════════════════════════════════════════════════════════════
sep()
h1('PARTE A — PLAN DE MEJORAS AL REPOSITORIO TYPESCRIPT (7 PASOS)')
body(
    'Estos pasos mejoran la calculadora que ya tienes en GitHub. '
    'Cada paso es independiente — puedes implementarlos en cualquier orden, '
    'aunque el orden recomendado maximiza el impacto.'
)
doc.add_paragraph('')

# ──────────────────────────────
# PASO A-1
# ──────────────────────────────
h2('PASO A-1 — String Sizing (Dimensionado eléctrico de strings) [PRIORIDAD ALTA]')
badge('⏱', 'Tiempo estimado:', '2-3 horas', AZUL)
badge('📁', 'Archivo nuevo a crear:', 'client/src/lib/stringSizing.ts', VERDE)
badge('📁', 'Archivo a modificar:', 'client/src/components/EnergyProductionSimulator.tsx', NARANJA)
doc.add_paragraph('')
body('OBJETIVO: Validar que la combinación de paneles en serie/paralelo es compatible con el inversor seleccionado. Sin esta validación, el usuario puede configurar un sistema que dañe el inversor.')
doc.add_paragraph('')
body('PASOS:')
bullet('1. Crear el archivo stringSizing.ts con la función calcularStringSizing() del documento de diagnóstico')
bullet('2. Agregar en EnergyProductionSimulator.tsx una nueva sección "Dimensionado Eléctrico" con campos para: Voc, Vmp, Isc, coef. temperatura, Vdc_max del inversor, rango MPPT, temperatura mínima del lugar')
bullet('3. Mostrar los resultados: N paneles en serie (mín/ópt/máx), N strings en paralelo, Voc en frío, Vmp caliente, ratio DC/AC, semáforo verde/rojo')
bullet('4. Si hay advertencias (tensión fuera de rango, ratio DC/AC inválido), mostrar alerta roja con texto explicativo')
doc.add_paragraph('')
badge('✅', 'Verificación:', 'Ingresar un sistema con Voc=40V, 20 paneles en serie → debe dar alerta si supera Vdc_max=800V del inversor.', VERDE)
doc.add_paragraph('')

# ──────────────────────────────
# PASO A-2
# ──────────────────────────────
h2('PASO A-2 — Factor de cobertura activa [PRIORIDAD ALTA]')
badge('⏱', 'Tiempo estimado:', '1 hora', AZUL)
badge('📁', 'Archivo a modificar:', 'client/src/lib/iamSoilingEngine.ts', NARANJA)
doc.add_paragraph('')
body('OBJETIVO: Agregar el parámetro factorCobertura (0-1) a BIPVSimulationConfig para que la potencia se calcule sobre el área activa real con celdas, no el área total de la fachada.')
doc.add_paragraph('')
body('PASOS:')
bullet('1. Agregar factorCobertura?: number al interface BIPVSimulationConfig (valor por defecto 0.85)')
bullet('2. Modificar la función que calcula la potencia DC: multiplicar areaM2 × factorCobertura antes de calcular')
bullet('3. Agregar en BIPVGlassSimulator.tsx un slider "Factor de cobertura activa" con valores sugeridos: Fachada continua (0.90), Muro cortina con marcos (0.80), Ventana integrada (0.65)')
bullet('4. Mostrar junto al slider el área activa resultante en m² para que el usuario entienda el impacto')
doc.add_paragraph('')
badge('✅', 'Verificación:', 'Con área=100m² y factorCobertura=0.80, la potencia pico debe ser 20% menor que con factorCobertura=1.00.', VERDE)
doc.add_paragraph('')

# ──────────────────────────────
# PASO A-3
# ──────────────────────────────
h2('PASO A-3 — Balance energético mensual + Clasificación A+/A/B/C/D [PRIORIDAD ALTA]')
badge('⏱', 'Tiempo estimado:', '3-4 horas', AZUL)
badge('📁', 'Archivo nuevo a crear:', 'client/src/lib/energyBalance.ts', VERDE)
badge('📁', 'Archivo a modificar:', 'client/src/components/BIPVROIOptimizer.tsx', NARANJA)
doc.add_paragraph('')
body('OBJETIVO: Reemplazar el porcentaje fijo de autoconsumo por un balance real mes a mes que compare la generación BIPV con el consumo real del edificio.')
doc.add_paragraph('')
body('PASOS:')
bullet('1. Crear energyBalance.ts con calcularBalanceEnergetico() del documento de diagnóstico')
bullet('2. En BIPVROIOptimizer.tsx agregar nueva pestaña "Balance Mensual" junto a los tabs existentes')
bullet('3. Agregar formulario de entrada: consumo mensual del edificio (12 campos de kWh/mes, o un valor anual dividido automáticamente con variación estacional)')
bullet('4. Mostrar tabla de 12 meses con columnas: Generación / Consumo / Autoconsumo / Excedente / Compra red / % Autosuficiencia')
bullet('5. Mostrar barra de progreso por mes coloreada: verde si el BIPV cubre > 80%, amarillo 40-80%, rojo < 40%')
bullet('6. Agregar tarjeta grande con la Clasificación Energética (A+/A/B/C/D) basada en la autosuficiencia anual promedio')
bullet('7. Actualizar el cálculo del VAN en el ROI para usar los valores reales mes a mes en vez del autoconsumo % fijo')
doc.add_paragraph('')
badge('✅', 'Verificación:', 'Con 10.000 kWh/año de generación y 15.000 kWh/año de consumo → clasificación C (67% autosuficiencia).', VERDE)
doc.add_paragraph('')

# ──────────────────────────────
# PASO A-4
# ──────────────────────────────
h2('PASO A-4 — Gráfico y tabla de producción con degradación en 25 años [PRIORIDAD MEDIA]')
badge('⏱', 'Tiempo estimado:', '2 horas', AZUL)
badge('📁', 'Archivo a modificar:', 'client/src/lib/bipvROIOptimizer.ts + BIPVROIOptimizer.tsx', NARANJA)
doc.add_paragraph('')
body('OBJETIVO: Agregar la función calcularProyeccion25Anos() y mostrar al usuario la curva de producción con degradación año a año, el flujo de caja acumulado, y el CO₂ evitado.')
doc.add_paragraph('')
body('PASOS:')
bullet('1. Agregar calcularProyeccion25Anos() al final de bipvROIOptimizer.ts (código en documento de diagnóstico)')
bullet('2. Crear nueva pestaña "Proyección 25 años" en BIPVROIOptimizer.tsx')
bullet('3. Gráfico 1 (Recharts LineChart): eje X = año 1-25, eje Y = kWh/año. Línea azul descendente con área rellena. Agregar línea horizontal en la producción del año 1 como referencia')
bullet('4. Gráfico 2 (Recharts BarChart): flujo de caja acumulado año a año. Barras rojas = antes del payback, barras verdes = después del payback. La barra donde el color cambia es el año de retorno')
bullet('5. Tarjeta de resumen: CO₂ evitado total en 25 años (toneladas), equivalente en árboles plantados (1 árbol ≈ 20 kg CO₂/año), kWh totales generados')
body('   Nota: el parámetro factorCO2 debe usar la función co2EvitadoTon() del Paso A-5 para que coincida con el país del proyecto.')
doc.add_paragraph('')
badge('✅', 'Verificación:', 'Sistema de 50 kWp, degradación 0.5%/año: la producción del año 25 debe ser ~88% de la del año 1 (0.995^24 ≈ 0.887).', VERDE)
doc.add_paragraph('')

# ──────────────────────────────
# PASO A-5
# ──────────────────────────────
h2('PASO A-5 — Convertidor de unidades + CO₂ por país [PRIORIDAD MEDIA]')
badge('⏱', 'Tiempo estimado:', '2 horas', AZUL)
badge('📁', 'Archivo nuevo a crear:', 'client/src/lib/unitConversions.ts', VERDE)
badge('📁', 'Archivo a modificar:', 'client/src/components/EnergyProductionSimulator.tsx', NARANJA)
doc.add_paragraph('')
body('OBJETIVO: Agregar un convertidor de unidades interactivo (°C/°F, W/kW/HP, Wh/kWh/MWh, m²/ft², HSP) y una tabla de factores de emisión CO₂ por país.')
doc.add_paragraph('')
body('PASOS:')
bullet('1. Crear unitConversions.ts con todas las funciones del documento de diagnóstico')
bullet('2. Agregar en EnergyProductionSimulator.tsx un panel colapsable "Convertidor de Unidades" al inicio del formulario')
bullet('3. El convertidor tiene 4 secciones: Temperatura / Potencia / Energía / Irradiancia-HSP')
bullet('4. En cada sección: dos campos de entrada, flecha ⇄ en el centro, conversión en tiempo real con onChange')
bullet('5. Agregar selector de país en los parámetros del proyecto (Chile, Colombia, México, Argentina, Perú, España, Global)')
bullet('6. Usar el factor de CO₂ del país seleccionado en la función co2EvitadoTon() para los reportes')
bullet('7. Mostrar el CO₂ evitado anual en la página principal del simulador (tarjeta pequeña junto a la producción anual)')
doc.add_paragraph('')
badge('✅', 'Verificación:', '1000 Wh/m²/día → HSP = 1.0. 25°C → 77°F. 1 HP → 745.7 W.', VERDE)
doc.add_paragraph('')

# ──────────────────────────────
# PASO A-6
# ──────────────────────────────
h2('PASO A-6 — Comparativa de impacto térmico (ventilado vs confinado) [PRIORIDAD MEDIA]')
badge('⏱', 'Tiempo estimado:', '1 hora', AZUL)
badge('📁', 'Archivo a modificar:', 'client/src/lib/iamSoilingEngine.ts + BIPVGlassSimulator.tsx', NARANJA)
doc.add_paragraph('')
body('OBJETIVO: Mostrar al usuario el impacto de la temperatura de celda en fachadas confinadas (sin ventilación trasera) vs ventiladas. Esto justifica diseños con cámara de aire.')
doc.add_paragraph('')
body('PASOS:')
bullet('1. Agregar compararImpactoTermico() a iamSoilingEngine.ts (código en documento de diagnóstico)')
bullet('2. En BIPVGlassSimulator.tsx, cuando el usuario seleccione tipo de montaje "fachada_confinada" o kBipv=1.3, mostrar una tarjeta de alerta naranja')
bullet('3. La tarjeta muestra: T celda ventilado (k=1.0) / T celda confinado (k=1.3) / Diferencia °C / Pérdida de producción %')
bullet('4. Agregar un botón "¿Cómo mejorar?" que muestre un tooltip: "Diseñar cámara de ventilación trasera de al menos 10 cm reduce la temperatura 10-15°C y recupera el 5-8% de producción perdida"')
doc.add_paragraph('')
badge('✅', 'Verificación:', 'Irradiancia=800 W/m², T_amb=30°C, NOCT=43°C: T_celda ventilado ≈ 52°C, T_celda confinado ≈ 61.8°C, diferencia ≈ 9.8°C.', VERDE)
doc.add_paragraph('')

# ──────────────────────────────
# PASO A-7
# ──────────────────────────────
h2('PASO A-7 — Banco de baterías (sistemas híbridos) [PRIORIDAD BAJA]')
badge('⏱', 'Tiempo estimado:', '1-2 horas', AZUL)
badge('📁', 'Archivo a modificar:', 'client/src/lib/stringSizing.ts (agregar al final)', NARANJA)
doc.add_paragraph('')
body('OBJETIVO: Permitir al usuario dimensionar el banco de baterías para sistemas BIPV con almacenamiento (energía + autonomía requerida).')
doc.add_paragraph('')
body('PASOS:')
bullet('1. Agregar dimensionarBaterias() al final de stringSizing.ts (código en documento de diagnóstico)')
bullet('2. Agregar nueva pestaña "Almacenamiento" en EnergyProductionSimulator.tsx (solo visible si el usuario activa "sistema con baterías")')
bullet('3. Campos de entrada: consumo diario kWh, días de autonomía (1-3), voltaje del banco (12/24/48V), tecnología (LFP/NMC/plomo/flujo)')
bullet('4. Resultados: Capacidad útil kWh, Capacidad bruta kWh, Capacidad Ah, Número de baterías unitarias, DOD aplicado')
bullet('5. Agregar al reporte PDF la sección de baterías si está configurada')
doc.add_paragraph('')
badge('✅', 'Verificación:', '20 kWh/día, 1 día autonomía, LFP (DOD=90%): capacidad bruta ≈ 23.4 kWh, a 48V → ≈ 487 Ah → 5 baterías de 100 Ah.', VERDE)
doc.add_paragraph('')

# ═══════════════════════════════════════════════════════════════
# PARTE B — CALCULADORA PYTHON CON STREAMLIT
# ═══════════════════════════════════════════════════════════════
sep()
h1('PARTE B — CALCULADORA BIPV COMPLETA EN PYTHON CON STREAMLIT (10 PASOS)')
body(
    'Esta parte construye desde cero la calculadora BIPV en Python puro usando Streamlit. '
    'Es ideal para aprender Python fotovoltaico y para uso en ingeniería de proyectos. '
    'Se puede usar de forma independiente o como complemento a la calculadora web TypeScript.'
)
doc.add_paragraph('')
badge('📦', 'Librerías necesarias:', 'streamlit, pvlib, numpy, pandas, matplotlib, plotly, python-docx, openpyxl, scipy, numpy-financial', AZUL)
doc.add_paragraph('')

codigo(
    '# Instalar todas las librerías necesarias:\n'
    'pip install streamlit pvlib numpy pandas matplotlib plotly\n'
    'pip install python-docx openpyxl scipy numpy-financial requests'
)

# ──────────────────────────────
# PASO B-1
# ──────────────────────────────
h2('PASO B-1 — Estructura del proyecto y configuración de Streamlit')
badge('⏱', 'Tiempo estimado:', '30 minutos', AZUL)
doc.add_paragraph('')
body('OBJETIVO: Crear la estructura de carpetas del proyecto Python y la navegación principal con páginas de Streamlit.')
doc.add_paragraph('')
body('Estructura de carpetas a crear:')
codigo(
    'calculadora_bipv_python/\n'
    '├── app.py                    # Punto de entrada principal\n'
    '├── pages/\n'
    '│   ├── 1_🌞_Recurso_Solar.py\n'
    '│   ├── 2_🔧_Modulo_BIPV.py\n'
    '│   ├── 3_⚡_Produccion.py\n'
    '│   ├── 4_🌑_Sombras.py\n'
    '│   ├── 5_🔌_Electrico.py\n'
    '│   ├── 6_📊_Balance.py\n'
    '│   ├── 7_💰_Financiero.py\n'
    '│   └── 8_📄_Reporte.py\n'
    '├── calculos/\n'
    '│   ├── solar.py              # pvlib: posición solar, transposición\n'
    '│   ├── termica.py            # Temperatura celda, modelo BIPV\n'
    '│   ├── electrica.py          # String sizing, baterías\n'
    '│   ├── energia.py            # Producción IEC 61724, PR\n'
    '│   ├── sombras.py            # Análisis de sombras\n'
    '│   ├── financiero.py         # VAN, TIR, LCOE, payback\n'
    '│   └── conversiones.py       # Unidades, HSP, CO2\n'
    '├── datos/\n'
    '│   ├── tecnologias.py        # Catálogo de módulos BIPV\n'
    '│   └── regiones.py           # Ciudades, irradiancia, CO2 por país\n'
    '└── reportes/\n'
    '    ├── word_report.py        # Genera Word con python-docx\n'
    '    └── excel_report.py       # Genera Excel con openpyxl'
)

body('Archivo app.py principal:')
codigo(
    'import streamlit as st\n\n'
    'st.set_page_config(\n'
    '    page_title="Calculadora BIPV Python",\n'
    '    page_icon="🌞",\n'
    '    layout="wide",\n'
    '    initial_sidebar_state="expanded"\n'
    ')\n\n'
    'st.title("🌞 Calculadora BIPV — Python Edition")\n'
    'st.markdown("""\n'
    'Calculadora fotovoltaica para sistemas BIPV (Building-Integrated Photovoltaics).\n'
    'Navega por las páginas en el menú lateral para ingresar los datos paso a paso.\n'
    '""")\n\n'
    '# Estado global compartido entre páginas\n'
    'if "proyecto" not in st.session_state:\n'
    '    st.session_state["proyecto"] = {\n'
    '        "nombre": "", "latitud": -33.45, "longitud": -70.65,\n'
    '        "pais": "Chile", "ciudad": "Santiago"\n'
    '    }'
)
doc.add_paragraph('')

# ──────────────────────────────
# PASO B-2
# ──────────────────────────────
h2('PASO B-2 — Página de Recurso Solar (pvlib integrado)')
badge('⏱', 'Tiempo estimado:', '2 horas', AZUL)
doc.add_paragraph('')
body('OBJETIVO: Obtener irradiancia real del lugar usando pvlib con datos TMY de PVGIS o EPW. El usuario ingresa latitud/longitud y obtiene el perfil solar mensual y el diagrama de trayectoria solar.')
doc.add_paragraph('')
body('Archivo: pages/1_🌞_Recurso_Solar.py')
codigo(
    'import streamlit as st\n'
    'import pvlib\n'
    'import pandas as pd\n'
    'import plotly.express as px\n\n'
    'st.header("🌞 Paso 1 — Recurso Solar del Lugar")\n\n'
    'col1, col2, col3 = st.columns(3)\n'
    'with col1:\n'
    '    lat = st.number_input("Latitud (°)", value=-33.45, step=0.01)\n'
    'with col2:\n'
    '    lon = st.number_input("Longitud (°)", value=-70.65, step=0.01)\n'
    'with col3:\n'
    '    alt = st.number_input("Altitud (m s.n.m.)", value=567, step=10)\n\n'
    'if st.button("📡 Obtener datos PVGIS"):\n'
    '    with st.spinner("Consultando base de datos PVGIS (EU JRC)..."):\n'
    '        # pvlib obtiene datos TMY directamente de la API de PVGIS\n'
    '        tmy_data, months_selected, inputs, metadata = pvlib.iotools.get_pvgis_tmy(\n'
    '            latitude=lat, longitude=lon, outputformat="json",\n'
    '            usehorizon=True, startyear=2005, endyear=2020\n'
    '        )\n'
    '        st.session_state["tmy"] = tmy_data\n'
    '        st.success(f"Datos obtenidos: {len(tmy_data)} horas/año")\n\n'
    '        # HSP mensual promedio\n'
    '        tmy_data["mes"] = tmy_data.index.month\n'
    '        hsp_mensual = tmy_data.groupby("mes")["ghi"].sum() / 1000  # kWh/m²/mes\n'
    '        hsp_diaria  = hsp_mensual / tmy_data.groupby("mes").size() * 24 / 1000\n\n'
    '        fig = px.bar(x=["Ene","Feb","Mar","Abr","May","Jun",\n'
    '                        "Jul","Ago","Sep","Oct","Nov","Dic"],\n'
    '                     y=hsp_mensual.values,\n'
    '                     labels={"x":"Mes","y":"GHI (kWh/m²/mes)"},\n'
    '                     title="Irradiación Global Horizontal por mes",\n'
    '                     color_discrete_sequence=["#F4A020"])\n'
    '        st.plotly_chart(fig, use_container_width=True)'
)
doc.add_paragraph('')

# ──────────────────────────────
# PASO B-3
# ──────────────────────────────
h2('PASO B-3 — Página de Módulo BIPV (catálogo + configuración)')
badge('⏱', 'Tiempo estimado:', '1.5 horas', AZUL)
doc.add_paragraph('')
body('OBJETIVO: El usuario selecciona la tecnología de vidrio BIPV (CdTe, mono-Si, CIGS, HJT), el nivel de transparencia, el área y el tipo de montaje. Se calculan la potencia pico y la eficiencia real.')
doc.add_paragraph('')
body('Archivo: calculos/termica.py + pages/2_🔧_Modulo_BIPV.py')
codigo(
    '# calculos/termica.py\n'
    'def temperatura_celda(t_amb, irradiancia, noct, k_bipv=1.0):\n'
    '    """Temperatura de celda con penalidad BIPV.\n'
    '    k_bipv = 1.0 → ventilado (convencional)\n'
    '    k_bipv = 1.3 → confinado (fachada sin ventilación trasera)\n'
    '    """\n'
    '    return t_amb + k_bipv * ((noct - 20) / 800) * irradiancia\n\n'
    'def eficiencia_real(efic_stc, coef_temp, t_celda, t_ref=25.0):\n'
    '    """Eficiencia corregida por temperatura."""\n'
    '    return efic_stc * (1 + (coef_temp / 100) * (t_celda - t_ref))\n\n'
    'def potencia_bipv(irradiancia, area_m2, efic_real, factor_cobertura=0.85,\n'
    '                  transparencia=0.20, factor_iam=1.0, factor_soiling=1.0):\n'
    '    """Potencia DC generada [W]."""\n'
    '    area_activa  = area_m2 * factor_cobertura\n'
    '    frac_opaca   = 1 - transparencia\n'
    '    return irradiancia * area_activa * efic_real * frac_opaca * factor_iam * factor_soiling'
)
doc.add_paragraph('')

# ──────────────────────────────
# PASO B-4
# ──────────────────────────────
h2('PASO B-4 — Página de Producción Energética (IEC 61724 en Python)')
badge('⏱', 'Tiempo estimado:', '2 horas', AZUL)
doc.add_paragraph('')
body('OBJETIVO: Implementar el cálculo de producción anual con la cadena completa de pérdidas (IEC 61724-1:2021) usando pvlib para la parte solar.')
doc.add_paragraph('')
body('Archivo: calculos/energia.py')
codigo(
    '# calculos/energia.py\n'
    'import pvlib\n'
    'import numpy as np\n\n'
    'def calcular_produccion_anual(\n'
    '    tmy_data,          # DataFrame TMY de pvlib\n'
    '    potencia_kWp,      # kWp del sistema\n'
    '    inclinacion,       # grados\n'
    '    azimut,            # grados (0=Norte, 90=Este, 180=Sur, 270=Oeste)\n'
    '    lat, lon,\n'
    '    perdidas_config    # dict con factores de pérdida\n'
    '):\n'
    '    loc = pvlib.location.Location(lat, lon, altitude=0)\n\n'
    '    # Posición solar hora a hora\n'
    '    solar_pos = loc.get_solarposition(tmy_data.index)\n\n'
    '    # Irradiancia en el plano inclinado (modelo de Perez)\n'
    '    poa = pvlib.irradiance.get_total_irradiance(\n'
    '        surface_tilt=inclinacion, surface_azimuth=azimut,\n'
    '        solar_zenith=solar_pos["apparent_zenith"],\n'
    '        solar_azimuth=solar_pos["azimuth"],\n'
    '        dni=tmy_data["dni"], ghi=tmy_data["ghi"], dhi=tmy_data["dhi"],\n'
    '        model="perez"\n'
    '    )["poa_global"]\n\n'
    '    # Rendimiento de referencia Yr (kWh/kWp)\n'
    '    Yr = poa.sum() / 1000\n\n'
    '    # Temperatura de celda (modelo pvlib Faiman)\n'
    '    params = pvlib.temperature.TEMPERATURE_MODEL_PARAMETERS["sapm"]["open_rack_glass_polymer"]\n'
    '    t_cell = pvlib.temperature.sapm_cell(\n'
    '        poa, tmy_data["temp_air"], tmy_data["wind_speed"], **params\n'
    '    )\n\n'
    '    # Cadena de pérdidas\n'
    '    L_suciedad   = perdidas_config.get("soiling", 0.02)   # 2%\n'
    '    L_sombras    = perdidas_config.get("shading", 0.03)   # 3%\n'
    '    L_temperatura= perdidas_config.get("temperatura", 0.04)\n'
    '    L_inversor   = perdidas_config.get("inversor", 0.035)\n'
    '    L_cableado   = perdidas_config.get("cableado", 0.015)\n'
    '    L_otros      = perdidas_config.get("otros", 0.02)\n\n'
    '    # Factor de rendimiento PR\n'
    '    PR = (1 - L_suciedad) * (1 - L_sombras) * (1 - L_temperatura) * \\\n'
    '         (1 - L_inversor) * (1 - L_cableado) * (1 - L_otros)\n\n'
    '    # Rendimiento final Yf\n'
    '    Yf = Yr * PR\n\n'
    '    # Energía anual\n'
    '    energia_anual_kwh = Yf * potencia_kWp\n\n'
    '    return {\n'
    '        "Yr_kWh_kWp": round(Yr, 1),\n'
    '        "Yf_kWh_kWp": round(Yf, 1),\n'
    '        "PR": round(PR, 3),\n'
    '        "energia_anual_kwh": round(energia_anual_kwh, 1),\n'
    '        "CF": round(Yf / 8760, 3),\n'
    '    }'
)
doc.add_paragraph('')

# ──────────────────────────────
# PASO B-5
# ──────────────────────────────
h2('PASO B-5 — Página de Análisis de Sombras (diagrama solar 2D)')
badge('⏱', 'Tiempo estimado:', '2 horas', AZUL)
doc.add_paragraph('')
body('OBJETIVO: Generar el diagrama de trayectoria solar del lugar y permitir que el usuario defina ángulos de obstáculos (edificios vecinos, cornisas) para calcular la pérdida por sombras.')
doc.add_paragraph('')
body('Archivo: calculos/sombras.py + pages/4_🌑_Sombras.py')
codigo(
    '# calculos/sombras.py\n'
    'import pvlib\n'
    'import pandas as pd\n'
    'import plotly.graph_objects as go\n\n'
    'def generar_diagrama_solar(lat, lon, obstaculos=None):\n'
    '    """\n'
    '    Genera diagrama solar (elevación vs azimut) con trayectoria de cada mes.\n'
    '    obstaculos: lista de dict {azimut_ini, azimut_fin, elevacion_max}\n'
    '    """\n'
    '    import numpy as np\n'
    '    times = pd.date_range("2024-01-01", "2024-12-31", freq="1h", tz="America/Santiago")\n'
    '    loc   = pvlib.location.Location(lat, lon)\n'
    '    pos   = loc.get_solarposition(times)\n'
    '    pos   = pos[pos["apparent_elevation"] > 0]\n\n'
    '    fig = go.Figure()\n'
    '    colores = ["#E63946","#F4A261","#2A9D8F","#264653","#A8DADC",\n'
    '               "#457B9D","#1D3557","#F1FAEE","#E9C46A","#F77F00","#D62828","#023E8A"]\n\n'
    '    for mes in range(1, 13):\n'
    '        data_mes = pos[pos.index.month == mes]\n'
    '        fig.add_trace(go.Scatter(\n'
    '            x=data_mes["azimuth"], y=data_mes["apparent_elevation"],\n'
    '            mode="lines", name=f"Mes {mes}",\n'
    '            line=dict(color=colores[mes-1], width=1.5)\n'
    '        ))\n\n'
    '    # Dibujar obstáculos como rectángulos de sombra\n'
    '    if obstaculos:\n'
    '        for obs in obstaculos:\n'
    '            fig.add_shape(type="rect",\n'
    '                x0=obs["azimut_ini"], x1=obs["azimut_fin"],\n'
    '                y0=0, y1=obs["elevacion_max"],\n'
    '                fillcolor="rgba(200,0,0,0.3)", line_color="red")\n\n'
    '    fig.update_layout(\n'
    '        title="Diagrama de Trayectoria Solar — Obstáculos de Sombra",\n'
    '        xaxis_title="Azimut (°)", yaxis_title="Elevación solar (°)"\n'
    '    )\n'
    '    return fig\n\n'
    'def calcular_factor_sombra_anual(lat, lon, obstaculos):\n'
    '    """Fracción de horas de sol bloqueadas por obstáculos (0=sin sombra, 1=sombra total)."""\n'
    '    import numpy as np\n'
    '    times = pd.date_range("2024-01-01", "2024-12-31", freq="1h", tz="America/Santiago")\n'
    '    loc   = pvlib.location.Location(lat, lon)\n'
    '    pos   = loc.get_solarposition(times)\n'
    '    sol   = pos[pos["apparent_elevation"] > 0]\n'
    '    total = len(sol)\n'
    '    sombradas = 0\n'
    '    for _, row in sol.iterrows():\n'
    '        for obs in obstaculos:\n'
    '            if (obs["azimut_ini"] <= row["azimuth"] <= obs["azimut_fin"]\n'
    '                    and row["apparent_elevation"] <= obs["elevacion_max"]):\n'
    '                sombradas += 1\n'
    '                break\n'
    '    return round(sombradas / total, 3) if total > 0 else 0.0'
)
doc.add_paragraph('')

# ──────────────────────────────
# PASO B-6
# ──────────────────────────────
h2('PASO B-6 — Página Eléctrica (string sizing + baterías)')
badge('⏱', 'Tiempo estimado:', '1.5 horas', AZUL)
doc.add_paragraph('')
body('OBJETIVO: Implementar el dimensionado eléctrico de strings (validación Voc/MPPT/ratio DC-AC) y el sizing del banco de baterías. Reutiliza las funciones de la Parte A.')
doc.add_paragraph('')
body('Archivo: calculos/electrica.py')
codigo(
    '# calculos/electrica.py\n'
    'import math\n\n'
    'def string_sizing(Voc_stc, Vmp_stc, Isc_stc, coef_v,\n'
    '                  Vdc_max, Vmppt_min, Vmppt_max, Idc_max,\n'
    '                  T_min=-5, T_max=70):\n'
    '    Voc_frio     = Voc_stc * (1 + (coef_v/100) * (T_min - 25))\n'
    '    Vmp_caliente = Vmp_stc * (1 + (coef_v/100) * (T_max - 25))\n'
    '    n_serie_max  = int(Vdc_max   / Voc_frio)\n'
    '    n_serie_min  = math.ceil(Vmppt_min / Vmp_caliente)\n'
    '    n_serie_opt  = int(Vmppt_max  / Vmp_stc)\n'
    '    n_par_max    = int(Idc_max    / Isc_stc)\n'
    '    return {\n'
    '        "n_serie_min": n_serie_min, "n_serie_max": n_serie_max,\n'
    '        "n_serie_opt": n_serie_opt, "n_paralelo_max": n_par_max,\n'
    '        "Voc_frio": round(Voc_frio, 1),\n'
    '        "Vmp_caliente": round(Vmp_caliente, 1),\n'
    '        "valido": n_serie_min <= n_serie_opt <= n_serie_max\n'
    '    }\n\n'
    'DOD = {"litio-LFP":0.90, "litio-NMC":0.80, "plomo-acido":0.50, "flujo":0.80}\n\n'
    'def banco_baterias(consumo_kwh_dia, dias_autonomia, voltaje_v,\n'
    '                   tecnologia="litio-LFP", cap_bateria_ah=100):\n'
    '    dod      = DOD[tecnologia]\n'
    '    bruta    = consumo_kwh_dia * dias_autonomia / (dod * 0.95)\n'
    '    cap_ah   = bruta * 1000 / voltaje_v\n'
    '    n_bat    = math.ceil(cap_ah / cap_bateria_ah)\n'
    '    return {"capacidad_kwh": round(consumo_kwh_dia*dias_autonomia, 1),\n'
    '            "capacidad_bruta_kwh": round(bruta, 1),\n'
    '            "capacidad_ah": round(cap_ah), "n_baterias": n_bat,\n'
    '            "dod": dod, "tecnologia": tecnologia}'
)
doc.add_paragraph('')

# ──────────────────────────────
# PASO B-7
# ──────────────────────────────
h2('PASO B-7 — Página Balance Energético (consumo vs generación)')
badge('⏱', 'Tiempo estimado:', '1.5 horas', AZUL)
doc.add_paragraph('')
body('OBJETIVO: Comparar mes a mes la energía generada por el BIPV contra el consumo del edificio. Calcular autoconsumo, excedente, compra de red, y clasificación energética A+/A/B/C/D.')
doc.add_paragraph('')
body('Archivo: pages/6_📊_Balance.py')
codigo(
    'import streamlit as st\n'
    'import pandas as pd\n'
    'import plotly.graph_objects as go\n\n'
    'st.header("📊 Paso 6 — Balance Energético Mensual")\n\n'
    'generacion = st.session_state.get("produccion_mensual", [0]*12)\n\n'
    'st.subheader("Consumo del edificio por mes (kWh)")\n'
    'consumo = []\n'
    'cols = st.columns(6)\n'
    'meses = ["Ene","Feb","Mar","Abr","May","Jun","Jul","Ago","Sep","Oct","Nov","Dic"]\n'
    'for i in range(12):\n'
    '    with cols[i % 6]:\n'
    '        consumo.append(st.number_input(meses[i], min_value=0.0, value=1500.0, key=f"c{i}"))\n\n'
    'tarifa_compra = st.number_input("Tarifa compra $/kWh", value=120.0)\n'
    'tarifa_venta  = st.number_input("Tarifa inyección $/kWh", value=60.0)\n\n'
    'if st.button("Calcular balance"):\n'
    '    filas = []\n'
    '    for i in range(12):\n'
    '        g, c = generacion[i], consumo[i]\n'
    '        auto = min(g, c)\n'
    '        exc  = max(0, g - c)\n'
    '        comp = max(0, c - g)\n'
    '        ahorro = auto*tarifa_compra + exc*tarifa_venta\n'
    '        filas.append({"Mes":meses[i],"Gen(kWh)":round(g),"Cons(kWh)":round(c),\n'
    '                      "Autoconsumo":round(auto),"Excedente":round(exc),\n'
    '                      "Compra red":round(comp),"Ahorro $":round(ahorro),\n'
    '                      "%Auto":round(auto/g*100 if g>0 else 0,1)})\n'
    '    df = pd.DataFrame(filas)\n'
    '    st.dataframe(df, use_container_width=True)\n\n'
    '    auto_pct = sum(min(g,c) for g,c in zip(generacion,consumo)) / sum(consumo) * 100\n'
    '    clase = "A+" if auto_pct>=80 else "A" if auto_pct>=60 else "B" if auto_pct>=40 else "C" if auto_pct>=20 else "D"\n'
    '    st.metric("Autosuficiencia anual", f"{auto_pct:.1f}%", delta=f"Clase {clase}")'
)
doc.add_paragraph('')

# ──────────────────────────────
# PASO B-8
# ──────────────────────────────
h2('PASO B-8 — Página Financiero (VAN, TIR, LCOE, CO₂)')
badge('⏱', 'Tiempo estimado:', '2 horas', AZUL)
doc.add_paragraph('')
body('OBJETIVO: Implementar el análisis financiero completo con numpy-financial para VAN (NPV) y TIR (IRR). Incluir tabla de flujo de caja año a año y tabla de CO₂ evitado por país.')
doc.add_paragraph('')
body('Archivo: calculos/financiero.py')
codigo(
    '# calculos/financiero.py\n'
    'import numpy as np\n'
    'import numpy_financial as npf\n\n'
    'FACTOR_CO2 = {"Chile":0.294,"Colombia":0.126,"México":0.398,\n'
    '              "Argentina":0.341,"Perú":0.214,"España":0.187,"Global":0.436}\n\n'
    'def analisis_financiero(\n'
    '    costo_total,           # $ inversión inicial\n'
    '    energia_ano1_kwh,      # kWh generados año 1\n'
    '    tarifa_kwh,            # $/kWh tarifa de compra\n'
    '    degradacion=0.005,     # 0.5%/año\n'
    '    crecimiento_tarifa=0.035,\n'
    '    tasa_descuento=0.08,\n'
    '    mantenimiento_anual=0.01,  # % del costo total por año\n'
    '    horizonte=25,\n'
    '    pais="Chile"\n'
    '):\n'
    '    flujos = [-costo_total]\n'
    '    for n in range(1, horizonte+1):\n'
    '        prod   = energia_ano1_kwh * (1-degradacion)**(n-1)\n'
    '        ingreso= prod * tarifa_kwh * (1+crecimiento_tarifa)**(n-1)\n'
    '        mant   = costo_total * mantenimiento_anual\n'
    '        flujos.append(ingreso - mant)\n\n'
    '    van    = npf.npv(tasa_descuento, flujos)\n'
    '    tir    = npf.irr(flujos) * 100\n\n'
    '    # LCOE ($/kWh)\n'
    '    kwh_totales = sum(energia_ano1_kwh*(1-degradacion)**n for n in range(horizonte))\n'
    '    gastos_pv   = sum((costo_total*mantenimiento_anual)/((1+tasa_descuento)**n)\n'
    '                       for n in range(1,horizonte+1))\n'
    '    lcoe = (costo_total + gastos_pv) / kwh_totales if kwh_totales>0 else 0\n\n'
    '    # Payback simple\n'
    '    acum, payback = -costo_total, None\n'
    '    for n, f in enumerate(flujos[1:], 1):\n'
    '        acum += f\n'
    '        if acum >= 0 and payback is None:\n'
    '            payback = n\n\n'
    '    # CO2\n'
    '    factor = FACTOR_CO2.get(pais, FACTOR_CO2["Global"])\n'
    '    co2_total_ton = kwh_totales * factor / 1000\n\n'
    '    return {"VAN_USD": round(van, 0), "TIR_pct": round(tir, 2),\n'
    '            "LCOE_USD_kWh": round(lcoe, 4), "payback_anos": payback,\n'
    '            "co2_evitado_ton": round(co2_total_ton, 1),\n'
    '            "flujos": flujos}'
)
doc.add_paragraph('')

# ──────────────────────────────
# PASO B-9
# ──────────────────────────────
h2('PASO B-9 — Página de Reporte (Word + Excel descargables)')
badge('⏱', 'Tiempo estimado:', '2 horas', AZUL)
doc.add_paragraph('')
body('OBJETIVO: Generar automáticamente un informe Word con todos los resultados del proyecto y un Excel con las tablas de datos, disponibles para descarga directa desde Streamlit.')
doc.add_paragraph('')
body('Archivo: reportes/word_report.py + pages/8_📄_Reporte.py')
codigo(
    '# pages/8_📄_Reporte.py\n'
    'import streamlit as st\n'
    'from reportes.word_report import generar_word\n'
    'from reportes.excel_report import generar_excel\n'
    'import io\n\n'
    'st.header("📄 Paso 8 — Generar Informe del Proyecto")\n\n'
    'col1, col2 = st.columns(2)\n\n'
    'with col1:\n'
    '    if st.button("📝 Generar Informe Word"):\n'
    '        datos = st.session_state  # todos los datos del proyecto\n'
    '        buffer = io.BytesIO()\n'
    '        generar_word(datos, buffer)\n'
    '        buffer.seek(0)\n'
    '        st.download_button(\n'
    '            label="⬇️ Descargar Word (.docx)",\n'
    '            data=buffer,\n'
    '            file_name=f"Informe_BIPV_{datos.get(\'proyecto\',{}).get(\'nombre\',\'proyecto\')}.docx",\n'
    '            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"\n'
    '        )\n\n'
    'with col2:\n'
    '    if st.button("📊 Generar Excel con datos"):\n'
    '        datos = st.session_state\n'
    '        buffer = io.BytesIO()\n'
    '        generar_excel(datos, buffer)\n'
    '        buffer.seek(0)\n'
    '        st.download_button(\n'
    '            label="⬇️ Descargar Excel (.xlsx)",\n'
    '            data=buffer,\n'
    '            file_name="Datos_BIPV.xlsx",\n'
    '            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"\n'
    '        )'
)
doc.add_paragraph('')

# ──────────────────────────────
# PASO B-10
# ──────────────────────────────
h2('PASO B-10 — Despliegue (Streamlit Community Cloud — gratis)')
badge('⏱', 'Tiempo estimado:', '30 minutos', AZUL)
doc.add_paragraph('')
body('OBJETIVO: Publicar la calculadora Python en internet de forma gratuita usando Streamlit Community Cloud. El resultado es una URL pública que cualquier persona puede usar.')
doc.add_paragraph('')
body('PASOS:')
bullet('1. Subir el proyecto a GitHub (nuevo repositorio, ej: calculadora-bipv-python)')
bullet('2. Crear el archivo requirements.txt con todas las librerías:')
codigo(
    '# requirements.txt\n'
    'streamlit>=1.35\n'
    'pvlib>=0.10\n'
    'numpy>=1.26\n'
    'pandas>=2.0\n'
    'plotly>=5.20\n'
    'matplotlib>=3.8\n'
    'python-docx>=1.1\n'
    'openpyxl>=3.1\n'
    'scipy>=1.12\n'
    'numpy-financial>=1.0\n'
    'requests>=2.31'
)
bullet('3. Ir a share.streamlit.io → "New app" → seleccionar el repositorio → branch main → archivo app.py')
bullet('4. Streamlit Cloud instala automáticamente requirements.txt y despliega la app')
bullet('5. La URL resultante es: https://[tu-usuario]-calculadora-bipv-python-app-[hash].streamlit.app')
bullet('6. Cada vez que se hace git push al repositorio, la app se actualiza automáticamente')
doc.add_paragraph('')
badge('✅', 'Verificación final:', 'La calculadora debe funcionar en el navegador sin instalar nada. Probar el flujo completo: ingresar latitud Santiago (-33.45, -70.65) → obtener PVGIS → definir fachada 50m² CdTe → calcular → generar reporte Word.', VERDE)
doc.add_paragraph('')

# ═══════════════════════════════════════════════════════════════
# CRONOGRAMA
# ═══════════════════════════════════════════════════════════════
sep()
h1('CRONOGRAMA SUGERIDO — 4 SEMANAS')
doc.add_paragraph('')

tabla_cron = doc.add_table(rows=1, cols=4)
tabla_cron.style = 'Table Grid'
hdrs = ['Semana', 'Parte A (TypeScript)', 'Parte B (Python)', 'Objetivo de la semana']
for i, h_txt in enumerate(hdrs):
    tabla_cron.rows[0].cells[i].text = h_txt
    cell_shade(tabla_cron.rows[0].cells[i], '1A5C8A')
    for par in tabla_cron.rows[0].cells[i].paragraphs:
        for run in par.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            run.font.size = Pt(9)

semanas = [
    ('Semana 1', 'A-1: String Sizing\nA-2: Factor cobertura', 'B-1: Estructura\nB-2: Recurso solar', 'Cimientos — dimensionado eléctrico y recurso solar en Python'),
    ('Semana 2', 'A-3: Balance mensual\nA-5: Convertidor de unidades', 'B-3: Módulo BIPV\nB-4: Producción IEC 61724', 'Cálculos centrales — producción y balance energético'),
    ('Semana 3', 'A-4: Proyección 25 años\nA-6: Impacto térmico', 'B-5: Sombras\nB-6: Eléctrico\nB-7: Balance', 'Análisis avanzado — sombras, degradación y balance mensual'),
    ('Semana 4', 'A-7: Banco baterías\nRevisión y pruebas', 'B-8: Financiero\nB-9: Reportes\nB-10: Deploy', 'Producción — análisis financiero, reportes y despliegue'),
]

for fila in semanas:
    row = tabla_cron.add_row().cells
    for i, val in enumerate(fila):
        row[i].text = val
        for par in row[i].paragraphs:
            for run in par.runs:
                run.font.size = Pt(9)

doc.add_paragraph('')

# Pie
footer = doc.add_paragraph('Plan elaborado sobre análisis de código real — github.com/ventas108/calculadora-bipv  |  Python Edition — Streamlit + pvlib  |  2026')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.runs[0].font.size = Pt(9)
footer.runs[0].font.color.rgb = GRIS
footer.runs[0].italic = True

doc.save('Plan_Maestro_Calculadora_BIPV.docx')
print("Documento generado correctamente.")
