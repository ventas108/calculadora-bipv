# -*- coding: utf-8 -*-
"""
Genera la Ficha Técnica comercial de la Calculadora BIPV (versión extendida,
con datos extraídos del Manual de Operación v3 — agosto 2026).
Salida: entregables/FICHA_TECNICA_CALCULADORA_BIPV_agosto2026.docx
Ejecutar:  python3 scripts/generar_ficha_tecnica.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "entregables/FICHA_TECNICA_CALCULADORA_BIPV_agosto2026.docx"

C_VERDE_OSC = RGBColor(0x1B, 0x5E, 0x20)
C_VERDE_CLR = RGBColor(0xE8, 0xF5, 0xE9)
C_AZUL      = RGBColor(0x0D, 0x47, 0xA1)
C_AZUL_CLR  = RGBColor(0xE3, 0xF2, 0xFD)
C_NARANJA   = RGBColor(0xE6, 0x51, 0x00)
C_NARANJA_CLR = RGBColor(0xFF, 0xF3, 0xE0)
C_GRIS      = RGBColor(0x37, 0x47, 0x4F)
C_GRIS_CLR  = RGBColor(0xEC, 0xEF, 0xF1)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Cm(1.5)
    s.left_margin = s.right_margin = Cm(1.7)
st = doc.styles["Normal"]
st.font.name = "Calibri"; st.font.size = Pt(10)

def hexs(rgb): return f"{int(rgb[0]):02X}{int(rgb[1]):02X}{int(rgb[2]):02X}"
def cell_bg(cell, rgb):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexs(rgb))
    cell._tc.get_or_add_tcPr().append(shd)

def para(text="", size=10, bold=False, color=None, align=None, space_after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align: p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold
    if color: r.font.color.rgb = color
    return p

def h(text, color=C_AZUL, size=13): para(text, size=size, bold=True, color=color, space_after=4)
def sub(text): para(text, bold=True, color=C_VERDE_OSC, space_after=2)

def box(text, bg, fg, bold_first=None):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]; cell_bg(c, bg)
    p = c.paragraphs[0]
    if bold_first:
        r = p.add_run(bold_first + " "); r.bold = True; r.font.color.rgb = fg; r.font.size = Pt(10)
    r = p.add_run(text); r.font.color.rgb = fg; r.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def tabla(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, htxt in enumerate(headers):
        c = t.rows[0].cells[j]; cell_bg(c, C_VERDE_OSC)
        r = c.paragraphs[0].add_run(htxt); r.bold = True
        r.font.color.rgb = WHITE; r.font.size = Pt(9.5)
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            if i % 2 == 1: cell_bg(cells[j], C_VERDE_CLR)
            r = cells[j].paragraphs[0].add_run(val); r.font.size = Pt(9)
            if j == 0: r.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

# ══════════ PORTADA ══════════
para("FICHA TÉCNICA", size=22, bold=True, color=C_VERDE_OSC,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para("Calculadora BIPV — Plataforma de simulación fotovoltaica integrada en edificios y agrivoltaica",
     size=12.5, bold=True, color=C_GRIS, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para("Innovación Química · calc.innovacionquimica.com.co · Versión agosto 2026",
     size=9, color=C_GRIS, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

box("Simula 8 760 horas del año con datos climáticos reales del sitio, resuelve la curva "
    "corriente-voltaje de cada panel, modela la activación de los diodos de bypass bajo "
    "sombra parcial y entrega TIR, VPN, Payback y LCOE en pesos colombianos con la Ley "
    "1715/2014 aplicada. Cada número del reporte es trazable y auditable — del rayo de sol "
    "al flujo de caja.", C_AZUL_CLR, C_AZUL, "¿Qué es?")

# ══════════ 1. COBERTURA ══════════
h("1. Seis tecnologías de integración, cada una con su física correcta")
para("La Calculadora no usa un modelo genérico: cada tipo de instalación carga automáticamente "
     "su densidad de potencia, su Performance Ratio típico y su inclinación físicamente correcta, "
     "con alertas reactivas si un valor sale del rango recomendado.")
tabla(["Tipo de instalación", "Densidad (W/m²)", "PR típico", "Tilt defecto", "Modelo térmico"],
[
 ("🏢 Fachada BIPV", "80–180", "0.65", "90°", "Confinado (k=1.3)"),
 ("🏠 Techo inclinado BIPV", "100–200", "0.75", "15°", "Ventilado (k=1.0)"),
 ("⛱️ Pérgola BIPV", "60–150", "0.70", "10°", "Semi-ventilado (k=1.1)"),
 ("Marquesina BIPV", "70–160", "0.68", "30°", "Semi-ventilado (k=1.1)"),
 ("🏚️ Techo plano", "120–220", "0.78", "10°", "Ventilado (k=1.0)"),
 ("🌱 Granja FV / agrivoltaica", "130–250", "0.80", "15° (≈5°N Colombia)", "Ventilado + bifacial"),
])
box("Modo agrivoltaico completo: factor de ocupación configurable (5–100%) que dimensiona "
    "paneles y presupuesto sobre el área útil, sincroniza la separación entre filas (GCR) y "
    "muestra en Vista 3D las filas elevadas a 3 m sobre el cultivo. Validado con proyectos "
    "reales de cultivo bajo paneles en Colombia.", C_VERDE_CLR, C_VERDE_OSC, "🌱")

# ══════════ 2. MOTORES ══════════
h("2. Motores de cálculo — la diferencia entre estimar y simular")

sub("☀️ Recurso solar hora a hora (8 760 iteraciones/año)")
para("Archivos climáticos TMY/EPW reales del sitio + transposición Hay-Davies al plano exacto "
     "de instalación. Para bifaciales y agrivoltaica, modelo infinite_sheds de pvlib (estándar "
     "NREL) que calcula la luz que llega a la cara trasera según la separación real entre filas. "
     "El reporte declara la ganancia bifacial anual (%) con altura de montaje y albedo — el "
     "banco ve de dónde sale cada kWh extra.")

sub("⚡ Motor IV — curva corriente-voltaje real, no regla de tres")
para("Resuelve el modelo de diodo único del panel en cada condición de irradiancia y "
     "temperatura, con Voc, Isc, Vmp, Imp y coeficientes térmicos de la ficha real del "
     "fabricante. Se activa automáticamente con ficha completa e incluye triple defensa "
     "contra el error más común del mercado: el conteo de celdas en paneles half-cut "
     "(verificación física Ns ≈ Voc/0.74).")

sub("🌗 Diodos de bypass — el cálculo que casi nadie hace")
para("Cuando un obstáculo sombrea 2 de 8 módulos de un string, la pérdida real NO es 25%: "
     "puede ser 40–60% de la producción en esas horas, porque los diodos de bypass eliminan "
     "la tensión completa de los módulos sombreados. La Calculadora resuelve el circuito IV "
     "del string hora a hora con el factor de sombra geométrico del sitio (CSV de la "
     "Calculadora de Sombreado 3D) y entrega pérdida anual, horas de bypass activo y tabla "
     "mensual. Referencia técnica: Deline et al. 2013 (NREL).")
tabla(["Pérdida bypass anual", "Diagnóstico", "Acción sugerida"],
[
 ("< 2%", "🟢 Bajo", "Sombras leves — diseño aprobado"),
 ("2–5%", "🟡 Moderado", "Considerar redistribución de strings"),
 ("5–10%", "🔴 Alto", "Evaluar cambio de orientación o layout"),
 ("> 10%", "⛔ Muy alto", "Rediseño del sombreado — detectado ANTES de construir"),
])

sub("🔬 Motor Óptico BIPV")
para("Cascada completa de pérdidas que solo existen en integración arquitectónica: reflexión "
     "por ángulo de incidencia (IAM), suciedad (soiling) y efecto térmico del confinamiento "
     "en fachada (el panel confinado opera ~30% más caliente que en rack ventilado, k=1.3).")

sub("🏢 Multi-superficie — el edificio completo en una simulación")
para("Combina fachada sur + techo plano + pérgola + marquesina en un solo sistema: POA y "
     "producción por cada orientación, bypass individual por superficie y una E_ac total "
     "trazable que alimenta el financiero, las baterías y el cálculo de CO₂ evitado.")

box("El Performance Ratio no se asume: se obtiene. La cadena óptica → térmica → eléctrica "
    "produce el PR como resultado físico del sitio. En Bogotá o Medellín puede superar el "
    "100% (altitud y baja temperatura) — y la herramienta lo demuestra, capa por capa.",
    C_VERDE_CLR, C_VERDE_OSC, "🎯 Argumento clave:")

# ══════════ 3. FINANCIERO ══════════
h("3. Análisis financiero bancable — Colombia, no genérico")
para("TIR, VPN, Payback y LCOE calculados sobre flujo de caja en COP con TRM en línea, "
     "precio real del inversor del catálogo y los tres beneficios de la Ley 1715/2014: "
     "deducción de renta (Art. 11), exclusión de IVA sobre equipos (Art. 12) y depreciación "
     "acelerada (Art. 14).")
tabla(["Indicador", "🟢 Atractivo", "🟡 Aceptable", "🔴 Revisar"],
[
 ("TIR", "> 12%", "8–12%", "< 8%"),
 ("VPN", "> 0 USD", "≈ 0", "< 0 USD"),
 ("Payback simple", "< 10 años", "10–15 años", "> 15 años"),
 ("LCOE", "< tarifa red", "≈ tarifa red", "> tarifa red"),
])
sub("Estimación Rápida — rango de inversión en menos de 2 minutos")
para("Benchmarks del mercado colombiano (julio 2026) por tipo de instalación, escenario "
     "económico (optimista/base/conservador) y factor de zona geográfica (Bogotá ×1.00 hasta "
     "Urabá/Chocó ×1.17). Desglosa CAPEX en equipos (55–65%), construcción (20–28%), costos "
     "blandos (8–14%) y contingencias, más OPEX anual de 5 componentes (O&M, limpieza, fondo "
     "de reposición de inversor, monitoreo y seguro). Precisión declarada honestamente: "
     "±25–35% — perfecta para decidir la factibilidad; las cotizaciones reales siempre tienen "
     "prioridad automática cuando se ingresan.")
tabla(["Tipo", "CAPEX referencia (USD/Wp)", "OPEX (USD/kWp·año)"],
[
 ("Granja FV en campo", "0.70–1.10", "8–12"),
 ("Techo industrial", "0.90–1.45", "6–12"),
 ("BIPV fachada / pérgola", "1.60–3.20", "10–16"),
])

# ══════════ 4. CATÁLOGOS ══════════
h("4. Catálogos inteligentes — cero transcripción manual")
para("Paneles, inversores y baterías se cargan arrastrando la ficha técnica PDF del "
     "fabricante. El extractor lee potencias, tensiones MPPT, corrientes por tracker, "
     "coeficientes térmicos, celdas y bifacialidad — incluso de fichas escaneadas (OCR), "
     "en español o inglés, con tablas multi-modelo (un PDF, varios equipos).")
tabla(["Capacidad", "Detalle"],
[
 ("Fabricantes de inversores reconocidos", "20+: Growatt, Deye, Solis, Huawei, SMA, Fronius, GoodWe, Sungrow, Victron, SolarEdge, MUST, SolaX, LuxPower, POWEST y más"),
 ("Control de calidad automático", "Validador físico que rechaza valores imposibles + detección de fallos silenciosos: ningún dato incompleto entra al catálogo sin aviso"),
 ("Verificación permanente", "Banco de regresión de 84 fichas reales de paneles + 26 casos de inversores que corren antes de cada mejora"),
 ("Emparejamiento panel-inversor", "Motor automático que verifica ventanas MPPT, ratios DC/AC y compatibilidad eléctrica del arreglo"),
])

# ══════════ 5. ENTREGABLES ══════════
h("5. Entregables con trazabilidad total")
para("El Reporte PDF ejecutivo incluye hasta 11 secciones seleccionables: proyecto, recurso "
     "solar, motor óptico, producción con diagnóstico PR real vs esperado, bypass diodes con "
     "semáforo de impacto, desglose multi-superficie, financiero, costos con KPIs de "
     "bancabilidad (USD/Wp, USD/m², OPEX/CAPEX), balance con baterías y CO₂ evitado.")
box("Cada reporte declara explícitamente qué energía se usó y por qué — por ejemplo: "
    "\"E_ac usada: 88 150 kWh/año (corregida por bypass diodes; pérdida descontada: "
    "2 850 kWh/año)\". El cliente, el banco o la UPME pueden verificar que las cifras son "
    "conservadoras, no optimistas.", C_NARANJA_CLR, C_NARANJA, "🔍 Trazabilidad:")

# ══════════ 6. ESPECIFICACIONES ══════════
h("6. Especificaciones generales")
tabla(["Característica", "Detalle"],
[
 ("Resolución temporal", "Horaria — 8 760 horas por año meteorológico típico (TMY/EPW)"),
 ("Tipos de instalación", "6: fachada, techo inclinado, techo plano, pérgola, marquesina, granja FV/agrivoltaica"),
 ("Multi-superficie", "Ilimitadas orientaciones combinadas con bypass individual por superficie"),
 ("Baterías", "Balance horario consumo vs generación: autogeneración, autosuficiencia, excedentes y dimensionamiento recomendado"),
 ("Moneda y normativa", "COP con TRM en línea · Ley 1715/2014 (Arts. 11, 12, 14) · RETIE/UPME en costos blandos"),
 ("Base científica", "pvlib (NREL, estándar de la industria) + modelos propios BIPV + Deline et al. 2013 para bypass"),
 ("Acceso", "Aplicación web, sin instalación — calc.innovacionquimica.com.co"),
 ("Ecosistema", "Integrada con la Calculadora de Sombreado 3D (bipv.innovacionquimica.com.co)"),
])

box("Para desarrolladores inmobiliarios, arquitectos, EPCs, banca verde y proyectos "
    "agrivoltaicos que necesitan cifras defendibles ante un comité de inversión — no "
    "promesas de folleto.", C_GRIS_CLR, C_GRIS, "👥 Diseñada para:")

para("Ficha técnica — agosto de 2026 · Innovación Química · calc.innovacionquimica.com.co",
     size=8.5, color=C_GRIS, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

doc.save(OUT)
print("OK —", OUT)
