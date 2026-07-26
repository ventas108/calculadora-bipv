
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Título
title = doc.add_heading('ADENDA: Reglas Principales de Python', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1A, 0x5C, 0x8A)

subtitle = doc.add_paragraph('Aplicadas a Cálculos Fotovoltaicos')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(13)
subtitle.runs[0].font.bold = True
subtitle.runs[0].font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)

doc.add_paragraph('')


def add_regla(numero, titulo, descripcion, sintaxis, ejemplo_desc, ejemplo_codigo):
    h = doc.add_heading(f'{numero}. {titulo}', level=1)
    h.runs[0].font.size = Pt(13)
    h.runs[0].font.color.rgb = RGBColor(0x1A, 0x5C, 0x8A)

    p = doc.add_paragraph()
    r = p.add_run('¿Qué hace? ')
    r.bold = True
    p.add_run(descripcion)

    p2 = doc.add_paragraph()
    r2 = p2.add_run('Sintaxis:  ')
    r2.bold = True
    cr = p2.add_run(sintaxis)
    cr.font.name = 'Courier New'
    cr.font.size = Pt(10)
    cr.font.color.rgb = RGBColor(0x17, 0x6B, 0x17)

    p3 = doc.add_paragraph()
    r3 = p3.add_run('Ejemplo FV:  ')
    r3.bold = True
    p3.add_run(ejemplo_desc + '  ')
    er = p3.add_run(ejemplo_codigo)
    er.font.name = 'Courier New'
    er.font.size = Pt(10)
    er.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    doc.add_paragraph('')


def add_seccion_titulo(texto):
    doc.add_paragraph('')
    h = doc.add_heading(texto, level=2)
    h.runs[0].font.size = Pt(12)
    h.runs[0].font.color.rgb = RGBColor(0x1A, 0x6B, 0x3C)
    doc.add_paragraph('')


def add_conversion(numero, nombre_from, nombre_to, formula_explicada, formula_codigo, ejemplo_desc, ejemplo_codigo):
    """Bloque especial para conversiones de métricas."""
    h = doc.add_heading(f'{numero}. Conversión: {nombre_from}  →  {nombre_to}', level=1)
    h.runs[0].font.size = Pt(13)
    h.runs[0].font.color.rgb = RGBColor(0x6E, 0x27, 0x94)  # morado para conversiones

    p = doc.add_paragraph()
    r = p.add_run('Fórmula:  ')
    r.bold = True
    p.add_run(formula_explicada)

    p2 = doc.add_paragraph()
    r2 = p2.add_run('En Python:  ')
    r2.bold = True
    cr = p2.add_run(formula_codigo)
    cr.font.name = 'Courier New'
    cr.font.size = Pt(10)
    cr.font.color.rgb = RGBColor(0x17, 0x6B, 0x17)

    p3 = doc.add_paragraph()
    r3 = p3.add_run('Ejemplo FV:  ')
    r3.bold = True
    p3.add_run(ejemplo_desc + '  ')
    er = p3.add_run(ejemplo_codigo)
    er.font.name = 'Courier New'
    er.font.size = Pt(10)
    er.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    doc.add_paragraph('')


# ─────────────────────────────────────────────
#  PARTE 1 — REGLAS BÁSICAS DE PYTHON
# ─────────────────────────────────────────────
add_seccion_titulo('PARTE 1 — REGLAS BÁSICAS DE PYTHON')

add_regla(1, 'Variable — Guardar un dato',
    'Almacena un valor en memoria con un nombre para usarlo después.',
    'nombre = valor',
    'Guardar la potencia de un panel solar:',
    'potencia_panel = 400   # watts')

add_regla(2, 'print() — Mostrar resultado',
    'Muestra un mensaje o resultado en la consola.',
    'print("texto", variable)',
    'Mostrar energía diaria generada:',
    'print("Energía diaria:", energia, "Wh")')

add_regla(3, 'input() — Pedir dato al usuario',
    'Pausa el programa y espera que el usuario escriba un valor.',
    'variable = input("mensaje")',
    'Pedir horas de sol del lugar:',
    'horas = input("Horas de sol al día: ")')

add_regla(4, 'float() — Convertir a número decimal',
    'Convierte texto a número decimal para realizar cálculos.',
    'variable = float(variable)',
    'Convertir horas de sol ingresadas para calcular:',
    'horas = float(horas)')

add_regla(5, 'int() — Convertir a número entero',
    'Convierte texto o decimal a número entero sin decimales.',
    'variable = int(variable)',
    'Convertir cantidad de paneles del usuario:',
    'paneles = int(input("Cantidad de paneles: "))')

add_regla(6, 'Operaciones aritméticas ( + - * / )',
    'Realiza suma, resta, multiplicación y división entre valores.',
    'resultado = valor1 * valor2 / valor3',
    'Calcular energía total del sistema solar:',
    'energia = potencia * horas * num_paneles')

add_regla(7, 'round() — Redondear decimales',
    'Redondea un número a la cantidad de decimales que indiques.',
    'round(numero, decimales)',
    'Redondear kWh mensuales generados:',
    'kwh_mes = round(energia * 30 / 1000, 2)')

add_regla(8, 'if / else — Condición',
    'Ejecuta acciones diferentes según si se cumple o no una condición.',
    'if condicion:\n      accion1\nelse:\n      accion2',
    'Verificar si el sistema cubre el consumo del hogar:',
    'if energia >= consumo:\n      print("Sistema suficiente")\nelse:\n      print("Faltan paneles")')

add_regla(9, 'for — Bucle (repetición)',
    'Repite un bloque de código para cada elemento de una lista o rango.',
    'for item in lista:\n      accion',
    'Calcular energía de varios modelos de paneles:',
    'for p in [300, 400, 500]:\n      print("Panel", p, "W →", p*5, "Wh/día")')

add_regla(10, 'range() — Generar secuencia de números',
    'Crea una secuencia de números desde un inicio hasta un fin.',
    'range(inicio, fin)',
    'Simular producción mensual durante 12 meses:',
    'for mes in range(1, 13):\n      print("Mes", mes, "→", kwh, "kWh")')

add_regla(11, 'Lista — Guardar varios valores',
    'Almacena múltiples valores en una sola variable.',
    'lista = [valor1, valor2, valor3]',
    'Guardar potencias de paneles disponibles:',
    'paneles = [300, 370, 400, 450, 550]')

add_regla(12, 'len() — Contar elementos de una lista',
    'Devuelve cuántos elementos tiene una lista.',
    'len(lista)',
    'Contar cuántos modelos de paneles hay:',
    'print("Modelos disponibles:", len(paneles))')

add_regla(13, 'Comentario con #',
    'Agrega una nota explicativa que Python ignora al ejecutar el código.',
    '# texto explicativo',
    'Documentar la fórmula usada:',
    '# Energía(Wh) = Potencia(W) x Horas_sol x N_paneles')

add_regla(14, 'type() — Ver tipo de dato',
    'Muestra si un dato es texto (str), entero (int) o decimal (float).',
    'print(type(variable))',
    'Verificar que horas_sol es número antes de calcular:',
    'print(type(horas_sol))   # debe ser <class float>')

add_regla(15, 'Fórmula con paréntesis — Control de orden',
    'Usa paréntesis para controlar el orden de las operaciones matemáticas.',
    'resultado = (a - b) * c / d',
    'Calcular retorno de inversión del sistema solar:',
    'roi = costo_sistema / (ahorro_anual)')


# ─────────────────────────────────────────────
#  PARTE 2 — CONVERSIONES DE MÉTRICAS
# ─────────────────────────────────────────────
add_seccion_titulo('PARTE 2 — CONVERSIONES DE MÉTRICAS EN PYTHON')

p_intro = doc.add_paragraph(
    'En ingeniería fotovoltaica es frecuente recibir datos en una unidad '
    'y necesitar calcular en otra. Esta sección muestra cómo escribir esas '
    'conversiones directamente en Python, con la fórmula, el código y un ejemplo aplicado.'
)
p_intro.runs[0].italic = True
doc.add_paragraph('')

# ── TEMPERATURA ──
add_seccion_titulo('  A) Conversiones de Temperatura')

add_conversion(16,
    'Celsius (°C)', 'Fahrenheit (°F)',
    '°F = (°C × 9/5) + 32',
    'tf = (tc * 9/5) + 32',
    'Convertir temperatura de operación del panel (25 °C estándar):',
    'tc = 25\ntf = (tc * 9/5) + 32\nprint("Temperatura:", tf, "°F")   # → 77.0 °F')

add_conversion(17,
    'Fahrenheit (°F)', 'Celsius (°C)',
    '°C = (°F − 32) × 5/9',
    'tc = (tf - 32) * 5/9',
    'Convertir dato climático de estación en °F a °C para calcular pérdida:',
    'tf = 95\ntc = (tf - 32) * 5/9\nprint("Temperatura:", round(tc,1), "°C")   # → 35.0 °C')

add_conversion(18,
    'Celsius (°C)', 'Kelvin (K)',
    'K = °C + 273.15',
    'tk = tc + 273.15',
    'Calcular eficiencia real del panel con temperatura en Kelvin:',
    'tc = 45\ntk = tc + 273.15\nprint("Temperatura en Kelvin:", tk, "K")   # → 318.15 K')

# ── POTENCIA ──
add_seccion_titulo('  B) Conversiones de Potencia')

add_conversion(19,
    'Watts (W)', 'Kilowatts (kW)',
    'kW = W ÷ 1 000',
    'kw = w / 1000',
    'Convertir potencia total de sistema de 3 200 W a kW:',
    'w = 3200\nkw = w / 1000\nprint("Potencia:", kw, "kW")   # → 3.2 kW')

add_conversion(20,
    'Kilowatts (kW)', 'Watts (W)',
    'W = kW × 1 000',
    'w = kw * 1000',
    'Expresar en W un panel nominado en 0.40 kW:',
    'kw = 0.40\nw = kw * 1000\nprint("Potencia:", w, "W")   # → 400.0 W')

add_conversion(21,
    'Watts (W)', 'Horsepower — HP (CV)',
    '1 HP = 745.7 W  →  HP = W ÷ 745.7',
    'hp = w / 745.7',
    'Comparar potencia de inversor (1 500 W) en HP:',
    'w = 1500\nhp = round(w / 745.7, 2)\nprint("Potencia:", hp, "HP")   # → 2.01 HP')

add_conversion(22,
    'Horsepower (HP)', 'Watts (W)',
    'W = HP × 745.7',
    'w = hp * 745.7',
    'Calcular carga en vatios de bomba de 2 HP alimentada por panel:',
    'hp = 2\nw = hp * 745.7\nprint("Carga bomba:", w, "W")   # → 1491.4 W')

# ── ENERGÍA ──
add_seccion_titulo('  C) Conversiones de Energía')

add_conversion(23,
    'Watt-hora (Wh)', 'Kilowatt-hora (kWh)',
    'kWh = Wh ÷ 1 000',
    'kwh = wh / 1000',
    'Convertir producción diaria de 4 800 Wh a kWh para factura:',
    'wh = 4800\nkwh = wh / 1000\nprint("Producción:", kwh, "kWh")   # → 4.8 kWh')

add_conversion(24,
    'Kilowatt-hora (kWh)', 'Watt-hora (Wh)',
    'Wh = kWh × 1 000',
    'wh = kwh * 1000',
    'Expresar en Wh el consumo mensual de 180 kWh para dimensionar banco de baterías:',
    'kwh = 180\nwh = kwh * 1000\nprint("Consumo:", wh, "Wh")   # → 180000 Wh')

add_conversion(25,
    'Kilowatt-hora (kWh)', 'Megawatt-hora (MWh)',
    'MWh = kWh ÷ 1 000',
    'mwh = kwh / 1000',
    'Calcular producción anual de un parque solar en MWh:',
    'kwh_anual = 52000\nmwh = kwh_anual / 1000\nprint("Producción anual:", mwh, "MWh")   # → 52.0 MWh')

add_conversion(26,
    'Watt-hora (Wh)', 'Joules (J)',
    '1 Wh = 3 600 J  →  J = Wh × 3 600',
    'j = wh * 3600',
    'Convertir energía de batería de 100 Wh a Joules:',
    'wh = 100\nj = wh * 3600\nprint("Energía:", j, "J")   # → 360000 J')

# ── ÁREA / SUPERFICIE ──
add_seccion_titulo('  D) Conversiones de Área')

add_conversion(27,
    'Metros cuadrados (m²)', 'Pies cuadrados (ft²)',
    '1 m² = 10.764 ft²  →  ft² = m² × 10.764',
    'ft2 = m2 * 10.764',
    'Calcular área de techo en ft² para clientes que usan sistema imperial:',
    'm2 = 30\nft2 = round(m2 * 10.764, 2)\nprint("Área:", ft2, "ft²")   # → 322.92 ft²')

add_conversion(28,
    'Pies cuadrados (ft²)', 'Metros cuadrados (m²)',
    '1 ft² = 0.0929 m²  →  m² = ft² × 0.0929',
    'm2 = ft2 * 0.0929',
    'Convertir techo de 400 ft² a m² para calcular cuántos paneles caben:',
    'ft2 = 400\nm2 = round(ft2 * 0.0929, 2)\nprint("Área:", m2, "m²")   # → 37.16 m²')

add_conversion(29,
    'Metros cuadrados (m²)', 'Hectáreas (ha)',
    '1 ha = 10 000 m²  →  ha = m² ÷ 10 000',
    'ha = m2 / 10000',
    'Expresar superficie de un parque solar de 25 000 m² en hectáreas:',
    'm2 = 25000\nha = m2 / 10000\nprint("Superficie:", ha, "ha")   # → 2.5 ha')

add_conversion(30,
    'Kilómetros cuadrados (km²)', 'Metros cuadrados (m²)',
    '1 km² = 1 000 000 m²  →  m² = km² × 1 000 000',
    'm2 = km2 * 1000000',
    'Convertir irradiación global en km² a m² para proyectar potencia instalable:',
    'km2 = 0.05\nm2 = km2 * 1000000\nprint("Área disponible:", m2, "m²")   # → 50000.0 m²')

# ─────────────────────────────────────────────
#  PARTE 3 — HORAS DE SOL PICO E IRRADIANCIA
# ─────────────────────────────────────────────
add_seccion_titulo('PARTE 3 — HORAS DE SOL PICO (HSP) E IRRADIANCIA SOLAR')

p_def = doc.add_paragraph()
r_def = p_def.add_run('Conceptos clave antes de programar:')
r_def.bold = True
r_def.font.size = Pt(11)
r_def.font.color.rgb = RGBColor(0x6E, 0x27, 0x94)
doc.add_paragraph('')

# Definición HSP
h_hsp = doc.add_heading('¿Qué es la Hora de Sol Pico (HSP)?', level=2)
h_hsp.runs[0].font.size = Pt(11)
h_hsp.runs[0].font.color.rgb = RGBColor(0x6E, 0x27, 0x94)

p_hsp = doc.add_paragraph(
    'Una HSP equivale a 1 hora de irradiancia constante a 1 000 W/m² (condición estándar). '
    'Es decir, si tu zona recibe 5 HSP al día, significa que la energía solar acumulada '
    'de ese día es equivalente a tener el sol al máximo durante 5 horas seguidas. '
    'Se obtiene de bases de datos como NASA POWER, PVGIS o SolarGIS, ingresando la '
    'latitud/longitud del proyecto.'
)
p_hsp.runs[0].font.size = Pt(10)
doc.add_paragraph('')

# Definición Irradiancia
h_irr = doc.add_heading('¿Qué es la Irradiancia?', level=2)
h_irr.runs[0].font.size = Pt(11)
h_irr.runs[0].font.color.rgb = RGBColor(0x6E, 0x27, 0x94)

p_irr = doc.add_paragraph(
    'La irradiancia es la potencia del sol que llega a 1 m² de superficie en un instante, '
    'medida en W/m². La irradiación es la energía acumulada en un período (Wh/m² o kWh/m²). '
    'La relación fundamental es: Irradiación (kWh/m²) = HSP × 1 kW/m².'
)
p_irr.runs[0].font.size = Pt(10)
doc.add_paragraph('')

add_seccion_titulo('  E) Cálculos y Conversiones de HSP e Irradiancia')

# ── Conversión 31: kWh/m²/día → HSP
add_conversion(31,
    'Irradiación diaria (kWh/m²/día)', 'HSP del lugar',
    'HSP = kWh/m²/día  (son numéricamente iguales, solo cambia la interpretación)',
    'hsp = irradiacion_kwh_m2',
    'El atlas solar indica 5.2 kWh/m²/día para tu ciudad — ese valor ES las HSP:',
    'irradiacion_kwh_m2 = 5.2   # dato de NASA POWER o PVGIS\n'
    'hsp = irradiacion_kwh_m2\n'
    'print("HSP del lugar:", hsp, "h/día")   # → 5.2 h/día')

# ── Conversión 32: Wh/m²/día → HSP
add_conversion(32,
    'Irradiación en Wh/m²/día', 'HSP',
    'HSP = Wh/m²/día ÷ 1 000',
    'hsp = irradiacion_wh_m2 / 1000',
    'Un sensor registró 5 200 Wh/m²/día. Convertir a HSP:',
    'irradiacion_wh_m2 = 5200\n'
    'hsp = irradiacion_wh_m2 / 1000\n'
    'print("HSP:", hsp, "h/día")   # → 5.2 h/día')

# ── Conversión 33: HSP → Energía generada (Wh)
add_conversion(33,
    'HSP + Potencia del sistema (W)', 'Energía diaria generada (Wh)',
    'Energía (Wh) = Potencia_pico (W) × HSP',
    'energia_wh = potencia_pico_w * hsp',
    'Sistema de 2 000 W en zona con 5.2 HSP. ¿Cuánto genera al día?',
    'potencia_pico_w = 2000   # 5 paneles × 400 W\n'
    'hsp = 5.2\n'
    'energia_wh = potencia_pico_w * hsp\n'
    'print("Energía diaria:", energia_wh, "Wh")   # → 10400 Wh → 10.4 kWh')

# ── Conversión 34: Energía real con pérdidas (PR)
add_conversion(34,
    'Energía ideal (Wh) + Factor de rendimiento (PR)', 'Energía real entregada (Wh)',
    'Energía_real = Potencia_pico × HSP × PR    (PR típico: 0.75 a 0.85)',
    'energia_real_wh = potencia_pico_w * hsp * pr',
    'Calcular energía real considerando pérdidas del sistema (PR = 0.80):',
    'potencia_pico_w = 2000\n'
    'hsp = 5.2\n'
    'pr  = 0.80   # 80% de rendimiento real (pérdidas por calor, cableado, inversor)\n'
    'energia_real_wh = potencia_pico_w * hsp * pr\n'
    'print("Energía real:", energia_real_wh, "Wh")   # → 8320 Wh')

# ── Conversión 35: Potencia pico necesaria
add_conversion(35,
    'Consumo diario (kWh) + HSP + PR', 'Potencia pico necesaria (kWp)',
    'Potencia_pico (kWp) = Consumo_diario (kWh) ÷ (HSP × PR)',
    'potencia_kWp = consumo_kwh / (hsp * pr)',
    'Hogar con 18 kWh/día de consumo, zona con 5 HSP y PR 0.78. ¿Cuántos kWp instalar?',
    'consumo_kwh = 18\n'
    'hsp = 5.0\n'
    'pr  = 0.78\n'
    'potencia_kWp = consumo_kwh / (hsp * pr)\n'
    'print("Potencia necesaria:", round(potencia_kWp, 2), "kWp")   # → 4.62 kWp')

# ── Conversión 36: Número de paneles necesarios
add_conversion(36,
    'Potencia pico (kWp) + Potencia por panel (W)', 'Número de paneles',
    'N_paneles = Potencia_pico (W) ÷ Potencia_panel (W)  → redondear hacia arriba',
    'import math\nn_paneles = math.ceil(potencia_w / potencia_panel_w)',
    'Sistema de 4.62 kWp con paneles de 400 W. ¿Cuántos paneles se necesitan?',
    'import math\n'
    'potencia_w = 4620        # 4.62 kWp convertido a W\n'
    'potencia_panel_w = 400\n'
    'n_paneles = math.ceil(potencia_w / potencia_panel_w)\n'
    'print("Paneles necesarios:", n_paneles)   # → 12 paneles')

# ── Conversión 37: Irradiancia W/m² → kW/m²
add_conversion(37,
    'Irradiancia (W/m²)', 'Irradiancia (kW/m²)',
    'kW/m² = W/m² ÷ 1 000',
    'irr_kw_m2 = irr_w_m2 / 1000',
    'Sensor registra 850 W/m² en el mediodía. Convertir a kW/m²:',
    'irr_w_m2 = 850\n'
    'irr_kw_m2 = irr_w_m2 / 1000\n'
    'print("Irradiancia:", irr_kw_m2, "kW/m²")   # → 0.85 kW/m²')

# ── Conversión 38: Eficiencia real del panel según temperatura
add_conversion(38,
    'Eficiencia nominal + temperatura real (°C)', 'Eficiencia real del panel (%)',
    'Efic_real = Efic_nom × (1 + Coef_temp × (T_real − 25))\n'
    '   Coef_temp típico: −0.0035 por °C  (pérdida del 0.35 % por cada °C sobre 25°C)',
    'efic_real = efic_nom * (1 + coef_temp * (temp_c - 25))',
    'Panel con eficiencia nominal 20 %, operando a 45 °C. ¿Cuál es su eficiencia real?',
    'efic_nom  = 0.20        # 20% de eficiencia en condición estándar\n'
    'coef_temp = -0.0035     # coeficiente de temperatura del fabricante\n'
    'temp_c    = 45          # temperatura real de operación\n'
    'efic_real = efic_nom * (1 + coef_temp * (temp_c - 25))\n'
    'print("Eficiencia real:", round(efic_real * 100, 2), "%")   # → 18.6 %')

# ── Conversión 39: Irradiación mensual → producción mensual
add_conversion(39,
    'Irradiación mensual (kWh/m²/mes) + Potencia pico + PR', 'Producción mensual (kWh)',
    'Producción_mes (kWh) = Potencia_pico (kWp) × Irradiación_mes (kWh/m²) × PR',
    'prod_mes_kwh = potencia_kWp * irradiacion_mes * pr',
    'Sistema de 3 kWp, mes con 150 kWh/m² de irradiación, PR 0.80:',
    'potencia_kWp     = 3.0\n'
    'irradiacion_mes  = 150   # kWh/m²/mes (dato de PVGIS o NASA)\n'
    'pr               = 0.80\n'
    'prod_mes_kwh = potencia_kWp * irradiacion_mes * pr\n'
    'print("Producción mensual:", prod_mes_kwh, "kWh")   # → 360.0 kWh')

# ── Conversión 40: Producción anual estimada
add_conversion(40,
    'Producción mensual (kWh/mes)', 'Producción anual (kWh/año)',
    'Producción_año = suma de los 12 meses  (o simplificada: promedio × 12)',
    'prod_anual = sum(producciones_por_mes)',
    'Calcular producción anual con irradiación mensual variable por estación:',
    '# Irradiación mensual típica de una ciudad (kWh/m²/mes)\n'
    'irradiacion_meses = [130, 140, 160, 170, 180, 175, 185, 178, 165, 150, 135, 125]\n'
    'potencia_kWp = 3.0\n'
    'pr = 0.80\n'
    'producciones = [potencia_kWp * irr * pr for irr in irradiacion_meses]\n'
    'prod_anual = round(sum(producciones), 1)\n'
    'print("Producción anual:", prod_anual, "kWh/año")')

# ── EJEMPLO INTEGRADOR COMPLETO HSP ──
add_seccion_titulo('EJEMPLO INTEGRADOR COMPLETO — HSP + Irradiancia + Temperatura + Dimensionado')

p_ej2 = doc.add_paragraph()
r_t2 = p_ej2.add_run('Caso: Dimensionar un sistema solar desde cero usando todos los conceptos de HSP')
r_t2.bold = True
r_t2.font.size = Pt(11)
r_t2.font.color.rgb = RGBColor(0x6E, 0x27, 0x94)

codigo_hsp = (
    'import math\n\n'
    '# ── DATOS DEL PROYECTO ──\n'
    'consumo_kwh_dia   = 20.0    # consumo diario del hogar (kWh/día)\n'
    'hsp               = 5.2     # HSP del lugar (dato de NASA POWER o PVGIS)\n'
    'pr                = 0.80    # rendimiento del sistema (80%)\n'
    'potencia_panel_w  = 400     # potencia nominal de cada panel (W)\n'
    'efic_nom          = 0.20    # eficiencia nominal del panel (20%)\n'
    'coef_temp         = -0.0035 # coeficiente de temperatura (/°C)\n'
    'temp_f            = 95      # temperatura ambiente del lugar (°F)\n\n'
    '# ── PASO 1: Convertir temperatura ──\n'
    'temp_c = (temp_f - 32) * 5/9\n'
    'print("Temperatura de operación:", round(temp_c, 1), "°C")\n\n'
    '# ── PASO 2: Eficiencia real según temperatura ──\n'
    'efic_real = efic_nom * (1 + coef_temp * (temp_c - 25))\n'
    'print("Eficiencia real del panel:", round(efic_real * 100, 2), "%")\n\n'
    '# ── PASO 3: Potencia pico necesaria ──\n'
    'potencia_kWp = consumo_kwh_dia / (hsp * pr)\n'
    'print("Potencia pico necesaria:", round(potencia_kWp, 2), "kWp")\n\n'
    '# ── PASO 4: Número de paneles ──\n'
    'potencia_w = potencia_kWp * 1000\n'
    'n_paneles  = math.ceil(potencia_w / potencia_panel_w)\n'
    'print("Paneles necesarios:", n_paneles, "unidades")\n\n'
    '# ── PASO 5: Energía real que generará el sistema ──\n'
    'potencia_real_w  = n_paneles * potencia_panel_w\n'
    'energia_real_wh  = potencia_real_w * hsp * pr\n'
    'energia_real_kwh = energia_real_wh / 1000\n'
    'print("Energía real generada:", round(energia_real_kwh, 2), "kWh/día")\n\n'
    '# ── PASO 6: ¿Cubre el consumo? ──\n'
    'if energia_real_kwh >= consumo_kwh_dia:\n'
    '    excedente = round(energia_real_kwh - consumo_kwh_dia, 2)\n'
    '    print("Sistema suficiente. Excedente:", excedente, "kWh/día")\n'
    'else:\n'
    '    faltante = round(consumo_kwh_dia - energia_real_kwh, 2)\n'
    '    print("Sistema insuficiente. Faltan:", faltante, "kWh/día")\n'
)

p_c2 = doc.add_paragraph()
cr2 = p_c2.add_run(codigo_hsp)
cr2.font.name = 'Courier New'
cr2.font.size = Pt(9.5)
cr2.font.color.rgb = RGBColor(0x10, 0x10, 0x60)

doc.add_paragraph('')

# ── EJEMPLO INTEGRADOR ──
add_seccion_titulo('EJEMPLO INTEGRADOR — Uso combinado de conversiones')

p_ej = doc.add_paragraph()
r_titulo = p_ej.add_run('Caso: Calcular producción y verificar suficiencia de un sistema solar')
r_titulo.bold = True
r_titulo.font.size = Pt(11)
r_titulo.font.color.rgb = RGBColor(0x1A, 0x5C, 0x8A)

codigo_integrador = (
    '# Datos de entrada\n'
    'potencia_w     = 400          # potencia de cada panel en W\n'
    'num_paneles    = 8            # cantidad de paneles\n'
    'horas_sol      = 5.5          # horas de sol pico (HSP) del lugar\n'
    'consumo_kwh    = 20           # consumo diario del hogar en kWh\n'
    'temp_f         = 95           # temperatura ambiente en °F\n\n'
    '# Conversión de temperatura\n'
    'temp_c = (temp_f - 32) * 5/9\n'
    'print("Temperatura de operación:", round(temp_c,1), "°C")\n\n'
    '# Energía generada en Wh y kWh\n'
    'energia_wh  = potencia_w * num_paneles * horas_sol\n'
    'energia_kwh = energia_wh / 1000\n'
    'print("Energía generada:", energia_kwh, "kWh/día")\n\n'
    '# ¿Cubre el consumo?\n'
    'if energia_kwh >= consumo_kwh:\n'
    '    print("Sistema suficiente")\n'
    'else:\n'
    '    faltante = consumo_kwh - energia_kwh\n'
    '    print("Faltan", faltante, "kWh — se necesitan más paneles")\n'
)

p_code = doc.add_paragraph()
cr = p_code.add_run(codigo_integrador)
cr.font.name = 'Courier New'
cr.font.size = Pt(9.5)
cr.font.color.rgb = RGBColor(0x10, 0x10, 0x60)

doc.add_paragraph('')

# ─────────────────────────────────────────────
#  PARTE 4 — DIMENSIONADO ELÉCTRICO (como PVsyst)
# ─────────────────────────────────────────────
add_seccion_titulo('PARTE 4 — DIMENSIONADO ELÉCTRICO DEL SISTEMA')

p_intro4 = doc.add_paragraph(
    'PVsyst calcula la arquitectura eléctrica del sistema: cuántos paneles van en serie '
    '(string), cuántos strings en paralelo, si el inversor es compatible en tensión y '
    'corriente, y cuál es la temperatura real de la celda. Estos cálculos evitan errores '
    'graves de diseño. A continuación, los mismos cálculos en Python.'
)
p_intro4.runs[0].font.size = Pt(10)
p_intro4.runs[0].italic = True
doc.add_paragraph('')

add_conversion(41,
    'Temperatura ambiente (°C) + NOCT', 'Temperatura real de la celda (°C)',
    'T_celda = T_amb + (NOCT − 20) × Irradiancia / 800\n'
    '   NOCT típico: 45°C  |  Irradiancia de referencia: 800 W/m²',
    'T_celda = T_amb + ((NOCT - 20) / 800) * irradiancia_w_m2',
    'Panel con NOCT 45°C, T_amb 30°C, irradiancia al mediodía 950 W/m²:',
    'T_amb          = 30\n'
    'NOCT           = 45\n'
    'irradiancia    = 950   # W/m²\n'
    'T_celda = T_amb + ((NOCT - 20) / 800) * irradiancia\n'
    'print("Temperatura de celda:", round(T_celda, 1), "°C")   # → 56.6 °C')

add_conversion(42,
    'Tensión Voc nominal + temperatura de celda', 'Voc real del panel (V)',
    'Voc_real = Voc_nom × (1 + CoefV × (T_celda − 25))\n'
    '   CoefV típico: −0.0029 /°C  (pérdida del 0.29 % por cada °C sobre 25°C)',
    'Voc_real = Voc_nom * (1 + coef_v * (T_celda - 25))',
    'Panel con Voc 49 V, CoefV −0.0029, celda a 56.6°C:',
    'Voc_nom  = 49.0\n'
    'coef_v   = -0.0029\n'
    'T_celda  = 56.6\n'
    'Voc_real = Voc_nom * (1 + coef_v * (T_celda - 25))\n'
    'print("Voc real:", round(Voc_real, 2), "V")   # → 43.47 V')

add_conversion(43,
    'Tensión Voc real + Rango de tensión del inversor (V)', 'Número de paneles en serie (string)',
    'N_serie_max = Vdc_max_inv ÷ Voc_real_min_temp\n'
    'N_serie_min = Vmppt_min_inv ÷ Vmp_real_max_temp\n'
    '   Elegir N_serie dentro de ese rango (valor entero)',
    'import math\n'
    'n_max = math.floor(Vdc_max / Voc_fria)\n'
    'n_min = math.ceil(Vmppt_min / Vmp_caliente)',
    'Inversor con Vdc_max 1000V, MPPT_min 200V. Voc en frío 52V, Vmp en caliente 36V:',
    'import math\n'
    'Vdc_max     = 1000   # tensión máxima del inversor\n'
    'Vmppt_min   = 200    # tensión mínima del rango MPPT\n'
    'Voc_fria    = 52.0   # Voc del panel a temperatura mínima (-10°C)\n'
    'Vmp_caliente= 36.0   # Vmp del panel a temperatura máxima (70°C)\n'
    'n_max = math.floor(Vdc_max / Voc_fria)\n'
    'n_min = math.ceil(Vmppt_min / Vmp_caliente)\n'
    'print(f"Paneles en serie: entre {n_min} y {n_max}")   # → entre 6 y 19')

add_conversion(44,
    'Corriente Isc del string + Corriente máx del inversor (A)', 'Número de strings en paralelo',
    'N_strings = floor(Idc_max_inv ÷ Isc_string)\n'
    '   Isc_string ≈ Isc_panel  (para un string; se multiplica por strings en paralelo)',
    'import math\n'
    'n_strings = math.floor(Idc_max_inv / Isc_panel)',
    'Inversor con Idc_max 30 A, panel con Isc 10.5 A. ¿Cuántos strings en paralelo?',
    'import math\n'
    'Idc_max_inv = 30.0   # corriente máxima DC del inversor\n'
    'Isc_panel   = 10.5   # corriente de cortocircuito del panel\n'
    'n_strings = math.floor(Idc_max_inv / Isc_panel)\n'
    'print("Strings en paralelo:", n_strings)   # → 2 strings')

add_conversion(45,
    'Potencia DC total (kWp) + Potencia AC del inversor (kW)', 'Ratio DC/AC (Sizing Ratio)',
    'Ratio_DC_AC = Potencia_DC (kWp) ÷ Potencia_AC (kW)\n'
    '   Rango ideal PVsyst: 1.10 a 1.30  (sobredimensionar DC compensa pérdidas)',
    'ratio_dc_ac = potencia_dc_kWp / potencia_ac_kW',
    'Sistema de 4.8 kWp con inversor de 4 kW AC:',
    'potencia_dc_kWp = 4.8\n'
    'potencia_ac_kW  = 4.0\n'
    'ratio_dc_ac = potencia_dc_kWp / potencia_ac_kW\n'
    'print("Ratio DC/AC:", round(ratio_dc_ac, 2))   # → 1.2  ✓ dentro del rango ideal')

add_conversion(46,
    'Potencia DC del sistema + Eficiencia del inversor (%)', 'Potencia AC entregada a la red (kW)',
    'P_AC = P_DC × Eficiencia_inversor\n'
    '   Eficiencia típica de inversores modernos: 97–98.5 %',
    'P_ac_kW = P_dc_kW * efic_inv',
    'Inversor recibe 3.6 kW DC con eficiencia del 97%:',
    'P_dc_kW  = 3.6\n'
    'efic_inv = 0.97\n'
    'P_ac_kW  = P_dc_kW * efic_inv\n'
    'print("Potencia AC entregada:", round(P_ac_kW, 3), "kW")   # → 3.492 kW')

add_conversion(47,
    'Energía diaria + Voltaje del banco de baterías + DOD', 'Capacidad del banco de baterías (Ah)',
    'C_Ah = Energía_Wh ÷ (V_banco × DOD)\n'
    '   DOD: profundidad de descarga (0.5 para plomo-ácido, 0.8 para litio)',
    'C_ah = energia_wh_autonomia / (v_banco * dod)',
    'Consumo de autonomía 2 días = 10 000 Wh, banco 48V, baterías litio (DOD 80%):',
    'energia_wh_autonomia = 10000   # Wh para N días de autonomía\n'
    'v_banco              = 48      # voltaje del banco de baterías\n'
    'dod                  = 0.80    # profundidad de descarga litio\n'
    'C_ah = energia_wh_autonomia / (v_banco * dod)\n'
    'print("Capacidad banco:", round(C_ah, 1), "Ah")   # → 260.4 Ah')

# ─────────────────────────────────────────────
#  PARTE 5 — PÉRDIDAS DEL SISTEMA (Loss Diagram PVsyst)
# ─────────────────────────────────────────────
add_seccion_titulo('PARTE 5 — PÉRDIDAS DEL SISTEMA (Loss Diagram de PVsyst en Python)')

p_intro5 = doc.add_paragraph(
    'El diagrama de pérdidas (Loss Diagram) es uno de los resultados más importantes de PVsyst. '
    'Muestra cómo la energía solar se va reduciendo en cada etapa hasta llegar al contador. '
    'Cada factor de pérdida se aplica como un coeficiente multiplicador. En Python se puede '
    'replicar este cálculo paso a paso, partiendo de la energía teórica y aplicando cada pérdida.'
)
p_intro5.runs[0].font.size = Pt(10)
p_intro5.runs[0].italic = True
doc.add_paragraph('')

add_conversion(48,
    'Energía ideal (sin pérdidas)', 'Energía real paso a paso aplicando cada pérdida',
    'E_real = E_ideal × (1 − pérdida_1) × (1 − pérdida_2) × ... × (1 − pérdida_N)',
    'e = e_ideal * (1-p_temp) * (1-p_suciedad) * (1-p_sombra) * (1-p_mismatch) * (1-p_cableado) * efic_inv',
    'Calcular energía entregada a la red aplicando el loss diagram completo:',
    '# Energía ideal (STC)\n'
    'e_ideal_kwh  = 20.0    # kWh/día en condición ideal\n\n'
    '# Factores de pérdida (valores típicos de PVsyst)\n'
    'p_temperatura = 0.048  # 4.8% pérdida por temperatura\n'
    'p_suciedad    = 0.030  # 3.0% suciedad/polvo sobre el panel\n'
    'p_sombra      = 0.020  # 2.0% sombras lejanas (horizonte)\n'
    'p_mismatch    = 0.015  # 1.5% desajuste entre paneles del string\n'
    'p_cableado    = 0.015  # 1.5% resistencia del cableado DC\n'
    'efic_inversor = 0.970  # 97.0% eficiencia del inversor\n\n'
    'e_real = (e_ideal_kwh\n'
    '         * (1 - p_temperatura)\n'
    '         * (1 - p_suciedad)\n'
    '         * (1 - p_sombra)\n'
    '         * (1 - p_mismatch)\n'
    '         * (1 - p_cableado)\n'
    '         * efic_inversor)\n'
    'print("Energía entregada a la red:", round(e_real, 2), "kWh/día")   # → 16.67 kWh')

add_conversion(49,
    'Energía real + Energía ideal', 'Performance Ratio — PR (%)',
    'PR = Energía_real ÷ Energía_ideal × 100\n'
    '   PR óptimo según PVsyst: 75–85 %',
    'pr_real = (e_real / e_ideal_kwh) * 100',
    'Calcular el PR del sistema con los datos del loss diagram anterior:',
    'e_ideal_kwh = 20.0\n'
    'e_real      = 16.67\n'
    'pr_real = (e_real / e_ideal_kwh) * 100\n'
    'print("Performance Ratio:", round(pr_real, 1), "%")   # → 83.4 %')

add_conversion(50,
    'Energía anual generada (kWh/año) + Potencia pico (kWp)', 'Rendimiento específico (kWh/kWp/año)',
    'Rendimiento_esp = Energía_anual ÷ Potencia_pico\n'
    '   Valor de referencia PVsyst: 1 200–1 800 kWh/kWp/año según zona',
    'rendimiento_esp = energia_anual_kwh / potencia_kWp',
    'Sistema de 5 kWp que genera 7 800 kWh al año:',
    'energia_anual_kwh = 7800\n'
    'potencia_kWp      = 5.0\n'
    'rendimiento_esp   = energia_anual_kwh / potencia_kWp\n'
    'print("Rendimiento específico:", rendimiento_esp, "kWh/kWp/año")   # → 1560.0')

add_conversion(51,
    'Energía real + Potencia instalada + Horas del período', 'Factor de Capacidad — CF (%)',
    'CF = Energía_real ÷ (Potencia_instalada × Horas_período) × 100\n'
    '   CF solar típico: 15–25 %',
    'cf = (energia_anual_kwh / (potencia_kWp * 8760)) * 100',
    'Sistema de 5 kWp que generó 7 800 kWh en 1 año (8 760 horas):',
    'energia_anual_kwh = 7800\n'
    'potencia_kWp      = 5.0\n'
    'horas_año         = 8760\n'
    'cf = (energia_anual_kwh / (potencia_kWp * horas_año)) * 100\n'
    'print("Factor de Capacidad:", round(cf, 1), "%")   # → 17.8 %')

add_conversion(52,
    'Producción año 1 + Tasa de degradación anual (%)', 'Producción en año N (con degradación)',
    'Prod_año_N = Prod_año_1 × (1 − degradacion) ^ (N − 1)\n'
    '   Degradación típica según PVsyst: 0.5 % anual (módulos monocristalinos)',
    'prod_n = prod_año1 * ((1 - degradacion) ** (n - 1))',
    'Sistema genera 7 800 kWh en año 1. ¿Cuánto genera en el año 25?',
    'prod_año1   = 7800   # kWh en el primer año\n'
    'degradacion = 0.005  # 0.5% de pérdida por año\n'
    'n           = 25     # año a calcular\n'
    'prod_n = prod_año1 * ((1 - degradacion) ** (n - 1))\n'
    'print("Producción año 25:", round(prod_n, 1), "kWh")   # → 6920.5 kWh\n\n'
    '# Ver degradación para todos los 25 años\n'
    'for año in range(1, 26):\n'
    '    p = round(prod_año1 * ((1 - degradacion) ** (año - 1)), 1)\n'
    '    print(f"Año {año:2d}: {p} kWh")')

# ─────────────────────────────────────────────
#  PARTE 6 — ANÁLISIS FINANCIERO
# ─────────────────────────────────────────────
add_seccion_titulo('PARTE 6 — ANÁLISIS FINANCIERO (PVsyst Economic Analysis)')

p_intro6 = doc.add_paragraph(
    'PVsyst incluye un módulo de análisis económico que calcula el retorno de la inversión. '
    'Los indicadores clave son: Payback simple, Valor Actual Neto (VAN/NPV), Tasa Interna de '
    'Retorno (TIR/IRR) y el LCOE (Costo Nivelado de Energía). Estos son los mismos cálculos '
    'en Python, explicados paso a paso.'
)
p_intro6.runs[0].font.size = Pt(10)
p_intro6.runs[0].italic = True
doc.add_paragraph('')

add_conversion(53,
    'Costo total del sistema + Ahorro anual en cuenta eléctrica', 'Payback simple (años)',
    'Payback = Costo_total ÷ Ahorro_anual',
    'payback = costo_total / ahorro_anual',
    'Sistema que cuesta $6 000 000 y ahorra $900 000 al año en electricidad:',
    'costo_total  = 6000000   # pesos o dólares\n'
    'ahorro_anual =  900000\n'
    'payback = costo_total / ahorro_anual\n'
    'print("Payback simple:", round(payback, 1), "años")   # → 6.7 años')

add_conversion(54,
    'Energía generada en 25 años (kWh totales) + Costo total del sistema', 'LCOE ($/kWh)',
    'LCOE = Costo_total ÷ Energía_total_25_años\n'
    '   Incluye degradación y costo de O&M (Operación y Mantenimiento)',
    'lcoe = (costo_total + costo_om_total) / energia_total_25',
    'Sistema $6 000 000, O&M $50 000/año, genera ~175 000 kWh en 25 años:',
    'costo_total      = 6000000\n'
    'costo_om_anual   =   50000\n'
    'prod_año1        =    7800   # kWh/año\n'
    'degradacion      =   0.005\n\n'
    'energia_total_25 = sum(prod_año1 * ((1-degradacion)**(a)) for a in range(25))\n'
    'costo_om_total   = costo_om_anual * 25\n'
    'lcoe = (costo_total + costo_om_total) / energia_total_25\n'
    'print("LCOE:", round(lcoe, 1), "$/kWh")   # → 44.8 $/kWh (o pesos/kWh)')

add_conversion(55,
    'Flujos de caja anuales + Tasa de descuento', 'VAN / NPV (Valor Actual Neto)',
    'VAN = −Inversión_inicial + Σ (Flujo_año_n ÷ (1 + tasa)^n)\n'
    '   VAN > 0 → el proyecto es rentable  |  tasa típica: 6–10 %',
    'van = -inversion + sum(flujo / (1+tasa)**n for n, flujo in enumerate(flujos, 1))',
    'Inversión $6 000 000, ahorro creciente 3% anual, tasa 7%, horizonte 25 años:',
    'inversion    = 6000000\n'
    'ahorro_año1  =  900000\n'
    'crecimiento  =    0.03   # tarifa eléctrica sube 3% al año\n'
    'tasa         =    0.07   # tasa de descuento (costo del dinero)\n\n'
    'flujos = [ahorro_año1 * ((1 + crecimiento) ** (n-1)) for n in range(1, 26)]\n'
    'van = -inversion + sum(f / (1+tasa)**n for n, f in enumerate(flujos, 1))\n'
    'print("VAN:", round(van, 0), "$")   # VAN > 0 → proyecto rentable')

add_conversion(56,
    'Flujos de caja anuales + Inversión inicial', 'TIR / IRR (Tasa Interna de Retorno, %)',
    'TIR = tasa que hace VAN = 0  →  se calcula por iteración (bisección)\n'
    '   TIR > tasa de descuento → el proyecto es rentable',
    '# Método de bisección para encontrar la TIR\n'
    'tasa_baja, tasa_alta = 0.0, 1.0\n'
    'for _ in range(100): ...',
    'Calcular la TIR del mismo proyecto anterior con bisección en Python:',
    'inversion   = 6000000\n'
    'ahorro_año1 =  900000\n'
    'crecimiento =   0.03\n'
    'flujos = [-inversion] + [ahorro_año1*((1+crecimiento)**(n)) for n in range(25)]\n\n'
    '# Bisección para TIR\n'
    'tasa_baja, tasa_alta = 0.0, 1.0\n'
    'for _ in range(200):\n'
    '    tasa_mid = (tasa_baja + tasa_alta) / 2\n'
    '    van_mid  = sum(f/(1+tasa_mid)**n for n,f in enumerate(flujos))\n'
    '    if van_mid > 0: tasa_baja = tasa_mid\n'
    '    else:           tasa_alta = tasa_mid\n'
    'tir = tasa_mid * 100\n'
    'print("TIR:", round(tir, 2), "%")   # → TIR estimada del proyecto')

add_conversion(57,
    'Energía generada anual (kWh/año)', 'CO₂ evitado (toneladas/año)',
    'CO2_evitado = Energía_kWh × Factor_emisión\n'
    '   Factor de emisión eléctrica chilena: ≈ 0.294 kg CO₂/kWh (SEN 2024)\n'
    '   Factor promedio Latinoamérica: ≈ 0.36 kg CO₂/kWh',
    'co2_kg = energia_kwh_anual * factor_emision\nco2_ton = co2_kg / 1000',
    'Sistema que genera 7 800 kWh/año en Chile. ¿Cuánto CO₂ evita?',
    'energia_kwh_anual = 7800\n'
    'factor_emision    = 0.294   # kg CO2/kWh — factor SEN Chile 2024\n'
    'co2_kg  = energia_kwh_anual * factor_emision\n'
    'co2_ton = co2_kg / 1000\n'
    'print("CO2 evitado:", round(co2_ton, 2), "ton CO2/año")   # → 2.29 ton CO2/año\n'
    'print("En 25 años:", round(co2_ton * 25, 1), "ton CO2")   # → 57.3 ton CO2')

# ── EJEMPLO INTEGRADOR FINAL COMPLETO ──
add_seccion_titulo('EJEMPLO INTEGRADOR FINAL — Simulación completa tipo PVsyst en Python')

p_ej3 = doc.add_paragraph()
r_t3 = p_ej3.add_run(
    'Caso: Análisis técnico-económico completo de un sistema solar residencial de 5 kWp'
)
r_t3.bold = True
r_t3.font.size = Pt(11)
r_t3.font.color.rgb = RGBColor(0x6E, 0x27, 0x94)

codigo_final = (
    'import math\n\n'
    '# ══════════════════════════════════════════════\n'
    '# DATOS DE ENTRADA DEL PROYECTO\n'
    '# ══════════════════════════════════════════════\n'
    'potencia_panel_w   = 400      # W por panel\n'
    'n_paneles          = 12       # cantidad de paneles (string × paralelo)\n'
    'hsp                = 5.2      # Horas de Sol Pico del lugar\n'
    'temp_amb_c         = 30       # temperatura ambiente media (°C)\n'
    'NOCT               = 45       # temperatura nominal de operación del panel\n'
    'efic_nom           = 0.20     # eficiencia nominal del panel\n'
    'coef_temp          = -0.0035  # coeficiente de temperatura (/°C)\n'
    'efic_inversor      = 0.97     # eficiencia del inversor\n'
    'degradacion_anual  = 0.005    # 0.5% de pérdida por año\n'
    'costo_sistema      = 6000000  # costo total instalado ($)\n'
    'tarifa_kwh         = 120      # precio del kWh en la tarifa eléctrica ($)\n'
    'crecimiento_tarifa = 0.03     # alza anual de la tarifa eléctrica (3%)\n'
    'tasa_descuento     = 0.07     # tasa de descuento para VAN (7%)\n'
    'factor_co2         = 0.294    # kg CO2/kWh evitado\n\n'
    '# ── PASO 1: Temperatura real de la celda ──\n'
    'irradiancia   = hsp * 1000 / hsp   # simplificado: 1 000 W/m² al pico\n'
    'T_celda = temp_amb_c + ((NOCT - 20) / 800) * 950\n'
    'efic_real = efic_nom * (1 + coef_temp * (T_celda - 25))\n'
    'print(f"Temperatura celda: {T_celda:.1f} °C | Efic. real: {efic_real*100:.2f} %")\n\n'
    '# ── PASO 2: Energía ideal y real con pérdidas ──\n'
    'potencia_dc_w = potencia_panel_w * n_paneles\n'
    'e_ideal_wh    = potencia_dc_w * hsp\n'
    'p_temp        = 1 - efic_real / efic_nom   # pérdida real por temperatura\n'
    'e_real_kwh = (e_ideal_wh\n'
    '              * (1 - p_temp)\n'
    '              * (1 - 0.030)   # suciedad\n'
    '              * (1 - 0.020)   # sombras\n'
    '              * (1 - 0.015)   # mismatch\n'
    '              * (1 - 0.015)   # cableado\n'
    '              * efic_inversor) / 1000\n'
    'pr = e_real_kwh / (e_ideal_wh / 1000) * 100\n'
    'print(f"Energía real: {e_real_kwh:.2f} kWh/día | PR: {pr:.1f} %")\n\n'
    '# ── PASO 3: Producción anual con degradación ──\n'
    'e_año1     = e_real_kwh * 365\n'
    'prod_25    = [round(e_año1 * ((1-degradacion_anual)**(a)), 1) for a in range(25)]\n'
    'total_25   = round(sum(prod_25), 0)\n'
    'print(f"Producción año 1: {e_año1:.0f} kWh | Total 25 años: {total_25:.0f} kWh")\n\n'
    '# ── PASO 4: Análisis financiero ──\n'
    'ahorro_año1 = e_año1 * tarifa_kwh\n'
    'payback     = costo_sistema / ahorro_año1\n'
    'flujos_van  = [ahorro_año1*((1+crecimiento_tarifa)**n) for n in range(25)]\n'
    'van         = -costo_sistema + sum(f/(1+tasa_descuento)**(n+1) for n,f in enumerate(flujos_van))\n'
    'lcoe        = (costo_sistema + 50000*25) / total_25\n'
    'print(f"Ahorro año 1: ${ahorro_año1:,.0f}")\n'
    'print(f"Payback: {payback:.1f} años | VAN: ${van:,.0f} | LCOE: ${lcoe:.1f}/kWh")\n\n'
    '# ── PASO 5: CO₂ evitado ──\n'
    'co2_25_ton = round(total_25 * factor_co2 / 1000, 1)\n'
    'print(f"CO2 evitado en 25 años: {co2_25_ton} toneladas")\n'
)

p_c3 = doc.add_paragraph()
cr3 = p_c3.add_run(codigo_final)
cr3.font.name = 'Courier New'
cr3.font.size = Pt(9)
cr3.font.color.rgb = RGBColor(0x10, 0x10, 0x60)

doc.add_paragraph('')

# ─────────────────────────────────────────────
#  PARTE 7 — GRÁFICOS CON MATPLOTLIB
# ─────────────────────────────────────────────
add_seccion_titulo('PARTE 7 — GENERACIÓN DE GRÁFICOS CON MATPLOTLIB')

p_intro7 = doc.add_paragraph(
    'Matplotlib es la librería estándar de Python para visualizar resultados. '
    'Para usarla sin conflictos con los cálculos ya escritos, existen reglas de '
    'orden y estructura que SIEMPRE debes respetar. Se explican primero las reglas '
    'obligatorias y luego cada tipo de gráfico aplicado a fotovoltaica.'
)
p_intro7.runs[0].font.size = Pt(10)
p_intro7.runs[0].italic = True
doc.add_paragraph('')

# ── REGLAS OBLIGATORIAS ──
h_reglas = doc.add_heading('REGLAS OBLIGATORIAS — Cómo evitar conflictos al graficar', level=2)
h_reglas.runs[0].font.size = Pt(11)
h_reglas.runs[0].font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

reglas_texto = [
    ('REGLA 1 — Importar siempre al principio del archivo',
     'El import de matplotlib debe ir en la PRIMERA línea del archivo, '
     'antes de cualquier cálculo. Si lo pones en medio del código, puede '
     'generar errores de estado o sobrescribir variables ya definidas.',
     'import matplotlib\nmatplotlib.use("Agg")        # modo sin ventana emergente (ideal para scripts)\nimport matplotlib.pyplot as plt\n# A partir de aquí van tus variables y cálculos...'),

    ('REGLA 2 — Abrir y cerrar cada gráfico correctamente',
     'Cada gráfico debe abrirse con plt.figure() y cerrarse con plt.close() '
     'al terminar. Si no cierras, el siguiente gráfico se dibuja encima del '
     'anterior y los resultados se mezclan.',
     'plt.figure()           # abre un lienzo nuevo\n# ... código del gráfico ...\nplt.savefig("grafico.png")  # guarda el archivo\nplt.close()            # cierra y libera memoria — OBLIGATORIO'),

    ('REGLA 3 — Usar savefig() en lugar de show()',
     'plt.show() abre una ventana emergente y CONGELA el programa hasta que '
     'el usuario la cierre. En scripts de cálculo fotovoltaico eso interrumpe '
     'el flujo. Usa siempre savefig() para guardar la imagen sin pausar.',
     'plt.savefig("produccion_mensual.png", dpi=150, bbox_inches="tight")\n# dpi=150 → calidad aceptable | bbox_inches="tight" → sin recortes'),

    ('REGLA 4 — Nombrar los ejes y el título siempre',
     'Un gráfico sin etiquetas no sirve para informes ni para interpretar '
     'resultados. PVsyst siempre etiqueta sus gráficos — sigue ese estándar.',
     'plt.title("Producción mensual del sistema solar")\nplt.xlabel("Mes del año")\nplt.ylabel("Energía generada (kWh)")\nplt.grid(True, linestyle="--", alpha=0.5)  # grilla suave'),

    ('REGLA 5 — plt.close("all") para limpiar si hay muchos gráficos',
     'Si tu script genera varios gráficos en secuencia, usa plt.close("all") '
     'antes de empezar el bloque de gráficos para asegurarte de que no haya '
     'ningún lienzo residual de una ejecución anterior.',
     'plt.close("all")   # limpia todos los gráficos abiertos\n# Ahora genera tus gráficos con seguridad'),
]

for titulo_r, desc_r, cod_r in reglas_texto:
    p_rt = doc.add_paragraph()
    rr = p_rt.add_run(f'▶  {titulo_r}')
    rr.bold = True
    rr.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    p_rd = doc.add_paragraph(desc_r)
    p_rd.runs[0].font.size = Pt(10)

    p_rc = doc.add_paragraph()
    cr_r = p_rc.add_run(cod_r)
    cr_r.font.name = 'Courier New'
    cr_r.font.size = Pt(9.5)
    cr_r.font.color.rgb = RGBColor(0x17, 0x6B, 0x17)
    doc.add_paragraph('')

add_seccion_titulo('  F) Tipos de Gráficos Aplicados a Fotovoltaica')

# Helper gráfico
def add_grafico(numero, tipo, descripcion, cuando_usarlo, ejemplo_desc, ejemplo_codigo):
    h = doc.add_heading(f'{numero}. Gráfico de {tipo}', level=1)
    h.runs[0].font.size = Pt(13)
    h.runs[0].font.color.rgb = RGBColor(0xD4, 0x7A, 0x00)

    p = doc.add_paragraph()
    r1 = p.add_run('¿Qué muestra? ')
    r1.bold = True
    p.add_run(descripcion)

    p2 = doc.add_paragraph()
    r2 = p2.add_run('¿Cuándo usarlo en FV? ')
    r2.bold = True
    p2.add_run(cuando_usarlo)

    p3 = doc.add_paragraph()
    r3 = p3.add_run('Ejemplo FV:  ')
    r3.bold = True
    p3.add_run(ejemplo_desc)

    p4 = doc.add_paragraph()
    cr4 = p4.add_run(ejemplo_codigo)
    cr4.font.name = 'Courier New'
    cr4.font.size = Pt(9.5)
    cr4.font.color.rgb = RGBColor(0x10, 0x10, 0x60)
    doc.add_paragraph('')


add_grafico(58,
    'Barras — Producción mensual',
    'Compara cantidades entre categorías o períodos. Cada barra representa un valor.',
    'Producción kWh por mes, comparar meses con más/menos sol, consumo vs generación.',
    'Producción mensual de un sistema de 3 kWp con irradiación variable:',
    'import matplotlib\nmatplotlib.use("Agg")\nimport matplotlib.pyplot as plt\n\n'
    'meses = ["Ene","Feb","Mar","Abr","May","Jun",\n'
    '         "Jul","Ago","Sep","Oct","Nov","Dic"]\n'
    'produccion_kwh = [520,480,450,390,340,310,320,360,400,460,500,530]\n\n'
    'plt.figure(figsize=(10, 5))\n'
    'plt.bar(meses, produccion_kwh, color="orange", edgecolor="black")\n'
    'plt.title("Producción mensual del sistema solar — 3 kWp")\n'
    'plt.xlabel("Mes")\n'
    'plt.ylabel("Energía generada (kWh)")\n'
    'plt.grid(axis="y", linestyle="--", alpha=0.5)\n'
    'plt.savefig("produccion_mensual.png", dpi=150, bbox_inches="tight")\n'
    'plt.close()\n'
    'print("Gráfico guardado como produccion_mensual.png")')

add_grafico(59,
    'Línea — Degradación de producción en 25 años',
    'Muestra la evolución de un valor a lo largo del tiempo. Ideal para tendencias.',
    'Producción año a año con degradación, evolución del ahorro acumulado, curva del VAN.',
    'Producción anual de 7 800 kWh en año 1, degradación 0.5% anual — 25 años:',
    'import matplotlib\nmatplotlib.use("Agg")\nimport matplotlib.pyplot as plt\n\n'
    'prod_año1   = 7800\n'
    'degradacion = 0.005\n'
    'años        = list(range(1, 26))\n'
    'produccion  = [round(prod_año1 * ((1-degradacion)**(a-1)), 1) for a in años]\n\n'
    'plt.figure(figsize=(10, 5))\n'
    'plt.plot(años, produccion, marker="o", color="steelblue", linewidth=2)\n'
    'plt.fill_between(años, produccion, alpha=0.15, color="steelblue")\n'
    'plt.title("Producción anual con degradación — 25 años de vida útil")\n'
    'plt.xlabel("Año del sistema")\n'
    'plt.ylabel("Energía generada (kWh/año)")\n'
    'plt.grid(True, linestyle="--", alpha=0.5)\n'
    'plt.savefig("degradacion_25_años.png", dpi=150, bbox_inches="tight")\n'
    'plt.close()\n'
    'print("Gráfico guardado como degradacion_25_años.png")')

add_grafico(60,
    'Torta (Pie) — Pérdidas del sistema (Loss Diagram)',
    'Muestra la proporción de cada parte respecto al total. Ideal para porcentajes.',
    'Desglose del loss diagram: qué porcentaje se pierde en temperatura, suciedad, '
    'sombras, mismatch, cableado e inversor. Réplica visual del gráfico de PVsyst.',
    'Loss diagram de un sistema con PR 82 % — desglose visual de pérdidas:',
    'import matplotlib\nmatplotlib.use("Agg")\nimport matplotlib.pyplot as plt\n\n'
    'conceptos = ["Energía útil (PR 82%)", "Pérd. temperatura",\n'
    '             "Suciedad", "Sombras", "Mismatch",\n'
    '             "Cableado DC", "Inversor"]\n'
    'porcentajes = [82.0, 4.8, 3.0, 2.0, 1.5, 1.5, 3.0, 2.2]\n'
    'porcentajes = [82.0, 4.8, 3.0, 2.0, 1.5, 1.5, 5.2]  # suma 100%\n'
    'colores = ["#2ecc71","#e74c3c","#e67e22","#f39c12","#9b59b6","#3498db","#95a5a6"]\n'
    'explotar = [0.05, 0, 0, 0, 0, 0, 0]  # destaca la porción útil\n\n'
    'plt.figure(figsize=(8, 8))\n'
    'plt.pie(porcentajes, labels=conceptos, colors=colores,\n'
    '        autopct="%1.1f%%", explode=explotar, startangle=140)\n'
    'plt.title("Loss Diagram — Distribución de pérdidas del sistema solar")\n'
    'plt.savefig("loss_diagram.png", dpi=150, bbox_inches="tight")\n'
    'plt.close()\n'
    'print("Gráfico guardado como loss_diagram.png")')

add_grafico(61,
    'Barras dobles — Generación vs Consumo mensual',
    'Compara dos series de datos lado a lado en el mismo gráfico.',
    'Ver mes a mes si el sistema cubre el consumo o hay déficit/excedente. '
    'Equivale al gráfico de balance energético de PVsyst.',
    'Generación solar vs consumo del hogar, mes a mes:',
    'import matplotlib\nmatplotlib.use("Agg")\nimport matplotlib.pyplot as plt\nimport numpy as np\n\n'
    'meses      = ["Ene","Feb","Mar","Abr","May","Jun",\n'
    '              "Jul","Ago","Sep","Oct","Nov","Dic"]\n'
    'generacion = [520, 480, 450, 390, 340, 310, 320, 360, 400, 460, 500, 530]\n'
    'consumo    = [400, 380, 420, 440, 460, 500, 490, 470, 430, 410, 390, 420]\n\n'
    'x      = np.arange(len(meses))\n'
    'ancho  = 0.35\n\n'
    'plt.figure(figsize=(12, 5))\n'
    'plt.bar(x - ancho/2, generacion, ancho, label="Generación FV", color="orange")\n'
    'plt.bar(x + ancho/2, consumo,    ancho, label="Consumo hogar", color="steelblue")\n'
    'plt.xticks(x, meses)\n'
    'plt.title("Balance energético mensual — Generación vs Consumo")\n'
    'plt.xlabel("Mes")\n'
    'plt.ylabel("Energía (kWh)")\n'
    'plt.legend()\n'
    'plt.grid(axis="y", linestyle="--", alpha=0.5)\n'
    'plt.savefig("balance_energetico.png", dpi=150, bbox_inches="tight")\n'
    'plt.close()\n'
    'print("Gráfico guardado como balance_energetico.png")')

add_grafico(62,
    'Línea doble — HSP mensual e Irradiancia',
    'Muestra dos variables relacionadas sobre el mismo eje temporal para compararlas.',
    'Comparar la HSP mensual con la irradiancia media, ver correlación entre temperatura '
    'ambiente y eficiencia del panel a lo largo del año.',
    'HSP mensual e irradiancia media mensual — recurso solar del lugar:',
    'import matplotlib\nmatplotlib.use("Agg")\nimport matplotlib.pyplot as plt\n\n'
    'meses       = ["Ene","Feb","Mar","Abr","May","Jun",\n'
    '               "Jul","Ago","Sep","Oct","Nov","Dic"]\n'
    'hsp_mensual = [6.2, 5.9, 5.5, 4.8, 4.1, 3.8, 4.0, 4.5, 5.0, 5.6, 6.0, 6.3]\n'
    'irr_media   = [6200, 5900, 5500, 4800, 4100, 3800,\n'
    '               4000, 4500, 5000, 5600, 6000, 6300]  # Wh/m²/día\n\n'
    'fig, ax1 = plt.subplots(figsize=(11, 5))\n\n'
    '# Eje izquierdo: HSP\n'
    'ax1.plot(meses, hsp_mensual, color="orange", marker="o", linewidth=2, label="HSP (h/día)")\n'
    'ax1.set_ylabel("HSP (h/día)", color="orange")\n'
    'ax1.tick_params(axis="y", labelcolor="orange")\n\n'
    '# Eje derecho: Irradiancia\n'
    'ax2 = ax1.twinx()\n'
    'ax2.plot(meses, irr_media, color="steelblue", marker="s",\n'
    '         linewidth=2, linestyle="--", label="Irradiancia (Wh/m²/día)")\n'
    'ax2.set_ylabel("Irradiancia (Wh/m²/día)", color="steelblue")\n'
    'ax2.tick_params(axis="y", labelcolor="steelblue")\n\n'
    'ax1.set_title("Recurso solar mensual — HSP e Irradiancia")\n'
    'ax1.set_xlabel("Mes")\n'
    'ax1.grid(True, linestyle="--", alpha=0.4)\n'
    'plt.savefig("hsp_irradiancia.png", dpi=150, bbox_inches="tight")\n'
    'plt.close()\n'
    'print("Gráfico guardado como hsp_irradiancia.png")')

add_grafico(63,
    'Línea acumulada — Flujo de caja y Payback',
    'Muestra cómo se acumula el ahorro año a año hasta recuperar la inversión inicial.',
    'Visualizar el punto exacto de payback: dónde la línea cruza el eje cero. '
    'Equivale al gráfico de flujo de caja acumulado del análisis económico de PVsyst.',
    'Inversión $6 000 000 recuperada con ahorro anual creciente al 3%:',
    'import matplotlib\nmatplotlib.use("Agg")\nimport matplotlib.pyplot as plt\n\n'
    'inversion    = 6000000\n'
    'ahorro_año1  = 900000\n'
    'crecimiento  = 0.03\n'
    'años         = list(range(0, 26))\n\n'
    '# Flujo acumulado: empieza negativo (inversión) y sube con cada ahorro\n'
    'flujo_acum = [-inversion]\n'
    'for n in range(1, 26):\n'
    '    ahorro_n = ahorro_año1 * ((1 + crecimiento) ** (n-1))\n'
    '    flujo_acum.append(flujo_acum[-1] + ahorro_n)\n\n'
    'colores_barras = ["#e74c3c" if v < 0 else "#2ecc71" for v in flujo_acum]\n\n'
    'plt.figure(figsize=(11, 5))\n'
    'plt.bar(años, flujo_acum, color=colores_barras, edgecolor="gray", linewidth=0.5)\n'
    'plt.axhline(0, color="black", linewidth=1.2, linestyle="--")\n'
    'plt.title("Flujo de caja acumulado — Punto de Payback del sistema solar")\n'
    'plt.xlabel("Año")\n'
    'plt.ylabel("Flujo acumulado ($)")\n'
    'plt.grid(axis="y", linestyle="--", alpha=0.4)\n'
    'plt.savefig("payback_flujo.png", dpi=150, bbox_inches="tight")\n'
    'plt.close()\n'
    'print("Gráfico guardado como payback_flujo.png")')

add_grafico(64,
    'Subplots — Múltiples gráficos en una sola imagen',
    'Organiza varios gráficos en una cuadrícula dentro de una sola imagen. '
    'Permite presentar un informe visual completo en un solo archivo.',
    'Generar un panel de 4 gráficos en una imagen: producción, pérdidas, '
    'payback y CO₂. Equivale a la hoja de resumen ejecutivo de PVsyst.',
    'Panel de 4 resultados principales del proyecto en una sola imagen:',
    'import matplotlib\nmatplotlib.use("Agg")\nimport matplotlib.pyplot as plt\n\n'
    '# ── DATOS ──\n'
    'meses       = ["E","F","M","A","M","J","J","A","S","O","N","D"]\n'
    'produccion  = [520,480,450,390,340,310,320,360,400,460,500,530]\n'
    'perdidas    = [82.0, 4.8, 3.0, 2.0, 1.5, 1.5, 5.2]\n'
    'etiq_perd   = ["Útil","Temp","Suciedad","Sombra","Mismatch","Cable","Inversor"]\n'
    'años        = list(range(1, 26))\n'
    'prod_anual  = [7800 * (0.995**(a-1)) for a in años]\n'
    'co2_acum    = [round(p * 0.294 / 1000, 2) for p in prod_anual]\n\n'
    '# ── FIGURA CON 4 SUBGRÁFICOS ──\n'
    'fig, ((ax1, ax2), (ax3, ax4)) = plt.subplots(2, 2, figsize=(14, 9))\n'
    'fig.suptitle("Resumen ejecutivo — Sistema solar 3 kWp", fontsize=14, fontweight="bold")\n\n'
    '# Gráfico 1: Producción mensual\n'
    'ax1.bar(meses, produccion, color="orange", edgecolor="black")\n'
    'ax1.set_title("Producción mensual (kWh)")\n'
    'ax1.set_ylabel("kWh")\n'
    'ax1.grid(axis="y", linestyle="--", alpha=0.5)\n\n'
    '# Gráfico 2: Loss diagram\n'
    'ax2.pie(perdidas, labels=etiq_perd, autopct="%1.0f%%", startangle=140,\n'
    '        colors=["#2ecc71","#e74c3c","#e67e22","#f39c12","#9b59b6","#3498db","#95a5a6"])\n'
    'ax2.set_title("Loss Diagram (%)")\n\n'
    '# Gráfico 3: Degradación 25 años\n'
    'ax3.plot(años, prod_anual, color="steelblue", marker=".", linewidth=1.8)\n'
    'ax3.fill_between(años, prod_anual, alpha=0.12, color="steelblue")\n'
    'ax3.set_title("Producción anual con degradación (kWh)")\n'
    'ax3.set_xlabel("Año")\n'
    'ax3.grid(True, linestyle="--", alpha=0.5)\n\n'
    '# Gráfico 4: CO₂ evitado acumulado\n'
    'ax4.plot(años, [sum(co2_acum[:i+1]) for i in range(25)],\n'
    '         color="#27ae60", marker=".", linewidth=1.8)\n'
    'ax4.set_title("CO₂ evitado acumulado (ton)")\n'
    'ax4.set_xlabel("Año")\n'
    'ax4.grid(True, linestyle="--", alpha=0.5)\n\n'
    'plt.tight_layout()\n'
    'plt.savefig("resumen_ejecutivo.png", dpi=150, bbox_inches="tight")\n'
    'plt.close()\n'
    'print("Panel de 4 gráficos guardado como resumen_ejecutivo.png")')

# ── ESTRUCTURA CORRECTA DE ARCHIVO ──
add_seccion_titulo('REGLA DE ORO — Estructura correcta de un archivo Python con cálculos Y gráficos')

p_oro = doc.add_paragraph(
    'La causa más común de errores al mezclar cálculos con gráficos es el ORDEN de '
    'las instrucciones. La estructura correcta es siempre la misma: primero los '
    'imports, luego los cálculos, luego los gráficos. Nunca al revés ni mezclados.'
)
p_oro.runs[0].font.size = Pt(10)
p_oro.runs[0].italic = True
doc.add_paragraph('')

codigo_estructura = (
    '# ══════════════════════════════════════════════════════════\n'
    '# ESTRUCTURA CORRECTA — Archivo Python con cálculos y gráficos\n'
    '# ══════════════════════════════════════════════════════════\n\n'
    '# ── BLOQUE 1: IMPORTS (siempre al inicio, nunca en medio del código) ──\n'
    'import math\n'
    'import matplotlib\n'
    'matplotlib.use("Agg")           # evita ventanas emergentes\n'
    'import matplotlib.pyplot as plt\n'
    'import numpy as np              # solo si usas barras dobles o arrays\n\n'
    '# ── BLOQUE 2: CÁLCULOS (variables, fórmulas, conversiones) ──\n'
    'potencia_w   = 400\n'
    'n_paneles    = 12\n'
    'hsp          = 5.2\n'
    'energia_kwh  = potencia_w * n_paneles * hsp / 1000\n'
    'pr           = 0.80\n'
    'energia_real = round(energia_kwh * pr, 2)\n'
    'print("Energía real:", energia_real, "kWh/día")\n\n'
    '# ── BLOQUE 3: GRÁFICOS (siempre al final, después de todos los cálculos) ──\n'
    'plt.close("all")                # limpia lienzos residuales\n\n'
    'plt.figure(figsize=(8, 4))\n'
    'plt.bar(["Energía ideal", "Energía real"],\n'
    '        [energia_kwh, energia_real],\n'
    '        color=["steelblue", "orange"])\n'
    'plt.title("Comparación: energía ideal vs real")\n'
    'plt.ylabel("kWh/día")\n'
    'plt.grid(axis="y", linestyle="--", alpha=0.5)\n'
    'plt.savefig("comparacion_energia.png", dpi=150, bbox_inches="tight")\n'
    'plt.close()                     # cierra siempre al terminar cada gráfico\n'
    'print("Gráfico guardado correctamente.")\n\n'
    '# ══ RESUMEN DE REGLAS ═══════════════════════════════════════\n'
    '# 1. import matplotlib + matplotlib.use("Agg") → primeras líneas del archivo\n'
    '# 2. Todos los cálculos ANTES de los gráficos\n'
    '# 3. plt.figure()  →  código del gráfico  →  plt.savefig()  →  plt.close()\n'
    '# 4. NUNCA plt.show() en scripts — usa siempre plt.savefig()\n'
    '# 5. plt.close("all") al inicio del bloque de gráficos para limpiar\n'
    '# ════════════════════════════════════════════════════════════\n'
)

p_est = doc.add_paragraph()
cr_est = p_est.add_run(codigo_estructura)
cr_est.font.name = 'Courier New'
cr_est.font.size = Pt(9)
cr_est.font.color.rgb = RGBColor(0x10, 0x10, 0x60)

doc.add_paragraph('')

# Pie de página
footer = doc.add_paragraph('Elaborado para curso de Python aplicado a Energía Solar Fotovoltaica — 2026')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.runs[0].font.size = Pt(9)
footer.runs[0].font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
footer.runs[0].italic = True

doc.save('Adenda_Python_Fotovoltaico.docx')
print("Documento creado correctamente.")
