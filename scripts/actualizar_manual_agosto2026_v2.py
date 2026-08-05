"""
Actualiza entregables/MANUAL_CALCULADORA_BIPV_v3_agosto2026.docx con las
últimas tareas aplicadas (5-ago-2026): agrivoltaica (factor de ocupación,
GCR sincronizado, Vista 3D granja), Motor IV automático, precio real del
inversor en Financiero y detección Ns half-cut. Mismo formato del manual.

Ejecutar:  python3 scripts/actualizar_manual_agosto2026_v2.py
"""
import copy, shutil
from docx import Document
from docx.shared import RGBColor

SRC = "entregables/MANUAL_CALCULADORA_BIPV_v3_agosto2026.docx"
shutil.copy(SRC, SRC + ".bak2")
doc = Document(SRC)

SEP = "────────────────────────────────────────────────────────────"
AZUL = RGBColor(0x1A, 0x53, 0x76)

def find_para(pred):
    for p in doc.paragraphs:
        if pred(p.text.strip()):
            return p
    raise RuntimeError("párrafo no encontrado")

# ── 1. TOC: nueva entrada después de "Preguntas frecuentes" ────────────────
toc_faq = find_para(lambda t: t == "Preguntas frecuentes")
new_toc = copy.deepcopy(toc_faq._p)
toc_faq._p.addnext(new_toc)
from docx.text.paragraph import Paragraph
p_toc = Paragraph(new_toc, toc_faq._parent)
for r in list(p_toc.runs):
    r._r.getparent().remove(r._r)
r = p_toc.add_run("Anexo — Actualizaciones 5 de agosto 2026 (agrivoltaica y Motor IV)  ")
rn = p_toc.add_run("NUEVO")
rn.bold = True
rn.font.color.rgb = RGBColor(0xE6, 0x51, 0x00)

# ── 2. Anexo: insertar antes del bloque de cierre ("Manual actualizado…") ──
cierre = find_para(lambda t: t.startswith("Manual actualizado el"))
prev = cierre._p.getprevious()
if prev is not None and prev.tag.endswith("}p") and Paragraph(prev, cierre._parent).text.strip() == SEP:
    anchor = Paragraph(prev, cierre._parent)
else:
    anchor = cierre

def before(text="", style=None, bold=False, color=None, size=None):
    p = anchor.insert_paragraph_before(text, style=style)
    if p.runs:
        r = p.runs[0]
        r.bold = bold
        if color: r.font.color.rgb = color
    return p

def sep():          before(SEP)
def h2(t):
    p = before(t, style="Heading 2")
    if p.runs: p.runs[0].font.color.rgb = AZUL
def sub(t):         before(t, bold=True)
def body(t):        before(t)
def bullet(t):      before(t, style="List Bullet")
def warn(t):
    p = before(t)
    if p.runs:
        p.runs[0].bold = True
        p.runs[0].font.color.rgb = RGBColor(0xE6, 0x51, 0x00)
def tip(t):
    p = before(t)
    if p.runs: p.runs[0].font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

sep()
h2("18. Anexo — Actualizaciones del 5 de agosto de 2026")
body("Esta entrega incorpora el soporte completo para proyectos agrivoltaicos "
     "(cultivo bajo paneles), la activación automática del Motor IV y el uso "
     "del precio real del inversor en el análisis financiero.")

sep()
sub("18.1 Página 1 — Proyecto: Factor de ocupación con paneles (agrivoltaica)  ACTUALIZADO")
body("Nuevo campo \"Factor de ocupación con paneles (%)\" (rango 5–100, por defecto 100). "
     "Define qué porcentaje del terreno bruto queda cubierto por paneles; el resto queda "
     "libre para el cultivo. La app calcula el área útil = área × factor/100 y la usa en "
     "Dimensionamiento, Presupuesto y Reporte PDF.")
bullet("Ejemplo: terreno de 3 000 m² con factor 30% → 900 m² útiles para paneles "
       "(~340 paneles en vez de ~1 040) y 2 100 m² libres para el cultivo.")
bullet("El factor se guarda en el proyecto (proyecto_actual.json) y se restaura al recargarlo.")
warn("⚠️ Para no cometer errores: si tu proyecto es una granja con cultivo, ajusta el factor "
     "ANTES de pasar a Dimensionamiento. Si lo dejas en 100%, el conteo de paneles, el "
     "presupuesto y el USD/m² se calcularán como si todo el terreno llevara paneles.")

sep()
sub("18.2 Página 2 — Recurso Solar: GCR sincronizado con el factor de ocupación  ACTUALIZADO")
body("El control \"GCR (cobertura del suelo)\" del modelo bifacial ahora arranca con el mismo "
     "valor que el factor de ocupación definido en Proyecto (factor 30% → GCR 0.30), porque "
     "ambos representan la misma fracción de suelo cubierta por paneles. La sincronización "
     "solo aplica la primera vez; después se respeta el valor que guardes aquí.")
bullet("Si el GCR y el factor difieren en más de 15 puntos, aparece una alerta con el valor "
       "sugerido — el sombreado entre filas y el conteo de paneles estarían usando supuestos distintos.")
warn("⚠️ Para no cometer errores: no subas el GCR \"para producir más\" sin cambiar también el "
     "factor de ocupación en Proyecto. Un GCR alto junta las filas (más sombra mutua y menos "
     "luz al cultivo); mantén ambos valores alineados.")

sep()
sub("18.3 Página 9 — Vista 3D: modo Granja agrivoltaica  NUEVO")
body("Cuando el tipo de instalación es \"Granja fotovoltaica\", la Vista 3D ya no dibuja un "
     "edificio: muestra el terreno completo en verde (el cultivo) con las filas de paneles "
     "inclinadas y elevadas a 3 m cuando el factor de ocupación es menor a 100%.")
bullet("La separación entre filas se calcula como ancho del colector ÷ GCR "
       "(con 30% de ocupación, ~3.7 m entre ejes de filas).")
bullet("Un resumen indica número de filas, separación, altura de montaje y porcentaje "
       "de suelo libre para el cultivo.")
tip("Nota: la vista muestra las filas como bandas continuas, no módulos individuales. "
    "El conteo oficial de paneles sigue siendo el de la Página 4 — Dimensionamiento.")

sep()
sub("18.4 Página 7 — Financiero: precio real del inversor del catálogo  ACTUALIZADO")
body("El precio del inversor seleccionado en el catálogo fluye automáticamente al análisis "
     "financiero, en lugar de usar un estimado genérico por kW.")
warn("⚠️ Para no cometer errores: verifica que la ficha del inversor en el catálogo tenga "
     "precio cargado. Si el campo está vacío, revisa el valor que aparece en Financiero "
     "antes de generar el presupuesto.")

sep()
sub("18.5 Página 3 — Motor IV: activación automática y defensa Ns half-cut  ACTUALIZADO")
body("Si el panel seleccionado tiene su ficha técnica completa (Voc, Isc, Vmp, Imp, Ns y "
     "coeficientes térmicos), el Motor IV se activa automáticamente sin configuración manual.")
bullet("Nueva defensa para paneles half-cut: el extractor de fichas infiere el número de "
       "celdas en serie (Ns) desde el conteo de semiceldas y lo verifica contra el Voc "
       "(regla práctica: Ns ≈ Voc / 0.74 en paneles HJT).")
bullet("Si el Ns auto-extraído es inconsistente, el Motor IV lo corrige al vuelo y el "
       "validador físico lo marca en el diagnóstico.")
warn("⚠️ Para no cometer errores: en paneles half-cut la ficha suele reportar el TOTAL de "
     "semiceldas (p. ej. 144); el Ns eléctrico es la mitad (72) porque hay dos strings en "
     "paralelo. Si cargas un panel a mano, verifica el Ns con la regla Voc/0.74.")

sep()
sub("18.6 Flujo recomendado para proyectos agrivoltaicos")
body("1 Proyecto (tipo Granja fotovoltaica + factor de ocupación) → 2 Recurso Solar "
     "(verificar GCR sincronizado → Calcular POA) → 4 Dimensionamiento (área útil) → "
     "9 Vista 3D (verificación visual de filas y cultivo) → 6 Producción → 7 Financiero → "
     "8 Presupuesto → 10 Reporte PDF.")
tip("Regla de oro agrivoltaica: factor de ocupación (Proyecto) = GCR (Recurso Solar). "
    "Si cambias uno, revisa el otro.")

# ── 3. Pie: fecha y novedades ──────────────────────────────────────────────
cierre.runs[0].text = "Manual actualizado el 5 de agosto de 2026"
for p in doc.paragraphs:
    if p.text.startswith("Novedades de esta versión:"):
        p.runs[0].text = ("Novedades de esta versión: modo agrivoltaico completo (factor de "
                          "ocupación, GCR sincronizado, Vista 3D de granja), Motor IV "
                          "automático con ficha completa, precio real del inversor en "
                          "Financiero y defensa Ns half-cut.")
        for r in p.runs[1:]:
            r.text = ""
        break

doc.save(SRC)
print("OK — manual actualizado:", SRC)
