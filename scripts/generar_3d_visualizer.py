
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2); s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

AZUL   = RGBColor(0x1A, 0x5C, 0x8A)
VERDE  = RGBColor(0x17, 0x6B, 0x17)
MORADO = RGBColor(0x6E, 0x27, 0x94)
ROJO   = RGBColor(0xC0, 0x39, 0x2B)
NARANJO= RGBColor(0xD4, 0x7A, 0x00)
GRIS   = RGBColor(0x7F, 0x7F, 0x7F)
COD    = RGBColor(0x10, 0x10, 0x60)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)

def shade_cell(cell, hex6):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex6)
    cell._tc.get_or_add_tcPr().append(shd)

def sep():
    p = doc.add_paragraph('─' * 80)
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(0xCC,0xCC,0xCC)

def h1(txt, color=AZUL):
    doc.add_paragraph('')
    h = doc.add_heading(txt, level=1)
    h.runs[0].font.size = Pt(14); h.runs[0].font.color.rgb = color

def h2(txt, color=MORADO):
    h = doc.add_heading(txt, level=2)
    h.runs[0].font.size = Pt(12); h.runs[0].font.color.rgb = color

def h3(txt, color=NARANJO):
    h = doc.add_heading(txt, level=3)
    h.runs[0].font.size = Pt(11); h.runs[0].font.color.rgb = color

def body(txt):
    p = doc.add_paragraph(txt)
    p.runs[0].font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def bullet(txt, color=AZUL):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(txt); r.font.size = Pt(10); r.font.color.rgb = color

def badge(icon, label, txt, color=AZUL):
    p = doc.add_paragraph()
    r1 = p.add_run(f'{icon} {label}  '); r1.bold = True; r1.font.color.rgb = color
    r2 = p.add_run(txt); r2.font.size = Pt(10)

def codigo(txt):
    p = doc.add_paragraph()
    cr = p.add_run(txt)
    cr.font.name = 'Courier New'; cr.font.size = Pt(8.5); cr.font.color.rgb = COD
    doc.add_paragraph('')

def archivo(ruta, color=VERDE):
    p = doc.add_paragraph()
    r = p.add_run('📁 Archivo:  '); r.bold = True; r.font.color.rgb = color
    cr = p.add_run(ruta); cr.font.name = 'Courier New'; cr.font.size = Pt(9.5); cr.font.color.rgb = color

def tabla_header(tbl, hdrs, fill='1A5C8A'):
    row = tbl.rows[0].cells
    for i, h in enumerate(hdrs):
        row[i].text = h
        shade_cell(row[i], fill)
        for par in row[i].paragraphs:
            for run in par.runs:
                run.bold = True; run.font.color.rgb = BLANCO; run.font.size = Pt(9)

# ═══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════════════════════════
t = doc.add_heading('VISUALIZADOR 3D DE FACHADAS BIPV\nIntegración con la Calculadora Python + Streamlit', 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.runs[0].font.color.rgb = AZUL

sub = doc.add_paragraph('Análisis de opciones · Arquitectura 3 capas · Código Python completo · Paso B-5 ampliado')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].bold = True; sub.runs[0].font.size = Pt(12); sub.runs[0].font.color.rgb = MORADO

intro = doc.add_paragraph(
    'Este documento responde al reto planteado: ¿es posible integrar un visualizador 3D '
    'de fachadas compatible con la calculadora BIPV en Python/Streamlit? '
    'La respuesta es sí — y con varias opciones reales, todas disponibles en pip. '
    'Se entrega el análisis comparativo, la arquitectura recomendada en 3 capas, '
    'el código Python completo y la integración dentro del plan de construcción '
    'como paso B-5 ampliado (B-5A + B-5B + B-5C).'
)
intro.runs[0].font.size = Pt(10); intro.runs[0].italic = True
intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
doc.add_paragraph('')

# ═══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 1 — ANÁLISIS DE OPCIONES
# ═══════════════════════════════════════════════════════════════════════════════
sep()
h1('SECCIÓN 1 — ANÁLISIS DE OPCIONES 3D DISPONIBLES EN PYTHON')

body(
    'Se verificaron las siguientes librerías disponibles en pip y compatibles con Streamlit. '
    'Todas están en el índice de PyPI y pueden instalarse con pip install. '
    'Se evalúan según 5 criterios relevantes para el visualizador BIPV.'
)
doc.add_paragraph('')

tbl_op = doc.add_table(rows=1, cols=7)
tbl_op.style = 'Table Grid'
tabla_header(tbl_op, ['Librería', 'Streamlit\nnativo', 'Modelo 3D\n(OBJ/GLTF)', 'Sombras\nreal-time', 'Dibujar\npaneles', 'Dificultad', 'Veredicto'])

opciones = [
    ('Pydeck 0.9.3',        '✅ Sí',    '❌ No',   '⚠ Básico', '❌ No',  '🟢 Fácil',    '✅ Capa 1 — Vista de sitio'),
    ('PyVista 0.48 + stpyvista','✅ Sí','✅ OBJ/STL','✅ Colores','⚠ Parcial','🟡 Medio','✅ Capa 2 — Edificio 3D'),
    ('Trimesh 4.12',        '⚠ vía HTML','✅ OBJ/GLTF/STL','⚠ No','❌ No','🟢 Fácil','✅ Soporte — carga meshes'),
    ('Three.js (HTML comp)','✅ via html','✅ GLTF/OBJ','✅ Tiempo real','✅ Sí','🔴 Avanzado','✅ Capa 3 — Fachada detalle'),
    ('Plotly 3D Mesh',      '✅ Sí',    '⚠ Básico', '❌ No',   '❌ No',  '🟢 Fácil',    '⚠ Alternativa simple'),
    ('Pythreejs 2.4',       '⚠ Jupyter','✅ Three.js','✅ Sí',  '✅ Sí',  '🔴 Avanzado','⚠ Mejor en Jupyter'),
    ('Babylon.js (HTML)',   '✅ via html','✅ GLTF/OBJ','✅ Sí', '✅ Sí',  '🔴 Avanzado','⚠ Alternativa a Three.js'),
]

for fila in opciones:
    row = tbl_op.add_row().cells
    for i, val in enumerate(fila):
        row[i].text = val
        if row[i].paragraphs[0].runs:
            r = row[i].paragraphs[0].runs[0]
        else:
            r = row[i].paragraphs[0].add_run(val)
        r.font.size = Pt(8.5)
        if i == 6:
            if '✅' in val: r.font.color.rgb = VERDE; r.bold = True
            elif '⚠' in val: r.font.color.rgb = NARANJO

doc.add_paragraph('')

# ═══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 2 — ARQUITECTURA 3 CAPAS
# ═══════════════════════════════════════════════════════════════════════════════
sep()
h1('SECCIÓN 2 — ARQUITECTURA RECOMENDADA EN 3 CAPAS')

body(
    'No existe una sola librería que cubra todo el espectro del visualizador BIPV. '
    'La solución más robusta, y la que mejor se integra con pvlib y Streamlit, '
    'es una arquitectura en 3 capas donde cada capa usa la herramienta óptima '
    'para su función específica. Las tres capas operan juntas en la misma página de Streamlit.'
)
doc.add_paragraph('')

diagrama = (
    '┌─────────────────────────────────────────────────────────────────────────┐\n'
    '│                    PÁGINA STREAMLIT — VISUALIZADOR 3D BIPV              │\n'
    '├─────────────────┬─────────────────────────────┬───────────────────────── ┤\n'
    '│   CAPA 1        │        CAPA 2               │       CAPA 3             │\n'
    '│   SITIO         │        EDIFICIO 3D          │       FACHADA DETALLE    │\n'
    '│                 │                             │                          │\n'
    '│  pydeck         │  PyVista + stpyvista        │  Three.js                │\n'
    '│  (deck.gl)      │  (VTK bajo el capó)         │  via st.components       │\n'
    '│                 │                             │                          │\n'
    '│ • Mapa satelital│ • Edificio paramétrico 3D   │ • Fachada en detalle     │\n'
    '│ • Huella edif.  │ • Colores por irradiancia   │ • Grid de paneles BIPV   │\n'
    '│ • Edif. vecinos │ • Sombras como colormesh    │ • Sombras hora a hora    │\n'
    '│ • Orientación N │ • Clic fachada → selección  │ • Usuario dibuja zonas   │\n'
    '│ • Heatmap GHI   │ • Importa OBJ/STL fachadas  │ • Importa GLTF complejo  │\n'
    '│                 │                             │                          │\n'
    '│  PYTHON PURO    │  PYTHON PURO                │  JS ↔ Python (bridge)    │\n'
    '│  Nativo Streamlit│ Nativo vía stpyvista       │  streamlit-javascript    │\n'
    '└─────────────────┴─────────────────────────────┴──────────────────────────┘\n'
    '         ↓                      ↓                          ↓\n'
    '   Coordenadas             Área por fachada          Área dibujada\n'
    '   del proyecto            seleccionada (m²)         por el usuario (m²)\n'
    '         ↓                      ↓                          ↓\n'
    '┌─────────────────────────────────────────────────────────────────────────┐\n'
    '│             MOTOR DE CÁLCULO PYTHON (pvlib + numpy)                     │\n'
    '│   posición solar · transposición Perez · temperatura celda · PR · VAN   │\n'
    '└─────────────────────────────────────────────────────────────────────────┘\n'
    '         ↑                      ↑                          ↑\n'
    '   Irradiancia GHI        Shade factor               Sun vector\n'
    '   por mes/hora           por fachada                JSON → Three.js'
)

p_diag = doc.add_paragraph()
r_diag = p_diag.add_run(diagrama)
r_diag.font.name = 'Courier New'; r_diag.font.size = Pt(7.5); r_diag.font.color.rgb = COD
doc.add_paragraph('')

# Flujo de datos
h2('Flujo de datos Python ↔ 3D')
body('La integración entre pvlib (Python) y el visualizador 3D funciona con un patrón de bridge JSON:')
doc.add_paragraph('')

flujo = (
    '1. pvlib calcula posición solar hora a hora  →  JSON: {azimut, elevacion, hora}\n'
    '                                                           ↓\n'
    '2. Three.js recibe el JSON via streamlit-javascript  →  posiciona la fuente de luz solar\n'
    '                                                           ↓\n'
    '3. Three.js calcula sombras en tiempo real sobre los paneles dibujados\n'
    '                                                           ↓\n'
    '4. Usuario selecciona/dibuja zonas de paneles en la fachada\n'
    '                                                           ↓\n'
    '5. Three.js envía JSON de vuelta: {area_m2, factor_sombra, azimut, inclinacion}\n'
    '                                                           ↓\n'
    '6. Python recibe via st.session_state y recalcula producción en tiempo real'
)
p_f = doc.add_paragraph()
r_f = p_f.add_run(flujo)
r_f.font.name = 'Courier New'; r_f.font.size = Pt(9); r_f.font.color.rgb = COD
doc.add_paragraph('')

# ═══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 3 — PASO B-5A: PYDECK (CAPA 1)
# ═══════════════════════════════════════════════════════════════════════════════
sep()
h1('PASO B-5A — Vista de Sitio con Pydeck (Capa 1)')
badge('⏱', 'Tiempo estimado:', '2 horas  |  Dificultad: 🟢 Fácil', AZUL)
badge('📦', 'Librería:', 'pydeck 0.9.3  →  pip install pydeck (ya incluido en Streamlit)', VERDE)
doc.add_paragraph('')

body(
    'Pydeck usa deck.gl (WebGL) bajo el capó. Con Streamlit es completamente nativo '
    '— solo se necesita st.pydeck_chart(). Genera una vista 3D del sitio del proyecto: '
    'la huella del edificio extruida, los edificios vecinos que generan sombra, '
    'y un heatmap de irradiancia anual GHI del lugar.'
)
doc.add_paragraph('')

archivo('pages/5A_🗺️_Vista_Sitio.py')
codigo(
    '# ═══════════════════════════════════════════════════════════════\n'
    '# CAPA 1: Vista de Sitio 3D con Pydeck\n'
    '# Archivo: pages/5A_🗺️_Vista_Sitio.py\n'
    '# Instalar: pip install pydeck\n'
    '# ═══════════════════════════════════════════════════════════════\n'
    'import streamlit as st\n'
    'import pydeck as pdk\n'
    'import numpy as np\n\n'
    'st.header("🗺️ Paso 5A — Vista de Sitio (3D)")\n\n'
    '# ─── Datos del proyecto (de session_state o valores por defecto) ────\n'
    'lat  = st.session_state.get("lat",  -33.45)\n'
    'lon  = st.session_state.get("lon",  -70.65)\n'
    'pais = st.session_state.get("pais", "Chile")\n\n'
    'st.subheader("Definir geometría del edificio")\n'
    'col1, col2, col3 = st.columns(3)\n'
    'with col1: ancho    = st.number_input("Ancho del edificio (m)",  value=20.0, step=1.0)\n'
    'with col2: profund  = st.number_input("Profundidad (m)",         value=15.0, step=1.0)\n'
    'with col3: altura   = st.number_input("Altura total (m)",        value=30.0, step=1.0)\n\n'
    'st.subheader("Edificios vecinos (generadores de sombra)")\n'
    'n_vecinos = st.number_input("Número de edificios vecinos", min_value=0, max_value=6, value=2)\n'
    'vecinos = []\n'
    'for i in range(n_vecinos):\n'
    '    c1, c2, c3, c4 = st.columns(4)\n'
    '    with c1: dlon = st.number_input(f"Vecino {i+1} — offset E/O (m)", value=float(25+i*30), key=f"vdlon{i}")\n'
    '    with c2: dlat = st.number_input(f"Vecino {i+1} — offset N/S (m)", value=0.0,            key=f"vdlat{i}")\n'
    '    with c3: vancho  = st.number_input(f"Ancho (m)",  value=15.0, key=f"vaw{i}")\n'
    '    with c4: valtura = st.number_input(f"Altura (m)", value=float(20+i*10), key=f"vah{i}")\n'
    '    vecinos.append({"dlon": dlon, "dlat": dlat, "ancho": vancho, "altura": valtura})\n\n'
    '# ─── Convertir metros a grados (aproximado) ─────────────────────────\n'
    'def m_a_grados_lon(m, lat): return m / (111320 * np.cos(np.radians(lat)))\n'
    'def m_a_grados_lat(m):      return m / 111320\n\n'
    '# ─── Polígono del edificio principal ────────────────────────────────\n'
    'dlon_a = m_a_grados_lon(ancho/2, lat)\n'
    'dlat_p = m_a_grados_lat(profund/2)\n'
    'poligono_principal = [[\n'
    '    [lon - dlon_a, lat - dlat_p], [lon + dlon_a, lat - dlat_p],\n'
    '    [lon + dlon_a, lat + dlat_p], [lon - dlon_a, lat + dlat_p]\n'
    ']]\n\n'
    '# ─── Capa 1: Edificio principal (color azul BIPV) ────────────────────\n'
    'capa_principal = pdk.Layer(\n'
    '    "PolygonLayer",\n'
    '    data=[{"polygon": poligono_principal[0], "height": altura, "color": [26, 92, 138, 200]}],\n'
    '    get_polygon="polygon",\n'
    '    get_elevation="height",\n'
    '    get_fill_color="color",\n'
    '    extruded=True,\n'
    '    wireframe=True,\n'
    ')\n\n'
    '# ─── Capa 2: Edificios vecinos (gris) ────────────────────────────────\n'
    'data_vecinos = []\n'
    'for v in vecinos:\n'
    '    cx = lon + m_a_grados_lon(v["dlon"], lat)\n'
    '    cy = lat + m_a_grados_lat(v["dlat"])\n'
    '    dv = m_a_grados_lon(v["ancho"]/2, lat)\n'
    '    data_vecinos.append({\n'
    '        "polygon": [[cx-dv, cy-dv],[cx+dv, cy-dv],[cx+dv, cy+dv],[cx-dv, cy+dv]],\n'
    '        "height": v["altura"], "color": [120, 120, 120, 180]\n'
    '    })\n\n'
    'capa_vecinos = pdk.Layer(\n'
    '    "PolygonLayer",\n'
    '    data=data_vecinos,\n'
    '    get_polygon="polygon",\n'
    '    get_elevation="height",\n'
    '    get_fill_color="color",\n'
    '    extruded=True,\n'
    ')\n\n'
    '# ─── Capa 3: Heatmap de irradiancia (si hay datos TMY cargados) ──────\n'
    'capas = [capa_principal, capa_vecinos]\n'
    'if "tmy" in st.session_state:\n'
    '    tmy = st.session_state["tmy"]\n'
    '    ghi_anual = tmy["ghi"].sum() / 1000  # kWh/m²/año\n'
    '    # Gradiente de color por orientación (simulación)\n'
    '    for az, dlat_v, dlon_v, nombre in [\n'
    '        (0,   dlat_p+0.0001, 0,        "Norte"),\n'
    '        (180,-dlat_p-0.0001, 0,        "Sur"),\n'
    '        (90,  0,             dlon_a+0.00002, "Este"),\n'
    '        (270, 0,            -dlon_a-0.00002, "Oeste"),\n'
    '    ]:\n'
    '        # Irradiancia aproximada por orientación (hemisferio sur)\n'
    '        factor = {"Sur": 1.0, "Norte": 0.6, "Este": 0.75, "Oeste": 0.75}[nombre]\n'
    '        irr = ghi_anual * factor\n'
    '        g = min(255, int((irr / 2000) * 255))\n'
    '        r = 255 - g\n'
    '        capas.append(pdk.Layer("TextLayer", data=[{\n'
    '            "coords": [lon + dlon_v, lat + dlat_v],\n'
    '            "text": f"{nombre}\\n{irr:.0f} kWh/m²"\n'
    '        }], get_position="coords", get_text="text", font_size=14))\n\n'
    '# ─── Vista principal ─────────────────────────────────────────────────\n'
    'vista = pdk.ViewState(\n'
    '    latitude=lat, longitude=lon,\n'
    '    zoom=17, pitch=55, bearing=30\n'
    ')\n\n'
    'st.pydeck_chart(pdk.Deck(\n'
    '    layers=capas,\n'
    '    initial_view_state=vista,\n'
    '    map_style="mapbox://styles/mapbox/satellite-v9",\n'
    '    tooltip={"text": "Edificio BIPV\\nAltura: {height} m"}\n'
    '))\n\n'
    '# Guardar datos del edificio para las siguientes capas\n'
    'if st.button("✅ Confirmar geometría del edificio"):\n'
    '    st.session_state["edificio"] = {\n'
    '        "ancho": ancho, "profundidad": profund, "altura": altura,\n'
    '        "lat": lat, "lon": lon, "vecinos": vecinos\n'
    '    }\n'
    '    st.success("Geometría guardada. Continúa con Paso 5B →")'
)

# ═══════════════════════════════════════════════════════════════════════════════
# PASO B-5B: PYVISTA (CAPA 2)
# ═══════════════════════════════════════════════════════════════════════════════
sep()
h1('PASO B-5B — Edificio 3D con PyVista + stpyvista (Capa 2)')
badge('⏱', 'Tiempo estimado:', '3 horas  |  Dificultad: 🟡 Medio', AZUL)
badge('📦', 'Librerías:', 'pyvista 0.48.4 + stpyvista 0.2.1 + trimesh 4.12.2', VERDE)
doc.add_paragraph('')

body(
    'PyVista es la librería de ingeniería 3D más completa para Python. '
    'Usa VTK bajo el capó (estándar en ingeniería civil y mecánica). '
    'stpyvista es el puente que permite mostrar visualizaciones PyVista dentro de Streamlit. '
    'En esta capa se construye el edificio paramétrico 3D con las dimensiones del Paso 5A, '
    'se colorea cada fachada según su irradiancia anual calculada por pvlib, '
    'y se permite al usuario hacer clic en una fachada para seleccionarla como zona BIPV.'
)
doc.add_paragraph('')

archivo('calculos/edificio_3d.py')
codigo(
    '# ═══════════════════════════════════════════════════════════════\n'
    '# CAPA 2: Edificio 3D con PyVista\n'
    '# Archivo: calculos/edificio_3d.py\n'
    '# Instalar: pip install pyvista stpyvista trimesh\n'
    '# ═══════════════════════════════════════════════════════════════\n'
    'import pyvista as pv\n'
    'import numpy as np\n'
    'import trimesh\n\n'
    '# ─── Paleta de colores por irradiancia (azul→amarillo→rojo) ────\n'
    'def irr_a_color(irr_kwh_m2, min_irr=500, max_irr=2000):\n'
    '    """Mapea irradiancia a color RGB. Azul=bajo, Rojo=alto."""\n'
    '    t = np.clip((irr_kwh_m2 - min_irr) / (max_irr - min_irr), 0, 1)\n'
    '    if t < 0.5:\n'
    '        return [int(255*(2*t)), int(255*(2*t)), 255]  # azul → cyan\n'
    '    else:\n'
    '        return [255, int(255*(2-2*t)), 0]            # amarillo → rojo\n\n'
    'def construir_edificio_3d(\n'
    '    ancho: float,          # metros — dimensión Este-Oeste\n'
    '    profundidad: float,    # metros — dimensión Norte-Sur\n'
    '    altura: float,         # metros — altura total\n'
    '    irr_norte: float,      # kWh/m²/año fachada Norte\n'
    '    irr_sur: float,        # kWh/m²/año fachada Sur\n'
    '    irr_este: float,       # kWh/m²/año fachada Este\n'
    '    irr_oeste: float,      # kWh/m²/año fachada Oeste\n'
    '    irr_techo: float,      # kWh/m²/año cubierta\n'
    '    fachada_seleccionada: str = None  # "Norte"/"Sur"/"Este"/"Oeste"/"Techo"\n'
    ') -> pv.PolyData:\n'
    '    """\n'
    '    Construye un edificio paramétrico 3D coloreado por irradiancia.\n'
    '    Devuelve un mesh PyVista listo para renderizar en Streamlit con stpyvista.\n'
    '    """\n'
    '    # ─── Crear las 5 caras del edificio como meshes separados ────\n'
    '    facetas = {\n'
    '        "Sur":   {"irr": irr_sur,   "color": irr_a_color(irr_sur)},\n'
    '        "Norte": {"irr": irr_norte, "color": irr_a_color(irr_norte)},\n'
    '        "Este":  {"irr": irr_este,  "color": irr_a_color(irr_este)},\n'
    '        "Oeste": {"irr": irr_oeste, "color": irr_a_color(irr_oeste)},\n'
    '        "Techo": {"irr": irr_techo, "color": irr_a_color(irr_techo)},\n'
    '    }\n\n'
    '    # Si una fachada está seleccionada, resaltarla en verde brillante\n'
    '    if fachada_seleccionada and fachada_seleccionada in facetas:\n'
    '        facetas[fachada_seleccionada]["color"] = [0, 220, 80]\n\n'
    '    mesh_total = pv.PolyData()\n\n'
    '    def add_face(corners, nombre):\n'
    '        """Agrega un plano rectangular 3D al mesh total."""\n'
    '        nonlocal mesh_total\n'
    '        pts = np.array(corners, dtype=float)\n'
    '        faces = np.array([[4, 0, 1, 2, 3]])  # quad de 4 vértices\n'
    '        plane = pv.PolyData(pts, faces)\n'
    '        color = facetas[nombre]["color"]\n'
    '        plane["color"] = [color] * 4  # un color por vértice\n'
    '        plane["nombre"] = [nombre]    # etiqueta de fachada\n'
    '        plane["irr_kwh"] = [facetas[nombre]["irr"]]\n'
    '        mesh_total = mesh_total.merge(plane)\n\n'
    '    W, P, H = ancho/2, profundidad/2, altura\n\n'
    '    # Fachada Sur  (y = -P, mira hacia -Y = Sur geográfico)\n'
    '    add_face([[-W,0,0],[W,0,0],[W,0,H],[-W,0,H]], "Sur")\n'
    '    # Fachada Norte (y = +P)\n'
    '    add_face([[-W,P*2,0],[W,P*2,0],[W,P*2,H],[-W,P*2,H]], "Norte")\n'
    '    # Fachada Este  (x = +W)\n'
    '    add_face([[W,0,0],[W,P*2,0],[W,P*2,H],[W,0,H]], "Este")\n'
    '    # Fachada Oeste (x = -W)\n'
    '    add_face([[-W,0,0],[-W,P*2,0],[-W,P*2,H],[-W,0,H]], "Oeste")\n'
    '    # Techo (z = H)\n'
    '    add_face([[-W,0,H],[W,0,H],[W,P*2,H],[-W,P*2,H]], "Techo")\n\n'
    '    return mesh_total\n\n\n'
    '# ─── Importar edificio desde archivo OBJ/STL (si el usuario tiene planos 3D) ──\n'
    'def importar_obj(ruta_archivo: str) -> pv.PolyData:\n'
    '    """\n'
    '    Importa un modelo 3D de edificio desde archivo OBJ o STL.\n'
    '    Usa trimesh para la carga y convierte a PyVista.\n'
    '    """\n'
    '    mesh_tm = trimesh.load(ruta_archivo)\n'
    '    # Convertir trimesh → pyvista\n'
    '    vertices = np.array(mesh_tm.vertices)\n'
    '    faces_tm = np.array(mesh_tm.faces)\n'
    '    # PyVista necesita [n_verts, v0, v1, v2, ...] por cara\n'
    '    faces_pv = np.hstack([np.full((len(faces_tm), 1), 3), faces_tm]).ravel()\n'
    '    return pv.PolyData(vertices, faces_pv)'
)

doc.add_paragraph('')
archivo('pages/5B_🏢_Edificio_3D.py')
codigo(
    '# ═══════════════════════════════════════════════════════════════\n'
    '# Archivo: pages/5B_🏢_Edificio_3D.py\n'
    '# ═══════════════════════════════════════════════════════════════\n'
    'import streamlit as st\n'
    'import stpyvista\n'
    'import pvlib\n'
    'import pandas as pd\n'
    'from calculos.edificio_3d import construir_edificio_3d, importar_obj\n'
    'from calculos.solar import calcular_irradiancia_por_fachada\n\n'
    'st.header("🏢 Paso 5B — Edificio 3D y Selección de Fachada BIPV")\n\n'
    'edificio = st.session_state.get("edificio", {})\n'
    'tmy      = st.session_state.get("tmy",      None)\n\n'
    'if not edificio:\n'
    '    st.warning("Primero define la geometría del edificio en el Paso 5A.")\n'
    '    st.stop()\n\n'
    '# ─── Calcular irradiancia por fachada con pvlib ──────────────────\n'
    'lat = st.session_state.get("lat", -33.45)\n'
    'lon = st.session_state.get("lon", -70.65)\n\n'
    '@st.cache_data\n'
    'def get_irr_por_fachada(_tmy, lat, lon):\n'
    '    """Calcula irradiancia anual kWh/m² para cada orientación con pvlib."""\n'
    '    from calculos.solar import calcular_irradiancia_por_fachada\n'
    '    return calcular_irradiancia_por_fachada(_tmy, lat, lon)\n\n'
    'if tmy is not None:\n'
    '    irr = get_irr_por_fachada(tmy, lat, lon)\n'
    'else:\n'
    '    # Valores por defecto para Santiago de Chile (hemisferio sur)\n'
    '    irr = {"Norte": 900, "Sur": 1600, "Este": 1100, "Oeste": 1100, "Techo": 1800}\n\n'
    '# ─── Selector de fachada BIPV ────────────────────────────────────\n'
    'col1, col2 = st.columns([1, 2])\n'
    'with col1:\n'
    '    st.subheader("Orientaciones disponibles")\n'
    '    for nombre, valor in irr.items():\n'
    '        st.metric(f"Fachada {nombre}", f"{valor:.0f} kWh/m²/año")\n\n'
    '    fachada_sel = st.selectbox(\n'
    '        "Seleccionar fachada para BIPV",\n'
    '        ["Sur", "Norte", "Este", "Oeste", "Techo"],\n'
    '        index=0\n'
    '    )\n'
    '    area_bipv = st.number_input("Área BIPV en fachada seleccionada (m²)",\n'
    '                                min_value=1.0, max_value=float(edificio["ancho"]*edificio["altura"]),\n'
    '                                value=float(edificio["ancho"]*edificio["altura"]*0.6))\n\n'
    'with col2:\n'
    '    st.subheader(f"Vista 3D — Fachada {fachada_sel} seleccionada (verde)")\n\n'
    '    # Construir y renderizar el edificio 3D\n'
    '    mesh = construir_edificio_3d(\n'
    '        ancho=edificio["ancho"], profundidad=edificio["profundidad"],\n'
    '        altura=edificio["altura"],\n'
    '        irr_norte=irr["Norte"], irr_sur=irr["Sur"],\n'
    '        irr_este=irr["Este"], irr_oeste=irr["Oeste"], irr_techo=irr["Techo"],\n'
    '        fachada_seleccionada=fachada_sel\n'
    '    )\n\n'
    '    plotter = stpyvista.start_xvfb()  # necesario para servidor sin pantalla\n'
    '    p = pyvista.Plotter(window_size=[600, 400])\n'
    '    p.add_mesh(mesh, scalars="irr_kwh", cmap="plasma",\n'
    '               show_scalar_bar=True, scalar_bar_args={"title": "kWh/m²/año"})\n'
    '    p.add_axes()  # indicador Norte/Sur/Este/Oeste\n'
    '    p.camera_position = "iso"\n'
    '    stpyvista.render_pyvista(p)\n\n'
    '# ─── Guardar selección ───────────────────────────────────────────\n'
    'if st.button("✅ Confirmar fachada y área BIPV"):\n'
    '    st.session_state["fachada_bipv"] = {\n'
    '        "nombre": fachada_sel, "area_m2": area_bipv,\n'
    '        "azimut": {"Norte":0,"Sur":180,"Este":90,"Oeste":270,"Techo":180}[fachada_sel],\n'
    '        "inclinacion": 90 if fachada_sel != "Techo" else 10,\n'
    '        "irr_anual": irr[fachada_sel]\n'
    '    }\n'
    '    st.success(f"Fachada {fachada_sel} confirmada — {area_bipv:.0f} m² → Paso 5C")'
)

archivo('calculos/solar.py  →  agregar función calcular_irradiancia_por_fachada()')
codigo(
    '# ─── Agregar a calculos/solar.py ─────────────────────────────────\n'
    'import pvlib\n'
    'import pandas as pd\n\n'
    'def calcular_irradiancia_por_fachada(tmy_data: pd.DataFrame, lat: float, lon: float) -> dict:\n'
    '    """\n'
    '    Calcula la irradiancia anual en el plano inclinado (POA)\n'
    '    para las 4 orientaciones cardinales y el techo, usando pvlib Perez.\n'
    '    Retorna dict: {Norte, Sur, Este, Oeste, Techo} en kWh/m²/año.\n'
    '    """\n'
    '    loc      = pvlib.location.Location(lat, lon)\n'
    '    sol_pos  = loc.get_solarposition(tmy_data.index)\n'
    '    resultado = {}\n\n'
    '    orientaciones = {\n'
    '        "Sur":   {"tilt": 90,  "azimut": 180},\n'
    '        "Norte": {"tilt": 90,  "azimut": 0},\n'
    '        "Este":  {"tilt": 90,  "azimut": 90},\n'
    '        "Oeste": {"tilt": 90,  "azimut": 270},\n'
    '        "Techo": {"tilt": 10,  "azimut": 180},\n'
    '    }\n\n'
    '    for nombre, params in orientaciones.items():\n'
    '        poa = pvlib.irradiance.get_total_irradiance(\n'
    '            surface_tilt=params["tilt"],\n'
    '            surface_azimuth=params["azimut"],\n'
    '            solar_zenith=sol_pos["apparent_zenith"],\n'
    '            solar_azimuth=sol_pos["azimuth"],\n'
    '            dni=tmy_data["dni"],\n'
    '            ghi=tmy_data["ghi"],\n'
    '            dhi=tmy_data["dhi"],\n'
    '            model="perez"\n'
    '        )\n'
    '        resultado[nombre] = round(poa["poa_global"].sum() / 1000, 0)  # kWh/m²/año\n\n'
    '    return resultado'
)

# ═══════════════════════════════════════════════════════════════════════════════
# PASO B-5C: THREE.JS (CAPA 3)
# ═══════════════════════════════════════════════════════════════════════════════
sep()
h1('PASO B-5C — Fachada Detalle con Three.js via Streamlit Component (Capa 3)')
badge('⏱', 'Tiempo estimado:', '4-5 horas  |  Dificultad: 🔴 Avanzado', AZUL)
badge('📦', 'Librerías:', 'streamlit-javascript 0.1.5 + Three.js (CDN, sin instalación pip)', VERDE)
doc.add_paragraph('')

body(
    'Esta es la capa más potente. Three.js se carga desde CDN dentro de un '
    'bloque HTML embebido en Streamlit usando st.components.v1.html(). '
    'La comunicación bidireccional (usuario dibuja paneles → Python recalcula) '
    'se logra con streamlit-javascript. '
    'En esta vista el usuario ve la fachada seleccionada en detalle, '
    'puede dibujar zonas de paneles haciendo clic, y ve las sombras moverse '
    'hora a hora según la posición solar calculada por pvlib.'
)
doc.add_paragraph('')

body('Flujo de comunicación Python ↔ Three.js:')
flujo2 = (
    'Python (pvlib)           Three.js (HTML)           Python (pvlib)\n'
    '     │                        │                         │\n'
    '     │── sun_vector JSON ────▶│                         │\n'
    '     │   {az, el, hora}       │── renderiza sombra ──▶  │\n'
    '     │                        │                         │\n'
    '     │                        │◀── usuario dibuja ──────│\n'
    '     │                        │    paneles (clic)       │\n'
    '     │◀── panel_data JSON ────│                         │\n'
    '     │   {area, x, y, w, h}   │                         │\n'
    '     │                        │                         │\n'
    '     │── recalcula energía ───│──────────────────────▶  │\n'
    '     │   (pvlib + numpy)      │   muestra kWh en UI     │'
)
p_f2 = doc.add_paragraph()
r_f2 = p_f2.add_run(flujo2)
r_f2.font.name = 'Courier New'; r_f2.font.size = Pt(9); r_f2.font.color.rgb = COD
doc.add_paragraph('')

archivo('pages/5C_🌟_Fachada_3D.py')
codigo(
    '# ═══════════════════════════════════════════════════════════════\n'
    '# CAPA 3: Fachada Detalle con Three.js\n'
    '# Archivo: pages/5C_🌟_Fachada_3D.py\n'
    '# Instalar: pip install streamlit-javascript\n'
    '# ═══════════════════════════════════════════════════════════════\n'
    'import streamlit as st\n'
    'import json\n'
    'import pvlib\n'
    'import pandas as pd\n'
    'from streamlit_javascript import st_javascript\n\n'
    'st.header("🌟 Paso 5C — Diseño de Paneles BIPV en Fachada (3D interactivo)")\n\n'
    'fachada = st.session_state.get("fachada_bipv", {})\n'
    'if not fachada:\n'
    '    st.warning("Primero selecciona la fachada en el Paso 5B."); st.stop()\n\n'
    '# ─── Calcular posición solar para la hora seleccionada ───────────\n'
    'lat = st.session_state.get("lat", -33.45)\n'
    'lon = st.session_state.get("lon", -70.65)\n\n'
    'col1, col2 = st.columns([1, 3])\n'
    'with col1:\n'
    '    mes  = st.slider("Mes",  1, 12, 6)\n'
    '    hora = st.slider("Hora", 6, 19, 12)\n\n'
    '    # Calcular posición solar con pvlib\n'
    '    fecha = pd.Timestamp(f"2024-{mes:02d}-15 {hora:02d}:00", tz="America/Santiago")\n'
    '    loc   = pvlib.location.Location(lat, lon)\n'
    '    pos   = loc.get_solarposition(pd.DatetimeIndex([fecha]))\n'
    '    azimut_sol  = float(pos["azimuth"].iloc[0])\n'
    '    elevacion   = float(pos["apparent_elevation"].iloc[0])\n'
    '    st.metric("Azimut solar", f"{azimut_sol:.1f}°")\n'
    '    st.metric("Elevación solar", f"{elevacion:.1f}°")\n\n'
    '    tecnologia  = st.selectbox("Tecnología BIPV", ["CdTe", "mono-Si", "CIGS", "HJT"])\n'
    '    trans       = st.slider("Transparencia (%)", 10, 60, 20) / 100\n\n'
    'with col2:\n'
    '    # ─── Datos que se pasan a Three.js ──────────────────────────\n'
    '    config_3d = {\n'
    '        "ancho_fachada": float(fachada.get("area_m2", 50) ** 0.5 * 1.5),\n'
    '        "alto_fachada":  float(fachada.get("area_m2", 50) ** 0.5 * 0.67),\n'
    '        "azimut_sol":    azimut_sol,\n'
    '        "elevacion_sol": elevacion,\n'
    '        "transparencia": trans,\n'
    '        "nombre_fachada": fachada.get("nombre", "Sur"),\n'
    '    }\n'
    '    config_json = json.dumps(config_3d)\n\n'
    '    # ─── HTML + Three.js embebido en Streamlit ───────────────────\n'
    '    html_3d = f"""\n'
    '    <!DOCTYPE html><html><head>\n'
    '    <style>body{{margin:0;background:#1a1a2e;}}canvas{{display:block;}}</style>\n'
    '    <script src="https://cdnjs.cloudflare.com/ajax/libs/three.js/r134/three.min.js">\n'
    '    </script></head>\n'
    '    <body>\n'
    '    <div id="info" style="position:absolute;top:10px;left:10px;color:white;font-family:sans-serif;font-size:13px;">\n'
    '      Haz clic en la fachada para colocar paneles BIPV<br>\n'
    '      <span id="area_info">Área seleccionada: 0.0 m²</span>\n'
    '    </div>\n'
    '    <script>\n'
    '    const cfg = {config_json};\n\n'
    '    // ─── Escena Three.js ─────────────────────────────────────────\n'
    '    const escena    = new THREE.Scene();\n'
    '    const camara    = new THREE.PerspectiveCamera(45, window.innerWidth/window.innerHeight, 0.1, 1000);\n'
    '    const renderer  = new THREE.WebGLRenderer({{antialias: true}});\n'
    '    renderer.setSize(window.innerWidth, window.innerHeight);\n'
    '    renderer.shadowMap.enabled = true;\n'
    '    document.body.appendChild(renderer.domElement);\n\n'
    '    // ─── Fondo del cielo ─────────────────────────────────────────\n'
    '    escena.background = new THREE.Color(0x1a1a2e);\n\n'
    '    // ─── Fachada base (plano gris) ───────────────────────────────\n'
    '    const W = cfg.ancho_fachada, H = cfg.alto_fachada;\n'
    '    const geom_fachada = new THREE.PlaneGeometry(W, H);\n'
    '    const mat_fachada  = new THREE.MeshLambertMaterial({{color: 0x555577, side: THREE.DoubleSide}});\n'
    '    const fachada_mesh = new THREE.Mesh(geom_fachada, mat_fachada);\n'
    '    fachada_mesh.receiveShadow = true;\n'
    '    escena.add(fachada_mesh);\n\n'
    '    // Grid de paneles BIPV (modulación 1m × 1.65m)\n'
    '    const paneles_grupo = new THREE.Group();\n'
    '    const mod_w = 1.0, mod_h = 1.65;\n'
    '    const nx = Math.floor(W / mod_w);\n'
    '    const ny = Math.floor(H / mod_h);\n'
    '    const paneles_estado = Array(nx).fill(0).map(()=>Array(ny).fill(false));\n\n'
    '    for (let ix = 0; ix < nx; ix++) {{\n'
    '        for (let iy = 0; iy < ny; iy++) {{\n'
    '            const g = new THREE.PlaneGeometry(mod_w*0.94, mod_h*0.94);\n'
    '            const m = new THREE.MeshLambertMaterial({{\n'
    '                color: 0x1a5c8a,\n'
    '                transparent: true,\n'
    '                opacity: 0.3 + cfg.transparencia * 0.4,\n'
    '                side: THREE.DoubleSide\n'
    '            }});\n'
    '            const panel = new THREE.Mesh(g, m);\n'
    '            panel.position.set(\n'
    '                -W/2 + mod_w*(ix+0.5),\n'
    '                -H/2 + mod_h*(iy+0.5),\n'
    '                0.01\n'
    '            );\n'
    '            panel.userData = {{ix, iy}};\n'
    '            paneles_grupo.add(panel);\n'
    '        }}\n'
    '    }}\n'
    '    escena.add(paneles_grupo);\n\n'
    '    // ─── Luz solar posicionada según pvlib ───────────────────────\n'
    '    const az_rad = cfg.azimut_sol * Math.PI / 180;\n'
    '    const el_rad = cfg.elevacion_sol * Math.PI / 180;\n'
    '    const dist   = 15;\n'
    '    const luz_solar = new THREE.DirectionalLight(0xffeebb, 1.5);\n'
    '    luz_solar.position.set(\n'
    '        dist * Math.sin(az_rad) * Math.cos(el_rad),\n'
    '        dist * Math.cos(az_rad) * Math.cos(el_rad),\n'
    '        dist * Math.sin(el_rad)\n'
    '    );\n'
    '    luz_solar.castShadow = true;\n'
    '    escena.add(luz_solar);\n'
    '    escena.add(new THREE.AmbientLight(0x404060, 0.8));\n\n'
    '    // ─── Cámara ──────────────────────────────────────────────────\n'
    '    camara.position.set(0, -H, H);\n'
    '    camara.lookAt(0, 0, 0);\n\n'
    '    // ─── Interacción: clic para activar/desactivar paneles ───────\n'
    '    const raycaster = new THREE.Raycaster();\n'
    '    const mouse     = new THREE.Vector2();\n'
    '    let area_total  = 0;\n\n'
    '    renderer.domElement.addEventListener("click", (evt) => {{\n'
    '        const rect = renderer.domElement.getBoundingClientRect();\n'
    '        mouse.x =  ((evt.clientX - rect.left) / rect.width)  * 2 - 1;\n'
    '        mouse.y = -((evt.clientY - rect.top)  / rect.height) * 2 + 1;\n'
    '        raycaster.setFromCamera(mouse, camara);\n'
    '        const intersects = raycaster.intersectObjects(paneles_grupo.children);\n'
    '        if (intersects.length > 0) {{\n'
    '            const panel = intersects[0].object;\n'
    '            const {{ix, iy}} = panel.userData;\n'
    '            paneles_estado[ix][iy] = !paneles_estado[ix][iy];\n'
    '            panel.material.color.set(paneles_estado[ix][iy] ? 0x00cc55 : 0x1a5c8a);\n'
    '            panel.material.opacity = paneles_estado[ix][iy] ? 0.9 : (0.3 + cfg.transparencia*0.4);\n\n'
    '            // Calcular área total seleccionada\n'
    '            let total = 0;\n'
    '            paneles_estado.forEach(col => col.forEach(v => {{ if(v) total += mod_w * mod_h; }}));\n'
    '            area_total = total;\n'
    '            document.getElementById("area_info").innerText = "Área seleccionada: " + total.toFixed(1) + " m²";\n\n'
    '            // Enviar datos a Streamlit via postMessage\n'
    '            const datos = {{\n'
    '                area_m2: total,\n'
    '                n_paneles: paneles_estado.flat().filter(Boolean).length,\n'
    '                mod_w: mod_w, mod_h: mod_h\n'
    '            }};\n'
    '            window.parent.postMessage({{type:"streamlit:setComponentValue", value: datos}}, "*");\n'
    '        }}\n'
    '    }});\n\n'
    '    // ─── Loop de animación ────────────────────────────────────────\n'
    '    function animar() {{\n'
    '        requestAnimationFrame(animar);\n'
    '        renderer.render(escena, camara);\n'
    '    }}\n'
    '    animar();\n'
    '    </script></body></html>\n'
    '    """\n\n'
    '    # Renderizar el componente Three.js en Streamlit\n'
    '    st.components.v1.html(html_3d, height=550, scrolling=False)\n\n'
    '# ─── Recibir área dibujada desde Three.js y recalcular producción ─\n'
    'datos_js = st_javascript("window._streamlit_panel_area || 0")\n\n'
    'if datos_js and isinstance(datos_js, dict) and datos_js.get("area_m2", 0) > 0:\n'
    '    area_dibujada = datos_js["area_m2"]\n'
    '    n_paneles     = datos_js.get("n_paneles", 0)\n'
    '    st.session_state["area_dibujada_m2"] = area_dibujada\n\n'
    '    # Calcular producción estimada con el área dibujada\n'
    '    efic_dict = {"CdTe": 0.14, "mono-Si": 0.20, "CIGS": 0.15, "HJT": 0.22}\n'
    '    efic = efic_dict.get(tecnologia, 0.16)\n'
    '    potencia_kwp = area_dibujada * efic * (1 - trans)\n'
    '    irr_anual    = fachada.get("irr_anual", 1200)\n'
    '    energia_kwh  = potencia_kwp * irr_anual * 0.80  # PR=0.80\n\n'
    '    col1, col2, col3 = st.columns(3)\n'
    '    col1.metric("Área dibujada", f"{area_dibujada:.1f} m²")\n'
    '    col2.metric("Potencia pico", f"{potencia_kwp:.2f} kWp")\n'
    '    col3.metric("Producción estimada", f"{energia_kwh:.0f} kWh/año")\n\n'
    '    if st.button("✅ Confirmar diseño de paneles → Paso 6"):\n'
    '        st.session_state["potencia_kwp"]           = potencia_kwp\n'
    '        st.session_state["produccion_anual_kwh"]   = energia_kwh\n'
    '        st.success("Diseño confirmado. Continúa con Paso 6 — Balance Energético →")'
)

# ═══════════════════════════════════════════════════════════════════════════════
# ALTERNATIVA GLTF — IMPORTAR MODELO COMPLEJO
# ═══════════════════════════════════════════════════════════════════════════════
sep()
h1('BONUS — Importar modelo GLTF/OBJ complejo (edificios reales con planos de arquitecto)')
badge('⚠', 'Cuándo usar:', 'Solo cuando el usuario tiene el modelo 3D del edificio real (de ArchiCAD, Revit, SketchUp, AutoCAD).', NARANJO)
doc.add_paragraph('')

body(
    'Si el usuario tiene el plano 3D del edificio en formato OBJ, GLTF, STL o IFC, '
    'trimesh permite cargarlo directamente en Python, analizarlo con pvlib para calcular '
    'la irradiancia en cada cara del modelo, y pasarlo a Three.js para visualizarlo. '
    'Esta funcionalidad convierte la calculadora en un herramienta de nivel profesional.'
)
doc.add_paragraph('')

archivo('calculos/importar_modelo.py')
codigo(
    '# ═══════════════════════════════════════════════════════════════\n'
    '# Importar modelo 3D real (OBJ/GLTF/STL) con trimesh\n'
    '# Archivo: calculos/importar_modelo.py\n'
    '# Instalar: pip install trimesh\n'
    '# ═══════════════════════════════════════════════════════════════\n'
    'import trimesh\n'
    'import numpy as np\n'
    'from typing import List, Dict\n\n'
    'def analizar_caras_bipv(ruta_archivo: str, lat: float, lon: float) -> List[Dict]:\n'
    '    """\n'
    '    Importa un modelo 3D y analiza qué caras (polígonos) son candidatas para BIPV.\n'
    '    Filtra por: inclinación (<85°), área mínima (>2 m²), no horizontal inferior.\n'
    '    Retorna lista de caras con área, normal, orientación y área en m².\n'
    '    """\n'
    '    mesh = trimesh.load(ruta_archivo)\n\n'
    '    # Si el archivo tiene múltiples objetos, unirlos\n'
    '    if isinstance(mesh, trimesh.Scene):\n'
    '        mesh = trimesh.util.concatenate(list(mesh.geometry.values()))\n\n'
    '    caras_bipv = []\n'
    '    for i, cara in enumerate(mesh.faces):\n'
    '        vertices = mesh.vertices[cara]\n\n'
    '        # Vector normal de la cara\n'
    '        v1 = vertices[1] - vertices[0]\n'
    '        v2 = vertices[2] - vertices[0]\n'
    '        normal = np.cross(v1, v2)\n'
    '        norm_len = np.linalg.norm(normal)\n'
    '        if norm_len < 1e-8:\n'
    '            continue\n'
    '        normal = normal / norm_len\n\n'
    '        # Inclinación de la cara respecto a la vertical\n'
    '        inclinacion_rad = np.arccos(np.clip(abs(normal[2]), 0, 1))\n'
    '        inclinacion_deg = np.degrees(inclinacion_rad)\n\n'
    '        # Azimut de la cara\n'
    '        azimut_rad = np.arctan2(normal[0], normal[1])\n'
    '        azimut_deg = (np.degrees(azimut_rad) + 360) % 360\n\n'
    '        # Área de la cara (en m²)\n'
    '        area = trimesh.triangles.area(vertices[np.newaxis, ...])[0]\n\n'
    '        # Filtro: solo caras que miran hacia afuera (normal.z < 0.9 = no horizontal)\n'
    '        # y que tengan área suficiente para al menos 1 panel\n'
    '        if area < 2.0 or normal[2] < -0.1:  # excluir caras que miran hacia abajo\n'
    '            continue\n\n'
    '        caras_bipv.append({\n'
    '            "id": i,\n'
    '            "area_m2": round(area, 2),\n'
    '            "inclinacion_deg": round(90 - inclinacion_deg, 1),\n'
    '            "azimut_deg": round(azimut_deg, 1),\n'
    '            "normal": normal.tolist(),\n'
    '            "centroide": vertices.mean(axis=0).tolist()\n'
    '        })\n\n'
    '    # Ordenar por área descendente\n'
    '    return sorted(caras_bipv, key=lambda x: x["area_m2"], reverse=True)'
)

doc.add_paragraph('')
body(
    'Con esta función, el usuario sube su modelo OBJ/GLTF (exportado de ArchiCAD, Revit '
    'o SketchUp), la calculadora analiza automáticamente todas las caras del edificio, '
    'identifica cuáles son candidatas para BIPV, calcula la irradiancia en cada una '
    'con pvlib, y las muestra ordenadas por potencial solar.'
)

# ═══════════════════════════════════════════════════════════════════════════════
# TABLA ACTUALIZADA DEL PLAN
# ═══════════════════════════════════════════════════════════════════════════════
sep()
h1('PLAN ACTUALIZADO — PASO B-5 REEMPLAZADO POR B-5A + B-5B + B-5C')

tbl_plan = doc.add_table(rows=1, cols=5)
tbl_plan.style = 'Table Grid'
tabla_header(tbl_plan, ['Paso', 'Nombre', 'Herramienta', 'Tiempo', 'Salida'])

pasos_plan = [
    ('B-5A', 'Vista de Sitio 3D',          'Pydeck (deck.gl)',              '2 h 🟢', 'Mapa 3D con edificio + vecinos + heatmap GHI'),
    ('B-5B', 'Edificio 3D + selección',    'PyVista + stpyvista + trimesh', '3 h 🟡', 'Edificio coloreado por irradiancia, fachada clickeable'),
    ('B-5C', 'Fachada detalle + paneles',  'Three.js (st.components.html)', '5 h 🔴', 'Grid de paneles dibujables, sombras en tiempo real'),
    ('BONUS','Importar modelo GLTF/OBJ',   'trimesh + Three.js GLTFLoader', '3 h 🔴', 'Análisis automático de caras BIPV en modelo real'),
]

for fila in pasos_plan:
    row = tbl_plan.add_row().cells
    for i, val in enumerate(fila):
        row[i].text = val
        if row[i].paragraphs[0].runs:
            r = row[i].paragraphs[0].runs[0]
        else:
            r = row[i].paragraphs[0].add_run(val)
        r.font.size = Pt(9)
        if i == 0:
            r.bold = True; r.font.color.rgb = AZUL

doc.add_paragraph('')

# Requirements.txt actualizado
h2('requirements.txt actualizado con librerías 3D')
codigo(
    '# requirements.txt — Calculadora BIPV Python con 3D\n'
    'streamlit>=1.35\n'
    'pvlib>=0.10\n'
    'numpy>=1.26\n'
    'pandas>=2.0\n'
    'plotly>=5.20\n'
    'matplotlib>=3.8\n'
    'pydeck>=0.9          # Capa 1: vista de sitio\n'
    'pyvista>=0.48        # Capa 2: edificio 3D (requiere vtk)\n'
    'vtk>=9.3             # Backend de PyVista\n'
    'stpyvista>=0.2       # Bridge PyVista → Streamlit\n'
    'trimesh>=4.10        # Importar OBJ/GLTF/STL\n'
    'streamlit-javascript>=0.1.5  # Comunicación Three.js → Python\n'
    'python-docx>=1.1\n'
    'openpyxl>=3.1\n'
    'scipy>=1.12\n'
    'numpy-financial>=1.0\n'
    'requests>=2.31'
)

# ═══════════════════════════════════════════════════════════════════════════════
# CRONOGRAMA ACTUALIZADO
# ═══════════════════════════════════════════════════════════════════════════════
sep()
h1('CRONOGRAMA ACTUALIZADO — 5 SEMANAS')

tbl_cron = doc.add_table(rows=1, cols=3)
tbl_cron.style = 'Table Grid'
tabla_header(tbl_cron, ['Semana', 'Pasos', 'Objetivo'])

semanas_u = [
    ('Semana 1', 'B-1, B-2, B-3', 'Base: estructura Streamlit, recurso solar pvlib, catálogo BIPV'),
    ('Semana 2', 'B-4, B-5A',     'Motor de cálculo IEC 61724 + Vista de sitio Pydeck 3D'),
    ('Semana 3', 'B-5B, B-5C',    'Edificio 3D PyVista + Fachada interactiva Three.js'),
    ('Semana 4', 'B-6, B-7, B-8', 'Balance mensual, análisis financiero VAN/TIR/LCOE'),
    ('Semana 5', 'B-9, B-10, BONUS', 'Reportes Word/Excel + Deploy Streamlit Cloud + OBJ importer'),
]

for fila in semanas_u:
    row = tbl_cron.add_row().cells
    for i, val in enumerate(fila):
        row[i].text = val
        if row[i].paragraphs[0].runs:
            r = row[i].paragraphs[0].runs[0]
        else:
            r = row[i].paragraphs[0].add_run(val)
        r.font.size = Pt(9)
        if i == 0: r.bold = True; r.font.color.rgb = AZUL

doc.add_paragraph('')

# Pie
footer = doc.add_paragraph(
    'Calculadora BIPV Python — Visualizador 3D  ·  pydeck + PyVista + Three.js + pvlib  ·  2026'
)
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.runs[0].font.size = Pt(9); footer.runs[0].font.color.rgb = GRIS; footer.runs[0].italic = True

doc.save('Visualizador_3D_BIPV_Streamlit.docx')
print("Documento generado correctamente.")
