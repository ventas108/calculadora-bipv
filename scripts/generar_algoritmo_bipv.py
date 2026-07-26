
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ══════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════
t = doc.add_heading('ALGORITMO PASO A PASO — SOFTWARE BIPV', 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.runs[0].font.color.rgb = RGBColor(0x1A, 0x5C, 0x8A)

st = doc.add_paragraph('Building-Integrated Photovoltaics — Fotovoltaica Integrada en Edificios')
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
st.runs[0].font.size = Pt(12)
st.runs[0].font.bold = True
st.runs[0].font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)

doc.add_paragraph('')
intro = doc.add_paragraph(
    'Este documento describe el algoritmo completo, módulo a módulo, para programar '
    'un software de cálculo BIPV en Python. Cada etapa incluye: qué datos entran, '
    'qué se calcula, la fórmula o lógica utilizada, y el código Python exacto. '
    'El algoritmo sigue el mismo flujo de cálculo que usan herramientas profesionales '
    'como PVsyst, SAM (NREL) y DesignBuilder para sistemas BIPV.'
)
intro.runs[0].font.size = Pt(10)
intro.runs[0].italic = True
intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
doc.add_paragraph('')

# ══════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════
COLOR_MODULO   = RGBColor(0x1A, 0x5C, 0x8A)   # azul oscuro — módulos principales
COLOR_PASO     = RGBColor(0x6E, 0x27, 0x94)   # morado — pasos internos
COLOR_DECISION = RGBColor(0xC0, 0x39, 0x2B)   # rojo — decisiones
COLOR_CODIGO   = RGBColor(0x10, 0x10, 0x60)   # azul noche — código
COLOR_FORMULA  = RGBColor(0x17, 0x6B, 0x17)   # verde — fórmulas
COLOR_FLUJO    = RGBColor(0xD4, 0x7A, 0x00)   # naranja — flujograma

def titulo_modulo(texto):
    doc.add_paragraph('')
    h = doc.add_heading(texto, level=1)
    h.runs[0].font.size  = Pt(14)
    h.runs[0].font.color.rgb = COLOR_MODULO
    doc.add_paragraph('')

def titulo_paso(texto):
    h = doc.add_heading(texto, level=2)
    h.runs[0].font.size  = Pt(12)
    h.runs[0].font.color.rgb = COLOR_PASO

def bloque_entrada(texto):
    p = doc.add_paragraph()
    r = p.add_run('▶ DATOS DE ENTRADA:  ')
    r.bold = True
    r.font.color.rgb = COLOR_MODULO
    p.add_run(texto)

def bloque_salida(texto):
    p = doc.add_paragraph()
    r = p.add_run('◀ RESULTADO:  ')
    r.bold = True
    r.font.color.rgb = COLOR_FORMULA
    p.add_run(texto)

def bloque_formula(titulo_f, formula_f):
    p = doc.add_paragraph()
    r = p.add_run(f'  Fórmula — {titulo_f}:  ')
    r.bold = True
    cr = p.add_run(formula_f)
    cr.font.name  = 'Courier New'
    cr.font.size  = Pt(10)
    cr.font.color.rgb = COLOR_FORMULA

def bloque_decision(texto):
    p = doc.add_paragraph()
    r = p.add_run('◆ DECISIÓN:  ')
    r.bold = True
    r.font.color.rgb = COLOR_DECISION
    p.add_run(texto)

def bloque_codigo(codigo):
    p = doc.add_paragraph()
    cr = p.add_run(codigo)
    cr.font.name  = 'Courier New'
    cr.font.size  = Pt(9.5)
    cr.font.color.rgb = COLOR_CODIGO
    doc.add_paragraph('')

def flecha():
    p = doc.add_paragraph()
    r = p.add_run('                              ↓')
    r.font.size  = Pt(14)
    r.bold = True
    r.font.color.rgb = COLOR_FLUJO

def separador():
    p = doc.add_paragraph('─' * 78)
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

# ══════════════════════════════════════════
# DIAGRAMA DE FLUJO GENERAL (texto)
# ══════════════════════════════════════════
titulo_modulo('DIAGRAMA DE FLUJO GENERAL DEL SOFTWARE BIPV')

flujo = doc.add_paragraph()
r_f = flujo.add_run(
    '┌─────────────────────────────────────────────────────────────────┐\n'
    '│          INICIO DEL PROGRAMA BIPV                               │\n'
    '└──────────────────────────────┬──────────────────────────────────┘\n'
    '                               ↓\n'
    '┌──────────────────────────────────────────────────────────────────┐\n'
    '│  MÓDULO 1 — Datos del edificio y localización                   │\n'
    '│  (latitud, longitud, orientación fachadas, tipo de superficie)  │\n'
    '└──────────────────────────────┬──────────────────────────────────┘\n'
    '                               ↓\n'
    '┌──────────────────────────────────────────────────────────────────┐\n'
    '│  MÓDULO 2 — Recurso solar (irradiancia en plano inclinado)      │\n'
    '│  (HSP por fachada, transposición de irradiancia, horizonte)     │\n'
    '└──────────────────────────────┬──────────────────────────────────┘\n'
    '                               ↓\n'
    '┌──────────────────────────────────────────────────────────────────┐\n'
    '│  MÓDULO 3 — Caracterización del módulo BIPV                     │\n'
    '│  (tecnología, eficiencia, coeficientes, área activa, NOCT)      │\n'
    '└──────────────────────────────┬──────────────────────────────────┘\n'
    '                               ↓\n'
    '┌──────────────────────────────────────────────────────────────────┐\n'
    '│  MÓDULO 4 — Temperatura real de operación BIPV                  │\n'
    '│  (efecto NOCT modificado + acumulación de calor en fachada)     │\n'
    '└──────────────────────────────┬──────────────────────────────────┘\n'
    '                               ↓\n'
    '┌──────────────────────────────────────────────────────────────────┐\n'
    '│  MÓDULO 5 — Cálculo de energía generada                        │\n'
    '│  (energía ideal → pérdidas → energía real por fachada)         │\n'
    '└──────────────────────────────┬──────────────────────────────────┘\n'
    '                               ↓\n'
    '┌──────────────────────────────────────────────────────────────────┐\n'
    '│  MÓDULO 6 — Análisis de sombras BIPV                           │\n'
    '│  (sombra de aleros, voladizos, edificios vecinos, retranqueos)  │\n'
    '└──────────────────────────────┬──────────────────────────────────┘\n'
    '                               ↓\n'
    '┌──────────────────────────────────────────────────────────────────┐\n'
    '│  MÓDULO 7 — Dimensionado eléctrico del sistema BIPV             │\n'
    '│  (strings, inversor, ratio DC/AC, cableado, banco de baterías)  │\n'
    '└──────────────────────────────┬──────────────────────────────────┘\n'
    '                               ↓\n'
    '┌──────────────────────────────────────────────────────────────────┐\n'
    '│  MÓDULO 8 — Balance energético del edificio                     │\n'
    '│  (generación FV vs consumo edificio, autoconsumo, excedente)    │\n'
    '└──────────────────────────────┬──────────────────────────────────┘\n'
    '                               ↓\n'
    '┌──────────────────────────────────────────────────────────────────┐\n'
    '│  MÓDULO 9 — Análisis financiero y ambiental                     │\n'
    '│  (LCOE, Payback, VAN, TIR, CO₂ evitado, certificación energía)  │\n'
    '└──────────────────────────────┬──────────────────────────────────┘\n'
    '                               ↓\n'
    '┌──────────────────────────────────────────────────────────────────┐\n'
    '│  MÓDULO 10 — Generación de gráficos e informe final             │\n'
    '│  (gráficos matplotlib, resumen ejecutivo, archivo de resultados) │\n'
    '└──────────────────────────────┬──────────────────────────────────┘\n'
    '                               ↓\n'
    '┌──────────────────────────────────────────────────────────────────┐\n'
    '│                        FIN DEL PROGRAMA                         │\n'
    '└──────────────────────────────────────────────────────────────────┘'
)
r_f.font.name  = 'Courier New'
r_f.font.size  = Pt(8.5)
r_f.font.color.rgb = COLOR_FLUJO
doc.add_paragraph('')

# ══════════════════════════════════════════════════════════════
# MÓDULO 1 — EDIFICIO Y LOCALIZACIÓN
# ══════════════════════════════════════════════════════════════
separador()
titulo_modulo('MÓDULO 1 — Datos del edificio y localización')

bloque_entrada(
    'Latitud (°), Longitud (°), altitud (m), ciudad, orientación de cada fachada '
    '(azimut en °), inclinación de cada superficie (tilt en °), tipo de elemento BIPV '
    '(fachada ventilada, cubierta, vidrio, teja solar, pérgola).'
)
bloque_salida(
    'Diccionario con los parámetros geométricos del edificio listos para el cálculo solar.'
)
doc.add_paragraph('')

bloque_codigo(
    '# ── MÓDULO 1: Datos del edificio y localización ──\n\n'
    'edificio = {\n'
    '    "nombre"   : "Edificio Corporativo Norte",\n'
    '    "latitud"  : -33.45,      # grados decimales (negativo = hemisferio sur)\n'
    '    "longitud" : -70.66,      # grados decimales\n'
    '    "altitud_m": 520,         # metros sobre el nivel del mar\n'
    '    "zona_horaria": -4,       # UTC-4 para Chile\n'
    '}\n\n'
    '# Definir cada superficie BIPV del edificio\n'
    '# azimut: 0=Norte, 90=Este, 180=Sur, 270=Oeste\n'
    '# tilt: 0=horizontal, 90=vertical (fachada)\n'
    'superficies_bipv = [\n'
    '    {"id": "FAC_NORTE", "azimut": 0,   "tilt": 90, "area_m2": 80,\n'
    '     "tipo": "fachada_ventilada", "descripcion": "Fachada norte"},\n'
    '    {"id": "FAC_ESTE",  "azimut": 90,  "tilt": 90, "area_m2": 45,\n'
    '     "tipo": "fachada_ventilada", "descripcion": "Fachada este"},\n'
    '    {"id": "CUBIERTA",  "azimut": 0,   "tilt": 15, "area_m2": 120,\n'
    '     "tipo": "cubierta",           "descripcion": "Cubierta plana inclinada"},\n'
    '    {"id": "PERGOLA",   "azimut": 0,   "tilt": 30, "area_m2": 30,\n'
    '     "tipo": "pergola",            "descripcion": "Pérgola patio interior"},\n'
    ']\n\n'
    'print(f"Edificio: {edificio[\'nombre\']}")\n'
    'print(f"Ubicación: lat {edificio[\'latitud\']}° | lon {edificio[\'longitud\']}°")\n'
    'print(f"Superficies BIPV definidas: {len(superficies_bipv)}")'
)

flecha()

# ══════════════════════════════════════════════════════════════
# MÓDULO 2 — RECURSO SOLAR
# ══════════════════════════════════════════════════════════════
separador()
titulo_modulo('MÓDULO 2 — Recurso solar e irradiancia por superficie')

bloque_entrada(
    'Irradiación horizontal global (GHI en kWh/m²/día) de cada mes '
    '— obtenida de NASA POWER, PVGIS o SolarGIS. '
    'Azimut e inclinación (tilt) de cada superficie BIPV.'
)
bloque_salida(
    'Irradiación efectiva (kWh/m²/día) sobre cada superficie BIPV. '
    'HSP efectiva por orientación. Factor de corrección por orientación no óptima.'
)
doc.add_paragraph('')

titulo_paso('Paso 2.1 — Factor de corrección por inclinación y orientación (Kt)')
bloque_formula(
    'Factor Kt simplificado (Liu & Jordan)',
    'Kt ≈ cos(latitud - tilt) / cos(latitud)   [aproximación para azimut 0°=Norte]'
)
bloque_formula(
    'Irradiación en plano inclinado',
    'G_tilt = GHI × Kt_orientacion'
)
doc.add_paragraph('')

bloque_codigo(
    '# ── MÓDULO 2: Recurso solar e irradiancia por superficie ──\n'
    'import math\n\n'
    '# Irradiación horizontal global mensual (GHI) del lugar — kWh/m²/día\n'
    '# Fuente: NASA POWER / PVGIS ingresando latitud y longitud\n'
    'ghi_mensual = [5.2, 5.8, 6.1, 5.5, 4.8, 4.2, 4.5, 5.0, 5.4, 5.7, 5.3, 5.0]\n\n'
    '# Factores de corrección típicos por orientación y tilt en latitud ~33°S\n'
    '# (para cálculo preciso usar modelo de Hay-Davies o Perez)\n'
    'factor_kt = {\n'
    '    "fachada_ventilada": {\n'
    '        "azimut_0_norte": 0.72,   # fachada vertical norte — recibe buen sol\n'
    '        "azimut_90_este": 0.52,   # fachada vertical este — solo mañana\n'
    '        "azimut_180_sur": 0.28,   # fachada vertical sur — muy baja irradiancia\n'
    '        "azimut_270_oeste": 0.52, # fachada vertical oeste — solo tarde\n'
    '    },\n'
    '    "cubierta":  {"tilt_15": 0.98, "tilt_30": 1.03, "tilt_45": 0.95},\n'
    '    "pergola":   {"tilt_30": 1.03},\n'
    '}\n\n'
    '# Calcular HSP efectiva para cada superficie\n'
    'resultados_irradiancia = []\n\n'
    'for sup in superficies_bipv:\n'
    '    if sup["azimut"] == 0 and sup["tilt"] == 90:\n'
    '        kt = factor_kt["fachada_ventilada"]["azimut_0_norte"]\n'
    '    elif sup["azimut"] == 90 and sup["tilt"] == 90:\n'
    '        kt = factor_kt["fachada_ventilada"]["azimut_90_este"]\n'
    '    elif sup["tilt"] <= 20:\n'
    '        kt = factor_kt["cubierta"]["tilt_15"]\n'
    '    else:\n'
    '        kt = factor_kt["pergola"]["tilt_30"]\n\n'
    '    hsp_sup = [round(ghi * kt, 2) for ghi in ghi_mensual]\n'
    '    hsp_prom = round(sum(hsp_sup) / 12, 2)\n'
    '    resultados_irradiancia.append({\n'
    '        "id": sup["id"], "kt": kt,\n'
    '        "hsp_mensual": hsp_sup, "hsp_promedio": hsp_prom\n'
    '    })\n'
    '    print(f"{sup[\'id\']:12s} | Kt={kt:.2f} | HSP prom={hsp_prom} h/día")'
)

flecha()

# ══════════════════════════════════════════════════════════════
# MÓDULO 3 — CARACTERIZACIÓN DEL MÓDULO BIPV
# ══════════════════════════════════════════════════════════════
separador()
titulo_modulo('MÓDULO 3 — Caracterización del módulo BIPV')

bloque_entrada(
    'Tecnología del módulo (mono-Si, poli-Si, CdTe, CIGS, OPV, perovskita), '
    'potencia pico (Wp), eficiencia nominal (%), coeficiente de temperatura (%/°C), '
    'NOCT (°C), área total vs área activa (factor de relleno), transparencia si es vidrio BIPV.'
)
bloque_salida(
    'Potencia real por m² de fachada. Factor de cobertura. Parámetros de degradación.'
)
doc.add_paragraph('')

bloque_codigo(
    '# ── MÓDULO 3: Caracterización del módulo BIPV ──\n\n'
    '# Eficiencias típicas por tecnología BIPV (en condición estándar STC)\n'
    'tecnologias_bipv = {\n'
    '    "mono-Si":     {"efic_nom": 0.20, "coef_temp": -0.0035, "NOCT": 45,\n'
    '                   "degradacion": 0.005, "vida_util": 25},\n'
    '    "poli-Si":     {"efic_nom": 0.17, "coef_temp": -0.0038, "NOCT": 47,\n'
    '                   "degradacion": 0.006, "vida_util": 25},\n'
    '    "CdTe":        {"efic_nom": 0.16, "coef_temp": -0.0021, "NOCT": 44,\n'
    '                   "degradacion": 0.005, "vida_util": 25},\n'
    '    "CIGS":        {"efic_nom": 0.14, "coef_temp": -0.0025, "NOCT": 45,\n'
    '                   "degradacion": 0.006, "vida_util": 20},\n'
    '    "vidrio_BIPV": {"efic_nom": 0.10, "coef_temp": -0.0030, "NOCT": 50,\n'
    '                   "degradacion": 0.007, "vida_util": 20,\n'
    '                   "transparencia": 0.30},  # 30% de luz pasa a través\n'
    '}\n\n'
    '# Selección del módulo para el proyecto\n'
    'tecnologia_elegida = "mono-Si"\n'
    'modulo = tecnologias_bipv[tecnologia_elegida]\n\n'
    '# Factor de cobertura: no toda el área de fachada es celda activa\n'
    'factor_cobertura = 0.85   # 85% del área total es área activa con celdas\n\n'
    '# Potencia instalable por superficie\n'
    'for sup in superficies_bipv:\n'
    '    area_activa  = sup["area_m2"] * factor_cobertura\n'
    '    potencia_kWp = area_activa * modulo["efic_nom"]  # kWp (1kW/m² × efic × m²)\n'
    '    sup["area_activa_m2"] = round(area_activa, 1)\n'
    '    sup["potencia_kWp"]   = round(potencia_kWp, 2)\n'
    '    print(f"{sup[\'id\']:12s} | Área activa={area_activa:.1f} m² | {potencia_kWp:.2f} kWp")\n\n'
    'potencia_total_kWp = round(sum(s["potencia_kWp"] for s in superficies_bipv), 2)\n'
    'print(f"\\nPotencia BIPV total del edificio: {potencia_total_kWp} kWp")'
)

flecha()

# ══════════════════════════════════════════════════════════════
# MÓDULO 4 — TEMPERATURA DE OPERACIÓN BIPV
# ══════════════════════════════════════════════════════════════
separador()
titulo_modulo('MÓDULO 4 — Temperatura real de operación BIPV (efecto fachada)')

bloque_entrada(
    'NOCT del módulo, temperatura ambiente media mensual (°C), irradiancia efectiva (W/m²), '
    'tipo de montaje BIPV (ventilado o no ventilado).'
)
bloque_salida(
    'Temperatura real de celda por mes. Eficiencia real mensual. '
    'Factor de pérdida por temperatura mensual.'
)
doc.add_paragraph('')

p_nota = doc.add_paragraph()
rn = p_nota.add_run(
    '⚠ DIFERENCIA CLAVE BIPV vs FV estándar: '
)
rn.bold = True
rn.font.color.rgb = COLOR_DECISION
p_nota.add_run(
    'Los módulos BIPV integrados en fachadas SIN ventilación trasera alcanzan '
    'temperaturas 10–20°C más altas que un panel convencional con aire circulando. '
    'Esto reduce significativamente la eficiencia y debe modelarse por separado según '
    'el tipo de montaje.'
)
doc.add_paragraph('')

bloque_formula(
    'T_celda — montaje ventilado (fachada con cámara de aire)',
    'T_celda = T_amb + ((NOCT - 20) / 800) × Irradiancia_W_m2'
)
bloque_formula(
    'T_celda — montaje NO ventilado (vidrio BIPV pegado a estructura)',
    'T_celda = T_amb + ((NOCT + 10 - 20) / 800) × Irradiancia_W_m2  [+10°C de penalidad]'
)
doc.add_paragraph('')

bloque_codigo(
    '# ── MÓDULO 4: Temperatura real de operación BIPV ──\n\n'
    '# Temperatura ambiente media mensual del lugar (°C)\n'
    'temp_amb_mensual = [20,21,19,15,12,9,8,10,13,16,18,20]\n\n'
    '# Irradiancia de referencia para NOCT\n'
    'IRRADIANCIA_REF = 800   # W/m²\n\n'
    '# Penalidad de temperatura según tipo de montaje BIPV\n'
    'penalidad_noct = {\n'
    '    "fachada_ventilada": 0,    # cámara de aire → ventilación normal\n'
    '    "cubierta":          5,    # algo de acumulación de calor\n'
    '    "pergola":           0,    # ventilado por ambas caras\n'
    '    "vidrio_pegado":    15,    # sin ventilación → máxima penalidad\n'
    '}\n\n'
    'NOCT = modulo["NOCT"]\n\n'
    'for sup in superficies_bipv:\n'
    '    tipo = sup["tipo"]\n'
    '    pen  = penalidad_noct.get(tipo, 5)\n'
    '    irr_res = resultados_irradiancia\n'
    '    hsp_mes = next(r["hsp_mensual"] for r in irr_res if r["id"] == sup["id"])\n\n'
    '    t_celda_mes = []\n'
    '    efic_real_mes = []\n'
    '    for mes_i, (t_amb, hsp) in enumerate(zip(temp_amb_mensual, hsp_mes)):\n'
    '        irr_w_m2 = hsp * 1000 / 8   # W/m² promedio del día (aproximación)\n'
    '        T_cel = t_amb + ((NOCT + pen - 20) / IRRADIANCIA_REF) * irr_w_m2\n'
    '        efic  = modulo["efic_nom"] * (1 + modulo["coef_temp"] * (T_cel - 25))\n'
    '        t_celda_mes.append(round(T_cel, 1))\n'
    '        efic_real_mes.append(round(efic, 4))\n\n'
    '    sup["t_celda_mensual"]  = t_celda_mes\n'
    '    sup["efic_real_mensual"] = efic_real_mes\n'
    '    t_prom = round(sum(t_celda_mes) / 12, 1)\n'
    '    print(f"{sup[\'id\']:12s} | T_celda prom={t_prom}°C | penalidad={pen}°C")'
)

flecha()

# ══════════════════════════════════════════════════════════════
# MÓDULO 5 — ENERGÍA GENERADA
# ══════════════════════════════════════════════════════════════
separador()
titulo_modulo('MÓDULO 5 — Cálculo de energía generada por superficie BIPV')

bloque_entrada(
    'Potencia pico de cada superficie (kWp), HSP mensual efectiva, eficiencia real mensual, '
    'factores de pérdida (suciedad, sombras, mismatch, cableado, inversor).'
)
bloque_salida(
    'Energía generada mensual y anual por cada superficie. Energía total del edificio. '
    'Performance Ratio por superficie.'
)
doc.add_paragraph('')

bloque_formula(
    'Energía mensual real por superficie',
    'E_mes (kWh) = Potencia_kWp × HSP_media × Días_mes × PR_sistema'
)
doc.add_paragraph('')

bloque_codigo(
    '# ── MÓDULO 5: Cálculo de energía generada ──\n\n'
    '# Factores de pérdida — Loss Diagram BIPV\n'
    '# BIPV tiene mayor pérdida por temperatura que FV convencional\n'
    'perdidas = {\n'
    '    "suciedad_fachada" : 0.020,  # 2.0% — fachadas acumulan menos polvo que cubiertas\n'
    '    "sombras_propias"  : 0.030,  # 3.0% — aleros, balcones, voladizos\n'
    '    "mismatch"         : 0.020,  # 2.0% — mezcla de orientaciones en mismo inversor\n'
    '    "cableado_dc"      : 0.015,  # 1.5% — cables más largos en edificio\n'
    '    "efic_inversor"    : 0.030,  # 3.0% — pérdida del inversor (97% efic.)\n'
    '}\n\n'
    'dias_por_mes = [31,28,31,30,31,30,31,31,30,31,30,31]\n\n'
    '# PR general del sistema (multiplicar todos los (1-pérdida))\n'
    'pr_sistema = 1.0\n'
    'for p in perdidas.values():\n'
    '    pr_sistema *= (1 - p)\n'
    'pr_sistema = round(pr_sistema, 4)\n'
    'print(f"PR del sistema BIPV: {pr_sistema*100:.1f} %\\n")\n\n'
    'energia_total_kwh_anual = 0\n\n'
    'for sup in superficies_bipv:\n'
    '    irr_data = next(r for r in resultados_irradiancia if r["id"] == sup["id"])\n'
    '    energia_anual = 0\n'
    '    for mes_i in range(12):\n'
    '        hsp    = irr_data["hsp_mensual"][mes_i]\n'
    '        dias   = dias_por_mes[mes_i]\n'
    '        # Pérdida por temperatura ya incluida en efic_real\n'
    '        p_temp = 1 - (sup["efic_real_mensual"][mes_i] / modulo["efic_nom"])\n'
    '        e_mes  = (sup["potencia_kWp"] * hsp * dias\n'
    '                  * (1 - p_temp)\n'
    '                  * (1 - perdidas["suciedad_fachada"])\n'
    '                  * (1 - perdidas["sombras_propias"])\n'
    '                  * (1 - perdidas["mismatch"])\n'
    '                  * (1 - perdidas["cableado_dc"])\n'
    '                  * (1 - perdidas["efic_inversor"]))\n'
    '        energia_anual += e_mes\n'
    '    sup["energia_kwh_anual"] = round(energia_anual, 1)\n'
    '    energia_total_kwh_anual += sup["energia_kwh_anual"]\n'
    '    print(f"{sup[\'id\']:12s} | {sup[\'potencia_kWp\']} kWp | "\n'
    '          f"{sup[\'energia_kwh_anual\']:7.1f} kWh/año")\n\n'
    'energia_total_kwh_anual = round(energia_total_kwh_anual, 1)\n'
    'print(f"\\nEnergía BIPV total del edificio: {energia_total_kwh_anual} kWh/año")'
)

flecha()

# ══════════════════════════════════════════════════════════════
# MÓDULO 6 — ANÁLISIS DE SOMBRAS BIPV
# ══════════════════════════════════════════════════════════════
separador()
titulo_modulo('MÓDULO 6 — Análisis de sombras BIPV (específico de edificios)')

bloque_entrada(
    'Geometría del edificio: altura de aleros, profundidad de voladizos, retranqueos, '
    'altura de edificios vecinos, ángulo de horizonte. Azimut e inclinación de cada superficie.'
)
bloque_salida(
    'Factor de sombra mensual por superficie (0 = sin sombra, 1 = sombra total). '
    'Horas de sombra diarias por mes. Pérdida de energía por sombras en %.'
)
doc.add_paragraph('')

p_nota6 = doc.add_paragraph()
rn6 = p_nota6.add_run('⚠ DIFERENCIA CLAVE BIPV: ')
rn6.bold = True
rn6.font.color.rgb = COLOR_DECISION
p_nota6.add_run(
    'En BIPV, los módulos están integrados al edificio y sufren sombras propias '
    'de su arquitectura: aleros, voladizos, balcones y retranqueos. '
    'Este módulo calcula el ángulo de sombra del alero y determina qué meses '
    'la superficie queda parcialmente sombreada.'
)
doc.add_paragraph('')

bloque_formula(
    'Ángulo solar mínimo para evitar sombra de alero',
    'α_min = arctan(profundidad_alero / altura_desde_alero_a_panel)   [en grados]'
)
bloque_formula(
    'Pérdida por sombra de alero',
    'Factor_sombra = horas_sombreadas / horas_sol_total   → pérdida ≈ factor × 0.5'
)
doc.add_paragraph('')

bloque_codigo(
    '# ── MÓDULO 6: Análisis de sombras BIPV ──\n\n'
    '# Geometría de elementos de sombra del edificio\n'
    'sombras_edificio = {\n'
    '    "alero_profundidad_m"  : 0.60,   # vuelo del alero sobre la fachada (m)\n'
    '    "altura_panel_bajo_alero_m": 1.20, # distancia vertical del borde del panel al alero\n'
    '    "altura_edificio_vecino_m" : 12.0, # altura del edificio más próximo\n'
    '    "distancia_edificio_vecino_m": 15.0, # distancia al edificio vecino\n'
    '}\n\n'
    '# Ángulo mínimo del sol para que no haya sombra del alero\n'
    'prof  = sombras_edificio["alero_profundidad_m"]\n'
    'h_pan = sombras_edificio["altura_panel_bajo_alero_m"]\n'
    'angulo_critico_alero = math.degrees(math.atan(prof / h_pan))\n'
    'print(f"Ángulo crítico del alero: {angulo_critico_alero:.1f}°")\n'
    '# → Si el sol está por debajo de este ángulo, el alero sombrea el panel\n\n'
    '# Ángulo de sombra del edificio vecino\n'
    'h_vec = sombras_edificio["altura_edificio_vecino_m"]\n'
    'd_vec = sombras_edificio["distancia_edificio_vecino_m"]\n'
    'angulo_vecino = math.degrees(math.atan(h_vec / d_vec))\n'
    'print(f"Ángulo de sombra edificio vecino: {angulo_vecino:.1f}°")\n\n'
    '# Factor de sombra mensual por superficie (estimación simplificada)\n'
    '# En cálculo avanzado: usar modelo de trayectoria solar hora a hora\n'
    'factor_sombra_mensual = {\n'
    '    "FAC_NORTE": [0.05,0.04,0.03,0.02,0.01,0.01,0.01,0.02,0.03,0.04,0.05,0.06],\n'
    '    "FAC_ESTE":  [0.08,0.07,0.06,0.05,0.04,0.04,0.04,0.05,0.06,0.07,0.08,0.09],\n'
    '    "CUBIERTA":  [0.01,0.01,0.01,0.01,0.01,0.01,0.01,0.01,0.01,0.01,0.01,0.01],\n'
    '    "PERGOLA":   [0.10,0.09,0.08,0.07,0.06,0.06,0.06,0.07,0.08,0.09,0.10,0.11],\n'
    '}\n\n'
    'for sup in superficies_bipv:\n'
    '    fs = factor_sombra_mensual.get(sup["id"], [0.05]*12)\n'
    '    fs_prom = round(sum(fs)/12*100, 1)\n'
    '    sup["factor_sombra_mensual"] = fs\n'
    '    print(f"{sup[\'id\']:12s} | Pérdida media por sombra: {fs_prom}%")'
)

flecha()

# ══════════════════════════════════════════════════════════════
# MÓDULO 7 — DIMENSIONADO ELÉCTRICO
# ══════════════════════════════════════════════════════════════
separador()
titulo_modulo('MÓDULO 7 — Dimensionado eléctrico del sistema BIPV')

bloque_entrada(
    'Potencia total del sistema (kWp), parámetros eléctricos del módulo '
    '(Voc, Vmp, Isc, Imp), rango MPPT del inversor, temperatura mínima y máxima del lugar.'
)
bloque_salida(
    'Número de paneles en serie por string, número de strings en paralelo, '
    'modelo de inversor compatible, ratio DC/AC, esquema de cableado.'
)
doc.add_paragraph('')

bloque_decision(
    '¿El ratio DC/AC está entre 1.10 y 1.30? → SÍ: continuar. NO: ajustar número de módulos o cambiar inversor.'
)
doc.add_paragraph('')

bloque_codigo(
    '# ── MÓDULO 7: Dimensionado eléctrico BIPV ──\n\n'
    '# Parámetros eléctricos del módulo seleccionado (hoja técnica)\n'
    'parametros_electricos = {\n'
    '    "Voc_stc"  : 49.5,   # Tensión circuito abierto en STC (V)\n'
    '    "Vmp_stc"  : 41.2,   # Tensión máxima potencia en STC (V)\n'
    '    "Isc_stc"  : 10.2,   # Corriente cortocircuito en STC (A)\n'
    '    "Imp_stc"  :  9.7,   # Corriente máxima potencia en STC (A)\n'
    '    "coef_v"   : -0.0029, # Coeficiente tensión (%/°C)\n'
    '    "coef_i"   :  0.0005, # Coeficiente corriente (%/°C)\n'
    '}\n\n'
    '# Inversor seleccionado\n'
    'inversor = {\n'
    '    "modelo"    : "Fronius Symo 10.0",\n'
    '    "P_ac_kW"   : 10.0,\n'
    '    "Vdc_max"   : 1000,   # Tensión DC máxima admitida (V)\n'
    '    "Vmppt_min" : 200,    # Tensión mínima del rango MPPT (V)\n'
    '    "Vmppt_max" : 800,    # Tensión máxima del rango MPPT (V)\n'
    '    "Idc_max"   : 27.0,   # Corriente DC máxima admitida (A)\n'
    '    "efic"      : 0.975,  # Eficiencia del inversor\n'
    '}\n\n'
    '# Temperatura mínima y máxima del lugar (para Voc en frío y Vmp en caliente)\n'
    'T_min = -5     # temperatura mínima registrada (°C) — Voc máximo\n'
    'T_max = 70     # temperatura máxima de la celda en verano (°C) — Vmp mínimo\n\n'
    'Voc  = parametros_electricos["Voc_stc"]\n'
    'Vmp  = parametros_electricos["Vmp_stc"]\n'
    'cv   = parametros_electricos["coef_v"]\n\n'
    '# Voc en frío (temperatura mínima) → tensión máxima del string\n'
    'Voc_frio     = Voc * (1 + cv * (T_min - 25))\n'
    '# Vmp en caliente (temperatura máxima) → tensión mínima del string\n'
    'Vmp_caliente = Vmp * (1 + cv * (T_max - 25))\n\n'
    'n_serie_max = int(inversor["Vdc_max"]   / Voc_frio)\n'
    'n_serie_min = math.ceil(inversor["Vmppt_min"] / Vmp_caliente)\n'
    'n_serie_opt = int(inversor["Vmppt_max"] / Vmp)   # dentro del MPPT óptimo\n\n'
    'print(f"Paneles en serie: mín={n_serie_min}  máx={n_serie_max}  óptimo≈{n_serie_opt}")\n\n'
    'n_strings = int(inversor["Idc_max"] / parametros_electricos["Isc_stc"])\n'
    'print(f"Strings en paralelo: {n_strings}")\n\n'
    '# Verificación ratio DC/AC\n'
    'ratio = round(potencia_total_kWp / inversor["P_ac_kW"], 2)\n'
    'print(f"Ratio DC/AC: {ratio}")\n'
    'if 1.10 <= ratio <= 1.30:\n'
    '    print("✓ Ratio DC/AC dentro del rango óptimo (1.10–1.30)")\n'
    'else:\n'
    '    print("✗ Ajustar número de módulos o cambiar inversor")'
)

flecha()

# ══════════════════════════════════════════════════════════════
# MÓDULO 8 — BALANCE ENERGÉTICO
# ══════════════════════════════════════════════════════════════
separador()
titulo_modulo('MÓDULO 8 — Balance energético del edificio')

bloque_entrada(
    'Energía generada por el sistema BIPV (kWh/mes), perfil de consumo eléctrico '
    'del edificio por mes (kWh/mes), tarifa eléctrica, modalidad (autoconsumo, inyección a red).'
)
bloque_salida(
    'Autoconsumo mensual (%), excedente inyectado a la red, energía comprada a la red, '
    'ahorro efectivo mensual y anual.'
)
doc.add_parameter = None
doc.add_paragraph('')

bloque_decision(
    'Si generación > consumo → excedente se inyecta a red (o se almacena en batería). '
    'Si generación < consumo → déficit se cubre desde la red eléctrica.'
)
doc.add_paragraph('')

bloque_codigo(
    '# ── MÓDULO 8: Balance energético del edificio ──\n\n'
    '# Perfil de consumo eléctrico del edificio (kWh/mes)\n'
    'consumo_edificio_kwh = [1800,1650,1700,1750,1900,2100,\n'
    '                         2050,2000,1850,1780,1700,1820]\n\n'
    '# Distribución mensual de la generación total BIPV\n'
    '# (prorrateamos la energía anual con pesos por mes)\n'
    'pesos_mes = [0.09,0.08,0.09,0.08,0.07,0.07,0.08,0.08,0.08,0.09,0.09,0.10]\n'
    'gen_mensual = [round(energia_total_kwh_anual * w, 0) for w in pesos_mes]\n\n'
    'tarifa_kwh    = 120      # $/kWh precio de compra a la red\n'
    'precio_inyec  = 60       # $/kWh precio de venta de excedente a la red\n\n'
    'resumen_balance = []\n'
    'ahorro_total_anual = 0\n\n'
    'meses_nombres = ["Ene","Feb","Mar","Abr","May","Jun",\n'
    '                 "Jul","Ago","Sep","Oct","Nov","Dic"]\n\n'
    'for i in range(12):\n'
    '    gen  = gen_mensual[i]\n'
    '    cons = consumo_edificio_kwh[i]\n\n'
    '    if gen >= cons:\n'
    '        autoconsumo  = cons\n'
    '        excedente    = gen - cons\n'
    '        compra_red   = 0\n'
    '        ahorro_mes   = autoconsumo * tarifa_kwh + excedente * precio_inyec\n'
    '    else:\n'
    '        autoconsumo  = gen\n'
    '        excedente    = 0\n'
    '        compra_red   = cons - gen\n'
    '        ahorro_mes   = autoconsumo * tarifa_kwh\n\n'
    '    tasa_auto = round(autoconsumo / cons * 100, 1)\n'
    '    ahorro_total_anual += ahorro_mes\n'
    '    resumen_balance.append({\n'
    '        "mes": meses_nombres[i], "generacion": gen, "consumo": cons,\n'
    '        "autoconsumo_%": tasa_auto, "excedente": excedente,\n'
    '        "compra_red": compra_red, "ahorro_mes": round(ahorro_mes, 0)\n'
    '    })\n'
    '    print(f"{meses_nombres[i]}: Gen={gen:.0f} | Cons={cons} | "\n'
    '          f"Auto={tasa_auto}% | Ahorro=${ahorro_mes:,.0f}")\n\n'
    'ahorro_total_anual = round(ahorro_total_anual, 0)\n'
    'print(f"\\nAhorro total anual: ${ahorro_total_anual:,.0f}")'
)

flecha()

# ══════════════════════════════════════════════════════════════
# MÓDULO 9 — ANÁLISIS FINANCIERO Y AMBIENTAL
# ══════════════════════════════════════════════════════════════
separador()
titulo_modulo('MÓDULO 9 — Análisis financiero y ambiental')

bloque_entrada(
    'Costo total del sistema BIPV instalado (incluye módulos, integración arquitectónica, '
    'inversor, instalación eléctrica y permisos), ahorro anual, tasa de descuento, '
    'horizonte de evaluación (años), factor de emisión CO₂.'
)
bloque_salida(
    'Payback simple, VAN, TIR, LCOE, CO₂ evitado en 25 años, '
    'clasificación energética del edificio (A/B/C).'
)
doc.add_paragraph('')

bloque_codigo(
    '# ── MÓDULO 9: Análisis financiero y ambiental ──\n\n'
    '# Costos BIPV (son más altos que FV convencional por la integración arquitectónica)\n'
    'costos_bipv = {\n'
    '    "modulos_bipv"      : potencia_total_kWp * 1200000, # $/kWp — módulos premium\n'
    '    "integracion_arq"   : potencia_total_kWp *  300000, # $/kWp — mano de obra especial\n'
    '    "inversor_sistema"  : potencia_total_kWp *  150000, # $/kWp — inversor + cableado\n'
    '    "ingenieria_permisos": 500000,                      # costo fijo del proyecto\n'
    '}\n'
    'costo_total = sum(costos_bipv.values())\n'
    'costo_om_anual = costo_total * 0.01  # 1% anual de O&M\n\n'
    '# Parámetros financieros\n'
    'tasa_descuento     = 0.07    # 7% tasa de descuento\n'
    'crecimiento_tarifa = 0.04    # 4% alza anual de tarifa eléctrica\n'
    'horizonte_años     = 25\n'
    'factor_co2         = 0.294   # kg CO2/kWh — factor SEN Chile\n\n'
    '# Payback simple\n'
    'payback = round(costo_total / ahorro_total_anual, 1)\n'
    'print(f"Costo total BIPV: ${costo_total:,.0f}")\n'
    'print(f"Payback simple:   {payback} años")\n\n'
    '# VAN — Valor Actual Neto\n'
    'flujos = [ahorro_total_anual*((1+crecimiento_tarifa)**n) - costo_om_anual\n'
    '          for n in range(horizonte_años)]\n'
    'van = -costo_total + sum(f/(1+tasa_descuento)**(n+1) for n,f in enumerate(flujos))\n'
    'resultado_van = "RENTABLE" if van > 0 else "NO RENTABLE"\n'
    'print(f"VAN ({horizonte_años} años): ${van:,.0f} | {resultado_van}")\n\n'
    '# LCOE — Costo nivelado de energía\n'
    'energia_total_25 = sum(energia_total_kwh_anual*((1-0.005)**a) for a in range(25))\n'
    'lcoe = (costo_total + costo_om_anual*25) / energia_total_25\n'
    'print(f"LCOE: ${lcoe:.1f}/kWh")\n\n'
    '# TIR — Tasa Interna de Retorno (bisección)\n'
    'flujos_tir = [-costo_total] + flujos\n'
    'ta, tb = 0.0, 1.0\n'
    'for _ in range(200):\n'
    '    tm  = (ta + tb) / 2\n'
    '    van_m = sum(f/(1+tm)**n for n,f in enumerate(flujos_tir))\n'
    '    if van_m > 0: ta = tm\n'
    '    else:         tb = tm\n'
    'tir = round(tm * 100, 2)\n'
    'print(f"TIR: {tir}%")\n\n'
    '# CO₂ evitado en 25 años\n'
    'co2_ton = round(energia_total_25 * factor_co2 / 1000, 1)\n'
    'print(f"CO2 evitado en 25 años: {co2_ton} toneladas")\n\n'
    '# Clasificación energética básica del edificio (autoconsumo medio anual)\n'
    'autoconsumo_prom = round(\n'
    '    sum(b["autoconsumo_%"] for b in resumen_balance) / 12, 1)\n'
    'if   autoconsumo_prom >= 80: clase = "A+ (energía casi neta cero)"\n'
    'elif autoconsumo_prom >= 60: clase = "A  (alto autoconsumo)"\n'
    'elif autoconsumo_prom >= 40: clase = "B  (autoconsumo moderado)"\n'
    'else:                         clase = "C  (baja integración)"\n'
    'print(f"Clasificación BIPV del edificio: {clase}")'
)

flecha()

# ══════════════════════════════════════════════════════════════
# MÓDULO 10 — GRÁFICOS E INFORME
# ══════════════════════════════════════════════════════════════
separador()
titulo_modulo('MÓDULO 10 — Generación de gráficos e informe final')

bloque_entrada(
    'Todos los resultados calculados en los módulos 1–9.'
)
bloque_salida(
    'Archivos PNG con 6 gráficos. Archivo TXT o CSV con resumen de resultados. '
    'Opcionalmente: archivo Word o PDF con informe ejecutivo completo.'
)
doc.add_paragraph('')

bloque_codigo(
    '# ── MÓDULO 10: Gráficos e informe final ──\n\n'
    'import matplotlib\n'
    'matplotlib.use("Agg")    # sin ventana emergente\n'
    'import matplotlib.pyplot as plt\n'
    'import numpy as np\n\n'
    'plt.close("all")   # limpiar lienzos residuales antes de graficar\n\n'
    '# ─── GRÁFICO 1: Potencia instalada por superficie ───\n'
    'nombres_sup = [s["id"] for s in superficies_bipv]\n'
    'potencias   = [s["potencia_kWp"] for s in superficies_bipv]\n'
    'plt.figure(figsize=(8, 4))\n'
    'plt.bar(nombres_sup, potencias, color=["#2980b9","#e67e22","#27ae60","#8e44ad"])\n'
    'plt.title("Potencia BIPV instalada por superficie (kWp)")\n'
    'plt.ylabel("kWp")\n'
    'plt.grid(axis="y", linestyle="--", alpha=0.5)\n'
    'plt.savefig("01_potencia_por_superficie.png", dpi=150, bbox_inches="tight")\n'
    'plt.close()\n\n'
    '# ─── GRÁFICO 2: Balance mensual generación vs consumo ───\n'
    'x = np.arange(12)\n'
    'meses_n = [b["mes"] for b in resumen_balance]\n'
    'gen_g   = [b["generacion"] for b in resumen_balance]\n'
    'cons_g  = [b["consumo"]    for b in resumen_balance]\n'
    'plt.figure(figsize=(12, 5))\n'
    'plt.bar(x-0.2, gen_g,  0.4, label="Generación BIPV", color="orange")\n'
    'plt.bar(x+0.2, cons_g, 0.4, label="Consumo edificio",color="steelblue")\n'
    'plt.xticks(x, meses_n)\n'
    'plt.title("Balance energético mensual BIPV — Generación vs Consumo")\n'
    'plt.ylabel("kWh/mes")\n'
    'plt.legend()\n'
    'plt.grid(axis="y", linestyle="--", alpha=0.5)\n'
    'plt.savefig("02_balance_mensual.png", dpi=150, bbox_inches="tight")\n'
    'plt.close()\n\n'
    '# ─── GRÁFICO 3: Loss Diagram ───\n'
    'etiq_perd = ["Energía útil","Temperatura","Suciedad","Sombras","Mismatch","Cableado","Inversor"]\n'
    'vals_perd = [pr_sistema*100,\n'
    '             perdidas["suciedad_fachada"]*100,\n'
    '             perdidas["sombras_propias"]*100,\n'
    '             perdidas["mismatch"]*100,\n'
    '             perdidas["cableado_dc"]*100,\n'
    '             perdidas["efic_inversor"]*100,\n'
    '             100-(pr_sistema*100+sum(p*100 for p in perdidas.values()))]\n'
    '# Normalizar para que sume 100\n'
    'tot = sum(abs(v) for v in vals_perd)\n'
    'vals_norm = [abs(v)/tot*100 for v in vals_perd]\n'
    'plt.figure(figsize=(8, 8))\n'
    'plt.pie(vals_norm, labels=etiq_perd, autopct="%1.1f%%", startangle=140,\n'
    '        colors=["#2ecc71","#e74c3c","#e67e22","#f39c12","#9b59b6","#3498db","#95a5a6"])\n'
    'plt.title("Loss Diagram — Distribución de pérdidas BIPV")\n'
    'plt.savefig("03_loss_diagram.png", dpi=150, bbox_inches="tight")\n'
    'plt.close()\n\n'
    '# ─── GRÁFICO 4: Producción con degradación 25 años ───\n'
    'años_g  = list(range(1, 26))\n'
    'prod_g  = [round(energia_total_kwh_anual*(0.995**(a-1)),0) for a in años_g]\n'
    'plt.figure(figsize=(10, 5))\n'
    'plt.plot(años_g, prod_g, marker="o", color="steelblue", linewidth=2)\n'
    'plt.fill_between(años_g, prod_g, alpha=0.15, color="steelblue")\n'
    'plt.title("Producción BIPV anual con degradación — 25 años")\n'
    'plt.xlabel("Año")\n'
    'plt.ylabel("kWh/año")\n'
    'plt.grid(True, linestyle="--", alpha=0.5)\n'
    'plt.savefig("04_degradacion_25_años.png", dpi=150, bbox_inches="tight")\n'
    'plt.close()\n\n'
    '# ─── GRÁFICO 5: Flujo de caja acumulado (Payback) ───\n'
    'flujo_acum = [-costo_total]\n'
    'for n in range(25):\n'
    '    flujo_acum.append(flujo_acum[-1] + flujos[n])\n'
    'colores_pb = ["#e74c3c" if v<0 else "#2ecc71" for v in flujo_acum]\n'
    'plt.figure(figsize=(11, 5))\n'
    'plt.bar(range(26), flujo_acum, color=colores_pb, edgecolor="gray", linewidth=0.4)\n'
    'plt.axhline(0, color="black", linewidth=1.2, linestyle="--")\n'
    'plt.title("Flujo de caja acumulado — Punto de Payback BIPV")\n'
    'plt.xlabel("Año")\n'
    'plt.ylabel("$")\n'
    'plt.grid(axis="y", linestyle="--", alpha=0.4)\n'
    'plt.savefig("05_payback_flujo.png", dpi=150, bbox_inches="tight")\n'
    'plt.close()\n\n'
    '# ─── GRÁFICO 6: HSP por orientación ───\n'
    'meses_n2 = ["E","F","M","A","M","J","J","A","S","O","N","D"]\n'
    'colores_hsp = ["#2980b9","#e67e22","#27ae60","#8e44ad"]\n'
    'plt.figure(figsize=(11, 5))\n'
    'for idx, sup in enumerate(superficies_bipv):\n'
    '    irr_d = next(r for r in resultados_irradiancia if r["id"]==sup["id"])\n'
    '    plt.plot(meses_n2, irr_d["hsp_mensual"],\n'
    '             marker="o", label=sup["id"], color=colores_hsp[idx], linewidth=2)\n'
    'plt.title("HSP mensual efectiva por superficie BIPV")\n'
    'plt.xlabel("Mes")\n'
    'plt.ylabel("HSP (h/día)")\n'
    'plt.legend()\n'
    'plt.grid(True, linestyle="--", alpha=0.5)\n'
    'plt.savefig("06_hsp_por_superficie.png", dpi=150, bbox_inches="tight")\n'
    'plt.close()\n\n'
    '# ─── INFORME RESUMEN EN CONSOLA ───\n'
    'print("\\n" + "="*60)\n'
    'print("     RESUMEN EJECUTIVO — SISTEMA BIPV")\n'
    'print("="*60)\n'
    'print(f"Potencia total instalada : {potencia_total_kWp:.2f} kWp")\n'
    'print(f"Energía generada año 1   : {energia_total_kwh_anual:,.0f} kWh/año")\n'
    'print(f"Performance Ratio (PR)   : {pr_sistema*100:.1f} %")\n'
    'print(f"Autoconsumo promedio     : {autoconsumo_prom} %")\n'
    'print(f"Ahorro anual             : ${ahorro_total_anual:,.0f}")\n'
    'print(f"Costo total sistema      : ${costo_total:,.0f}")\n'
    'print(f"Payback                  : {payback} años")\n'
    'print(f"VAN ({horizonte_años} años)           : ${van:,.0f}")\n'
    'print(f"TIR                      : {tir} %")\n'
    'print(f"LCOE                     : ${lcoe:.1f}/kWh")\n'
    'print(f"CO2 evitado (25 años)    : {co2_ton} ton")\n'
    'print(f"Clasificación edificio   : {clase}")\n'
    'print("="*60)\n'
    'print("Gráficos generados: 6 archivos PNG en la carpeta del proyecto")'
)

# ══════════════════════════════════════════════════════════════
# TABLA RESUMEN
# ══════════════════════════════════════════════════════════════
separador()
doc.add_paragraph('')
h_tabla = doc.add_heading('TABLA RESUMEN — 10 MÓDULOS DEL SOFTWARE BIPV', level=1)
h_tabla.runs[0].font.size  = Pt(13)
h_tabla.runs[0].font.color.rgb = COLOR_MODULO

filas = [
    ('1', 'Edificio y localización',    'Latitud, longitud, azimut, tilt, área, tipo BIPV',
     'Diccionario de superficies'),
    ('2', 'Recurso solar',              'GHI mensual, factores Kt por orientación',
     'HSP efectiva por superficie'),
    ('3', 'Módulo BIPV',                'Tecnología, eficiencia, NOCT, coeficientes',
     'Potencia kWp por superficie'),
    ('4', 'Temperatura de operación',   'NOCT, T_amb mensual, tipo de montaje',
     'T_celda y eficiencia real mensual'),
    ('5', 'Energía generada',           'kWp, HSP, pérdidas, PR',
     'kWh/mes y kWh/año por superficie'),
    ('6', 'Análisis de sombras',        'Aleros, voladizos, edificios vecinos',
     'Factor de sombra mensual (%)'),
    ('7', 'Dimensionado eléctrico',     'Voc, Vmp, Isc, rango MPPT inversor',
     'N° paneles serie, strings, ratio DC/AC'),
    ('8', 'Balance energético',         'Generación vs consumo del edificio',
     'Autoconsumo, excedente, compra red'),
    ('9', 'Análisis financiero',        'Costo, tarifa, tasa descuento',
     'Payback, VAN, TIR, LCOE, CO₂'),
    ('10','Gráficos e informe',         'Todos los resultados anteriores',
     '6 gráficos PNG + resumen ejecutivo'),
]

tabla = doc.add_table(rows=1, cols=4)
tabla.style = 'Table Grid'
hdr = tabla.rows[0].cells
hdr[0].text = 'Módulo'
hdr[1].text = 'Nombre'
hdr[2].text = 'Datos de entrada principales'
hdr[3].text = 'Resultado'
for c in hdr:
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    c._tc.get_or_add_tcPr()
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE
    shd = _OE('w:shd')
    shd.set(_qn('w:val'),  'clear')
    shd.set(_qn('w:color'),'auto')
    shd.set(_qn('w:fill'), '1A5C8A')
    c._tc.get_or_add_tcPr().append(shd)

for fila in filas:
    row = tabla.add_row().cells
    for i, val in enumerate(fila):
        row[i].text = val
        row[i].paragraphs[0].runs[0].font.size = Pt(9)

doc.add_paragraph('')

# Nota final
nota_final = doc.add_paragraph(
    'NOTA DE IMPLEMENTACIÓN: Este algoritmo está diseñado para ejecutarse módulo a módulo '
    'en un único archivo Python, en el orden indicado. Cada módulo usa variables definidas '
    'en el anterior, por lo que el orden de ejecución es obligatorio. '
    'Para proyectos más complejos, cada módulo puede convertirse en una función (def) '
    'independiente que recibe sus datos de entrada como parámetros y retorna sus resultados.'
)
nota_final.runs[0].font.size = Pt(9)
nota_final.runs[0].italic = True
nota_final.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph('')

# Pie de página
footer = doc.add_paragraph(
    'Algoritmo BIPV — Python aplicado a Fotovoltaica Integrada en Edificios — 2026'
)
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.runs[0].font.size = Pt(9)
footer.runs[0].font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
footer.runs[0].italic = True

doc.save('Algoritmo_paso_a_paso_BIPV.docx')
print("Documento BIPV creado correctamente.")
