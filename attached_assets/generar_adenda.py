#!/usr/bin/env python3
"""Genera la Adenda Complemento FASE 9 en formato Word (.docx)"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Estilos generales ──────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    # fondo gris claro
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    p._p.get_or_add_pPr().append(shd)
    return p

def check(text):
    p = doc.add_paragraph()
    run = p.add_run('✅  ')
    run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
    run.bold = True
    p.add_run(text)
    return p

def warn(text):
    p = doc.add_paragraph()
    run = p.add_run('⚠️  ')
    p.add_run(text)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2E74B5')
        hdr[i].paragraphs[0]._p.get_or_add_pPr().append(shd)
        for run in hdr[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
    doc.add_paragraph()
    return table

# ══════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('ADENDA DE DESPLIEGUE — COMPLEMENTO')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run('Calculadora BIPV Python → Digital Ocean\n'
            'Post-Despliegue: Corrección Crítica Página 7 y Sincronización GitHub')

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = info.add_run('Mauricio Acevedo — Ingeniería BIPV Colombia\nFecha: 27 de julio de 2026')
run2.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# FASE 9
# ══════════════════════════════════════════════════════════════════════
h1('FASE 9 — Corrección de Error en Producción: Página 7 💰 Financiero')

h2('Contexto')
doc.add_paragraph(
    'Tras el despliegue exitoso descrito en la Adenda principal, se detectó un NameError '
    'en la Página 7 (Análisis Financiero — Ley 1715 de 2014) al intentar acceder a ella '
    'desde la app en producción.'
)

h2('Error detectado')
code_block('NameError: name \'tipo_cambio\' is not defined')
doc.add_paragraph(
    'Causa raíz: La variable tipo_cambio se utilizaba en la sección de cálculo CAPEX '
    '(línea ~120) pero se definía más abajo en la Sección 2 del formulario. '
    'Al ejecutarse el flujo lineal de Streamlit, Python intentaba usar la variable '
    'antes de que existiera en el scope.'
)

# ── 9.1 ───────────────────────────────────────────────────────────────
h3('PASO 9.1 — Diagnóstico del archivo en el servidor')
doc.add_paragraph('Verificación de la posición de tipo_cambio en el archivo en producción:')
code_block(
    '(venv) root@bipv-colombia:/var/www/bipv/calculadora-bipv#\n'
    'grep -n "tipo_cambio = float" bipv_python/pages/7_💰_Financiero.py'
)
doc.add_paragraph(
    'Resultado: tipo_cambio aparecía definida en la sección 2 del formulario (línea ~120) '
    'pero era referenciada antes en el bloque CAPEX → NameError confirmado.'
)

# ── 9.2 ───────────────────────────────────────────────────────────────
h3('PASO 9.2 — Aplicación del fix con nano')
doc.add_paragraph('Se abrió el archivo con el editor nano en el servidor:')
code_block(
    '(venv) root@bipv-colombia:/var/www/bipv/calculadora-bipv#\n'
    'nano bipv_python/pages/7_💰_Financiero.py'
)
doc.add_paragraph('Fix aplicado — líneas 44-45 añadidas ANTES de cualquier referencia a tipo_cambio:')

doc.add_paragraph('ANTES (estructura incorrecta):')
code_block(
    '42:     n_pan = st.number_input("Número de módulos", ...)\n'
    '43:\n'
    '        # ← tipo_cambio no existía aún; CAPEX la usaba y lanzaba NameError'
)

doc.add_paragraph('DESPUÉS (estructura corregida):')
code_block(
    '42:     n_pan = st.number_input("Número de módulos", ...)\n'
    '43:\n'
    '44: # TRM disponible desde el inicio (se actualiza en Sección 2)\n'
    '45: tipo_cambio = float(st.session_state.get("tipo_cambio", 3400.0))'
)

doc.add_paragraph('Guardado y cierre en nano:')
code_block('Ctrl + O  →  [nano: "Save modified buffer?"]  →  Y  →  Enter  →  Ctrl + X')

doc.add_paragraph('Verificación inmediata del fix:')
code_block(
    '(venv) root@bipv-colombia:/var/www/bipv/calculadora-bipv#\n'
    'grep -n "tipo_cambio = float" bipv_python/pages/7_💰_Financiero.py\n\n'
    '45:tipo_cambio = float(st.session_state.get("tipo_cambio", 3400.0))'
)
check('Fix confirmado en línea 45 del archivo en producción.')

# ── 9.3 ───────────────────────────────────────────────────────────────
h3('PASO 9.3 — Commit del fix en Git local')
code_block(
    '(venv) root@bipv-colombia:/var/www/bipv/calculadora-bipv#\n'
    'git add bipv_python/pages/7_💰_Financiero.py\n\n'
    'git commit -m "fix: tipo_cambio definido al inicio antes de CAPEX"'
)
doc.add_paragraph('Resultado:')
code_block(
    '[main 7dca2bf4] fix: tipo_cambio definido al inicio antes de CAPEX\n'
    ' 1 file changed, 1 insertion(+)'
)
check('Commit 7dca2bf4 creado exitosamente.')

# ── 9.4 ───────────────────────────────────────────────────────────────
h3('PASO 9.4 — Fallo en git push (token vencido) y diagnóstico')
code_block(
    '(venv) root@bipv-colombia:/var/www/bipv/calculadora-bipv#\n'
    'git push origin main'
)
doc.add_paragraph('ERROR obtenido:')
code_block(
    'remote: Invalid username or token.\n'
    'Password authentication is not supported for Git operations.\n'
    'fatal: Authentication failed for\n'
    "  'https://github.com/ventas108/calculadora-bipv.git/'"
)
doc.add_paragraph(
    'Causa: El Personal Access Token (PAT) de GitHub embebido en la URL del remote había '
    'vencido. GitHub eliminó la autenticación por contraseña en agosto 2021; '
    'requiere PAT o clave SSH.'
)

# ── 9.5 ───────────────────────────────────────────────────────────────
h3('PASO 9.5 — Generación de nuevo PAT en GitHub')
doc.add_paragraph('Procedimiento en el navegador:')
steps = [
    'Ir a: https://github.com/settings/tokens/new',
    'Note: servidor-bipv',
    'Expiration: 90 days',
    'Scope: ✅ repo  (acceso completo a repositorios privados)',
    'Clic "Generate token" → copiar el token generado.',
]
for s in steps:
    p = doc.add_paragraph(s, style='List Number')
warn('El token solo se muestra una vez al crearse. Guardarlo antes de cerrar la página.')

# ── 9.6 ───────────────────────────────────────────────────────────────
h3('PASO 9.6 — Actualización de la URL del remote y push exitoso')
doc.add_paragraph('Actualizar la URL del remote con el nuevo token:')
code_block(
    '(venv) root@bipv-colombia:/var/www/bipv/calculadora-bipv#\n'
    'git remote set-url origin \\\n'
    '  https://<NUEVO_TOKEN>@github.com/ventas108/calculadora-bipv.git\n\n'
    'git push origin main'
)
doc.add_paragraph('Resultado:')
code_block(
    'Enumerating objects: 14, done.\n'
    'Counting objects: 100% (14/14), done.\n'
    'Delta compression using up to 2 threads\n'
    'Compressing objects: 100% (10/10), done.\n'
    'Writing objects: 100% (10/10), 963 bytes | 963.00 KiB/s, done.\n'
    'Total 10 (delta 8), reused 0 (delta 0), pack-reused 0\n'
    'remote: Resolving deltas: 100% (8/8), completed with 4 local objects.\n'
    'To https://github.com/ventas108/calculadora-bipv.git\n'
    '   77112e3f..7dca2bf4  main -> main'
)
check('Push exitoso — rama main actualizada: 77112e3f → 7dca2bf4')

doc.add_paragraph('Verificación del historial de commits:')
code_block(
    'git log --oneline -3\n\n'
    '7dca2bf4 (HEAD -> main, origin/main, origin/HEAD)\n'
    '            fix: tipo_cambio definido al inicio antes de CAPEX\n'
    '289ac5c3    fix: tipo_cambio definido al inicio antes de CAPEX\n'
    '77112e3f    feat: Financiero — TRM default 3400, panel conversor USD→COP'
)
check('Repositorio GitHub sincronizado con el estado del servidor.')

# ── 9.7 ───────────────────────────────────────────────────────────────
h3('PASO 9.7 — Reinicio de PM2 para aplicar el fix')
code_block(
    '(venv) root@bipv-colombia:/var/www/bipv/calculadora-bipv#\n'
    'pm2 restart bipv-streamlit'
)
doc.add_paragraph('Resultado:')
code_block(
    'Use --update-env to update environment variables\n'
    '[PM2] Applying action restartProcessId on app [bipv-streamlit](ids: [1])\n'
    '[PM2] [bipv-streamlit](1) ✓\n\n'
    '┌────┬────────────────────┬──────────┬──────┬───────────┬──────────┐\n'
    '│ id │ name               │ mode     │  ↺   │ status    │ memory   │\n'
    '├────┼────────────────────┼──────────┼──────┼───────────┼──────────┤\n'
    '│  1 │ bipv-streamlit     │ fork     │  10  │ online    │ 6.4mb    │\n'
    '│  0 │ calculadora-bipv   │ fork     │  60  │ online    │ 183.1mb  │\n'
    '└────┴────────────────────┴──────────┴──────┴───────────┴──────────┘'
)
check('bipv-streamlit online — fix activo en producción.')
check('calculadora-bipv online — app Node.js sin cambios.')

# ══════════════════════════════════════════════════════════════════════
# RESULTADO FASE 9
# ══════════════════════════════════════════════════════════════════════
h2('🏆 RESULTADO FASE 9 — PÁGINA 7 OPERATIVA')
doc.add_paragraph('URL: https://calc.innovacionquimica.com.co  |  Página: 💰 Análisis Financiero — Ley 1715 de 2014')

add_table(
    ['✅ Logro', 'Detalle'],
    [
        ['Fix aplicado', 'tipo_cambio definida en línea 45 (scope global, antes de CAPEX)'],
        ['Commit', '7dca2bf4 en rama main'],
        ['GitHub sincronizado', '77112e3f → 7dca2bf4 (origin/main)'],
        ['PM2 reiniciado', 'bipv-streamlit online (restart #10)'],
        ['Página 7 operativa', 'Sin NameError; CAPEX, Art. 11/12/14, TIR, VPN, Payback y LCOE correctos'],
    ]
)

# ══════════════════════════════════════════════════════════════════════
# TABLA MÓDULOS
# ══════════════════════════════════════════════════════════════════════
h2('TABLA DE MÓDULOS — ESTADO ACTUALIZADO (27 jul 2026)')
add_table(
    ['Página', 'Función', 'Estado'],
    [
        ['🏠 Proyecto',          'Ciudad, área fachada, datos de sitio',   '✅ Funcional'],
        ['☀️ Recurso Solar',      'TMY + PVGIS + irradiancia POA',          '🔲 Pendiente'],
        ['🔬 Motor IV',          'Curva I-V + validación SDM',             '✅ Funcional'],
        ['📐 Dimensionamiento',  'String sizing con semáforo',             '✅ Funcional'],
        ['⚡ Mismatch',          'Análisis desajuste MPPT',                '🔲 Pendiente'],
        ['📊 Producción',        'Simulación IEC 61724 hora a hora',       '🔲 Pendiente'],
        ['💰 Financiero',        'VPN, TIR, Payback, Ley 1715/2014',       '✅ Funcional (corregida)'],
    ]
)

# ══════════════════════════════════════════════════════════════════════
# NOTA SEGURIDAD
# ══════════════════════════════════════════════════════════════════════
h2('NOTA DE SEGURIDAD — Buenas Prácticas con Tokens GitHub')
doc.add_paragraph(
    'Durante operaciones de emergencia en producción, los tokens pueden quedar expuestos '
    'en el historial de terminal. Medidas tras cada sesión:'
)
for s in [
    'Ir a https://github.com/settings/tokens',
    'Revocar el token utilizado en la sesión.',
    'Generar un nuevo token solo cuando sea necesario.',
]:
    doc.add_paragraph(s, style='List Number')

doc.add_paragraph('Alternativa recomendada a largo plazo — autenticación SSH (sin tokens):')
code_block(
    '# En el servidor Digital Ocean:\n'
    'ssh-keygen -t ed25519 -C "bipv-colombia-server"\n'
    'cat ~/.ssh/id_ed25519.pub\n'
    '# Agregar la clave en: GitHub → Settings → SSH and GPG keys → New SSH key\n\n'
    '# Cambiar remote a SSH:\n'
    'git remote set-url origin git@github.com:ventas108/calculadora-bipv.git'
)
doc.add_paragraph(
    'Con SSH, git push/pull no requiere token; la autenticación es por clave '
    'criptográfica almacenada en el servidor.'
)

# ── Pie ───────────────────────────────────────────────────────────────
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run(
    'Calculadora BIPV v1.0 — Motor SDM validado contra XLSM auditado '
    '(De Soto 2006). FF_max CdTe = 76.28% @ G=200 W/m²'
)
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

# ── Guardar ───────────────────────────────────────────────────────────
path = '/home/runner/workspace/attached_assets/Adenda_Complemento_FASE9_Financiero.docx'
doc.save(path)
print(f'OK — guardado en {path}')
